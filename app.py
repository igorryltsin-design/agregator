# REST API и маршруты импорта/экспорта перенесены в `routes.py` как Blueprint.
import os
import re
import json
import hashlib
import hmac
import logging
from logging.handlers import RotatingFileHandler
import time
import math
from collections import OrderedDict, defaultdict
from datetime import datetime, timedelta, timezone
from pathlib import Path
import io
import itertools
import heapq
import copy
import queue
import platform
import shutil
import uuid
import xml.etree.ElementTree as ET
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple, Mapping, cast
from urllib.parse import quote_plus

import click
from flask import Blueprint, Flask, request, redirect, url_for, jsonify, send_from_directory, send_file, Response, make_response, session, g, abort, current_app, has_app_context, stream_with_context
from functools import wraps
from werkzeug.utils import secure_filename
import sqlite3
import threading

try:
    from flask import copy_current_app_context
except ImportError:  # запасной вариант для старых версий Flask
    from flask import current_app

    def copy_current_app_context(func):
        app_obj = current_app._get_current_object()

        @wraps(func)
        def wrapper(*args, **kwargs):
            with app_obj.app_context():
                return func(*args, **kwargs)

        return wrapper
from sqlalchemy import func, and_, or_, exists, text, event, inspect as sa_inspect
from sqlalchemy.engine import Engine
from sqlalchemy.engine.url import make_url
from sqlalchemy.orm import aliased, joinedload
from models import (
    db,
    File,
    Tag,
    TagSchema,
    ChangeLog,
    upsert_tag,
    file_to_dict,
    Collection,
    User,
    CollectionMember,
    UserActionLog,
    UserPreference,
    TaskRecord,
    DocChatCache,
    LlmEndpoint,
    AiWordAccess,
    AiSearchSnippetCache,
    AiSearchKeywordFeedback,
    AiSearchFeedbackModel,
    AiSearchMetric,
    RagDocument,
    RagDocumentVersion,
    RagDocumentChunk,
    RagChunkEmbedding,
    RagIngestFailure,
    RagSession,
)

APP_START_TIME = time.time()
APP_STARTED_AT = datetime.utcnow()


@event.listens_for(Engine, "connect")
def _configure_sqlite_pragmas(dbapi_connection, connection_record):
    """Ensure required PRAGMA flags are enabled for SQLite connections."""
    try:
        import sqlite3 as _sqlite3
    except Exception:
        _sqlite3 = None
    if _sqlite3 is not None and isinstance(dbapi_connection, _sqlite3.Connection):
        try:
            import math as _math
            try:
                from agregator.rag.utils import bytes_to_vector as _bytes_to_vector  # type: ignore
            except Exception:
                _bytes_to_vector = None

            def _cosine_similarity_sqlite(blob_a, blob_b):
                if _bytes_to_vector is None or blob_a is None or blob_b is None:
                    return 0.0
                try:
                    vec_a = _bytes_to_vector(blob_a)
                    vec_b = _bytes_to_vector(blob_b)
                except Exception:
                    return 0.0
                if not vec_a or not vec_b or len(vec_a) != len(vec_b):
                    return 0.0
                dot = sum(a * b for a, b in zip(vec_a, vec_b))
                if dot == 0.0:
                    return 0.0
                norm_a = _math.sqrt(sum(a * a for a in vec_a))
                norm_b = _math.sqrt(sum(b * b for b in vec_b))
                if norm_a == 0.0 or norm_b == 0.0:
                    return 0.0
                return dot / (norm_a * norm_b)

            dbapi_connection.create_function("cosine_similarity", 2, _cosine_similarity_sqlite)
        except Exception:
            pass
        cursor = dbapi_connection.cursor()
        try:
            cursor.execute("PRAGMA foreign_keys = 1")
            cursor.execute("PRAGMA trusted_schema = 1")
            cursor.execute("PRAGMA recursive_triggers = 1")
        finally:
            cursor.close()


import fitz  # библиотека PyMuPDF
import requests
from agregator.config import AppConfig, load_app_config
from agregator.config_schema import CONFIG_FIELDS, field_api_key, field_snapshot_key
from agregator.runtime_settings import runtime_settings_store
from agregator.services import (
    HttpSettings,
    configure_http,
    configure_llm_cache,
    configure_search_cache,
    configure_llm_pool,
    configure_logging,
    get_browser_manager,
    get_http_settings,
    get_rotating_log_handler as svc_get_rotating_log_handler,
    get_task_queue,
    llm_cache_get,
    llm_cache_set,
    llm_cache_clear,
    llm_cache_stats,
    CachedLLMResponse,
    LlmPoolRejected,
    LlmPoolTimeout,
    FacetQueryParams,
    FacetService,
    SearchService,
    search_cache_get,
    search_cache_set,
    search_cache_clear,
    search_cache_stats,
    get_llm_pool,
    http_request,
    list_system_log_files as svc_list_system_log_files,
    resolve_log_name as svc_resolve_log_name,
    tail_log_file as svc_tail_log_file,
)
from agregator.osint import (
    OsintSearchService,
    SerpFetcher,
    SerpParser,
    SerpSettings,
    OsintRepositoryConfig,
    SUPPORTED_ENGINES,
    SUPPORTED_OSINT_LOCALES,
)
from agregator.osint.remote_browser import remote_browser_manager
from agregator.osint.proxy import fetch_proxied_page, rewrite_html_for_proxy
from agregator.osint.artifacts import resolve_artifact_path
from agregator.rag import (
    ChunkConfig,
    ContextSelector,
    KeywordRetriever,
    ContextSection,
    ValidationResult,
    RagIndexer,
    VectorRetriever,
    RetrievedChunk,
    EmbeddingBackend,
    HashEmbeddingBackend,
    load_embedding_backend,
    vector_to_bytes,
    build_system_prompt,
    build_user_prompt,
    fallback_answer,
    validate_answer,
    detect_language,
    extract_citations,
)
try:
    import docx
except ImportError:
    docx = None
try:
    from striprtf.striprtf import rtf_to_text
except ImportError:
    rtf_to_text = None
try:
    from ebooklib import epub
except ImportError:
    epub = None
try:
    import djvu.decode
except ImportError:
    djvu = None
try:
    import pytesseract  # необязательный OCR для PDF, состоящих из изображений
except ImportError:
    pytesseract = None
try:
    import mammoth
except ImportError:
    mammoth = None
try:
    from PIL import Image as PILImage
except Exception:
    PILImage = None
try:
    from faster_whisper import WhisperModel as FasterWhisperModel
except Exception:
    FasterWhisperModel = None
try:
    from huggingface_hub import snapshot_download as hf_snapshot_download
except Exception:
    hf_snapshot_download = None
import subprocess, shutil, wave, tempfile
import time
from html import escape

# ------------------- RU morphology (Natasha/pymorphy2) -------------------
try:
    from razdel import tokenize as _ru_tokenize
except Exception:
    _ru_tokenize = None
try:
    import pymorphy2
    _morph = pymorphy2.MorphAnalyzer()
except Exception:
    _morph = None

logger = logging.getLogger(__name__)

PIPELINE_API_KEY = (os.getenv("PIPELINE_API_KEY") or "").strip()


def _extract_pipeline_token() -> Optional[str]:
    header = request.headers.get('Authorization')
    token: Optional[str] = None
    if header:
        lowered = header.lower()
        if lowered.startswith('bearer '):
            token = header[7:].strip()
        else:
            token = header.strip()
    if not token:
        token = request.headers.get('X-Agregator-Key')
        if token:
            token = token.strip()
    if not token:
        token = request.headers.get('X-Agregator-Token')
        if token:
            token = token.strip()
    return token or None


def _check_pipeline_access() -> tuple[bool, Optional[Response]]:
    user = _load_current_user()
    if user and getattr(user, 'role', '') == 'admin':
        return True, None
    if PIPELINE_API_KEY:
        token = _extract_pipeline_token()
        if token and hmac.compare_digest(token, PIPELINE_API_KEY):
            return True, None
        resp = jsonify({'ok': False, 'error': 'Forbidden'})
        resp.status_code = 403
        return False, resp
    return True, None


def _rt():
    """Быстрый доступ к текущим runtime-настройкам."""
    return runtime_settings_store.current


def _lm_max_input_chars() -> int:
    """Return current input character limit for LLM calls."""
    try:
        value = int(getattr(_rt(), 'lm_max_input_chars', 4000) or 4000)
    except Exception:
        value = 4000
    return max(1, value)


def _lm_max_output_tokens() -> int:
    """Return current output token limit for LLM calls."""
    try:
        value = int(getattr(_rt(), 'lm_max_output_tokens', 256) or 256)
    except Exception:
        value = 256
    return max(1, value)


def _lm_safe_context_tokens() -> int:
    """Return the soft limit of tokens we keep within the prompt window."""
    try:
        value = int(getattr(_rt(), 'lm_context_token_limit', 3200) or 3200)
    except Exception:
        value = 3200
    return max(512, value)


def _guess_tokenizer_for_model(model_name: Optional[str], provider: Optional[str]) -> str:
    """Return a human-readable hint for the tokenizer based on model/provider."""
    name = (model_name or '').strip().lower()
    prov = (provider or '').strip().lower()

    if not name and not prov:
        return 'Авто (по модели)'

    def _has(substrings: Iterable[str]) -> bool:
        return any(sub in name for sub in substrings)

    if prov in {'openai', 'azure_openai', 'openrouter'}:
        if _has(('gpt-4o', 'o1', 'o3')):
            return 'tiktoken · o200k_base'
        if _has(('gpt-4.1', 'gpt-4-turbo', 'gpt-4', 'gpt-3.5', 'gpt-4o-mini', 'text-embedding-3')):
            return 'tiktoken · cl100k_base'
        if _has(('text-embedding-ada-002', 'ada-002')):
            return 'tiktoken · cl100k_base'
        return 'tiktoken (авто)'

    if prov == 'ollama':
        if _has(('llama', 'mixtral', 'mistral', 'gemma', 'phi', 'qwen', 'zephyr')):
            return 'SentencePiece / GGUF'
        return 'Встроенный GGUF'

    if _has(('gemma', 'mistral', 'mixtral', 'llama', 'phi', 'qwen', 'zephyr', 'command', 'falcon', 'aya', 'vicuna')):
        return 'SentencePiece / BPE'

    return 'Авто (по модели)'


def _always_ocr_first_page_dissertation() -> bool:
    """Return flag for forcing OCR on the first page of dissertations."""
    try:
        return bool(getattr(_rt(), 'always_ocr_first_page_dissertation', False))
    except Exception:
        return False


def _refresh_runtime_globals() -> None:
    """Expose frequently used runtime flags as module globals for legacy code."""
    runtime = runtime_settings_store.current
    for key, attr in _RUNTIME_ATTR_MAP.items():
        try:
            value = getattr(runtime, attr)
        except Exception:
            value = None
        if key in {'LM_MAX_INPUT_CHARS', 'LM_MAX_OUTPUT_TOKENS'}:
            try:
                value = int(value)
            except Exception:
                value = 0
            if value <= 0:
                value = 4000 if key == 'LM_MAX_INPUT_CHARS' else 256
        if key == 'FEEDBACK_TRAIN_INTERVAL_HOURS':
            try:
                value = max(0.0, float(value or 0.0))
            except Exception:
                value = 0.0
        if key == 'FEEDBACK_TRAIN_CUTOFF_DAYS':
            try:
                value = max(1, int(value or 1))
            except Exception:
                value = 90
        if key in {
            'TRANSCRIBE_ENABLED',
            'SUMMARIZE_AUDIO',
            'AUDIO_KEYWORDS_LLM',
            'IMAGES_VISION_ENABLED',
            'KEYWORDS_TO_TAGS_ENABLED',
            'TYPE_LLM_OVERRIDE',
            'MOVE_ON_RENAME',
            'COLLECTIONS_IN_SEPARATE_DIRS',
            'COLLECTION_TYPE_SUBDIRS',
            'DEFAULT_USE_LLM',
            'DEFAULT_PRUNE',
            'AI_RERANK_LLM',
            'AI_RAG_RETRY_ENABLED',
            'LLM_CACHE_ENABLED',
            'LLM_CACHE_ONLY_MODE',
            'SEARCH_CACHE_ENABLED',
            'ALWAYS_OCR_FIRST_PAGE_DISSERTATION',
            'SEARCH_FACET_INCLUDE_TYPES',
        }:
            value = bool(value)
        if key == 'PROMPTS':
            value = dict(value or {})
        globals()[key] = value


_RUNTIME_ATTR_MAP = {
    'SCAN_ROOT': 'scan_root',
    'EXTRACT_TEXT': 'extract_text',
    'LMSTUDIO_API_BASE': 'lmstudio_api_base',
    'LMSTUDIO_MODEL': 'lmstudio_model',
    'LMSTUDIO_API_KEY': 'lmstudio_api_key',
    'LM_DEFAULT_PROVIDER': 'lm_default_provider',
    'TRANSCRIBE_ENABLED': 'transcribe_enabled',
    'TRANSCRIBE_BACKEND': 'transcribe_backend',
    'TRANSCRIBE_MODEL_PATH': 'transcribe_model_path',
    'TRANSCRIBE_LANGUAGE': 'transcribe_language',
    'SUMMARIZE_AUDIO': 'summarize_audio',
    'AUDIO_KEYWORDS_LLM': 'audio_keywords_llm',
    'IMAGES_VISION_ENABLED': 'images_vision_enabled',
    'KEYWORDS_TO_TAGS_ENABLED': 'keywords_to_tags_enabled',
    'TYPE_DETECT_FLOW': 'type_detect_flow',
    'TYPE_LLM_OVERRIDE': 'type_llm_override',
    'IMPORT_SUBDIR': 'import_subdir',
    'MOVE_ON_RENAME': 'move_on_rename',
    'COLLECTIONS_IN_SEPARATE_DIRS': 'collections_in_separate_dirs',
    'COLLECTION_TYPE_SUBDIRS': 'collection_type_subdirs',
    'TYPE_DIRS': 'type_dirs',
    'DEFAULT_USE_LLM': 'default_use_llm',
    'DEFAULT_PRUNE': 'default_prune',
    'OCR_LANGS_CFG': 'ocr_langs_cfg',
    'PDF_OCR_PAGES_CFG': 'pdf_ocr_pages_cfg',
    'ALWAYS_OCR_FIRST_PAGE_DISSERTATION': 'always_ocr_first_page_dissertation',
    'PROMPTS': 'prompts',
    'AI_RERANK_LLM': 'ai_rerank_llm',
    'AI_RAG_RETRY_ENABLED': 'ai_rag_retry_enabled',
    'AI_RAG_RETRY_THRESHOLD': 'ai_rag_retry_threshold',
    'FEEDBACK_TRAIN_INTERVAL_HOURS': 'feedback_train_interval_hours',
    'FEEDBACK_TRAIN_CUTOFF_DAYS': 'feedback_train_cutoff_days',
    'LLM_CACHE_ENABLED': 'llm_cache_enabled',
    'LLM_CACHE_TTL_SECONDS': 'llm_cache_ttl_seconds',
    'LLM_CACHE_MAX_ITEMS': 'llm_cache_max_items',
    'LLM_CACHE_ONLY_MODE': 'llm_cache_only_mode',
    'SEARCH_CACHE_ENABLED': 'search_cache_enabled',
    'SEARCH_CACHE_TTL_SECONDS': 'search_cache_ttl_seconds',
    'SEARCH_CACHE_MAX_ITEMS': 'search_cache_max_items',
    'LM_MAX_INPUT_CHARS': 'lm_max_input_chars',
    'LM_MAX_OUTPUT_TOKENS': 'lm_max_output_tokens',
    'AZURE_OPENAI_API_VERSION': 'azure_openai_api_version',
    'AI_QUERY_VARIANTS_MAX': 'ai_query_variants_max',
    'SEARCH_FACET_TAG_KEYS': 'search_facet_tag_keys',
    'GRAPH_FACET_TAG_KEYS': 'graph_facet_tag_keys',
    'SEARCH_FACET_INCLUDE_TYPES': 'search_facet_include_types',
}

LM_MAX_INPUT_CHARS = 4000
LM_MAX_OUTPUT_TOKENS = 256
ALWAYS_OCR_FIRST_PAGE_DISSERTATION = False
PROMPTS = {}

for _GLOBAL_NAME in _RUNTIME_ATTR_MAP:
    globals().setdefault(_GLOBAL_NAME, None)


def _scan_root_path() -> Path:
    return _rt().scan_root


def _import_subdir_value() -> str:
    return (_rt().import_subdir or '').strip()


def _collections_in_separate_dirs() -> bool:
    return bool(_rt().collections_in_separate_dirs)


def _ensure_rag_embedding_defaults() -> None:
    runtime = runtime_settings_store.current
    updates: dict[str, object] = {}
    backend_raw = (runtime.rag_embedding_backend or '').strip().lower()
    if backend_raw in ('', 'auto'):
        updates['RAG_EMBEDDING_BACKEND'] = 'lm-studio'
    model_raw = (runtime.rag_embedding_model or '').strip()
    if not model_raw or model_raw == 'intfloat/multilingual-e5-large':
        updates['RAG_EMBEDDING_MODEL'] = 'nomic-ai/nomic-embed-text-v1.5-GGUF'
        if (runtime.rag_embedding_dim or 0) < 256:
            updates['RAG_EMBEDDING_DIM'] = 768
    if not runtime.rag_embedding_endpoint:
        updates['RAG_EMBEDDING_ENDPOINT'] = runtime.lmstudio_api_base or ''
    if runtime.rag_embedding_api_key is None:
        updates['RAG_EMBEDDING_API_KEY'] = ''
    if updates:
        runtime_settings_store.apply_updates(updates)


_RUNTIME_ATTR_MAP = {
    'SCAN_ROOT': 'scan_root',
    'EXTRACT_TEXT': 'extract_text',
    'LMSTUDIO_API_BASE': 'lmstudio_api_base',
    'LMSTUDIO_MODEL': 'lmstudio_model',
    'LMSTUDIO_API_KEY': 'lmstudio_api_key',
    'LM_DEFAULT_PROVIDER': 'lm_default_provider',
    'TRANSCRIBE_ENABLED': 'transcribe_enabled',
    'TRANSCRIBE_BACKEND': 'transcribe_backend',
    'TRANSCRIBE_MODEL_PATH': 'transcribe_model_path',
    'TRANSCRIBE_LANGUAGE': 'transcribe_language',
    'SUMMARIZE_AUDIO': 'summarize_audio',
    'AUDIO_KEYWORDS_LLM': 'audio_keywords_llm',
    'IMAGES_VISION_ENABLED': 'images_vision_enabled',
    'KEYWORDS_TO_TAGS_ENABLED': 'keywords_to_tags_enabled',
    'TYPE_DETECT_FLOW': 'type_detect_flow',
    'TYPE_LLM_OVERRIDE': 'type_llm_override',
    'IMPORT_SUBDIR': 'import_subdir',
    'MOVE_ON_RENAME': 'move_on_rename',
    'COLLECTIONS_IN_SEPARATE_DIRS': 'collections_in_separate_dirs',
    'COLLECTION_TYPE_SUBDIRS': 'collection_type_subdirs',
    'TYPE_DIRS': 'type_dirs',
    'DEFAULT_USE_LLM': 'default_use_llm',
    'DEFAULT_PRUNE': 'default_prune',
    'OCR_LANGS_CFG': 'ocr_langs_cfg',
    'PDF_OCR_PAGES_CFG': 'pdf_ocr_pages_cfg',
    'ALWAYS_OCR_FIRST_PAGE_DISSERTATION': 'always_ocr_first_page_dissertation',
    'PROMPTS': 'prompts',
    'AI_RERANK_LLM': 'ai_rerank_llm',
    'LLM_CACHE_ENABLED': 'llm_cache_enabled',
    'LLM_CACHE_TTL_SECONDS': 'llm_cache_ttl_seconds',
    'LLM_CACHE_MAX_ITEMS': 'llm_cache_max_items',
    'LLM_CACHE_ONLY_MODE': 'llm_cache_only_mode',
    'SEARCH_CACHE_ENABLED': 'search_cache_enabled',
    'SEARCH_CACHE_TTL_SECONDS': 'search_cache_ttl_seconds',
    'SEARCH_CACHE_MAX_ITEMS': 'search_cache_max_items',
    'LM_MAX_INPUT_CHARS': 'lm_max_input_chars',
    'LM_MAX_OUTPUT_TOKENS': 'lm_max_output_tokens',
    'AZURE_OPENAI_API_VERSION': 'azure_openai_api_version',
    'AI_QUERY_VARIANTS_MAX': 'ai_query_variants_max',
    'SEARCH_FACET_TAG_KEYS': 'search_facet_tag_keys',
    'GRAPH_FACET_TAG_KEYS': 'graph_facet_tag_keys',
    'SEARCH_FACET_INCLUDE_TYPES': 'search_facet_include_types',
}


def __getattr__(name: str):
    attr = _RUNTIME_ATTR_MAP.get(name)
    if attr is not None:
        return getattr(runtime_settings_store.current, attr)
    raise AttributeError(f"module 'app' has no attribute {name!r}")


ROLE_ADMIN = 'admin'
ROLE_EDITOR = 'editor'
ROLE_VIEWER = 'viewer'
_ROLE_ALIASES = {
    'user': ROLE_EDITOR,
    'editor': ROLE_EDITOR,
    'viewer': ROLE_VIEWER,
    'admin': ROLE_ADMIN,
}


def _normalize_user_role(value: str | None, *, default: str = ROLE_VIEWER) -> str:
    if not value:
        return default
    token = str(value).strip().lower()
    return _ROLE_ALIASES.get(token, default if token not in _ROLE_ALIASES else token)


def _user_role(user: User | None) -> str:
    if not user:
        return ROLE_VIEWER
    return _normalize_user_role(getattr(user, 'role', None), default=ROLE_EDITOR)

PIPELINE_API_KEY = (os.getenv("PIPELINE_API_KEY") or "").strip()


def _extract_pipeline_token() -> Optional[str]:
    header = request.headers.get('Authorization')
    token: Optional[str] = None
    if header:
        lowered = header.lower()
        if lowered.startswith('bearer '):
            token = header[7:].strip()
        else:
            token = header.strip()
    if not token:
        token = request.headers.get('X-Agregator-Key')
        if token:
            token = token.strip()
    if not token:
        token = request.headers.get('X-Agregator-Token')
        if token:
            token = token.strip()
    return token or None


def _check_pipeline_access() -> tuple[bool, Optional[Response]]:
    user = _load_current_user()
    if _user_role(user) == ROLE_ADMIN:
        return True, None
    if PIPELINE_API_KEY:
        token = _extract_pipeline_token()
        if token and hmac.compare_digest(token, PIPELINE_API_KEY):
            return True, None
        resp = jsonify({'ok': False, 'error': 'Forbidden'})
        resp.status_code = 403
        return False, resp
    return True, None


from agregator.caching import TimedCache  # extracted — single source of truth

FACET_CACHE = TimedCache(max_items=256, ttl=60)
facet_service = FacetService(FACET_CACHE)
search_service = SearchService(
    db=db,
    file_model=File,
    tag_model=Tag,
    cache_get=search_cache_get,
    cache_set=search_cache_set,
    logger=logging.getLogger('agregator.search'),
)

_FACET_CACHE_REBUILD_LOCK = threading.Lock()
_FACET_CACHE_REBUILD_PENDING = False


def _invalidate_facets_cache(reason: str | None = None) -> None:
    try:
        facet_service.invalidate(reason)
    except Exception:
        pass
    try:
        search_service.invalidate_cache(reason or 'facets')
    except Exception:
        pass
    try:
        _schedule_facet_cache_rebuild(reason)
    except Exception:
        pass


def _facet_default_contexts() -> list[FacetQueryParams]:
    cfg = current_app.config

    def _normalize_keys(value):
        if isinstance(value, (list, tuple)):
            return [str(v) for v in value]
        return None

    search_keys = _normalize_keys(cfg.get('SEARCH_FACET_TAG_KEYS'))
    graph_keys = _normalize_keys(cfg.get('GRAPH_FACET_TAG_KEYS'))
    include_types = bool(cfg.get('SEARCH_FACET_INCLUDE_TYPES', True))
    return [
        FacetQueryParams(
            query='',
            material_type='',
            context='search',
            include_types=include_types,
            tag_filters=[],
            collection_filter=None,
            allowed_scope=None,
            allowed_keys_list=search_keys,
            year_from='',
            year_to='',
            size_min='',
            size_max='',
            sources={'tags': True, 'authors': True, 'years': True},
            request_args=(),
        ),
        FacetQueryParams(
            query='',
            material_type='',
            context='graph',
            include_types=False,
            tag_filters=[],
            collection_filter=None,
            allowed_scope=None,
            allowed_keys_list=graph_keys,
            year_from='',
            year_to='',
            size_min='',
            size_max='',
            sources={'tags': True, 'authors': True, 'years': True},
            request_args=(('context', ('graph',)),),
        ),
    ]


def _schedule_facet_cache_rebuild(reason: str | None = None) -> None:
    if not has_app_context():
        return
    global _FACET_CACHE_REBUILD_PENDING
    with _FACET_CACHE_REBUILD_LOCK:
        if _FACET_CACHE_REBUILD_PENDING:
            return
        _FACET_CACHE_REBUILD_PENDING = True

    @copy_current_app_context
    def _run_rebuild():
        global _FACET_CACHE_REBUILD_PENDING
        try:
            contexts = _facet_default_contexts()
            for params in contexts:
                facet_service.get_facets(
                    params,
                    search_candidate_fn=_search_candidate_ids,
                    like_filter_fn=_apply_like_filter,
                )
            app.logger.debug('Facet cache prewarmed%s', f" ({reason})" if reason else '')
        except Exception as exc:
            app.logger.warning('Facet cache prewarm failed: %s', exc)
        finally:
            with _FACET_CACHE_REBUILD_LOCK:
                _FACET_CACHE_REBUILD_PENDING = False

    try:
        get_task_queue().submit(_run_rebuild, description='facet_cache_prewarm')
    except Exception as exc:
        app.logger.debug('Facet cache prewarm scheduling failed: %s', exc)
        with _FACET_CACHE_REBUILD_LOCK:
            _FACET_CACHE_REBUILD_PENDING = False


def _current_allowed_collections() -> set[int] | None:
    allowed = getattr(g, 'allowed_collection_ids', None)
    if allowed is None:
        return None
    try:
        return set(int(x) for x in allowed)
    except Exception:
        return None


def _ensure_search_support() -> None:
    search_service.ensure_support()


def _rebuild_files_fts(conn=None):
    search_service.rebuild_files(connection=conn)


def _rebuild_tags_fts(conn=None):
    search_service.rebuild_tags(connection=conn)


def _sync_file_to_fts(file_obj: File | None):
    search_service.sync_file(file_obj)


def _delete_file_from_fts(file_id: int | None):
    search_service.delete_file(file_id)


def _search_candidate_ids(query: str, limit: int = 4000) -> list[int] | None:
    return search_service.candidate_ids(query, limit=limit)

def _search_ranked_candidates(query: str, limit: int = 4000, profile: str | None = None) -> list[tuple[int, float]]:
    ranked = search_service.ranked_candidates(query, limit=limit, profile=profile)
    return [(int(fid), float(score)) for fid, score in ranked]

def _search_ranked_candidate_ids(query: str, limit: int = 4000, profile: str | None = None) -> list[int]:
    return [fid for fid, _score in _search_ranked_candidates(query, limit=limit, profile=profile)]


def _apply_like_filter(base_query, query: str):
    return search_service.apply_like_filter(base_query, query)


def _apply_text_search_filter(base_query, query: str):
    return search_service.apply_text_search_filter(base_query, query)


# Second TimedCache definition removed — using import from agregator.caching

from agregator.services.nlp import (  # extracted NLP module
    ru_tokens as _ru_tokens_impl,
    lemma as _lemma_impl,
    lemmas as _lemmas_impl,
    expand_synonyms as _expand_synonyms_impl,
    RU_SYNONYMS as _RU_SYNONYMS_EXT,
)
_RU_SYNONYMS = {
    'машина': ['автомобиль', 'авто', 'тачка', 'car'],
    'изображение': ['картинка', 'фото', 'фотография', 'image', 'picture', 'снимок', 'скриншот', 'screenshot'],
    'вуз': ['университет', 'институт', 'универ'],
    'статья': ['публикация', 'paper', 'publication', 'пейпер'],
    'диссертация': ['дисс', 'thesis', 'dissertation'],
    'журнал': ['journal', 'magazine', 'сборник'],
    'номер': ['выпуск', 'issue'],
    'страница': ['стр', 'страницы', 'pages', 'page'],
    'год': ['лет', 'г'],
}

# NLP functions delegate to agregator.services.nlp (extracted module)
def _ru_tokens(text: str) -> list[str]:
    return _ru_tokens_impl(text)

def _lemma(word: str) -> str:
    return _lemma_impl(word)

def _lemmas(text: str) -> list[str]:
    return _lemmas_impl(text)

def _expand_synonyms(lemmas: list[str]) -> set[str]:
    return _expand_synonyms_impl(lemmas)

LLM_ROUND_ROBIN: dict[str, itertools.cycle] = {}
LLM_ENDPOINT_SIGNATURE: dict[str, tuple[tuple[int, float, tuple[str, ...]], ...]] = {}
LLM_ENDPOINT_POOLS: dict[str, list[dict[str, str]]] = {}
LLM_ENDPOINT_UNIQUE: dict[str, list[dict[str, str]]] = {}
LLM_SCHEMA_READY = False
LMSTUDIO_IDLE_UNLOAD_THREAD_STARTED = False
LMSTUDIO_LAST_ACTIVITY_TS = time.time()
LMSTUDIO_IDLE_UNLOADED = False
LMSTUDIO_IDLE_LOCK = threading.Lock()
LLM_BUSY_HTTP_CODES = {409, 423, 429, 503}
LLM_BUSY_STATUS_VALUES = {'busy', 'processing', 'in_progress', 'queued', 'pending', 'rate_limited'}
LLM_PURPOSES = [
    {'id': 'summary', 'label': 'Резюме стенограмм'},
    {'id': 'keywords', 'label': 'Ключевые слова'},
    {'id': 'compose', 'label': 'Генерация ответов'},
    {'id': 'metadata', 'label': 'Извлечение метаданных'},
    {'id': 'vision', 'label': 'Анализ изображений'},
    {'id': 'doc_chat', 'label': 'Чат по документу'},
    {'id': 'rerank', 'label': 'Реранжирование поиска'},
    {'id': 'default', 'label': 'По умолчанию'},
]
LLM_PROVIDER_OPTIONS = [
    {'id': 'openai', 'label': 'OpenAI-совместимый (LM Studio, OpenAI, OpenRouter)'},
    {'id': 'azure_openai', 'label': 'Azure OpenAI'},
    {'id': 'ollama', 'label': 'Ollama (локальные модели)'},
]
TASK_RETENTION_WINDOW = timedelta(days=1)
TASK_FINAL_STATUSES = ('completed', 'error', 'cancelled')
TASK_STUCK_STATUSES = ('running', 'pending', 'queued', 'cancelling')

_RAG_RERANK_LOCK = threading.Lock()
_RAG_RERANK_CACHE: Dict[str, "CrossEncoderReranker"] = {}


def _reset_rag_rerank_cache(*, locked: bool = False) -> None:
    if not locked:
        with _RAG_RERANK_LOCK:
            _reset_rag_rerank_cache(locked=True)
        return
    for instance in _RAG_RERANK_CACHE.values():
        try:
            instance.close()
        except Exception:
            pass
    _RAG_RERANK_CACHE.clear()


def _get_rag_reranker() -> Optional["CrossEncoderReranker"]:
    runtime = _rt()
    backend_raw = str(getattr(runtime, "rag_rerank_backend", "none") or "").strip().lower()
    if backend_raw not in {"cross-encoder", "cross_encoder", "cross"}:
        if _RAG_RERANK_CACHE:
            _reset_rag_rerank_cache()
        return None
    model_name = str(getattr(runtime, "rag_rerank_model", "") or "").strip()
    if not model_name:
        return None
    device = getattr(runtime, "rag_rerank_device", None)
    batch_size = max(1, int(getattr(runtime, "rag_rerank_batch_size", 16) or 16))
    max_length = max(32, int(getattr(runtime, "rag_rerank_max_length", 512) or 512))
    max_chars = max(200, int(getattr(runtime, "rag_rerank_max_chars", 1200) or 1200))
    cache_key = f"{backend_raw}::{model_name}::{device or ''}::{batch_size}::{max_length}::{max_chars}"
    with _RAG_RERANK_LOCK:
        existing = _RAG_RERANK_CACHE.get(cache_key)
        if existing:
            return existing
        _reset_rag_rerank_cache(locked=True)
        try:
            from agregator.rag.rerank import CrossEncoderConfig, load_reranker

            reranker = load_reranker(
                backend_raw,
                config=CrossEncoderConfig(
                    model_name=model_name,
                    device=device,
                    batch_size=batch_size,
                    max_length=max_length,
                    truncate_chars=max_chars,
                ),
            )
        except Exception as exc:
            logger.warning("Не удалось инициализировать cross-encoder rerank (%s/%s): %s", backend_raw, model_name, exc)
            return None
        if reranker is None:
            return None
        _RAG_RERANK_CACHE[cache_key] = reranker
        return reranker

def _row_text_for_search(f: File) -> str:
    parts = [
        f.title or '', f.author or '', f.keywords or '', f.filename or '',
        getattr(f, 'text_excerpt', '') or '', getattr(f, 'abstract', '') or ''
    ]
    try:
        parts.extend([f"{t.key}={t.value}" for t in (f.tags or []) if t and t.key and t.value])
    except Exception:
        pass
    return ' '.join(parts)

# Простое кэширующее хранилище в памяти для расширения ключевых слов ИИ
AI_EXPAND_CACHE: dict[str, tuple[float, list[str], list[tuple[str, int]]]] = {}
AI_KEYWORD_PLAN_CACHE: dict[str, tuple[float, Dict[str, List[str]]]] = {}
QUERY_VARIANT_CACHE: dict[str, tuple[float, List[str]]] = {}
AuthorityCacheEntry = Dict[str, Any]
AUTHORITY_GRAPH_CACHE: dict[str, AuthorityCacheEntry] = {}
AI_FEEDBACK_MODEL_CACHE: tuple[float, Dict[int, Dict[str, Any]]] = (0.0, {})

try:
    QUERY_VARIANT_TTL = float(os.getenv("AI_QUERY_VARIANT_TTL", "600") or 600.0)
except Exception:
    QUERY_VARIANT_TTL = 600.0
try:
    AUTHORITY_GRAPH_TTL = float(os.getenv("AI_AUTHORITY_TTL", "1800") or 1800.0)
except Exception:
    AUTHORITY_GRAPH_TTL = 1800.0

# Весовые коэффициенты и параметры оценки (можно менять через переменные окружения)
def _getf(name: str, default: float) -> float:
    try:
        return float(os.getenv(name, str(default)))
    except Exception:
        return default

AI_SCORE_TITLE = _getf('AI_SCORE_TITLE', 2.5)
AI_SCORE_AUTHOR = _getf('AI_SCORE_AUTHOR', 1.5)
AI_SCORE_KEYWORDS = _getf('AI_SCORE_KEYWORDS', 1.2)
AI_SCORE_EXCERPT = _getf('AI_SCORE_EXCERPT', 1.0)
AI_SCORE_ABSTRACT = _getf('AI_SCORE_ABSTRACT', 1.0)
AI_SCORE_TAG = _getf('AI_SCORE_TAG', 1.0)
AI_SCORE_FTS = _getf('AI_SCORE_FTS', 1.8)
AI_SCORE_FTS_RAW = _getf('AI_SCORE_FTS_RAW', 2.2)
AI_BOOST_PHRASE = _getf('AI_BOOST_PHRASE', 3.0)
AI_BOOST_MULTI = _getf('AI_BOOST_MULTI', 0.6)  # дополнительный бонус за каждое уникальное слово
AI_BOOST_SNIPPET_COOCCUR = _getf('AI_BOOST_SNIPPET_COOCCUR', 0.8)
SNIPPET_CACHE_TTL_HOURS = int(os.getenv('AI_SNIPPET_CACHE_TTL_HOURS', '24'))
KEYWORD_IDF_MIN = float(os.getenv('KEYWORD_IDF_MIN', '1.25'))
CACHE_CLEANUP_INTERVAL_HOURS = int(os.getenv('AI_SNIPPET_CACHE_SWEEP_INTERVAL_HOURS', '24'))
try:
    FEEDBACK_MODEL_TTL = float(os.getenv("AI_FEEDBACK_TTL", "600") or 600.0)
except Exception:
    FEEDBACK_MODEL_TTL = 600.0
FEEDBACK_POS_PRIOR = _getf('AI_FEEDBACK_POS_PRIOR', 1.0)
FEEDBACK_NEG_PRIOR = _getf('AI_FEEDBACK_NEG_PRIOR', 1.0)
FEEDBACK_CLICK_WEIGHT = _getf('AI_FEEDBACK_CLICK_WEIGHT', 0.5)
FEEDBACK_MAX_WEIGHT = _getf('AI_FEEDBACK_MAX_WEIGHT', 2.5)
AI_FEEDBACK_WEIGHT_SCALE = _getf('AI_FEEDBACK_WEIGHT_SCALE', 0.8)
RAG_FEEDBACK_WEIGHT_SCALE = _getf('RAG_FEEDBACK_WEIGHT_SCALE', 0.5)
try:
    FEEDBACK_TRAIN_INTERVAL_HOURS = float(os.getenv("AI_FEEDBACK_TRAIN_INTERVAL_HOURS", "0") or 0.0)
except Exception:
    FEEDBACK_TRAIN_INTERVAL_HOURS = 0.0
try:
    FEEDBACK_TRAIN_CUTOFF_DAYS = int(os.getenv("AI_FEEDBACK_TRAIN_CUTOFF_DAYS", "90") or 90)
except Exception:
    FEEDBACK_TRAIN_CUTOFF_DAYS = 90

FEEDBACK_SCHEDULER_EVENT = threading.Event()
FEEDBACK_LAST_TRIGGER_AT = 0.0
FEEDBACK_ONLINE_TRAIN_LAST_AT = 0.0
try:
    FEEDBACK_ONLINE_TRAIN_MIN_INTERVAL_SEC = float(os.getenv("AI_FEEDBACK_ONLINE_TRAIN_MIN_INTERVAL_SEC", "900") or 900.0)
except Exception:
    FEEDBACK_ONLINE_TRAIN_MIN_INTERVAL_SEC = 900.0
try:
    FEEDBACK_ONLINE_TRAIN_EVERY_EVENTS = int(os.getenv("AI_FEEDBACK_ONLINE_TRAIN_EVERY_EVENTS", "25") or 25)
except Exception:
    FEEDBACK_ONLINE_TRAIN_EVERY_EVENTS = 25

def _now() -> float:
    return time.time()

def _sha256(s: str) -> str:
    try:
        return hashlib.sha256((s or "").encode("utf-8", errors="ignore")).hexdigest()
    except Exception:
        return hashlib.sha1((s or "").encode("utf-8", errors="ignore")).hexdigest()

def _query_fingerprint(query: str, payload: dict | None = None) -> str:
    payload = payload or {}
    allowed_scope = _current_allowed_collections()
    user = _load_current_user()
    scope_collections = None if allowed_scope is None else sorted(int(x) for x in allowed_scope)
    serialized = {
        'query': query.strip(),
        'user_id': getattr(user, 'id', None),
        'scope_collections': scope_collections,
        'top_k': payload.get('top_k'),
        'deep_search': bool(payload.get('deep_search')),
        'full_text': bool(payload.get('full_text')),
        'llm_snippets': bool(payload.get('llm_snippets')),
        'sources': payload.get('sources'),
        'filters': {
            'collection_ids': payload.get('collection_ids') or payload.get('collection_id'),
            'material_types': payload.get('material_types'),
            'year_from': payload.get('year_from'),
            'year_to': payload.get('year_to'),
            'tag_filters': payload.get('tag_filters'),
        }
    }
    try:
        data = json.dumps(serialized, sort_keys=True, ensure_ascii=False)
    except Exception:
        data = str(serialized)
    return _sha256(data)


def _cache_scope_identity() -> dict[str, Any]:
    user = _load_current_user()
    allowed_scope = _current_allowed_collections()
    try:
        scope_collections = None if allowed_scope is None else sorted(int(x) for x in allowed_scope)
    except Exception:
        scope_collections = None
    return {
        'user_id': getattr(user, 'id', None),
        'scope_collections': scope_collections,
    }

def _get_cached_snippet(file_id: int, query_hash: str, llm_variant: bool):
    if not file_id or not query_hash:
        return None
    try:
        entry = AiSearchSnippetCache.query.filter_by(file_id=file_id, query_hash=query_hash, llm_variant=llm_variant).first()
    except Exception:
        entry = None
    if not entry:
        return None
    if entry.expires_at and entry.expires_at < datetime.utcnow():
        try:
            db.session.delete(entry)
            db.session.commit()
        except Exception:
            db.session.rollback()
        return None
    return entry


def _store_snippet_cache(file_id: int, query_hash: str, llm_variant: bool, snippet: str, meta: dict | None = None, ttl_hours: int | None = None):
    if not file_id or not query_hash or not snippet:
        return
    ttl = ttl_hours if ttl_hours is not None else SNIPPET_CACHE_TTL_HOURS
    expires_at = datetime.utcnow() + timedelta(hours=max(1, ttl)) if ttl else None
    try:
        entry = AiSearchSnippetCache.query.filter_by(file_id=file_id, query_hash=query_hash, llm_variant=llm_variant).first()
        meta_json = json.dumps(meta, ensure_ascii=False) if meta else None
        if entry:
            entry.snippet = snippet
            entry.meta = meta_json
            entry.expires_at = expires_at
        else:
            entry = AiSearchSnippetCache(
                file_id=file_id,
                query_hash=query_hash,
                llm_variant=llm_variant,
                snippet=snippet,
                meta=meta_json,
                expires_at=expires_at,
            )
            db.session.add(entry)
        db.session.commit()
    except Exception:
        db.session.rollback()


def _prune_expired_snippet_cache() -> int:
    cutoff = datetime.utcnow()
    try:
        q = AiSearchSnippetCache.query.filter(
            AiSearchSnippetCache.expires_at.isnot(None),
            AiSearchSnippetCache.expires_at < cutoff
        )
        deleted = q.delete(synchronize_session=False)
        if deleted:
            db.session.commit()
            app.logger.info(f"[ai-search] snippet cache cleanup removed {deleted} rows")
        else:
            db.session.commit()
        return deleted or 0
    except Exception as exc:
        db.session.rollback()
        app.logger.warning(f"[ai-search] snippet cache cleanup failed: {exc}")
        return 0


_CLEANUP_THREAD_STARTED = False
_FEEDBACK_THREAD_STARTED = False


def _start_cache_cleanup_scheduler() -> None:
    global _CLEANUP_THREAD_STARTED
    if _CLEANUP_THREAD_STARTED or CACHE_CLEANUP_INTERVAL_HOURS <= 0:
        return

    interval = max(1, CACHE_CLEANUP_INTERVAL_HOURS) * 3600

    def loop():
        while True:
            try:
                with app.app_context():
                    _prune_expired_snippet_cache()
            except Exception:
                app.logger.exception("[ai-search] cache cleanup loop error")
            time.sleep(interval)

    threading.Thread(target=loop, name='ai-snippet-cache-cleaner', daemon=True).start()
    _CLEANUP_THREAD_STARTED = True


def _start_feedback_training_scheduler() -> None:
    global _FEEDBACK_THREAD_STARTED
    if _FEEDBACK_THREAD_STARTED:
        return

    def _queue_training(trigger: str, *, allow_duplicate: bool = False) -> None:
        global FEEDBACK_LAST_TRIGGER_AT
        try:
            with app.app_context():
                task_id, queued = _enqueue_feedback_training(
                    cutoff_days=FEEDBACK_TRAIN_CUTOFF_DAYS,
                    trigger=trigger,
                    submitted_by=None,
                    allow_duplicate=allow_duplicate,
                )
                if queued:
                    FEEDBACK_LAST_TRIGGER_AT = time.time()
                    app.logger.info("[ai-feedback] queued training task #%s (%s)", task_id, trigger)
                else:
                    app.logger.info("[ai-feedback] training already running (task #%s)", task_id)
        except Exception:
            app.logger.exception("[ai-feedback] failed to queue training (%s)", trigger)

    # Автозапуск даже при выключенном расписании
    _queue_training('initial', allow_duplicate=False)

    def loop():
        while True:
            interval_hours = 0.0
            try:
                interval_hours = max(0.0, float(FEEDBACK_TRAIN_INTERVAL_HOURS or 0.0))
            except Exception:
                interval_hours = 0.0
            if interval_hours <= 0.0:
                # Ждём сигнала о включении расписания
                FEEDBACK_SCHEDULER_EVENT.wait(timeout=300)
                FEEDBACK_SCHEDULER_EVENT.clear()
                continue
            wait_seconds = max(300.0, interval_hours * 3600.0)
            triggered = FEEDBACK_SCHEDULER_EVENT.wait(timeout=wait_seconds)
            if triggered:
                FEEDBACK_SCHEDULER_EVENT.clear()
                continue
            _queue_training('schedule')

    threading.Thread(target=loop, name='ai-feedback-trainer', daemon=True).start()
    _FEEDBACK_THREAD_STARTED = True

def _record_search_metric(query_hash: str, durations: dict[str, float], user: User | None, meta_extra: dict | None = None):
    try:
        known_fields = {'total', 'keywords', 'candidates', 'deep', 'llm_answer', 'llm_snippets'}
        extra_meta = {}
        for key, value in durations.items():
            if key not in known_fields and value is not None:
                extra_meta[key] = value
        if meta_extra:
            for key, value in meta_extra.items():
                if value is None:
                    continue
                extra_meta[key] = value
        metric = AiSearchMetric(
            query_hash=query_hash,
            user_id=getattr(user, 'id', None),
            total_ms=int(durations.get('total', 0) * 1000),
            keywords_ms=int(durations.get('keywords', 0) * 1000) if 'keywords' in durations else None,
            candidate_ms=int(durations.get('candidates', 0) * 1000) if 'candidates' in durations else None,
            deep_ms=int(durations.get('deep', 0) * 1000) if 'deep' in durations else None,
            llm_answer_ms=int(durations.get('llm_answer', 0) * 1000) if 'llm_answer' in durations else None,
            llm_snippet_ms=int(durations.get('llm_snippets', 0) * 1000) if 'llm_snippets' in durations else None,
            meta=json.dumps(extra_meta, ensure_ascii=False) if extra_meta else None,
        )
        db.session.add(metric)
        db.session.commit()
    except Exception:
        db.session.rollback()

# ------------------- Конфигурация -------------------

CONFIG = load_app_config()


def _apply_config_defaults(cfg: AppConfig) -> None:
    global BASE_DIR, LOGIN_BACKGROUNDS_DIR, RENAME_PATTERNS, DEFAULT_PROMPTS
    global FW_CACHE_DIR, SETTINGS_STORE_PATH, LOG_DIR, LOG_FILE_PATH
    global HTTP_DEFAULT_TIMEOUT, HTTP_CONNECT_TIMEOUT, HTTP_RETRIES, HTTP_BACKOFF_FACTOR
    global LOG_LEVEL, SENTRY_DSN, SENTRY_ENVIRONMENT

    runtime_settings_store.initialize(cfg)

    BASE_DIR = cfg.base_dir
    LOGIN_BACKGROUNDS_DIR = cfg.login_backgrounds_dir
    RENAME_PATTERNS = dict(cfg.rename_patterns)
    DEFAULT_PROMPTS = dict(cfg.default_prompts)
    FW_CACHE_DIR = cfg.fw_cache_dir
    SETTINGS_STORE_PATH = cfg.settings_store_path
    LOG_DIR = cfg.logs_dir
    LOG_FILE_PATH = cfg.log_file_path
    HTTP_DEFAULT_TIMEOUT = float(cfg.http_timeout)
    HTTP_CONNECT_TIMEOUT = float(cfg.http_connect_timeout)
    HTTP_RETRIES = int(cfg.http_retries)
    HTTP_BACKOFF_FACTOR = float(cfg.http_backoff_factor)
    LOG_LEVEL = cfg.log_level
    SENTRY_DSN = cfg.sentry_dsn
    SENTRY_ENVIRONMENT = cfg.sentry_environment

    _refresh_runtime_globals()
    _ensure_rag_embedding_defaults()


_apply_config_defaults(CONFIG)
LLM_PROVIDER_CHOICES = {'openai', 'openrouter', 'azure_openai', 'ollama'}
if runtime_settings_store.current.lm_default_provider not in LLM_PROVIDER_CHOICES:
    runtime_settings_store.current.lm_default_provider = 'openai'

# ------------------- Сохранение настроек во время работы -------------------

# ------------------- Сохранение настроек во время работы -------------------

# ------------------- Facet configuration (search & graph) -------------------

def _facet_config_state() -> dict:
    runtime = _rt()
    return {
        'search': {
            'include_types': bool(runtime.search_facet_include_types),
            'tag_keys': list(runtime.search_facet_tag_keys) if runtime.search_facet_tag_keys is not None else None,
        },
        'graph': {
            'tag_keys': list(runtime.graph_facet_tag_keys) if runtime.graph_facet_tag_keys is not None else None,
        }
    }


def _load_runtime_settings_from_disk() -> None:
    if not SETTINGS_STORE_PATH.exists():
        return
    try:
        data = json.loads(SETTINGS_STORE_PATH.read_text(encoding='utf-8'))
    except Exception as exc:
        logger.warning("Не удалось прочитать настройки: %s", exc)
        return
    runtime_settings_store.apply_updates(data)
    runtime_settings_store.current.apply_env()
    _reset_rag_rerank_cache()
    runtime = _rt()
    if runtime.lm_default_provider not in LLM_PROVIDER_CHOICES:
        runtime.lm_default_provider = 'openai'
    _refresh_runtime_globals()
    configure_llm_cache(
        enabled=runtime.llm_cache_enabled,
        max_items=runtime.llm_cache_max_items,
        ttl_seconds=runtime.llm_cache_ttl_seconds,
    )
    configure_search_cache(
        enabled=runtime.search_cache_enabled,
        max_items=runtime.search_cache_max_items,
        ttl_seconds=runtime.search_cache_ttl_seconds,
    )
    app_ref = globals().get('app')
    if app_ref:
        runtime.apply_to_flask_config(app_ref)


def _save_runtime_settings_to_disk() -> None:
    try:
        SETTINGS_STORE_PATH.parent.mkdir(parents=True, exist_ok=True)
        snapshot = runtime_settings_store.snapshot()
        SETTINGS_STORE_PATH.write_text(
            json.dumps(snapshot, ensure_ascii=False, indent=2),
            encoding='utf-8',
        )
    except Exception as exc:
        logger.warning("Не удалось сохранить настройки: %s", exc)


def _runtime_settings_snapshot() -> dict:
    return runtime_settings_store.snapshot()


def _apply_runtime_settings_to_config(app_obj: Flask) -> None:
    runtime_settings_store.current.apply_to_flask_config(app_obj)


# ------------------- Lightweight response helpers -------------------

# Error helpers delegate to agregator.middleware.errors (extracted module)
from agregator.middleware.errors import (
    json_error,
    html_page as _html_page,
    add_pipeline_cors_headers as _add_pipeline_cors_headers_impl,
    pipeline_cors_preflight as _pipeline_cors_preflight_impl,
)




_PREVIEW_HEAD = """
<style>
  :root { color-scheme: dark light; }
  body { margin:0; font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif; background:#0d1117; color:#c9d1d9; }
  a { color:#58a6ff; }
  .preview-shell { min-height:100vh; display:flex; flex-direction:column; }
  .preview-shell.embedded { min-height:auto; }
  .preview-header { padding:12px 16px; display:flex; align-items:center; gap:12px; justify-content:space-between; background:#161b22; border-bottom:1px solid #30363d; }
  .preview-title { font-size:16px; font-weight:600; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; }
  .preview-actions { display:flex; gap:8px; }
  .btn { display:inline-flex; align-items:center; justify-content:center; gap:6px; padding:8px 16px; border-radius:999px; font-size:13px;
         border:1px solid #30363d; background:#1f6feb; color:#fff; text-decoration:none; }
  .btn.secondary { background:transparent; color:#c9d1d9; }
  .preview-main { flex:1; overflow:auto; padding:16px; }
  .preview-shell.embedded .preview-main { padding:12px; }
  .preview-card { background:#161b22; border:1px solid #30363d; border-radius:14px; padding:16px; box-shadow:0 16px 48px rgba(0,0,0,0.35); }
  .preview-text { white-space:pre-wrap; max-height:70vh; overflow:auto; word-break:break-word; overflow-wrap:anywhere; background:#0d1117; padding:12px; border-radius:10px; border:1px solid #30363d; }
  .muted { color:#8b949e; font-size:13px; margin:12px 0 6px; }
  .section { margin-top:18px; }
  .section h3 { margin:0 0 8px; font-size:15px; font-weight:600; color:inherit; }
  img.responsive { max-width:100%; height:auto; border-radius:12px; border:1px solid #30363d; }
  .badge { display:inline-block; padding:4px 10px; border-radius:999px; background:#30363d; color:#c9d1d9; font-size:12px; margin-left:6px; }
  .tag { display:inline-block; padding:4px 10px; border-radius:999px; background:#30363d; color:#c9d1d9; font-size:12px; margin:2px 4px 2px 0; }
  mark { background:#bb8009; color:inherit; padding:0 2px; border-radius:3px; }
  embed, iframe { border:none; border-radius:10px; background:#0d1117; }
  audio { width:100%; }
  .docx-preview { background:#0d1117; border:1px solid #30363d; border-radius:10px; max-height:70vh; overflow:auto; padding:16px; line-height:1.55; }
  .docx-preview p { margin:0 0 0.9em; }
  .docx-preview h1, .docx-preview h2, .docx-preview h3, .docx-preview h4 { margin-top:1.2em; margin-bottom:0.6em; }
  .docx-preview ul, .docx-preview ol { padding-left:1.3em; }
  @media (prefers-color-scheme: light) {
    body { background:#f5f6f8; color:#101828; }
    .preview-header { background:#ffffff; border-color:rgba(15,23,42,0.09); }
    .preview-card { background:#ffffff; border-color:rgba(15,23,42,0.08); box-shadow:0 12px 36px rgba(15,23,42,0.2); }
    .preview-text { background:#f8fafc; border-color:rgba(15,23,42,0.08); }
    .btn { background:#1f6feb; color:#fff; border:1px solid #1f6feb; }
    .btn.secondary { background:#fff; color:#0f172a; border-color:rgba(15,23,42,0.14); }
    .muted { color:#64748b; }
    .badge, .tag { background:rgba(15,23,42,0.08); color:#0f172a; }
    .docx-preview { background:#f8fafc; border-color:rgba(15,23,42,0.08); }
  }
</style>
<script>
(function(){
  try {
    const params = new URLSearchParams(window.location.search);
    const mark = (params.get('mark') || '').trim();
    if (!mark) return;
    const terms = Array.from(new Set(mark.split(/[\\s,|]+/).filter(x => x && x.length >= 2))).slice(0, 5);
    if (!terms.length) return;
    const esc = s => s.replace(/[.*+?^${}()|[\\]\\]/g, '\\$&');
    const re = new RegExp('(' + terms.map(esc).join('|') + ')', 'gi');
    document.querySelectorAll('.preview-text, .docx-preview').forEach(el => {
      try { el.innerHTML = el.innerHTML.replace(re, '<mark>$1</mark>'); } catch (_) {}
    });
  } catch (err) {
    console.warn('highlight error', err);
  }
})();
</script>
""".strip()


def _sanitize_html_fragment(html_fragment: str) -> str:
    if not html_fragment:
        return ''
    cleaned = re.sub(r'(?is)<script.*?>.*?</script>', '', html_fragment)
    cleaned = re.sub(r'(?is)<style.*?>.*?</style>', '', cleaned)
    cleaned = re.sub(r'(?i) on[a-z]+="[^"]*"', '', cleaned)
    cleaned = re.sub(r"(?i) on[a-z]+='[^']*'", '', cleaned)
    cleaned = re.sub(r'(?i)javascript:', '', cleaned)
    return cleaned


def _docx_to_html(path: Path, limit_chars: int = 60000) -> str | None:
    if mammoth is None:
        return None
    try:
        with path.open('rb') as docx_file:
            result = mammoth.convert_to_html(docx_file)
        html_fragment = result.value or ''
        sanitized = _sanitize_html_fragment(html_fragment)
        return _limit_text_length(sanitized, limit_chars)
    except Exception as exc:
        app.logger.warning(f"DOCX preview conversion failed for {path}: {exc}")
        return None


def _render_preview(
    rel_url: str,
    *,
    is_pdf: bool,
    is_text: bool,
    is_audio: bool,
    is_image: bool,
    is_docx: bool,
    content: str,
    thumbnail_url: str | None,
    abstract: str,
    audio_url: str | None,
    duration: str | None,
    image_url: str | None,
    keywords: str | None,
    embedded: bool,
    docx_html: str | None
) -> Response:
    safe_rel = escape(rel_url)
    download_url = escape(url_for('download_file', rel_path=rel_url))
    view_url = escape(url_for('view_file', rel_path=rel_url))
    shell_class = "preview-shell embedded" if embedded else "preview-shell"

    sections: list[str] = []

    def section(title: str, html_content: str) -> None:
        sections.append(f'<div class="section"><h3>{escape(title)}</h3>{html_content}</div>')

    if is_pdf:
        sections.append(f'<embed src="{view_url}" type="application/pdf" width="100%" style="min-height:70vh;" />')
        if content:
            section('Фрагмент содержимого', f'<div class="preview-text">{escape(content)}</div>')
    elif is_audio:
        if audio_url:
            audio = f'<audio controls preload="metadata"><source src="{escape(audio_url)}" /></audio>'
            if duration:
                audio += f'<span class="badge">{escape(duration)}</span>'
            sections.append(audio)
        if abstract:
            section('Резюме', f'<div class="preview-text">{escape(abstract)}</div>')
        if keywords:
            chips = ''.join(f'<span class="tag">{escape(k.strip())}</span>' for k in keywords.split(',') if k.strip())
            if chips:
                section('Ключевые слова', chips)
        if content and not abstract:
            section('Фрагмент стенограммы', f'<div class="preview-text">{escape(content)}</div>')
    elif is_image:
        img_src = escape(image_url or url_for('media_file', rel_path=rel_url))
        sections.append(f'<img class="responsive" src="{img_src}" alt="image preview" />')
        if abstract:
            section('Описание', f'<div class="preview-text">{escape(abstract)}</div>')
        if keywords:
            chips = ''.join(f'<span class="tag">{escape(k.strip())}</span>' for k in keywords.split(',') if k.strip())
            if chips:
                section('Ключевые слова', chips)
    elif is_docx:
        if docx_html:
            section('Документ', f'<div class="docx-preview">{docx_html}</div>')
        elif content:
            section('Текст', f'<div class="preview-text">{escape(content)}</div>')
        else:
            sections.append('<div class="muted">Не удалось построить предпросмотр DOCX.</div>')
    elif is_text:
        section('Текст', f'<div class="preview-text">{escape(content)}</div>')
    elif thumbnail_url:
        sections.append(f'<img class="responsive" src="{escape(thumbnail_url)}" alt="thumbnail" />')
        if content:
            section('Фрагмент содержимого', f'<div class="preview-text">{escape(content)}</div>')
    else:
        if content:
            sections.append('<div class="muted">Предпросмотр для этого формата ограничен. Ниже показан доступный текст:</div>')
            sections.append(f'<div class="preview-text">{escape(content)}</div>')
        else:
            sections.append('<div class="muted">Предпросмотр не поддерживается для этого типа файла.</div>')

    if abstract and not any('Резюме' in s for s in sections):
        section('Резюме', f'<div class="preview-text">{escape(abstract)}</div>')
    if keywords and not any('Ключевые слова' in s for s in sections):
        chips = ''.join(f'<span class="tag">{escape(k.strip())}</span>' for k in keywords.split(',') if k.strip())
        if chips:
            section('Ключевые слова', chips)

    header_actions = f'<a class="btn" href="{download_url}" download>Скачать</a>'
    body = f"""
    <div class=\"{shell_class}\">
      <header class=\"preview-header\">
        <div class=\"preview-title\">Предпросмотр: {safe_rel}</div>
        <div class=\"preview-actions\">{header_actions}</div>
      </header>
      <main class=\"preview-main\">
        <div class=\"preview-card\">{''.join(sections)}</div>
      </main>
    </div>
    """
    return _html_page(f"Предпросмотр — {rel_url}", body, extra_head=_PREVIEW_HEAD)

# ------------------- Вспомогательные функции -------------------

def _normalize_author(val):
    """Convert various author representations (list/dict/etc.) to a string.
    Returns None for empty results.
    """
    try:
        if val is None:
            return None
        # список/кортеж имён или словарей
        if isinstance(val, (list, tuple, set)):
            out = []
            for x in val:
                if x is None:
                    continue
                if isinstance(x, dict):
                    if 'name' in x and x['name']:
                        s = str(x['name']).strip()
                    else:
                        s = " ".join(str(v).strip() for v in x.values() if v)
                else:
                    s = str(x).strip()
                if s:
                    out.append(s)
            return ", ".join(out) if out else None
        # один словарь вида {first,last} или {name}
        if isinstance(val, dict):
            if 'name' in val and val['name']:
                s = str(val['name']).strip()
                return s or None
            s = " ".join(str(v).strip() for v in val.values() if v)
            return s or None
        s = str(val).strip()
        return s or None
    except Exception:
        return None

def _normalize_year(val):
    """Normalize year to a short string; prefer 4-digit if present."""
    try:
        if val is None:
            return None
        if isinstance(val, (int, float)):
            return str(int(val))
        s = str(val).strip()
        m = re.search(r"\b(\d{4})\b", s)
        if m:
            return m.group(1)
        return s[:16] if s else None
    except Exception:
        return None


def _normalize_keywords_text(value: Any) -> Optional[str]:
    raw = str(value or "").strip()
    if not raw:
        return None
    parts = [p.strip() for p in re.split(r"[,;\n|]+", raw) if p and p.strip()]
    dedup: list[str] = []
    seen: set[str] = set()
    for item in parts:
        norm = item.lower()
        if norm in seen:
            continue
        seen.add(norm)
        dedup.append(item)
        if len(dedup) >= 50:
            break
    return ", ".join(dedup) if dedup else None


def _normalize_material_type_value(value: Any, *, ext: str = "", text_excerpt: str = "", filename: str = "") -> str:
    current = normalize_material_type(str(value or "").strip())
    allowed = _allowed_material_type_keys()
    if current in allowed and current:
        return current
    ext_l = (ext or "").lower()
    if ext_l in AUDIO_EXTS:
        return "audio"
    if ext_l in IMAGE_EXTS:
        return "image"
    cand = _detect_type_pre_llm(ext_l, text_excerpt or "", filename or "")
    if cand and cand in allowed:
        return cand
    guessed = guess_material_type(ext_l, text_excerpt or "", filename or "")
    guessed_norm = normalize_material_type(guessed)
    if guessed_norm in allowed and guessed_norm:
        return guessed_norm
    return "document"


def _apply_metadata_quality_rules(file_obj: File, *, ext: str = "", text_excerpt: str = "", filename: str = "") -> dict[str, Any]:
    changed: list[str] = []
    original_author = file_obj.author
    author_norm = _normalize_author(file_obj.author)
    if author_norm:
        author_norm = _normalize_author_name(author_norm)
    if author_norm and author_norm != (original_author or "").strip():
        file_obj.author = author_norm
        changed.append("author")

    original_year = file_obj.year
    year_norm = _normalize_year(file_obj.year)
    if year_norm:
        try:
            year_int = int(year_norm)
            if year_int < 1500 or year_int > 2100:
                year_norm = None
        except Exception:
            pass
    if year_norm and year_norm != (original_year or "").strip():
        file_obj.year = year_norm
        changed.append("year")

    original_keywords = file_obj.keywords
    keywords_norm = _normalize_keywords_text(file_obj.keywords)
    if keywords_norm and keywords_norm != (original_keywords or "").strip():
        file_obj.keywords = keywords_norm
        changed.append("keywords")

    original_type = file_obj.material_type
    type_norm = _normalize_material_type_value(
        file_obj.material_type,
        ext=ext or (file_obj.ext or ""),
        text_excerpt=text_excerpt or (file_obj.text_excerpt or ""),
        filename=filename or (file_obj.filename or ""),
    )
    if type_norm and type_norm != (original_type or "").strip():
        file_obj.material_type = type_norm
        changed.append("material_type")

    if not (file_obj.title or "").strip() and (filename or file_obj.filename):
        cleaned_title = re.sub(r"[_\-]+", " ", str(filename or file_obj.filename)).strip()
        if cleaned_title:
            file_obj.title = cleaned_title[:300]
            changed.append("title")

    quality = _metadata_quality(file_obj)
    return {"changed": changed, "quality": quality}

# --------- Расширенное извлечение тегов (общие и специфичные по типу) ---------
def extract_richer_tags(material_type: str, text: str, filename: str = "") -> dict:
    t = (text or "")
    tl = t.lower()
    tags: dict[str, str] = {}
    # предположение языка
    try:
        cyr = sum(1 for ch in t if ('а' <= ch.lower() <= 'я') or (ch in 'ёЁ'))
        lat = sum(1 for ch in t if 'a' <= ch.lower() <= 'z')
        if cyr + lat > 20:
            tags['lang'] = 'ru' if cyr >= lat else 'en'
    except Exception:
        pass
    # распространённые идентификаторы
    m = re.search(r"\b(10\.\d{4,9}\/[\w\-\.:;()\/[\]A-Za-z0-9]+)", t)
    if m:
        tags.setdefault('doi', m.group(1))
    m = re.search(r"\bISBN[:\s]*([0-9\- ]{10,20})", t, flags=re.I)
    if m:
        tags.setdefault('isbn', m.group(1).strip())
    m = re.search(r"\bISSN[:\s]*([0-9\- ]{8,15})\b", t, flags=re.I)
    if m:
        tags.setdefault('issn', m.group(1).strip())
    m = re.search(r"\bУДК[:\s]*([\d\.:\-]+)\b", t, flags=re.I)
    if m:
        tags.setdefault('udk', m.group(1))
    m = re.search(r"\bББК[:\s]*([A-ZА-Я0-9\.-/]+)\b", t, flags=re.I)
    if m:
        tags.setdefault('bbk', m.group(1))

    mt = (material_type or '').strip().lower()
    if mt in ("dissertation", "dissertation_abstract"):
        # код специальности вида 05.13.11
        m = re.search(r"\b(\d{2}\.\d{2}\.\d{2})\b", t)
        if m:
            tags.setdefault('specialty', m.group(1))
        if 'автореферат' in tl:
            tags.setdefault('kind', 'автореферат')
    elif mt == 'article':
        m = re.search(r"(journal|transactions|вестник|журнал)[:\s\-]+([^\n\r]{3,80})", tl, flags=re.I)
        if m:
            tags.setdefault('journal', m.group(2).strip().title())
        m = re.search(r"\b(\d+)\b.*?№\s*(\d+)\b.*?([\d]+)[\-–—]([\d]+)", t)
        if m:
            tags.setdefault('volume_issue', f"{m.group(1)}/{m.group(2)}")
            tags.setdefault('pages', f"{m.group(3)}–{m.group(4)}")
    elif mt == 'journal':
        m = re.search(r"(journal|transactions|вестник|журнал)[:\s\-]+([^\n\r]{3,80})", tl, flags=re.I)
        if m:
            tags.setdefault('journal', m.group(2).strip().title())
        m = re.search(r"\b(\d+)\b.*?№\s*(\d+)\b", t)
        if m:
            tags.setdefault('volume_issue', f"{m.group(1)}/{m.group(2)}")
        m = re.search(r"№\s*(\d+)\b", t)
        if m:
            tags.setdefault('number', m.group(1))
        toc_entries = _extract_journal_toc_entries(t)
        if toc_entries:
            formatted: list[str] = []
            for entry in toc_entries[:25]:
                title = entry.get('title', '').strip()
                if not title:
                    continue
                piece = title
                authors = entry.get('authors')
                page = entry.get('page')
                if authors:
                    piece = f"{authors}: {piece}"
                if page:
                    piece = f"{piece} (стр. {page})"
                formatted.append(piece)
            if formatted:
                tags.setdefault('toc', formatted)
    elif mt == 'textbook':
        m = re.search(r"учебн(?:ое|ик|ое\s+пособие)\s*[:\-]?\s*([^\n\r]{3,80})", tl)
        if m:
            tags.setdefault('discipline', m.group(1).strip().title())
    elif mt == 'monograph':
        m = re.search(r"(серия|series)\s*[:\-]?\s*([^\n\r]{3,80})", tl)
        if m:
            tags.setdefault('series', m.group(2).strip().title())
    elif mt == 'standard':
        pats = [
            r"\bГОСТ\s*R?\s*\d{1,5}(?:\.\d+)*[-–—]\d{2,4}\b",
            r"\bСТБ\s*\d{1,5}(?:\.\d+)*[-–—]\d{2,4}\b",
            r"\bСТО\s*[^\s]*\s*\d{1,6}[-–—]?\d{2,4}\b",
            r"\bISO\s*\d{3,5}(?:-\d+)*(?::\d{4})?\b",
            r"\bIEC\s*\d{3,5}(?:-\d+)*(?::\d{4})?\b",
            r"\bСП\s*\d{1,4}\.\d{1,4}-\d{4}\b",
            r"\bСанПиН\s*\d+[\.-]\d+[\.-]\d+\b",
            r"\bТУ\s*[A-Za-zА-Яа-я0-9\./-]+\b",
        ]
        for pat in pats:
            m = re.search(pat, t, flags=re.I)
            if m:
                tags.setdefault('standard', m.group(0))
                break
        if re.search(r"утратил[аи]?\s+силу|замен(ен|яет)|взамен", tl):
            tags.setdefault('status', 'replaced')
        elif re.search(r"введ(ен|ена)\s+впервые|действующ", tl):
            tags.setdefault('status', 'active')
    elif mt == 'proceedings':
        m = re.search(r"(материалы\s+конференции|proceedings\s+of\s+the|international\s+conference|symposium\s+on|workshop\s+on)[^\n\r]{0,120}", tl, flags=re.I)
        if m:
            tags.setdefault('conference', m.group(0).strip().title())
    elif mt == 'report':
        if re.search(r"техническое\s+задание\b|ТЗ\b", tl):
            tags.setdefault('doc_kind', 'Техническое задание')
        if re.search(r"пояснительная\s+записка\b", tl):
            tags.setdefault('doc_kind', 'Пояснительная записка')
    elif mt == 'patent':
        m = re.search(r"\b(?:RU|SU|US|WO|EP)\s?\d{4,10}[A-Z]?\d?\b", t)
        if m:
            tags.setdefault('patent_no', m.group(0))
        m = re.search(r"\b[A-H][0-9]{2}[A-Z]\s*\d+\/\d+\b", t)
        if m:
            tags.setdefault('ipc', m.group(0).replace(' ', ''))
    elif mt == 'presentation':
        if re.search(r"(слайды|slides|powerpoint|презентация)", tl):
            tags.setdefault('slides', 'yes')
    return tags

def _fw_alias_to_repo(ref: str) -> str | None:
    r = (ref or '').strip().lower()
    # Распространённые псевдонимы
    alias = {
        'tiny': 'Systran/faster-whisper-tiny',
        'base': 'Systran/faster-whisper-base',
        'small': 'Systran/faster-whisper-small',
        'medium': 'Systran/faster-whisper-medium',
        'large-v2': 'Systran/faster-whisper-large-v2',
        'large-v3': 'Systran/faster-whisper-large-v3',
        # Сжатые английские варианты
        'distil-small.en': 'Systran/faster-distil-whisper-small.en',
        'distil-medium.en': 'Systran/faster-distil-whisper-medium.en',
        'distil-large-v2': 'Systran/faster-distil-whisper-large-v2',
    }
    if r in alias:
        return alias[r]
    # Также принимаем форму org/name
    if '/' in r and len(r.split('/', 1)[0]) > 0:
        return ref
    return None

def _fw_model_dir_is_valid(path: Path) -> bool:
    """Проверить, что каталог модели содержит полноценные бинарники faster-whisper.
    Записываются модели в формате CTranslate2: требуется как минимум один .bin-файл весом > 1MB.
    Это помогает отфильтровать git-lfs плейсхолдеры (134 байта) и незавершённые выгрузки.
    """
    try:
        if not path or not path.exists() or not path.is_dir():
            return False
        for candidate in path.glob("*.bin*"):
            # Игнорируем плейсхолдеры git-lfs: у настоящего файла размер десятки мегабайт.
            if candidate.is_file() and candidate.stat().st_size >= 1_000_000:
                return True
    except Exception:
        return False
    return False


def _ensure_faster_whisper_model(model_ref: str) -> str:
    """Гарантировать локальный путь к модели faster-whisper.
    Если указан локальный каталог — используем его. Если это алиас или repo id — скачиваем в кэш.
    Возвращает путь к модели или пустую строку при ошибке.
    """
    try:
        if not model_ref:
            return ''
        p = Path(model_ref).expanduser()
        if p.exists() and p.is_dir():
            if _fw_model_dir_is_valid(p):
                return str(p)
            app.logger.warning(
                "Каталог модели faster-whisper найден (%s), но выглядит неполным — попробуем скачать заново.",
                p,
            )
        repo = _fw_alias_to_repo(model_ref)
        if not repo:
            return ''
        if hf_snapshot_download is None:
            app.logger.warning("Пакет huggingface_hub не установлен — автозагрузка модели faster-whisper недоступна")
            return ''
        # вычислить целевую директорию внутри кэша
        FW_CACHE_DIR.mkdir(parents=True, exist_ok=True)
        safe_name = repo.replace('/', '__')
        target_dir = FW_CACHE_DIR / safe_name
        if not _fw_model_dir_is_valid(target_dir):
            # загрузить/снять снимок в target_dir
            hf_snapshot_download(repo_id=repo, local_dir=str(target_dir), local_dir_use_symlinks=False, revision=None)
        if not _fw_model_dir_is_valid(target_dir):
            app.logger.warning(
                "Не удалось получить полноценную модель faster-whisper из репозитория %s — проверьте доступ к HuggingFace или загрузите модель вручную.",
                repo,
            )
            return ''
        return str(target_dir)
    except Exception as e:
        app.logger.warning(f"Failed to resolve faster-whisper model '{model_ref}': {e}")
        return ''

app = Flask(__name__)

DOC_CHAT_LOCK = threading.Lock()
DOC_CHAT_SESSIONS: dict[str, dict[str, Any]] = {}
DOC_CHAT_TTL_SECONDS = 6 * 3600  # 6 часов
DOC_CHAT_IMAGE_LIMIT = 0  # 0 — без ограничения количества изображений
DOC_CHAT_TEXT_LIMIT = 0  # 0 — индексация полного текста документа
DOC_CHAT_HISTORY_LIMIT = 12
DOC_CHAT_MODE_CONFIG: dict[str, dict[str, Any]] = {
    'default': {
        'label': 'Стандартный',
        'instruction': 'Используй сбалансированный режим: комбинируй краткие цитаты и пояснения, обязательно указывай ссылки вида «Текст N».',
        'max_chunks': 4,
        'score_threshold_factor': 0.45,
        'low_score_threshold_factor': 0.6,
        'min_threshold': 0.05,
        'fallback_chunks': 2,
        'image_limit': 3,
    },
    'quote': {
        'label': 'Цитата',
        'instruction': 'Режим «Цитата»: найди точный фрагмент текста, цитируй дословно и обязательно укажи страницу и ссылку на «Текст N». Если данных мало, честно укажи это.',
        'max_chunks': 3,
        'score_threshold_factor': 0.6,
        'low_score_threshold_factor': 0.75,
        'min_threshold': 0.1,
        'fallback_chunks': 2,
        'require_terms': True,
        'image_limit': 2,
    },
    'overview': {
        'label': 'Структурный обзор',
        'instruction': 'Режим «Структурный обзор»: сделай сжатый конспект по ключевым разделам, перечисли основные идеи и связи между ними.',
        'max_chunks': 6,
        'score_threshold_factor': 0.3,
        'low_score_threshold_factor': 0.45,
        'min_threshold': 0.0,
        'fallback_chunks': 4,
        'image_limit': 3,
    },
    'formulas': {
        'label': 'Формулы и графики',
        'instruction': 'Режим «Формулы и графики»: сосредоточься на формулах, графиках и других технических деталях. Если приводишь формулу, сохрани оригинальное написание.',
        'max_chunks': 5,
        'score_threshold_factor': 0.4,
        'low_score_threshold_factor': 0.55,
        'min_threshold': 0.02,
        'fallback_chunks': 3,
        'prefer_math_chunks': True,
        'prefer_technical_images': True,
        'image_limit': 2,
    },
}
DOC_CHAT_MODE_ALIASES: dict[str, str] = {
    '': 'default',
    'standard': 'default',
    'стандарт': 'default',
    'цитата': 'quote',
    'citation': 'quote',
    'quotes': 'quote',
    'overview': 'overview',
    'structure': 'overview',
    'структурный': 'overview',
    'формулы': 'formulas',
    'формулы_и_графики': 'formulas',
    'math': 'formulas',
    'graphs': 'formulas',
}
DOC_CHAT_FORMULA_KEYWORDS = (
    'формул', 'уравн', 'уравнение', 'уравнения', 'интеграл', 'сумма', '∑', '∫', '√', '≈', 'график', 'графика', 'графиков',
    'diagram', 'chart', 'plot', 'formula', 'equation', 'matrix', 'матриц', 'теорема', 'lemma', 'proof'
)
DOC_CHAT_FORMULA_PATTERN = re.compile(r"(?:\\[a-zA-Z]+|[А-Яа-яA-Za-z0-9]+?\s*(?:=|≥|≤|≈)\s*[А-Яа-яA-Za-z0-9]+?|[∑∫√±×÷])")
DOC_CHAT_TECH_IMAGE_KEYWORDS = (
    'формул', 'equation', 'formula', 'graph', 'график', 'diagram', 'chart', 'plot', 'figure', 'матриц', 'схема', 'schema'
)
DOC_CHAT_CITATION_RE = re.compile(r"(текст|text|изображение|image)\s+(\d+)", re.IGNORECASE)
DOC_CHAT_TONE_OPTIONS: dict[str, dict[str, Any]] = {
    'neutral': {
        'label': 'Нейтральный',
        'instruction': 'Сохраняй спокойный и официально-нейтральный тон, избегай эмоциональных оценок.',
    },
    'academic': {
        'label': 'Академичный',
        'instruction': 'Пиши академично: используй строгий стиль, подчёркивай термины и аккуратно цитируй источники.',
    },
    'friendly': {
        'label': 'Дружелюбный',
        'instruction': 'Объясняй дружелюбно и простыми словами, добавляй короткие контекстуальные пояснения.',
    },
}
DOC_CHAT_DETAIL_OPTIONS: dict[str, dict[str, Any]] = {
    'brief': {
        'label': 'Кратко',
        'instruction': 'Делай сжатый ответ: до 3 ключевых мыслей, без лишних подробностей.',
        'max_tokens': 650,
    },
    'balanced': {
        'label': 'Стандартно',
        'instruction': 'Держись баланса между краткостью и деталями, добавляй краткие цитаты и выводы.',
        'max_tokens': 900,
    },
    'deep': {
        'label': 'Подробно',
        'instruction': 'Раскрывай детали максимально полно, перечисляй ключевые аргументы и числовые данные.',
        'max_tokens': 1100,
    },
}
DOC_CHAT_LANGUAGE_OPTIONS: dict[str, dict[str, Any]] = {
    'ru': {
        'label': 'Русский',
        'instruction': 'Отвечай на русском языке, даже если вопрос задан иначе.',
    },
    'en': {
        'label': 'English',
        'instruction': 'Answer in English regardless of the question language.',
    },
    'auto': {
        'label': 'Авто',
        'instruction': 'Определи язык вопроса и отвечай на нём; если язык неясен, используй русский.',
    },
}
DOC_CHAT_PREFERENCE_DEFAULTS: dict[str, str] = {
    'tone': 'neutral',
    'detail': 'balanced',
    'language': 'ru',
}
DOC_CHAT_PREFERENCE_OPTIONS: dict[str, dict[str, dict[str, Any]]] = {
    'tone': DOC_CHAT_TONE_OPTIONS,
    'detail': DOC_CHAT_DETAIL_OPTIONS,
    'language': DOC_CHAT_LANGUAGE_OPTIONS,
}


def _doc_chat_now() -> float:
    return time.time()


def _doc_chat_cache_dir(session_id: str) -> Path:
    static_root = Path(app.static_folder or '.')
    target = static_root / 'cache' / 'doc_chat' / session_id
    try:
        target.mkdir(parents=True, exist_ok=True)
    except Exception:
        pass
    return target


def _doc_chat_per_file_dir(file_id: int) -> Path:
    static_root = Path(app.static_folder or '.')
    target = static_root / 'cache' / 'doc_chat' / 'files' / str(file_id)
    try:
        target.mkdir(parents=True, exist_ok=True)
    except Exception:
        pass
    return target


def _doc_chat_clear_per_file_dir(file_id: int) -> None:
    target = _doc_chat_per_file_dir(file_id)
    for child in target.iterdir():
        try:
            if child.is_dir():
                shutil.rmtree(child, ignore_errors=True)
            else:
                child.unlink(missing_ok=True)
        except Exception:
            continue


def _doc_chat_publish_images(file_obj: File, images: list[dict[str, Any]]) -> list[dict[str, Any]]:
    if not images:
        return []
    target_dir = _doc_chat_per_file_dir(file_obj.id)
    _doc_chat_clear_per_file_dir(file_obj.id)
    sanitized: list[dict[str, Any]] = []
    static_root = Path(app.static_folder or '.')
    for index, image in enumerate(images, start=1):
        entry = dict(image)
        candidates: list[Path] = []
        source_path = entry.get('path')
        if isinstance(source_path, str):
            candidates.append(Path(source_path))
        rel_path = entry.get('rel_path')
        if isinstance(rel_path, str):
            candidates.append(static_root / rel_path)
        dest_name = Path(rel_path or '').name or f"image_{index:02d}"
        dest_suffix = Path(dest_name).suffix or '.png'
        dest_file = target_dir / f"{index:02d}{dest_suffix}"
        copied = False
        for candidate in candidates:
            try:
                if candidate.exists():
                    shutil.copy2(candidate, dest_file)
                    copied = True
                    break
            except Exception:
                continue
        if not copied:
            continue
        rel_copy = f"cache/doc_chat/files/{file_obj.id}/{dest_file.name}"
        entry['rel_path'] = rel_copy
        entry['url'] = f"/static/{rel_copy}"
        entry.pop('path', None)
        sanitized.append(entry)
    return sanitized


def _doc_chat_prune(now: Optional[float] = None) -> None:
    now = now or _doc_chat_now()
    expired: list[tuple[str, Optional[str]]] = []
    with DOC_CHAT_LOCK:
        for sid, session in list(DOC_CHAT_SESSIONS.items()):
            updated = float(session.get('updated_at') or session.get('created_at') or now)
            if now - updated > DOC_CHAT_TTL_SECONDS:
                cache_path = session.get('cache_path')
                DOC_CHAT_SESSIONS.pop(sid, None)
                expired.append((sid, cache_path if isinstance(cache_path, str) else None))
    for _sid, cache_path in expired:
        if cache_path:
            try:
                shutil.rmtree(Path(cache_path), ignore_errors=True)
            except Exception:
                pass


def _doc_chat_get_cache_path(session_id: str) -> Optional[Path]:
    with DOC_CHAT_LOCK:
        session = DOC_CHAT_SESSIONS.get(session_id)
        cache_path = session.get('cache_path') if session else None
    if cache_path:
        try:
            path = Path(cache_path)
            path.mkdir(parents=True, exist_ok=True)
            return path
        except Exception:
            return None
    try:
        return _doc_chat_cache_dir(session_id)
    except Exception:
        return None


def _doc_chat_create_session(user: User, file_obj: File) -> dict[str, Any]:
    try:
        app.logger.info("[doc-chat] create_session start user=%s file=%s", user.id, file_obj.id)
    except Exception:
        pass
    session_id = uuid.uuid4().hex
    cache_dir = _doc_chat_cache_dir(session_id)
    try:
        app.logger.info("[doc-chat:%s] cache dir resolved %s", session_id[:8], cache_dir)
    except Exception:
        pass
    now = _doc_chat_now()
    collection = getattr(file_obj, 'collection', None)
    file_meta = {
        'id': file_obj.id,
        'title': (file_obj.title or file_obj.filename or file_obj.rel_path or f"Файл {file_obj.id}").strip(),
        'author': file_obj.author,
        'year': file_obj.year,
        'material_type': file_obj.material_type,
        'collection_id': getattr(collection, 'id', None),
        'collection_name': getattr(collection, 'name', None),
        'rel_path': file_obj.rel_path,
        'filename': file_obj.filename,
    }
    app.logger.info("[doc-chat:%s] building session payload", session_id[:8])
    session_payload: dict[str, Any] = {
        'id': session_id,
        'user_id': user.id,
        'file_id': file_obj.id,
        'file_meta': file_meta,
        'status': 'queued',
        'progress': [],
        'percent': 0.0,
        'error': None,
        'data': None,
        'history': [],
        'phase_history': [],
        'image_filter_summary': {},
        'created_at': now,
        'updated_at': now,
        'cache_path': str(cache_dir),
        'cache_rel': f"cache/doc_chat/{session_id}",
    }
    app.logger.info("[doc-chat:%s] payload prepared, pruning before lock", session_id[:8])
    _doc_chat_prune(now)
    app.logger.info("[doc-chat:%s] prune finished, acquiring lock", session_id[:8])
    with DOC_CHAT_LOCK:
        DOC_CHAT_SESSIONS[session_id] = session_payload
    app.logger.info("[doc-chat:%s] payload stored", session_id[:8])
    try:
        app.logger.info(
            "[doc-chat:%s] create_session done status=%s cache_path=%s",
            session_id[:8],
            session_payload['status'],
            session_payload.get('cache_path'),
        )
    except Exception:
        pass
    return session_payload


def _doc_chat_progress(session_id: str, line: str, *, percent: float | None = None) -> None:
    now = _doc_chat_now()
    with DOC_CHAT_LOCK:
        session = DOC_CHAT_SESSIONS.get(session_id)
        if not session:
            return
        progress = session.setdefault('progress', [])
        progress.append(str(line))
        if len(progress) > 400:
            session['progress'] = progress[-400:]
        if percent is not None:
            session['percent'] = max(0.0, min(100.0, float(percent)))
        session['updated_at'] = now
    try:
        label = f"[doc-chat:{session_id[:8]}]"
        if percent is not None:
            app.logger.info("%s %s (%.1f%%)", label, line, float(percent))
        else:
            app.logger.info("%s %s", label, line)
    except Exception:
        pass


def _doc_chat_record_phase(session_id: str, label: str, *, details: Optional[str] = None) -> None:
    now = _doc_chat_now()
    with DOC_CHAT_LOCK:
        session = DOC_CHAT_SESSIONS.get(session_id)
        if not session:
            return
        history = session.setdefault('phase_history', [])
        history.append({'label': label, 'details': details, 'ts': now})
        if len(history) > 6:
            session['phase_history'] = history[-6:]
        session['updated_at'] = now


def _doc_chat_set_status(session_id: str, status: str, *, error: Optional[str] = None, percent: float | None = None) -> None:
    now = _doc_chat_now()
    with DOC_CHAT_LOCK:
        session = DOC_CHAT_SESSIONS.get(session_id)
        if not session:
            return
        session['status'] = status
        session['error'] = error
        if percent is not None:
            session['percent'] = max(0.0, min(100.0, float(percent)))
        session['updated_at'] = now


def _doc_chat_store_data(session_id: str, data: dict[str, Any]) -> None:
    with DOC_CHAT_LOCK:
        session = DOC_CHAT_SESSIONS.get(session_id)
        if not session:
            return
        session['data'] = data
        session['updated_at'] = _doc_chat_now()


def _doc_chat_cache_is_valid(file_obj: File, cache: Optional[DocChatCache]) -> bool:
    if cache is None or not cache.data:
        return False
    try:
        file_sha1 = (file_obj.sha1 or '').strip()
        cache_sha1 = (cache.file_sha1 or '').strip()
        if file_sha1 and cache_sha1 and file_sha1 != cache_sha1:
            return False
        file_mtime = float(file_obj.mtime or 0.0)
        cache_mtime = float(cache.file_mtime or 0.0)
        if file_mtime and cache_mtime and file_mtime > cache_mtime + 1e-6:
            return False
    except Exception:
        return False
    payload = cache.data or {}
    chunk_total = int(payload.get('chunk_count') or cache.chunk_count or 0)
    if chunk_total <= 0:
        return False
    images = payload.get('images') or []
    static_root = Path(app.static_folder or '.')
    if images:
        for image in images:
            if not isinstance(image, dict):
                return False
            description = str(image.get('description') or '').strip()
            if not description:
                return False
            rel_path = image.get('rel_path')
            if rel_path:
                try:
                    if not (static_root / rel_path).exists():
                        return False
                except Exception:
                    return False
    embedding = payload.get('embedding') or {}
    if not embedding:
        return False
    if not cache.document_id:
        return False
    return True


def _doc_chat_load_cache_payload(file_obj: File) -> Optional[dict[str, Any]]:
    cache: Optional[DocChatCache] = getattr(file_obj, 'doc_chat_cache', None)
    if cache is None:
        cache = DocChatCache.query.get(file_obj.id)
    if not _doc_chat_cache_is_valid(file_obj, cache):
        return None
    payload = copy.deepcopy(cache.data or {})
    payload['chunk_count'] = int(payload.get('chunk_count') or cache.chunk_count or 0)
    payload['image_count'] = int(payload.get('image_count') or cache.image_count or len(payload.get('images') or []))
    if cache.text_size:
        payload['text_size'] = cache.text_size
    payload.setdefault('document_id', cache.document_id)
    static_root = Path(app.static_folder or '.')
    normalized_images: list[dict[str, Any]] = []
    for image in payload.get('images') or []:
        if not isinstance(image, dict):
            continue
        rel_path = image.get('rel_path')
        if rel_path:
            path = static_root / rel_path
            if not path.exists():
                return None
        normalized = dict(image)
        if rel_path:
            normalized['url'] = normalized.get('url') or f"/static/{rel_path}"
        normalized.pop('path', None)
        normalized_images.append(normalized)
    payload['images'] = normalized_images
    return payload


def _doc_chat_store_cache(file_obj: File, session_data: dict[str, Any], variant: dict[str, Any], *, vision_enabled: bool) -> None:
    payload = copy.deepcopy(session_data)
    images = payload.get('images') or []
    for image in images:
        if isinstance(image, dict):
            image.pop('path', None)
    record = DocChatCache.query.get(file_obj.id)
    if record is None:
        record = DocChatCache(file_id=file_obj.id, document_id=int(payload.get('document_id') or 0))
    record.document_id = int(payload.get('document_id') or 0)
    record.chunk_count = int(payload.get('chunk_count') or 0)
    record.image_count = int(payload.get('image_count') or len(images))
    record.text_size = int(payload.get('text_size') or 0)
    record.embedding_backend = (variant.get('backend') if isinstance(variant, dict) else None)
    record.embedding_model = (variant.get('model_name') if isinstance(variant, dict) else None)
    record.embedding_dim = None
    try:
        if isinstance(variant, dict) and variant.get('dim') is not None:
            record.embedding_dim = int(variant.get('dim') or 0)
    except Exception:
        record.embedding_dim = None
    record.vision_enabled = bool(vision_enabled)
    record.file_mtime = file_obj.mtime
    record.file_sha1 = file_obj.sha1
    record.data = payload
    try:
        db.session.merge(record)
        db.session.commit()
    except Exception as exc:
        db.session.rollback()
        app.logger.warning("doc-chat cache store failed for file %s: %s", file_obj.id, exc)


def _doc_chat_prepare_cache_for_file(file_id: int, user_id: Optional[int], allowed_ids: Optional[Sequence[int]]) -> bool:
    file_obj = File.query.get(file_id)
    if not file_obj:
        raise RuntimeError(f"Документ {file_id} не найден")
    if allowed_ids is not None and file_obj.collection_id not in allowed_ids:
        raise PermissionError("Нет доступа к документу")
    user_obj = User.query.get(user_id) if user_id else None
    if user_obj is None:
        raise RuntimeError("Не удалось определить пользователя для подготовки")
    session_payload = _doc_chat_create_session(user_obj, file_obj)
    try:
        _doc_chat_prepare_worker(
            session_payload['id'],
            file_obj.id,
            user_obj.id if getattr(user_obj, 'id', None) is not None else None,
            list(allowed_ids) if allowed_ids is not None else None,
        )
        return True
    finally:
        cache_path = session_payload.get('cache_path')
        with DOC_CHAT_LOCK:
            DOC_CHAT_SESSIONS.pop(session_payload['id'], None)
        if cache_path:
            try:
                shutil.rmtree(Path(cache_path), ignore_errors=True)
            except Exception:
                pass

def _doc_chat_extract_highlights(text: str, terms: Sequence[str], limit: int = 3) -> list[str]:
    """Return up to `limit` sentences from `text` that mention query terms."""
    text = (text or '').strip()
    if not text:
        return []
    sentences = re.split(r'(?<=[.!?])\s+', text)
    highlights: list[str] = []
    lowered_terms = [t for t in (terms or []) if len(t) >= 3]
    for sentence in sentences:
        normalized = sentence.strip()
        if not normalized:
            continue
        lower = normalized.lower()
        if lowered_terms:
            if any(term in lower for term in lowered_terms):
                highlights.append(normalized)
        elif len(normalized) > 40:
            highlights.append(normalized)
        if len(highlights) >= limit:
            break
    return highlights[:limit]


def _doc_chat_append_history(session_id: str, entry: dict[str, Any]) -> None:
    now = _doc_chat_now()
    with DOC_CHAT_LOCK:
        session = DOC_CHAT_SESSIONS.get(session_id)
        if not session:
            return
        history = session.setdefault('history', [])
        history.append(dict(entry))
        if len(history) > DOC_CHAT_HISTORY_LIMIT:
            session['history'] = history[-DOC_CHAT_HISTORY_LIMIT:]
        session['updated_at'] = now


def _doc_chat_snapshot(session_id: str) -> Optional[dict[str, Any]]:
    with DOC_CHAT_LOCK:
        session = DOC_CHAT_SESSIONS.get(session_id)
        if not session:
            return None
        snapshot = copy.deepcopy(session)
    data = snapshot.get('data')
    if isinstance(data, dict):
        images = data.get('images')
        if isinstance(images, list):
            for item in images:
                if isinstance(item, dict):
                    item.pop('vector', None)
                    item.pop('path', None)
    snapshot.pop('user_id', None)
    snapshot.pop('cache_path', None)
    return snapshot


def _doc_chat_public_session(session_id: str) -> Optional[dict[str, Any]]:
    return _doc_chat_snapshot(session_id)


def _doc_chat_internal_session(session_id: str) -> Optional[dict[str, Any]]:
    with DOC_CHAT_LOCK:
        session = DOC_CHAT_SESSIONS.get(session_id)
        if not session:
            return None
        return copy.deepcopy(session)


def _doc_chat_owner_id(session_id: str) -> Optional[int]:
    with DOC_CHAT_LOCK:
        session = DOC_CHAT_SESSIONS.get(session_id)
        if not session:
            return None
        try:
            return int(session.get('user_id'))
        except Exception:
            return None


def _doc_chat_resolve_mode(mode_name: Optional[str]) -> tuple[str, dict[str, Any]]:
    key = (mode_name or '').strip().lower()
    resolved = DOC_CHAT_MODE_ALIASES.get(key, key or 'default')
    if resolved not in DOC_CHAT_MODE_CONFIG:
        resolved = 'default'
    config = dict(DOC_CHAT_MODE_CONFIG[resolved])
    runtime = _rt()
    max_chunks_override = getattr(runtime, 'doc_chat_max_chunks', 0) or 0
    fallback_override = getattr(runtime, 'doc_chat_fallback_chunks', 0) or 0
    if max_chunks_override > 0:
        config['max_chunks'] = max(1, int(max_chunks_override))
    if fallback_override > 0:
        config['fallback_chunks'] = max(1, int(fallback_override))
    return resolved, config


def _doc_chat_trim_context(text: Optional[str], limit: int = 600) -> Optional[str]:
    if not text:
        return None
    cleaned = re.sub(r'\s+', ' ', text).strip()
    if not cleaned:
        return None
    if limit > 0 and len(cleaned) > limit:
        trimmed = cleaned[:limit]
        last_space = trimmed.rfind(' ')
        if last_space > 40:  # avoid chopping words too aggressively
            trimmed = trimmed[:last_space]
        cleaned = trimmed.rstrip() + '…'
    return cleaned


def _doc_chat_snippet_has_math(text: Optional[str]) -> bool:
    if not text:
        return False
    lowered = text.lower()
    if any(token in lowered for token in DOC_CHAT_FORMULA_KEYWORDS):
        return True
    return bool(DOC_CHAT_FORMULA_PATTERN.search(text))


def _doc_chat_safe_int(value: Any) -> Optional[int]:
    try:
        return int(value)
    except Exception:
        return None


def _doc_chat_image_min_settings() -> tuple[int, int]:
    runtime = _rt()
    min_width = max(0, int(getattr(runtime, 'doc_chat_image_min_width', 0) or 0))
    min_height = max(0, int(getattr(runtime, 'doc_chat_image_min_height', 0) or 0))
    return min_width, min_height


def _doc_chat_image_meets_min_size(width: Optional[int], height: Optional[int], min_width: int, min_height: int) -> bool:
    if min_width <= 0 and min_height <= 0:
        return True
    if width is None or height is None:
        return True
    if min_width > 0 and width < min_width:
        return False
    if min_height > 0 and height < min_height:
        return False
    return True


def _doc_chat_image_is_technical(image: dict[str, Any]) -> bool:
    description = str(image.get('description') or '').lower()
    keywords = ' '.join(str(k).lower() for k in (image.get('keywords') or []) if k)
    before_ctx = str(image.get('context_before') or '').lower()
    after_ctx = str(image.get('context_after') or '').lower()
    haystack = ' '.join((description, keywords, before_ctx, after_ctx))
    return any(token in haystack for token in DOC_CHAT_TECH_IMAGE_KEYWORDS)


def _doc_chat_history_context(history: Sequence[dict[str, Any]], *, max_pairs: int = 3, max_chars: int = 1200) -> str:
    if not history or max_pairs <= 0:
        return ""
    segments: list[str] = []
    user_turns = 0
    for turn in reversed(history):
        role = turn.get('role')
        if role not in {'user', 'assistant'}:
            continue
        content = str(turn.get('content') or '').strip()
        if not content:
            continue
        role_label = 'Пользователь' if role == 'user' else 'Помощник'
        mode_key = str(turn.get('mode') or '').strip()
        mode_label = DOC_CHAT_MODE_CONFIG.get(mode_key, {}).get('label')
        if mode_label:
            role_label = f"{role_label} ({mode_label})"
        segments.append(f"{role_label}: {content}")
        if role == 'user':
            user_turns += 1
            if user_turns >= max_pairs:
                break
    if not segments:
        return ""
    segments.reverse()
    text = "\n".join(segments)
    if len(text) > max_chars:
        text = text[-max_chars:]
        first_newline = text.find('\n')
        if 0 < first_newline < len(text) - 20:
            text = text[first_newline + 1 :]
    return text.strip()


def _doc_chat_recent_chunk_ids(history: Sequence[dict[str, Any]], *, limit: int = 4) -> list[int]:
    if not history or limit <= 0:
        return []
    chunk_ids: list[int] = []
    seen: set[int] = set()
    for turn in reversed(history):
        if turn.get('role') != 'assistant':
            continue
        sources = turn.get('sources') or {}
        text_sources = sources.get('texts') or []
        for source in text_sources:
            chunk_id = source.get('chunk_id')
            try:
                chunk_id_int = int(chunk_id)
            except Exception:
                continue
            if chunk_id_int in seen:
                continue
            seen.add(chunk_id_int)
            chunk_ids.append(chunk_id_int)
            if len(chunk_ids) >= limit:
                return chunk_ids
    return chunk_ids


def _doc_chat_pref_options_payload() -> dict[str, list[dict[str, str]]]:
    payload: dict[str, list[dict[str, str]]] = {}
    for key, options in DOC_CHAT_PREFERENCE_OPTIONS.items():
        entries: list[dict[str, str]] = []
        for value_key, meta in options.items():
            entries.append({
                'id': value_key,
                'label': meta.get('label', value_key.title()),
                'description': meta.get('description') or meta.get('instruction') or '',
            })
        payload[key] = entries
    return payload


def _doc_chat_normalize_pref_value(key: str, value: Any) -> Optional[str]:
    if key not in DOC_CHAT_PREFERENCE_OPTIONS:
        return None
    options = DOC_CHAT_PREFERENCE_OPTIONS[key]
    value_str = str(value or '').strip().lower()
    if not value_str:
        return DOC_CHAT_PREFERENCE_DEFAULTS.get(key)
    if value_str not in options:
        return DOC_CHAT_PREFERENCE_DEFAULTS.get(key)
    return value_str


def _doc_chat_get_user_preferences(user: Optional[User]) -> dict[str, str]:
    prefs = dict(DOC_CHAT_PREFERENCE_DEFAULTS)
    if not user:
        return prefs
    try:
        rows: List[UserPreference] = UserPreference.query.filter(UserPreference.user_id == user.id).all()
    except Exception:
        return prefs
    for row in rows:
        normalized = _doc_chat_normalize_pref_value(row.key, row.value)
        if normalized:
            prefs[row.key] = normalized
    return prefs


def _doc_chat_set_user_preferences(user: User, updates: Mapping[str, Any]) -> dict[str, str]:
    if not user:
        raise RuntimeError("user required for preferences update")
    changed = False
    for key, value in updates.items():
        if key not in DOC_CHAT_PREFERENCE_OPTIONS:
            continue
        normalized = _doc_chat_normalize_pref_value(key, value)
        if normalized is None:
            continue
        default_value = DOC_CHAT_PREFERENCE_DEFAULTS.get(key)
        current_pref = UserPreference.query.filter_by(user_id=user.id, key=key).one_or_none()
        if normalized == default_value:
            if current_pref:
                db.session.delete(current_pref)
                changed = True
            continue
        if current_pref:
            if current_pref.value != normalized:
                current_pref.value = normalized
                changed = True
        else:
            db.session.add(UserPreference(user_id=user.id, key=key, value=normalized))
            changed = True
    if changed:
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()
            raise
    return _doc_chat_get_user_preferences(user)


def _doc_chat_pdf_block_text(block: dict[str, Any]) -> str:
    if not block or block.get('type') != 0:
        return ''
    lines = block.get('lines') or []
    parts: list[str] = []
    for line in lines:
        spans = line.get('spans') or []
        span_text = ''.join(str(span.get('text') or '') for span in spans).strip()
        if span_text:
            parts.append(span_text)
    return '\n'.join(parts).strip()


def _doc_chat_pdf_neighbor_text(blocks: Sequence[dict[str, Any]], start_idx: int, direction: int) -> Optional[str]:
    idx = start_idx + direction
    while 0 <= idx < len(blocks):
        block = blocks[idx]
        if isinstance(block, dict) and block.get('type') == 0:
            text = _doc_chat_trim_context(_doc_chat_pdf_block_text(block))
            if text:
                return text
        idx += direction
    return None


def _doc_chat_docx_image_context(zf: "zipfile.ZipFile") -> dict[str, dict[str, Optional[str]]]:
    contexts: dict[str, dict[str, Optional[str]]] = {}
    try:
        document_xml = zf.read('word/document.xml')
        rels_xml = zf.read('word/_rels/document.xml.rels')
    except KeyError:
        return contexts
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'rel': 'http://schemas.openxmlformats.org/package/2006/relationships',
    }
    rels_root = ET.fromstring(rels_xml)
    rels_map: dict[str, str] = {}
    for rel in rels_root.findall('rel:Relationship', ns):
        rel_id = rel.attrib.get('Id')
        target = rel.attrib.get('Target')
        if not rel_id or not target:
            continue
        normalized = target.replace('\\', '/').lstrip('./')
        if not normalized.startswith('word/'):
            normalized = f"word/{normalized}"
        rels_map[rel_id] = normalized

    doc_root = ET.fromstring(document_xml)
    paragraphs = doc_root.findall('.//w:p', ns)
    paragraph_texts: list[str] = []
    paragraph_images: list[list[str]] = []
    for paragraph in paragraphs:
        texts: list[str] = []
        for node in paragraph.findall('.//w:t', ns):
            val = node.text
            if val:
                texts.append(val)
        paragraph_text = _doc_chat_trim_context(' '.join(texts), limit=600) or ''
        paragraph_texts.append(paragraph_text)
        images_in_paragraph: list[str] = []
        for blip in paragraph.findall('.//a:blip', ns):
            embed = blip.attrib.get(f"{{{ns['r']}}}embed")
            if not embed:
                continue
            target = rels_map.get(embed)
            if target:
                images_in_paragraph.append(target)
        paragraph_images.append(images_in_paragraph)

    total = len(paragraphs)
    for idx, targets in enumerate(paragraph_images):
        if not targets:
            continue
        before = None
        for j in range(idx - 1, -1, -1):
            if paragraph_texts[j]:
                before = paragraph_texts[j]
                break
        after = None
        for j in range(idx + 1, total):
            if paragraph_texts[j]:
                after = paragraph_texts[j]
                break
        for target in targets:
            contexts.setdefault(target, {'before': before, 'after': after})
    return contexts


def _doc_chat_extract_pdf_images(
    path: Path,
    session_id: str,
    limit: int,
    *,
    min_width: int,
    min_height: int,
    stats: dict[str, int],
) -> list[dict[str, Any]]:
    images: list[dict[str, Any]] = []
    try:
        doc = fitz.open(path)
    except Exception:
        return images
    try:
        seen: set[int] = set()
        context_map: dict[int, tuple[Optional[str], Optional[str]]] = {}
        for page_index, page in enumerate(doc):
            try:
                raw = page.get_text('rawdict') or {}
                page_blocks = raw.get('blocks') or []
            except Exception:
                page_blocks = []
            if not page_blocks:
                page_blocks = []
            for block_index, block in enumerate(page_blocks):
                if not isinstance(block, dict) or block.get('type') != 1:
                    continue
                try:
                    xref = int(block.get('image') or block.get('xref') or block.get('number') or 0)
                except Exception:
                    xref = 0
                stats['candidates'] += 1
                if not xref or xref in seen:
                    continue
                try:
                    base_image = doc.extract_image(xref) or {}
                except Exception:
                    continue
                image_bytes = base_image.get('image')
                if not image_bytes:
                    continue
                ext = str(base_image.get('ext') or 'png').lower()
                cache_dir = _doc_chat_get_cache_path(session_id)
                if cache_dir is None:
                    continue
                filename = f"page{page_index + 1:03d}_img{len(images) + 1:02d}.{ext}"
                file_path = cache_dir / filename
                try:
                    file_path.write_bytes(image_bytes)
                except Exception:
                    continue
                rel_path = f"cache/doc_chat/{session_id}/{filename}"
                before_text = _doc_chat_pdf_neighbor_text(page_blocks, block_index, -1)
                after_text = _doc_chat_pdf_neighbor_text(page_blocks, block_index, 1)
                context_map[xref] = (before_text, after_text)
                width = _doc_chat_safe_int(base_image.get('width'))
                height = _doc_chat_safe_int(base_image.get('height'))
                if not _doc_chat_image_meets_min_size(width, height, min_width, min_height):
                    stats['skipped_size'] += 1
                    seen.add(xref)
                    continue
                images.append({
                    'page': page_index + 1,
                    'ext': ext,
                    'rel_path': rel_path,
                    'path': str(file_path),
                    'width': width,
                    'height': height,
                    'context_before': before_text,
                    'context_after': after_text,
                })
                stats['kept'] += 1
                seen.add(xref)
                if limit > 0 and len(images) >= limit:
                    return images
            try:
                page_images = page.get_images(full=True) or []
            except Exception:
                page_images = []
            for img in page_images:
                try:
                    xref = int(img[0])
                except Exception:
                    continue
                stats['candidates'] += 1
                if not xref or xref in seen:
                    continue
                try:
                    base_image = doc.extract_image(xref) or {}
                except Exception:
                    continue
                image_bytes = base_image.get('image')
                if not image_bytes:
                    continue
                cache_dir = _doc_chat_get_cache_path(session_id)
                if cache_dir is None:
                    continue
                ext = str(base_image.get('ext') or 'png').lower()
                filename = f"page{page_index + 1:03d}_img{len(images) + 1:02d}.{ext}"
                file_path = cache_dir / filename
                try:
                    file_path.write_bytes(image_bytes)
                except Exception:
                    continue
                rel_path = f"cache/doc_chat/{session_id}/{filename}"
                before_text, after_text = context_map.get(xref, (None, None))
                width = _doc_chat_safe_int(base_image.get('width'))
                height = _doc_chat_safe_int(base_image.get('height'))
                if not _doc_chat_image_meets_min_size(width, height, min_width, min_height):
                    stats['skipped_size'] += 1
                    seen.add(xref)
                    continue
                images.append({
                    'page': page_index + 1,
                    'ext': ext,
                    'rel_path': rel_path,
                    'path': str(file_path),
                    'width': width,
                    'height': height,
                    'context_before': before_text,
                    'context_after': after_text,
                })
                stats['kept'] += 1
                seen.add(xref)
                if limit > 0 and len(images) >= limit:
                    return images
    finally:
        try:
            doc.close()
        except Exception:
            pass
    return images

def _doc_chat_extract_docx_images(
    path: Path,
    session_id: str,
    limit: int,
    *,
    min_width: int,
    min_height: int,
    stats: dict[str, int],
) -> list[dict[str, Any]]:
    images: list[dict[str, Any]] = []
    try:
        import zipfile
        with zipfile.ZipFile(path) as zf:
            context_map = _doc_chat_docx_image_context(zf)
            for name in zf.namelist():
                if not name.startswith('word/media/'):
                    continue
                data = zf.read(name)
                if not data:
                    continue
                ext = Path(name).suffix.lstrip('.').lower() or 'png'
                width = None
                height = None
                if PILImage is not None:
                    try:
                        with PILImage.open(io.BytesIO(data)) as im:
                            w, h = im.size
                            width = int(w)
                            height = int(h)
                    except Exception:
                        width = None
                        height = None
                stats['candidates'] += 1
                cache_dir = _doc_chat_get_cache_path(session_id)
                if cache_dir is None:
                    continue
                if not _doc_chat_image_meets_min_size(width, height, min_width, min_height):
                    continue
                filename = f"docx{len(images) + 1:02d}.{ext}"
                file_path = cache_dir / filename
                try:
                    file_path.write_bytes(data)
                except Exception:
                    continue
                rel_path = f"cache/doc_chat/{session_id}/{filename}"
                normalized_name = name.replace('\\', '/')
                ctx = context_map.get(normalized_name)
                images.append({
                    'page': None,
                    'ext': ext,
                    'rel_path': rel_path,
                    'path': str(file_path),
                    'width': width,
                    'height': height,
                    'context_before': ctx.get('before') if ctx else None,
                    'context_after': ctx.get('after') if ctx else None,
                })
                stats['kept'] += 1
                if limit > 0 and len(images) >= limit:
                    break
    except Exception:
        return images
    return images


def _doc_chat_extract_images(source_path: Optional[Path], session_id: str, limit: int = DOC_CHAT_IMAGE_LIMIT) -> tuple[list[dict[str, Any]], dict[str, int]]:
    stats_template = {'candidates': 0, 'kept': 0, 'skipped_size': 0}
    min_width, min_height = _doc_chat_image_min_settings()
    stats_template['min_width'] = min_width
    stats_template['min_height'] = min_height
    if not source_path or not source_path.exists():
        return [], stats_template
    ext = source_path.suffix.lower()
    if ext == '.pdf':
        result = _doc_chat_extract_pdf_images(
            source_path,
            session_id,
            limit,
            min_width=min_width,
            min_height=min_height,
            stats=stats_template,
        )
        return result, stats_template
    if ext == '.docx':
        result = _doc_chat_extract_docx_images(
            source_path,
            session_id,
            limit,
            min_width=min_width,
            min_height=min_height,
            stats=stats_template,
        )
        return result, stats_template
    return [], stats_template


def _doc_chat_extract_image_ocr(image_path: Path, *, words_limit: int = 200) -> dict[str, Any]:
    if pytesseract is None:
        return {}
    try:
        data = pytesseract.image_to_data(
            str(image_path),
            output_type=pytesseract.Output.DICT,
            lang=os.getenv('OCR_LANGS', 'rus+eng'),
        )
    except Exception as exc:
        app.logger.debug("doc-chat: image OCR failed for %s: %s", image_path, exc)
        return {}
    words: list[dict[str, Any]] = []
    text_parts: list[str] = []
    total = int(data.get('level') and len(data.get('level'))) if isinstance(data, dict) else 0
    if total <= 0:
        total = len(data.get('text') or []) if isinstance(data, dict) else 0
    for idx in range(total):
        word = ""
        try:
            word = str((data.get('text') or [])[idx] or "").strip()
        except Exception:
            word = ""
        if not word:
            continue
        try:
            conf_raw = data.get('conf')[idx] if isinstance(data.get('conf'), list) else None
            conf = float(conf_raw) if conf_raw is not None else None
        except Exception:
            conf = None
        try:
            left = int((data.get('left') or [0])[idx])
            top = int((data.get('top') or [0])[idx])
            width = int((data.get('width') or [0])[idx])
            height = int((data.get('height') or [0])[idx])
        except Exception:
            left = top = width = height = 0
        words.append({
            'text': word,
            'confidence': conf,
            'bbox': [left, top, width, height],
        })
        text_parts.append(word)
        if len(words) >= words_limit:
            break
    ocr_text = " ".join(text_parts).strip()
    if not ocr_text:
        return {}
    preview = ocr_text[:400].strip()
    has_math = _doc_chat_snippet_has_math(ocr_text)
    return {
        'text': ocr_text,
        'preview': preview,
        'words': words,
        'has_math': has_math,
    }


def _doc_chat_enrich_images(
    images: list[dict[str, Any]],
    *,
    vision_enabled: bool,
    ocr_enabled: bool,
    session_id: Optional[str] = None,
) -> list[dict[str, Any]]:
    enriched: list[dict[str, Any]] = []
    total = len(images)
    processed = 0
    for image in images:
        entry = dict(image)
        rel_path = entry.get('rel_path')
        entry['url'] = f"/static/{rel_path}" if rel_path else None
        entry['vision_enabled'] = bool(vision_enabled)
        entry.setdefault('keywords', [])
        entry.setdefault('description', '')
        if vision_enabled and entry.get('path'):
            try:
                vis = call_lmstudio_vision(Path(entry['path']), Path(entry['path']).name)
                if isinstance(vis, dict):
                    entry['description'] = (vis.get('description') or '').strip()
                    kws = vis.get('keywords') or []
                    if isinstance(kws, list):
                        entry['keywords'] = [str(k) for k in kws[:16]]
            except Exception as exc:
                entry['vision_error'] = str(exc)
        if not entry.get('description'):
            entry['description'] = 'Описание не получено.' if vision_enabled else 'Анализ изображений отключён.'
        if ocr_enabled and entry.get('path'):
            try:
                ocr_payload = _doc_chat_extract_image_ocr(Path(entry['path']))
            except Exception as exc:
                ocr_payload = {}
                entry['ocr_error'] = str(exc)
            if ocr_payload:
                entry['ocr_text'] = ocr_payload.get('text')
                entry['ocr_preview'] = ocr_payload.get('preview')
                entry['ocr_words'] = ocr_payload.get('words')
            entry['ocr_has_math'] = bool(ocr_payload.get('has_math'))
        enriched.append(entry)
        processed += 1
        if session_id and total > 0:
            percent = 70 + min(15, (processed / total) * 15)
            _doc_chat_progress(session_id, f"Обработано изображений: {processed}/{total}", percent=percent)
    return enriched


def _doc_chat_collect_chunk_previews(document_id: int, limit: int = 3) -> list[dict[str, Any]]:
    if not document_id:
        return []
    try:
        rows = (
            db.session.query(RagDocumentChunk)
            .filter(RagDocumentChunk.document_id == document_id)
            .order_by(RagDocumentChunk.ordinal.asc())
            .limit(limit)
            .all()
        )
    except Exception:
        return []
    previews: list[dict[str, Any]] = []
    for chunk in rows:
        text = (chunk.preview or chunk.content or '')[:400]
        previews.append({
            'ordinal': chunk.ordinal,
            'section_path': chunk.section_path,
            'preview': text,
        })
    return previews


def _doc_chat_resolve_embedding_variant(document_id: int) -> Optional[dict[str, Any]]:
    if not document_id:
        return None
    try:
        row = (
            db.session.query(
                RagChunkEmbedding.model_name,
                RagChunkEmbedding.model_version,
                func.max(RagChunkEmbedding.dim).label('dim'),
                func.count(RagChunkEmbedding.id).label('count'),
            )
            .join(RagDocumentChunk, RagChunkEmbedding.chunk_id == RagDocumentChunk.id)
            .filter(RagDocumentChunk.document_id == document_id)
            .group_by(RagChunkEmbedding.model_name, RagChunkEmbedding.model_version)
            .order_by(func.count(RagChunkEmbedding.id).desc())
            .first()
        )
    except Exception:
        return None
    if not row:
        return None
    model_name, model_version, dim, _count = row
    return {
        'model_name': model_name,
        'model_version': model_version,
        'dim': int(dim or 0),
    }


def _doc_chat_cosine(vec_a: Sequence[float], vec_b: Sequence[float]) -> float:
    if not vec_a or not vec_b or len(vec_a) != len(vec_b):
        return 0.0
    dot = sum(a * b for a, b in zip(vec_a, vec_b))
    if dot == 0:
        return 0.0
    norm_a = math.sqrt(sum(a * a for a in vec_a))
    norm_b = math.sqrt(sum(b * b for b in vec_b))
    if norm_a == 0.0 or norm_b == 0.0:
        return 0.0
    return dot / (norm_a * norm_b)


def _doc_chat_prepare_worker(session_id: str, file_id: int, user_id: int | None, allowed_ids: Optional[Sequence[int]]) -> None:
    allowed_set = set(int(x) for x in allowed_ids) if allowed_ids is not None else None
    with app.app_context():
        try:
            db.session.rollback()
        except Exception:
            db.session.remove()
        try:
            session_snapshot = _doc_chat_internal_session(session_id)
            if session_snapshot is None:
                app.logger.warning("[doc-chat:%s] session missing before подготовки", session_id[:8])
                return
            app.logger.info("[doc-chat:%s] start preparation for file_id=%s", session_id[:8], file_id)
            _doc_chat_set_status(session_id, 'processing', percent=5)
            _doc_chat_progress(session_id, "Подготовка документа начата", percent=5)

            file_obj = File.query.get(file_id)
            if not file_obj:
                _doc_chat_progress(session_id, "Ошибка: файл не найден")
                _doc_chat_set_status(session_id, 'error', error='Файл не найден', percent=100)
                return
            if allowed_set is not None and file_obj.collection_id not in allowed_set:
                _doc_chat_progress(session_id, "Ошибка: нет доступа к выбранному документу")
                _doc_chat_set_status(session_id, 'error', error='Нет доступа к документу', percent=100)
                return

            runtime = _rt()
            text, source_path = _collect_text_for_rag(file_obj, limit_chars=DOC_CHAT_TEXT_LIMIT)
            if not text.strip():
                raise RuntimeError("Не удалось извлечь текст из документа.")
            _doc_chat_progress(session_id, f"Извлечение текста завершено ({len(text)} символов)", percent=20)
            _doc_chat_record_phase(session_id, 'Извлечение текста', details=f"{len(text)} символов")

            chunk_max_tokens = runtime.doc_chat_chunk_max_tokens or 700
            chunk_overlap = runtime.doc_chat_chunk_overlap or 120
            chunk_min_tokens = runtime.doc_chat_chunk_min_tokens or 80
            chunk_config = ChunkConfig(
                max_tokens=max(16, int(chunk_max_tokens)),
                overlap=max(0, int(chunk_overlap)),
                min_tokens=max(1, int(chunk_min_tokens)),
            )
            indexer = RagIndexer(chunk_config=chunk_config, normalizer_version='v1')
            ingest_result = indexer.ingest_document(
                file_obj,
                text,
                metadata={'source': 'doc_chat', 'session_id': session_id},
                skip_if_unchanged=True,
                commit=True,
            )
            if ingest_result.get('skipped'):
                _doc_chat_progress(session_id, "RAG: документ уже проиндексирован", percent=35)
            else:
                chunks_count = int(ingest_result.get('chunks') or 0)
                _doc_chat_progress(session_id, f"RAG: создано чанков {chunks_count}", percent=35)
                _doc_chat_record_phase(session_id, 'Чанкование', details=f"{chunks_count} чанков")

            document = file_obj.rag_document
            if not document or not document.is_ready_for_rag:
                document = RagDocument.query.filter_by(file_id=file_obj.id, is_ready_for_rag=True).first()
            if not document:
                raise RuntimeError("RAG-документ не сформирован.")

            embedding_backend = (runtime.rag_embedding_backend or 'auto').strip().lower() or 'auto'
            embedding_model = runtime.rag_embedding_model or 'intfloat/multilingual-e5-large'
            embedding_dim = int(getattr(runtime, 'rag_embedding_dim', 384) or 384)
            embedding_batch = max(1, int(getattr(runtime, 'rag_embedding_batch_size', 32) or 32))
            embedding_device = getattr(runtime, 'rag_embedding_device', None)
            embedding_endpoint = getattr(runtime, 'rag_embedding_endpoint', None) or runtime.lmstudio_api_base
            embedding_api_key = getattr(runtime, 'rag_embedding_api_key', None) or runtime.lmstudio_api_key

            backend: EmbeddingBackend | None = None
            embedding_fallback_reason: Optional[str] = None
            try:
                backend = load_embedding_backend(
                    embedding_backend,
                    model_name=embedding_model,
                    dim=embedding_dim,
                    batch_size=embedding_batch,
                    device=embedding_device,
                    base_url=embedding_endpoint or None,
                    api_key=embedding_api_key or None,
                )
            except Exception as exc:
                embedding_fallback_reason = str(exc)
                _doc_chat_progress(session_id, f"Эмбеддинги недоступны ({embedding_fallback_reason}), используем hash fallback", percent=50)
                backend = HashEmbeddingBackend(dim=embedding_dim or 384, normalize=True, model_name="doc-chat-hash")

            try:
                created_embeddings = _embed_missing_chunks_for_document(
                    document.id,
                    backend,
                    batch_size=embedding_batch,
                    commit=True,
                )
                embedding_info = {
                    'backend': embedding_backend,
                    'model_name': getattr(backend, 'model_name', embedding_model),
                    'model_version': getattr(backend, 'model_version', None),
                    'dim': int(getattr(backend, 'dim', embedding_dim) or embedding_dim),
                }
                if embedding_fallback_reason:
                    embedding_info['fallback'] = embedding_fallback_reason
            finally:
                try:
                    backend.close()
                except Exception:
                    pass

            if created_embeddings:
                _doc_chat_progress(session_id, f"Эмбеддинги обновлены: {created_embeddings}", percent=55)
            else:
                _doc_chat_progress(session_id, "Эмбеддинги актуальны", percent=55)
            _doc_chat_record_phase(session_id, 'Эмбеддинги', details=f"{created_embeddings or 0} новых")

            variant = _doc_chat_resolve_embedding_variant(document.id) or {
                'model_name': embedding_info['model_name'],
                'model_version': embedding_info['model_version'],
                'dim': embedding_info['dim'],
            }
            variant['backend'] = embedding_backend

            chunk_count = int(
                db.session.query(func.count(RagDocumentChunk.id))
                .filter(RagDocumentChunk.document_id == document.id)
                .scalar() or 0
            )
            _doc_chat_progress(session_id, f"Документ готов: чанков {chunk_count}", percent=65)

            if not source_path:
                for candidate in _resolve_candidate_paths(file_obj):
                    if candidate.exists():
                        source_path = candidate
                        break

            images_raw, image_stats = _doc_chat_extract_images(
                source_path if isinstance(source_path, Path) else None,
                session_id,
                DOC_CHAT_IMAGE_LIMIT,
            )
            vision_enabled = bool(getattr(runtime, 'images_vision_enabled', False))
            ocr_enabled = bool(getattr(runtime, 'images_ocr_enabled', pytesseract is not None)) and pytesseract is not None
            images_enriched = _doc_chat_enrich_images(
                images_raw,
                vision_enabled=vision_enabled,
                ocr_enabled=ocr_enabled,
                session_id=session_id,
            ) if images_raw else []
            if images_enriched:
                _doc_chat_progress(session_id, f"Обработано изображений: {len(images_enriched)}", percent=80)
            else:
                _doc_chat_progress(session_id, "Изображения не найдены", percent=70)
            skip_count = image_stats.get('skipped_size', 0)
            phase_details = [f"обработано {len(images_enriched)}"]
            if skip_count:
                phase_details.append(f"пропущено {skip_count} по размеру")
            if image_stats.get('min_width') or image_stats.get('min_height'):
                phase_details.append(f"min={image_stats.get('min_width')}×{image_stats.get('min_height')}")
            _doc_chat_record_phase(session_id, 'Изображения', details=', '.join(phase_details))

            previews = _doc_chat_collect_chunk_previews(document.id)

            images_for_session = _doc_chat_publish_images(file_obj, images_enriched)

            session_data = {
                'document_id': document.id,
                'chunk_count': chunk_count,
                'language': document.lang_primary,
                'embedding': variant,
                'images': images_for_session,
                'image_count': len(images_for_session),
                'preview_chunks': previews,
                'text_size': len(text),
                'vision_enabled': vision_enabled,
                'ocr_enabled': ocr_enabled,
                'image_filter_summary': {
                    **image_stats,
                    'processed': len(images_enriched),
                } if image_stats else {},
            }
            if source_path:
                session_data['source_name'] = source_path.name

            _doc_chat_store_data(session_id, session_data)
            _doc_chat_progress(session_id, "Документ подготовлен", percent=95)
            _doc_chat_set_status(session_id, 'ready', percent=100)

            try:
                _doc_chat_store_cache(file_obj, session_data, variant, vision_enabled=vision_enabled)
            except Exception:
                pass

            actor = None
            if user_id:
                try:
                    actor = User.query.get(int(user_id))
                except Exception:
                    actor = None
            try:
                detail = json.dumps({
                    'session_id': session_id,
                    'file_id': file_obj.id,
                    'chunk_count': chunk_count,
                    'image_count': len(images_enriched),
                    'model': variant.get('model_name'),
                }, ensure_ascii=False)
                _log_user_action(actor, 'doc_chat_prepare', 'file', file_obj.id, detail=detail[:2000])
            except Exception:
                pass
        except Exception as exc:
            msg = str(exc)
            app.logger.exception("doc-chat preparation failed for session %s: %s", session_id, msg)
            _doc_chat_progress(session_id, f"Ошибка подготовки: {msg}")
            _doc_chat_set_status(session_id, 'error', error=msg, percent=100)
            db.session.rollback()
        finally:
            db.session.remove()


def _doc_chat_collection_job(task_id: int, file_ids: Sequence[int], user_id: Optional[int], allowed_ids: Optional[Sequence[int]]) -> None:
    try:
        db.session.rollback()
    except Exception:
        db.session.remove()

    task = TaskRecord.query.get(task_id)
    if not task:
        return

    task.status = 'running'
    task.started_at = datetime.utcnow()
    task.progress = 0.0
    task.error = None
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()

    summary: dict[str, Any] = {
        'total': len(file_ids),
        'processed': 0,
        'skipped': 0,
        'failures': [],
    }
    total = len(file_ids) or 1

    for index, file_id in enumerate(file_ids, start=1):
        try:
            file_obj = File.query.get(file_id)
            if not file_obj:
                summary['failures'].append({'file_id': file_id, 'reason': 'missing'})
            elif _doc_chat_cache_is_valid(file_obj, getattr(file_obj, 'doc_chat_cache', None)):
                summary['skipped'] += 1
            else:
                _doc_chat_prepare_cache_for_file(file_id, user_id, allowed_ids)
                summary['processed'] += 1
        except PermissionError as exc:
            db.session.rollback()
            summary['failures'].append({'file_id': file_id, 'reason': 'permission', 'message': str(exc)})
        except Exception as exc:
            db.session.rollback()
            app.logger.warning("[doc-chat] bulk cache build error for file %s: %s", file_id, exc)
            summary['failures'].append({'file_id': file_id, 'reason': 'error', 'message': str(exc)})

        progress = index / total
        task.progress = progress
        task.payload = json.dumps(summary, ensure_ascii=False)
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()

    task.progress = 1.0
    task.finished_at = datetime.utcnow()
    if summary['failures']:
        task.error = json.dumps(summary['failures'][:10], ensure_ascii=False)
    else:
        task.error = None
    task.status = 'completed'
    task.payload = json.dumps(summary, ensure_ascii=False)
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()


def _initialize_database_structures(app_obj: Flask) -> None:
    """Ensure ORM metadata, FTS tables and task queue state are ready."""
    with app_obj.app_context():
        db.create_all()
        _ensure_search_support()
        _reset_inflight_tasks()


def setup_app(config: AppConfig | None = None, *, ensure_database: bool = True) -> Flask:
    global CONFIG, LOG_DIR, LOG_FILE_PATH
    cfg = config or CONFIG
    if config is not None:
        CONFIG = cfg
        _apply_config_defaults(cfg)
        if runtime_settings_store.current.lm_default_provider not in LLM_PROVIDER_CHOICES:
            runtime_settings_store.current.lm_default_provider = 'openai'
    LOG_DIR = cfg.logs_dir
    LOG_FILE_PATH = cfg.log_file_path
    configure_http(
        HttpSettings(
            timeout=cfg.http_timeout,
            connect_timeout=cfg.http_connect_timeout,
            retries=cfg.http_retries,
            backoff_factor=cfg.http_backoff_factor,
        )
    )
    configure_llm_cache(
        enabled=cfg.llm_cache_enabled,
        max_items=cfg.llm_cache_max_items,
        ttl_seconds=cfg.llm_cache_ttl_seconds,
    )
    configure_search_cache(
        enabled=cfg.search_cache_enabled,
        max_items=cfg.search_cache_max_items,
        ttl_seconds=cfg.search_cache_ttl_seconds,
    )
    configure_llm_pool(
        workers=cfg.llm_pool_global_concurrency,
        per_user_limit=cfg.llm_pool_per_user_concurrency,
        max_queue=cfg.llm_queue_max_size,
    )
    try:
        LOG_DIR.mkdir(parents=True, exist_ok=True)
    except Exception:
        fallback_dir = cfg.base_dir
        fallback_dir.mkdir(parents=True, exist_ok=True)
        LOG_DIR = fallback_dir
        LOG_FILE_PATH = LOG_DIR / 'agregator.log'
    configure_logging(
        app,
        LOG_FILE_PATH,
        level=LOG_LEVEL,
        sentry_dsn=SENTRY_DSN or None,
        sentry_environment=SENTRY_ENVIRONMENT or None,
    )
    # Гарантируем запуск очереди фоновых задач
    get_task_queue().start()
    app.secret_key = cfg.flask_secret_key
    app.config.update(cfg.to_flask_config())
    app.config['MAX_CONTENT_LENGTH'] = cfg.max_content_length
    db.init_app(app)
    try:
        with app.app_context():
            if db.engine.dialect.name == 'sqlite':
                with db.engine.connect() as conn:
                    conn.execute(text("PRAGMA foreign_keys = 1"))
                    conn.execute(text("PRAGMA trusted_schema = 1"))
                    conn.execute(text("PRAGMA recursive_triggers = 1"))
                db.engine.dispose()
    except Exception as pragma_exc:
        logger.warning('Failed to apply SQLite PRAGMA settings: %s', pragma_exc)
    _load_runtime_settings_from_disk()
    _apply_runtime_settings_to_config(app)
    if 'admin_api' not in app.blueprints:
        app.register_blueprint(admin_bp)
    if 'users_api' not in app.blueprints:
        app.register_blueprint(users_bp)
    if 'osint_api' not in app.blueprints:
        app.register_blueprint(osint_bp)
    if ensure_database:
        _initialize_database_structures(app)
    with app.app_context():
        _schedule_facet_cache_rebuild('startup')
    return app


def _get_rotating_log_handler() -> RotatingFileHandler | None:
    return svc_get_rotating_log_handler(app, LOG_FILE_PATH)


def _list_system_log_files() -> list[dict[str, str | int | float | bool | None]]:
    files = []
    for entry in svc_list_system_log_files(LOG_DIR, LOG_FILE_PATH):
        files.append(
            {
                'name': entry.get('name'),
                'size': entry.get('size'),
                'modified_at': entry.get('modified_at'),
                'rotated': entry.get('rotated'),
            }
        )
    return files


def _tail_log_file(path: Path, max_lines: int = 200) -> list[str]:
    return svc_tail_log_file(path, max_lines=max_lines)


def _resolve_log_name(name: str | None) -> Path | None:
    return svc_resolve_log_name(LOG_DIR, LOG_FILE_PATH, name)
# ------------------- Утилиты -------------------

ALLOWED_EXTS = {".pdf", ".txt", ".md", ".docx", ".rtf", ".mp3", ".wav", ".m4a", ".flac", ".ogg",
                ".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"}
AUDIO_EXTS = {".mp3", ".wav", ".m4a", ".flac", ".ogg"}
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"}

# ------------------- Пользователи и доступ -------------------

DEFAULT_ADMIN_USER = CONFIG.default_admin_user
DEFAULT_ADMIN_PASSWORD = CONFIG.default_admin_password
LEGACY_ACCESS_CODE = CONFIG.legacy_access_code
SESSION_KEY = 'user_id'

_PUBLIC_PREFIXES = ('/static/', '/assets/', '/login-backgrounds/')
_PUBLIC_PATHS = {'/favicon.ico', '/api/auth/login', '/api/login-backgrounds'}

def _log_user_action(user: User | None, action: str, entity: str | None = None, entity_id: int | None = None, detail: str | None = None) -> None:
    try:
        rec = UserActionLog(user_id=user.id if user else None, action=action, entity=entity, entity_id=entity_id, detail=detail)
        db.session.add(rec)
        db.session.commit()
    except Exception:
        db.session.rollback()

def _user_to_payload(user: User) -> dict:
    role = _normalize_user_role(getattr(user, 'role', ROLE_VIEWER), default=ROLE_EDITOR)
    return {
        'id': user.id,
        'username': user.username,
        'role': role,
        'full_name': user.full_name,
        'created_at': user.created_at.isoformat() if user.created_at else None,
        'updated_at': user.updated_at.isoformat() if getattr(user, 'updated_at', None) else None,
        'aiword_access': _has_aiword_access(user),
        'can_upload': _user_can_upload(user),
        'can_import': _user_can_upload(user),
        'permissions': {
            'can_admin': role == ROLE_ADMIN,
            'can_edit': role in (ROLE_ADMIN, ROLE_EDITOR),
            'can_view': True,
        },
    }

def ensure_default_admin() -> None:
    try:
        with app.app_context():
            db.create_all()
            if User.query.count() == 0:
                password = DEFAULT_ADMIN_PASSWORD or LEGACY_ACCESS_CODE or 'admin123'
                user = User(username=DEFAULT_ADMIN_USER, role=ROLE_ADMIN)
                user.set_password(password)
                db.session.add(user)
                db.session.commit()
                app.logger.warning(
                    "Создан администратор по умолчанию '%s'. Рекомендуется немедленно сменить пароль.",
                    DEFAULT_ADMIN_USER,
                )
    except Exception as exc:
        app.logger.error("Не удалось создать администратора по умолчанию: %s", exc)

def _load_current_user() -> User | None:
    uid = session.get(SESSION_KEY)
    if not uid:
        g.current_user = None
        g.allowed_collection_ids = set()
        return None
    try:
        user = User.query.get(int(uid))
    except Exception:
        user = None
    if not user:
        session.pop(SESSION_KEY, None)
        g.current_user = None
        g.allowed_collection_ids = set()
        return None
    g.current_user = user
    _ensure_personal_collection(user)
    g.allowed_collection_ids = _compute_allowed_collection_ids(user)
    return user


def _compute_allowed_collection_ids(user: User | None):
    if not user:
        return set()
    if _user_role(user) == ROLE_ADMIN:
        return None
    ids: set[int] = set()
    try:
        own = db.session.query(Collection.id).filter(Collection.owner_id == user.id).all()
        ids.update(int(cid) for (cid,) in own)
    except Exception:
        pass
    try:
        membership = db.session.query(CollectionMember.collection_id, CollectionMember.role).filter(CollectionMember.user_id == user.id).all()
        ids.update(int(cid) for (cid, _role) in membership)
    except Exception:
        pass
    try:
        public = db.session.query(Collection.id).filter(Collection.is_private == False).all()
        ids.update(int(cid) for (cid,) in public)
    except Exception:
        pass
    return ids


def _ensure_personal_collection(user: User) -> None:
    try:
        personal = Collection.query.filter_by(owner_id=user.id, is_private=True).first()
        if not personal:
            name = f"Личная коллекция {user.username}"[:120]
            base_slug = _slugify(name)
            if not base_slug:
                base_slug = f"user-{user.id}"
            slug = base_slug
            i = 2
            while Collection.query.filter_by(slug=slug).first() is not None:
                slug = f"{base_slug}-{i}"
                i += 1
            personal = Collection(name=name, slug=slug, owner_id=user.id, is_private=True, searchable=True, graphable=True)
            db.session.add(personal)
            db.session.flush()
        member = CollectionMember.query.filter_by(collection_id=personal.id, user_id=user.id).first()
        if not member:
            db.session.add(CollectionMember(collection_id=personal.id, user_id=user.id, role='owner'))
        db.session.commit()
    except Exception:
        db.session.rollback()


def _has_collection_access(collection_id: int, write: bool = False) -> bool:
    user = _load_current_user()
    if not user:
        return False
    if _user_role(user) == ROLE_ADMIN:
        return True
    allowed = getattr(g, 'allowed_collection_ids', set())
    if allowed is None or collection_id in allowed:
        if not write:
            return True
        try:
            col = Collection.query.get(collection_id)
            if col and col.owner_id == user.id:
                return True
        except Exception:
            pass
        try:
            member = CollectionMember.query.filter_by(collection_id=collection_id, user_id=user.id).first()
            if member and member.role in ('owner', 'editor'):
                return True
        except Exception:
            pass
    return False


def _apply_file_access_filter(query):
    allowed = getattr(g, 'allowed_collection_ids', set())
    if allowed is None:
        return query
    if not allowed:
        return query.filter(File.collection_id == -1)
    return query.filter(File.collection_id.in_(allowed))


def _parse_collection_param(value: str | None, require_write: bool = False) -> int | None:
    if not value:
        return None
    try:
        cid = int(value)
    except Exception:
        abort(400)
    if require_write:
        if not _has_collection_access(cid, write=True):
            abort(403)
    else:
        if not _has_collection_access(cid):
            abort(403)
    return cid


def _aiword_allowed_user_ids() -> set[int]:
    try:
        return {row.user_id for row in AiWordAccess.query.all()}
    except Exception:
        return set()


def _ensure_aiword_access() -> User:
    user = _load_current_user()
    if not user:
        abort(401)
    if _user_role(user) == ROLE_ADMIN:
        return user
    allowed = _aiword_allowed_user_ids()
    if user.id in allowed:
        return user
    abort(403)


def _has_aiword_access(user: User | None) -> bool:
    if not user:
        return False
    if _user_role(user) == ROLE_ADMIN:
        return True
    try:
        return AiWordAccess.query.filter_by(user_id=user.id).first() is not None
    except Exception:
        return False


def _user_can_upload(user: User | None) -> bool:
    if not user:
        return False
    role = _normalize_user_role(getattr(user, 'role', None), default=ROLE_EDITOR)
    if role == ROLE_ADMIN:
        return True
    if role == ROLE_VIEWER:
        return False
    try:
        if Collection.query.filter(Collection.owner_id == user.id).limit(1).first():
            return True
    except Exception:
        pass
    try:
        member = CollectionMember.query.filter(
            CollectionMember.user_id == user.id,
            CollectionMember.role.in_(('owner', 'editor'))
        ).limit(1).first()
        if member:
            return True
    except Exception:
        pass
    return False


def _llm_parse_purposes(raw: str | None) -> list[str]:
    if raw is None:
        return ['default']
    parts = [p.strip().lower() for p in re.split(r'[;,\s]+', raw) if p.strip()]
    return parts or ['default']


def _llm_normalize_purposes(value) -> str | None:
    if value is None:
        return None
    if isinstance(value, str):
        parts = _llm_parse_purposes(value)
    elif isinstance(value, (list, tuple, set)):
        parts = []
        for item in value:
            token = str(item).strip().lower()
            if token:
                parts.append(token)
    else:
        return None
    seen: set[str] = set()
    ordered: list[str] = []
    for token in parts:
        if token not in seen:
            seen.add(token)
            ordered.append(token)
    return ','.join(ordered) if ordered else None


def _llm_pick_default_endpoint() -> LlmEndpoint | None:
    """Pick endpoint used as default route in common settings."""
    try:
        eps = LlmEndpoint.query.order_by(LlmEndpoint.weight.desc(), LlmEndpoint.created_at.desc()).all()
    except Exception:
        return None
    for ep in eps:
        if 'default' in _llm_parse_purposes(ep.purpose):
            return ep
    return None


def _sync_runtime_default_route_from_llm_endpoints() -> None:
    """Persist current default endpoint into runtime global LLM route."""
    ep = _llm_pick_default_endpoint()
    if ep is None:
        return
    provider = str(ep.provider or '').strip().lower() or 'openai'
    runtime_settings_store.apply_updates({
        'LMSTUDIO_API_BASE': ep.base_url,
        'LMSTUDIO_MODEL': ep.model,
        'LMSTUDIO_API_KEY': ep.api_key or '',
        'LM_DEFAULT_PROVIDER': provider,
    })
    _refresh_runtime_globals()
    _rt().apply_to_flask_config(app)
    _save_runtime_settings_to_disk()


def _invalidate_llm_cache() -> None:
    LLM_ROUND_ROBIN.clear()
    LLM_ENDPOINT_SIGNATURE.clear()
    LLM_ENDPOINT_POOLS.clear()
    LLM_ENDPOINT_UNIQUE.clear()


def _ensure_llm_pool(purpose: str | None) -> None:
    desired = (purpose or 'default').lower()
    _ensure_llm_schema_once()
    base_default = (LMSTUDIO_API_BASE or '').rstrip('/')
    model_default = LMSTUDIO_MODEL
    key_default = LMSTUDIO_API_KEY
    try:
        all_eps = LlmEndpoint.query.order_by(LlmEndpoint.created_at.asc()).all()
        candidates: list[tuple[LlmEndpoint, list[str]]] = []
        for ep in all_eps:
            purposes = _llm_parse_purposes(ep.purpose)
            if desired in purposes or (desired != 'default' and 'default' in purposes):
                candidates.append((ep, purposes))
        if not candidates and desired != 'default':
            for ep in all_eps:
                purposes = _llm_parse_purposes(ep.purpose)
                if 'default' in purposes:
                    candidates.append((ep, purposes))
        if candidates:
            sig = tuple(sorted((ep.id, float(ep.weight or 1.0), (ep.provider or ''), tuple(sorted(purposes))) for ep, purposes in candidates))
            if LLM_ENDPOINT_SIGNATURE.get(desired) != sig:
                pool: list[dict[str, str]] = []
                unique: list[dict[str, str]] = []
                seen: set[tuple] = set()
                for ep, _purposes in candidates:
                    provider_fallback = (_rt().lm_default_provider or 'openai')
                    entry = {
                        'id': ep.id,
                        'name': ep.name or f'endpoint-{ep.id}',
                        'base_url': (ep.base_url or base_default).rstrip('/'),
                        'model': ep.model or model_default,
                        'api_key': ep.api_key or key_default,
                        'provider': (ep.provider or provider_fallback).strip().lower() or provider_fallback,
                    }
                    weight = max(1, int(round(float(ep.weight or 1.0))))
                    for _ in range(weight):
                        pool.append(entry)
                    ident = (entry['id'], entry['base_url'], entry['model'], entry.get('api_key'), entry.get('provider'))
                    if ident not in seen:
                        seen.add(ident)
                        unique.append(entry)
                if not pool:
                    default_entry = {
                        'id': None,
                        'name': 'default',
                        'base_url': base_default,
                        'model': model_default,
                        'api_key': key_default,
                        'provider': (_rt().lm_default_provider or 'openai'),
                    }
                    pool = [default_entry]
                    unique = [default_entry]
                LLM_ENDPOINT_POOLS[desired] = pool
                LLM_ENDPOINT_UNIQUE[desired] = unique
                LLM_ROUND_ROBIN[desired] = itertools.cycle(pool)
                LLM_ENDPOINT_SIGNATURE[desired] = sig
            return
    except Exception:
        pass
    if desired not in LLM_ENDPOINT_POOLS:
        default_entry = {
            'id': None,
            'name': 'default',
            'base_url': base_default,
            'model': model_default,
            'api_key': key_default,
            'provider': (_rt().lm_default_provider or 'openai'),
        }
        LLM_ENDPOINT_POOLS[desired] = [default_entry]
        LLM_ENDPOINT_UNIQUE[desired] = [default_entry]
        LLM_ROUND_ROBIN[desired] = itertools.cycle(LLM_ENDPOINT_POOLS[desired])
        LLM_ENDPOINT_SIGNATURE[desired] = tuple()


def _select_llm_endpoint(purpose: str | None = None) -> tuple[str, str, str]:
    desired = (purpose or 'default').lower()
    _ensure_llm_pool(desired)
    pool = LLM_ENDPOINT_POOLS.get(desired) or []
    if not pool:
        return LMSTUDIO_API_BASE, LMSTUDIO_MODEL, LMSTUDIO_API_KEY
    rotation = LLM_ROUND_ROBIN.get(desired)
    if rotation is None:
        rotation = itertools.cycle(pool)
        LLM_ROUND_ROBIN[desired] = rotation
    choice = next(rotation)
    base = choice.get('base_url') or LMSTUDIO_API_BASE
    model = choice.get('model') or LMSTUDIO_MODEL
    key = choice.get('api_key') or LMSTUDIO_API_KEY
    return base, model, key


def _llm_iter_choices(purpose: str | None):
    desired = (purpose or 'default').lower()
    _ensure_llm_pool(desired)
    pool = LLM_ENDPOINT_POOLS.get(desired) or []
    if not pool:
        yield {
            'id': None,
            'name': 'default',
            'base_url': (LMSTUDIO_API_BASE or '').rstrip('/'),
            'model': LMSTUDIO_MODEL,
            'api_key': LMSTUDIO_API_KEY,
            'provider': (_rt().lm_default_provider or 'openai'),
        }
        return
    rotation = LLM_ROUND_ROBIN.get(desired)
    if rotation is None:
        rotation = itertools.cycle(pool)
        LLM_ROUND_ROBIN[desired] = rotation
    unique = LLM_ENDPOINT_UNIQUE.get(desired) or []
    max_iters = len(unique) if unique else len(pool) or 1
    seen: set[tuple] = set()
    for _ in range(max_iters):
        choice = next(rotation)
        ident = (choice.get('id'), choice.get('base_url'), choice.get('model'), choice.get('api_key'), choice.get('provider'))
        if ident in seen:
            continue
        seen.add(ident)
        yield choice


def _llm_choice_label(choice: dict) -> str:
    name = str(choice.get('name') or '').strip()
    base = str(choice.get('base_url') or '').strip()
    model = str(choice.get('model') or '').strip()
    ident = choice.get('id')
    parts = []
    if name:
        parts.append(name)
    if model:
        parts.append(model)
    if base:
        parts.append(base)
    if not parts:
        parts.append(f'id={ident}')
    return ' | '.join(parts)


def _llm_choice_provider(choice: dict) -> str:
    provider = str(choice.get('provider') or '').strip().lower()
    aliases = {
        'azure': 'azure_openai',
        'azure-openai': 'azure_openai',
    }
    provider = aliases.get(provider, provider)
    if provider not in LLM_PROVIDER_CHOICES:
        provider = runtime_settings_store.current.lm_default_provider
    if provider not in LLM_PROVIDER_CHOICES:
        provider = 'openai'
    return provider


def _llm_response_indicates_busy(response) -> bool:
    if response is None:
        return False
    try:
        status_code = int(response.status_code)
    except Exception:
        status_code = None
    if status_code in LLM_BUSY_HTTP_CODES:
        return True
    payload = None
    try:
        payload = response.json() if hasattr(response, 'json') else None
    except ValueError:
        payload = None
    except json.JSONDecodeError:
        payload = None
    if isinstance(payload, dict):
        status = str(payload.get('status') or '').lower()
        message = str(payload.get('message') or '').lower()
        if status in LLM_BUSY_STATUS_VALUES:
            return True
        if any(token in message for token in ('busy', 'processing', 'queue', 'rate')):
            return True
    text = str(getattr(response, 'text', '') or '').lower()
    if any(token in text for token in ('busy', 'processing', 'queue full')):
        return True
    return False


def _llm_choice_url(choice: dict) -> str:
    base = str(choice.get('base_url') or LMSTUDIO_API_BASE or '').strip()
    if not base:
        return ''
    base = base.rstrip('/')
    if not base:
        return ''
    provider = _llm_choice_provider(choice)
    if provider == 'ollama':
        if re.search(r"/api/(chat|generate)$", base):
            return base
        return f"{base}/api/chat"
    if provider == 'azure_openai':
        url = base
        if not re.search(r"/chat/completions(?:\?|$)", url):
            url = f"{url}/chat/completions"
        if 'api-version=' not in url:
            api_version = choice.get('api_version') or AZURE_OPENAI_API_VERSION
            sep = '&' if '?' in url else '?'
            url = f"{url}{sep}api-version={api_version}"
        return url
    if base.endswith('/chat/completions'):
        return base
    if base.endswith('/v1'):
        return f"{base}/chat/completions"
    return f"{base}/chat/completions"


def _llm_choice_headers(choice: dict) -> dict:
    headers = {"Content-Type": "application/json"}
    key = choice.get('api_key')
    if key:
        provider = _llm_choice_provider(choice)
        if provider == 'azure_openai':
            headers["api-key"] = key
        else:
            headers["Authorization"] = f"Bearer {key}"
    return headers


def _llm_timeout_pair(total_timeout: int | float | None) -> tuple[float, float]:
    try:
        base = float(total_timeout) if total_timeout is not None else 120.0
    except Exception:
        base = 120.0
    base = max(5.0, min(base, 300.0))
    connect = min(10.0, max(3.0, base * 0.25))
    read = max(connect + 2.0, base)
    return connect, read


def _llm_send_chat(
    choice: dict,
    messages: list[dict],
    *,
    temperature: float,
    max_tokens: int | None,
    top_p: float = 1.0,
    timeout: int = 120,
    extra_payload: dict | None = None,
    cache_bucket: str | None = None,
    cache_only: bool | None = None,
):
    _lmstudio_mark_activity()
    provider = _llm_choice_provider(choice)
    url = _llm_choice_url(choice)
    if not url:
        raise ValueError('empty_url')
    model = choice.get('model') or LMSTUDIO_MODEL
    if provider == 'ollama':
        payload = {
            'model': model,
            'messages': messages,
            'stream': False,
            'options': {
                'temperature': float(temperature),
                'top_p': float(top_p),
            },
        }
        if max_tokens is not None:
            try:
                payload['options']['num_predict'] = max(0, int(max_tokens))
            except Exception:
                pass
        if extra_payload:
            payload.update(extra_payload)
    else:
        payload = {
            'model': model,
            'messages': messages,
            'temperature': float(temperature),
            'top_p': float(top_p),
        }
        if max_tokens is not None:
            try:
                payload['max_tokens'] = max(0, int(max_tokens))
            except Exception:
                pass
        if extra_payload:
            payload.update(extra_payload)
    cache_enabled = bool(LLM_CACHE_ENABLED)
    cache_only = LLM_CACHE_ONLY_MODE if cache_only is None else bool(cache_only)
    cache_key: str | None = None
    if cache_enabled:
        cache_key_material = {
            'bucket': cache_bucket or (choice.get('name') or provider),
            'provider': provider,
            'url': url,
            'model': model,
            'payload': payload,
            'scope': _cache_scope_identity(),
        }
        cache_key = hashlib.sha256(json.dumps(cache_key_material, sort_keys=True, ensure_ascii=False).encode('utf-8')).hexdigest()
        cached_resp = llm_cache_get(cache_key)
        if cached_resp is not None:
            cached_resp.headers.setdefault('X-LLM-Cache', 'hit')
            return provider, cached_resp
        if cache_only:
            raise RuntimeError('llm_cache_only_mode')

    owner_id: int | None = None
    try:
        req_user = _load_current_user()
        if req_user is not None and getattr(req_user, 'id', None) is not None:
            owner_id = int(req_user.id)
    except Exception:
        owner_id = None
    retry_count = max(0, int(app.config.get('LLM_RETRY_COUNT', 1) or 0))
    retry_backoff_ms = max(0, int(app.config.get('LLM_RETRY_BACKOFF_MS', 500) or 0))
    pool_timeout = max(1, int(app.config.get('LLM_REQUEST_TIMEOUT_SEC', timeout) or timeout))
    per_user_queue_max = max(1, int(app.config.get('LLM_QUEUE_PER_USER_MAX', 10) or 10))
    response = None
    last_exc: Exception | None = None
    for attempt in range(retry_count + 1):
        try:
            response = get_llm_pool().submit_and_wait(
                owner_id=owner_id,
                timeout_sec=pool_timeout,
                per_user_queue_max=per_user_queue_max,
                fn=lambda: http_request(
                    'POST',
                    url,
                    headers=_llm_choice_headers(choice),
                    json=payload,
                    timeout=_llm_timeout_pair(timeout),
                    logger=app.logger,
                ),
            )
        except LlmPoolRejected as exc:
            raise RuntimeError(str(exc)) from exc
        except LlmPoolTimeout as exc:
            raise RuntimeError(str(exc)) from exc
        except Exception as exc:
            last_exc = exc
            response = None
        if response is not None and getattr(response, 'status_code', 0) < 500:
            break
        if attempt < retry_count and retry_backoff_ms > 0:
            time.sleep(float(retry_backoff_ms) / 1000.0)
    if response is None and last_exc is not None:
        raise RuntimeError(f'llm_request_failed:{last_exc}')

    if response is None:
        raise RuntimeError('llm_response_none')

    if response.status_code >= 400:
        try:
            preview = (response.text or '')[:200]
            app.logger.warning(
                "LLM HTTP %s (%s): %s",
                response.status_code,
                _llm_choice_label(choice),
                preview,
            )
        except Exception:
            pass

    if cache_enabled and cache_key and response.status_code < 400:
        try:
            json_payload = response.json()
        except Exception:
            json_payload = None
        cached = CachedLLMResponse(
            status_code=response.status_code,
            data=json_payload,
            text=response.text,
            headers=dict(response.headers or {}),
        )
        cached.headers.setdefault('X-LLM-Cache', 'store')
        llm_cache_set(cache_key, cached)
    _lmstudio_mark_activity()
    return provider, response


def _llm_extract_content(provider: str, data) -> str:
    if not isinstance(data, dict):
        return ''
    if provider == 'ollama':
        message = data.get('message')
        if isinstance(message, dict):
            content = message.get('content')
            if content:
                return str(content)
        for key in ('response', 'content', 'text'):
            value = data.get(key)
            if value:
                return str(value)
        return ''
    choices = data.get('choices')
    if isinstance(choices, list) and choices:
        message = choices[0].get('message') if isinstance(choices[0], dict) else None
        if isinstance(message, dict):
            content = message.get('content')
            if content:
                return str(content)
    return ''


def _is_public_path(path: str) -> bool:
    if any(path.startswith(prefix) for prefix in _PUBLIC_PREFIXES):
        return True
    if path in _PUBLIC_PATHS:
        return True
    if path == '/' or path.startswith('/access'):
        return True
    if request.method == 'GET' and path.startswith('/app'):
        return True
    return False

@app.before_request
def _access_gate():
    path = request.path or '/'
    user = _load_current_user()
    if _is_public_path(path):
        return None
    if path == '/api/training/problem-pipeline':
        if request.method == 'OPTIONS':
            return None
        allowed, failure_response = _check_pipeline_access()
        if allowed:
            return None
        return _add_pipeline_cors_headers(failure_response)
    if user:
        return None
    if path.startswith('/api/') or request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        resp = jsonify({'ok': False, 'error': 'Не авторизовано'})
        resp.status_code = 401
        return _add_pipeline_cors_headers(resp)
    return redirect('/app/login')

@app.route('/api/auth/login', methods=['POST'])
def api_auth_login():
    data = request.get_json(silent=True) or request.form or {}
    username = (data.get('username') or '').strip()
    password = data.get('password') or ''
    if not username or not password:
        return jsonify({'ok': False, 'error': 'Введите логин и пароль'}), 400
    try:
        user = User.query.filter(func.lower(User.username) == username.lower()).first()
    except Exception:
        user = None
    if not user or not user.check_password(password):
        session.pop(SESSION_KEY, None)
        return jsonify({'ok': False, 'error': 'Неверный логин или пароль'}), 401
    session[SESSION_KEY] = user.id
    session.permanent = True
    g.current_user = user
    _log_user_action(user, 'login')
    return jsonify({'ok': True, 'user': _user_to_payload(user)})

@app.route('/api/auth/logout', methods=['POST'])
def api_auth_logout():
    user = _load_current_user()
    session.pop(SESSION_KEY, None)
    g.current_user = None
    _log_user_action(user, 'logout')
    return jsonify({'ok': True})

@app.route('/api/auth/me')
def api_auth_me():
    user = _load_current_user()
    if not user:
        return jsonify({'ok': False}), 401
    return jsonify({'ok': True, 'user': _user_to_payload(user)})

@app.route('/api/profile', methods=['GET', 'POST'])
def api_profile():
    user = _load_current_user()
    if not user:
        return jsonify({'ok': False, 'error': 'Не авторизовано'}), 401
    if request.method == 'GET':
        return jsonify({'ok': True, 'user': _user_to_payload(user)})
    data = request.get_json(silent=True) or {}
    runtime = runtime_settings_store.current
    default_backend = (runtime.rag_embedding_backend or 'lm-studio').strip().lower() or 'lm-studio'
    default_model = runtime.rag_embedding_model or 'nomic-ai/nomic-embed-text-v1.5-GGUF'
    default_dim = max(8, int(runtime.rag_embedding_dim or 768))
    default_batch = max(1, int(runtime.rag_embedding_batch_size or 32))
    default_device = runtime.rag_embedding_device
    default_endpoint = runtime.rag_embedding_endpoint or runtime.lmstudio_api_base or ''
    default_api_key = runtime.rag_embedding_api_key or runtime.lmstudio_api_key or ''
    runtime = runtime_settings_store.current
    default_backend = (runtime.rag_embedding_backend or 'lm-studio').strip().lower() or 'lm-studio'
    default_model = runtime.rag_embedding_model or 'nomic-ai/nomic-embed-text-v1.5-GGUF'
    default_dim = max(8, int(runtime.rag_embedding_dim or 768))
    default_batch = max(1, int(runtime.rag_embedding_batch_size or 32))
    default_device = runtime.rag_embedding_device
    default_endpoint = runtime.rag_embedding_endpoint or runtime.lmstudio_api_base or ''
    default_api_key = runtime.rag_embedding_api_key or runtime.lmstudio_api_key or ''
    runtime = runtime_settings_store.current
    default_backend = (runtime.rag_embedding_backend or 'lm-studio').strip().lower() or 'lm-studio'
    default_model = runtime.rag_embedding_model or 'nomic-ai/nomic-embed-text-v1.5-GGUF'
    default_dim = max(8, int(runtime.rag_embedding_dim or 768))
    default_batch = max(1, int(runtime.rag_embedding_batch_size or 32))
    default_device = runtime.rag_embedding_device
    default_endpoint = runtime.rag_embedding_endpoint or runtime.lmstudio_api_base or ''
    default_api_key = runtime.rag_embedding_api_key or runtime.lmstudio_api_key or ''
    runtime = runtime_settings_store.current
    default_backend = (runtime.rag_embedding_backend or 'lm-studio').strip().lower() or 'lm-studio'
    default_model = runtime.rag_embedding_model or 'nomic-ai/nomic-embed-text-v1.5-GGUF'
    default_dim = max(8, int(runtime.rag_embedding_dim or 768))
    default_batch = max(1, int(runtime.rag_embedding_batch_size or 32))
    default_device = runtime.rag_embedding_device
    default_endpoint = runtime.rag_embedding_endpoint or runtime.lmstudio_api_base or ''
    default_api_key = runtime.rag_embedding_api_key or runtime.lmstudio_api_key or ''
    runtime = runtime_settings_store.current
    default_backend = (runtime.rag_embedding_backend or 'lm-studio').strip().lower() or 'lm-studio'
    default_model = runtime.rag_embedding_model or 'nomic-ai/nomic-embed-text-v1.5-GGUF'
    default_dim = max(8, int(runtime.rag_embedding_dim or 768))
    default_batch = max(1, int(runtime.rag_embedding_batch_size or 32))
    default_device = runtime.rag_embedding_device
    default_endpoint = runtime.rag_embedding_endpoint or runtime.lmstudio_api_base or ''
    default_api_key = runtime.rag_embedding_api_key or runtime.lmstudio_api_key or ''
    default_backend = (runtime.rag_embedding_backend or 'lm-studio').strip().lower() or 'lm-studio'
    default_model = runtime.rag_embedding_model or 'nomic-ai/nomic-embed-text-v1.5-GGUF'
    default_dim = max(8, int(runtime.rag_embedding_dim or 768))
    default_batch = max(1, int(runtime.rag_embedding_batch_size or 32))
    default_device = runtime.rag_embedding_device
    default_endpoint = runtime.rag_embedding_endpoint or runtime.lmstudio_api_base or ''
    default_api_key = runtime.rag_embedding_api_key or runtime.lmstudio_api_key or ''
    default_backend = (runtime.rag_embedding_backend or 'lm-studio').strip().lower() or 'lm-studio'
    default_model = runtime.rag_embedding_model or 'nomic-ai/nomic-embed-text-v1.5-GGUF'
    default_dim = max(8, int(runtime.rag_embedding_dim or 768))
    default_batch = max(1, int(runtime.rag_embedding_batch_size or 32))
    default_device = runtime.rag_embedding_device
    default_endpoint = runtime.rag_embedding_endpoint or runtime.lmstudio_api_base or ''
    default_api_key = runtime.rag_embedding_api_key or runtime.lmstudio_api_key or ''
    user.full_name = (data.get('full_name') or '').strip() or None
    db.session.commit()
    try:
        _log_user_action(user, 'profile_update', 'user', user.id, detail=json.dumps({'full_name': user.full_name}))
    except Exception:
        pass
    return jsonify({'ok': True, 'user': _user_to_payload(user)})

@app.route('/api/auth/password', methods=['POST'])
def api_auth_change_password():
    user = _load_current_user()
    if not user:
        return jsonify({'ok': False, 'error': 'Не авторизовано'}), 401
    data = request.get_json(silent=True) or {}
    current = data.get('current_password') or ''
    new_password = (data.get('new_password') or '').strip()
    if not user.check_password(current):
        return jsonify({'ok': False, 'error': 'Неверный текущий пароль'}), 400
    if len(new_password) < 6:
        return jsonify({'ok': False, 'error': 'Пароль должен содержать не менее 6 символов'}), 400
    user.set_password(new_password)
    db.session.commit()
    _log_user_action(user, 'change_password')
    return jsonify({'ok': True})

def _require_admin() -> User:
    user = _load_current_user()
    if not user or user.role != 'admin':
        abort(403)
    return user

def require_admin(fn):
    @wraps(fn)
    def wrapper(*args, **kwargs):
        _require_admin()
        return fn(*args, **kwargs)
    return wrapper

admin_bp = Blueprint('admin_api', __name__, url_prefix='/api/admin')
users_bp = Blueprint('users_api', __name__, url_prefix='/api/users')
osint_bp = Blueprint('osint_api', __name__, url_prefix='/api/osint')

_OSINT_SERVICE: OsintSearchService | None = None
_OSINT_LOCK = threading.Lock()


def _osint_env_bool(name: str, default: bool) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    token = str(raw).strip().lower()
    return token in {"1", "true", "yes", "on"}


def _osint_env_int(name: str, default: int, *, minimum: int | None = None, maximum: int | None = None) -> int:
    raw = os.getenv(name)
    value = default
    if raw is not None:
        text = str(raw).strip()
        if text:
            try:
                value = int(text)
            except ValueError:
                value = default
    if minimum is not None:
        value = max(minimum, value)
    if maximum is not None:
        value = min(maximum, value)
    return value


def _osint_sanitize_params(raw: Mapping[str, Any] | None) -> dict[str, str] | None:
    if not raw:
        return None
    sanitized: dict[str, str] = {}
    for key, value in raw.items():
        try:
            k = str(key).strip()
            if not k:
                continue
            sanitized[k] = str(value)
        except Exception:
            continue
    return sanitized or None


def _osint_llm_chat(messages: list[dict[str, str]]) -> dict:
    max_tokens = _osint_env_int("OSINT_LLM_MAX_TOKENS", 512, minimum=64, maximum=2048)
    timeout = _osint_env_int("OSINT_LLM_TIMEOUT_SECONDS", 120, minimum=30, maximum=300)
    last_error: Exception | None = None
    for choice in _llm_iter_choices("osint"):
        label = _llm_choice_label(choice)
        try:
            provider, response = _llm_send_chat(
                choice,
                messages,
                temperature=0.1,
                max_tokens=max_tokens,
                timeout=timeout,
                cache_bucket="osint-serp",
            )
            if _llm_response_indicates_busy(response):
                app.logger.info("OSINT LLM endpoint busy (%s)", label)
                last_error = RuntimeError("busy")
                continue
            response.raise_for_status()
            data = response.json()
            content = _llm_extract_content(provider, data)
            if content:
                return {
                    "content": content,
                    "model": choice.get("model") or provider,
                    "raw": getattr(response, "text", ""),
                }
            last_error = RuntimeError(f"empty response from {label}")
        except Exception as exc:  # noqa: BLE001
            last_error = exc
            app.logger.info("OSINT LLM attempt failed (%s): %s", label, exc)
            continue
    if last_error:
        raise last_error
    raise RuntimeError("osint_llm_no_response")


def _get_osint_service() -> OsintSearchService:
    global _OSINT_SERVICE
    if _OSINT_SERVICE is not None:
        return _OSINT_SERVICE
    with _OSINT_LOCK:
        if _OSINT_SERVICE is not None:
            return _OSINT_SERVICE
        llm_enabled = _osint_env_bool("OSINT_LLM_ENABLED", True)
        parser = SerpParser(
            llm_chat=_osint_llm_chat if llm_enabled else None,
            max_items=_osint_env_int("OSINT_MAX_ITEMS", 10, minimum=1, maximum=30),
            max_html_chars=_osint_env_int("OSINT_MAX_HTML_CHARS", 20000, minimum=2000, maximum=120000),
        )
        retry_agents_env = os.getenv("OSINT_RETRY_USER_AGENTS") or ""
        retry_agents = tuple(
            filter(
                None,
                (token.strip() for token in retry_agents_env.split(",") if token.strip()),
            )
        )
        retry_proxies_env = os.getenv("OSINT_RETRY_PROXIES") or ""
        retry_proxies = tuple(
            filter(
                None,
                (token.strip() for token in retry_proxies_env.replace(";", ",").split(",") if token.strip()),
            )
        )
        search_api_url = os.getenv("OSINT_SEARCH_API_URL") or None
        search_api_method = (os.getenv("OSINT_SEARCH_API_METHOD") or "POST").strip()
        search_api_key = os.getenv("OSINT_SEARCH_API_KEY") or None
        settings = SerpSettings(
            cache_enabled=_osint_env_bool("OSINT_CACHE_ENABLED", True),
            cache_ttl_seconds=_osint_env_int("OSINT_CACHE_TTL_SECONDS", 900, minimum=30, maximum=86400),
            cache_max_items=_osint_env_int("OSINT_CACHE_MAX_ITEMS", 128, minimum=16, maximum=512),
            wait_after_load_ms=_osint_env_int("OSINT_WAIT_AFTER_LOAD_MS", 1200, minimum=0, maximum=10000),
            wait_selector=os.getenv("OSINT_WAIT_SELECTOR") or None,
            navigation_timeout_ms=_osint_env_int("OSINT_NAVIGATION_TIMEOUT_MS", 45000, minimum=5000, maximum=240000),
            reuse_cache_on_block=_osint_env_bool("OSINT_REUSE_CACHE_ON_BLOCK", True),
            user_agent_override=os.getenv("OSINT_USER_AGENT") or None,
            retry_user_agents=retry_agents,
            retry_proxies=retry_proxies,
            search_api_url=search_api_url,
            search_api_method=search_api_method,
            search_api_key=search_api_key,
        )
        fetcher = SerpFetcher(settings=settings)
        repo_config = OsintRepositoryConfig(
            url=os.getenv("OSINT_DATABASE_URL") or None,
            echo=_osint_env_bool("OSINT_DB_ECHO", False),
        )
        _OSINT_SERVICE = OsintSearchService(
            fetcher=fetcher,
            parser=parser,
            repository_config=repo_config,
        )
        return _OSINT_SERVICE


@osint_bp.route('/engines', methods=['GET'])
def api_osint_engines():
    _require_admin()
    return jsonify({'ok': True, 'engines': list(SUPPORTED_ENGINES)})


@osint_bp.route('/history', methods=['GET'])
def api_osint_history():
    _require_admin()
    try:
        limit = int(request.args.get('limit', '10') or '10')
    except Exception:
        limit = 10
    limit = max(1, min(limit, 50))
    service = _get_osint_service()
    items = service.list_jobs(limit)
    return jsonify({'ok': True, 'items': items})


@osint_bp.route('/search', methods=['POST'])
def api_osint_search():
    user = _require_admin()
    data = request.get_json(silent=True) or {}
    query = (data.get('query') or '').strip()
    if not query:
        return jsonify({'ok': False, 'error': 'Укажите поисковый запрос'}), 400
    locale = (data.get('locale') or 'ru-RU').strip() or 'ru-RU'
    if locale not in SUPPORTED_OSINT_LOCALES:
        locale = SUPPORTED_OSINT_LOCALES[0]
    region = data.get('region')
    if region is not None:
        region = str(region).strip() or None
    safe = bool(data.get('safe'))
    include_html = bool(data.get('include_html'))
    include_llm = bool(data.get('include_llm_payload'))
    build_ontology = bool(data.get('build_ontology'))
    schedule_payload = data.get('schedule')
    if schedule_payload is not None and not isinstance(schedule_payload, Mapping):
        return jsonify({'ok': False, 'error': 'Поле schedule должно быть объектом'}), 400
    max_results_value = data.get('max_results')
    max_results: int | None = None
    if max_results_value not in (None, ''):
        try:
            max_results = int(max_results_value)
        except Exception:
            return jsonify({'ok': False, 'error': 'Некорректное значение max_results'}), 400
        max_results = max(1, min(max_results, 50))

    locales = OsintSearchService._normalize_requested_locales(data.get('locales'), locale)
    raw_sources = data.get('sources')
    sources: list[dict]
    if raw_sources is None:
        engine_raw = str(data.get('engine') or 'google').strip().lower()
        if engine_raw not in SUPPORTED_ENGINES:
            return jsonify({
                'ok': False,
                'error': 'Неподдерживаемая поисковая система',
                'supported': list(SUPPORTED_ENGINES),
            }), 400
        sources = [{'type': 'engine', 'engine': engine_raw}]
    elif isinstance(raw_sources, list):
        sources = []
        for entry in raw_sources:
            if isinstance(entry, str):
                engine_candidate = entry.strip().lower()
                if engine_candidate not in SUPPORTED_ENGINES and engine_candidate != 'local':
                    return jsonify({
                        'ok': False,
                        'error': f"Неподдерживаемый источник: {entry}",
                        'supported': list(SUPPORTED_ENGINES) + ['local'],
                    }), 400
                if engine_candidate == 'local':
                    sources.append({'type': 'local'})
                else:
                    sources.append({'type': 'engine', 'engine': engine_candidate})
            elif isinstance(entry, Mapping):
                obj = dict(entry)
                src_type = str(obj.get('type') or obj.get('source') or obj.get('engine') or '').lower()
                if src_type in SUPPORTED_ENGINES:
                    obj['type'] = 'engine'
                    obj['engine'] = src_type
                elif src_type in {'local', 'filesystem', 'catalog'}:
                    obj['type'] = 'local'
                    if src_type in {'filesystem', 'catalog'}:
                        options = dict(obj.get('options') or {})
                        options.setdefault('mode', src_type)
                        obj['options'] = options
                elif obj.get('engine'):
                    eng = str(obj['engine']).lower()
                    if eng not in SUPPORTED_ENGINES:
                        return jsonify({
                            'ok': False,
                            'error': f"Неподдерживаемая поисковая система: {eng}",
                            'supported': list(SUPPORTED_ENGINES),
                        }), 400
                    obj['engine'] = eng
                    obj['type'] = 'engine'
                else:
                    return jsonify({'ok': False, 'error': f"Неподдерживаемый источник: {entry}"}), 400
                sources.append(obj)
            else:
                return jsonify({'ok': False, 'error': f"Неподдерживаемый источник: {entry}"}), 400
    else:
        return jsonify({'ok': False, 'error': 'Поле sources должно быть массивом'}), 400

    for src in sources:
        if src.get('type') == 'engine':
            engine_name = str(src.get('engine') or '').lower()
            if engine_name not in SUPPORTED_ENGINES:
                return jsonify({
                    'ok': False,
                    'error': f"Неподдерживаемая поисковая система: {engine_name}",
                    'supported': list(SUPPORTED_ENGINES),
                }), 400
            src.setdefault('max_results', max_results)
            src.setdefault('force_refresh', bool(data.get('force_refresh')))
            extra_params_raw = src.get('extra_params')
            if extra_params_raw and isinstance(extra_params_raw, Mapping):
                src['extra_params'] = _osint_sanitize_params(extra_params_raw)
        elif src.get('type') == 'local':
            opts = dict(src.get('options') or {})
            if src.get('mode'):
                opts.setdefault('mode', src.get('mode'))
            src['options'] = opts

    service = _get_osint_service()
    try:
        job = service.start_job(
            query=query,
            locale=locale,
            region=region,
            safe=safe,
            sources=sources,
            params={
                "include_html": include_html,
                "include_llm_payload": include_llm,
                "max_results": max_results,
                "build_ontology": build_ontology,
                "locales": locales,
                **({"schedule": schedule_payload} if schedule_payload is not None else {}),
            },
            user_id=getattr(user, "id", None),
        )
    except ValueError as exc:
        return jsonify({'ok': False, 'error': str(exc)}), 400
    except Exception as exc:  # noqa: BLE001
        app.logger.exception("OSINT search failed: %s", exc)
        return jsonify({'ok': False, 'error': 'Не удалось выполнить поиск', 'details': str(exc)}), 500
    try:
        _log_user_action(
            user,
            'osint_search_start',
            'osint',
            job.get('id'),
            detail=json.dumps({'query': query, 'sources': sources}, ensure_ascii=False),
        )
    except Exception:
        pass
    return jsonify({'ok': True, 'job': job})


@osint_bp.route('/jobs/<int:job_id>/schedule', methods=['POST'])
def api_osint_job_schedule(job_id: int):
    user = _require_admin()
    payload = request.get_json(silent=True) or {}
    schedule_data = payload.get('schedule') if isinstance(payload.get('schedule'), Mapping) else payload
    service = _get_osint_service()
    try:
        job = service.update_schedule(job_id, schedule_data)
    except ValueError as exc:
        return jsonify({'ok': False, 'error': str(exc)}), 400
    except Exception as exc:  # noqa: BLE001
        app.logger.exception("OSINT schedule update failed: %s", exc)
        return jsonify({'ok': False, 'error': 'Не удалось обновить расписание', 'details': str(exc)}), 500
    try:
        _log_user_action(
            user,
            'osint_schedule_update',
            'osint',
            job_id,
            detail=json.dumps(schedule_data or {}, ensure_ascii=False),
        )
    except Exception:
        pass
    return jsonify({'ok': True, 'job': job})


@osint_bp.route('/artifacts/<path:relative_path>', methods=['GET'])
def api_osint_artifact(relative_path: str):
    _require_admin()
    try:
        artifact_path = resolve_artifact_path(relative_path)
    except ValueError:
        return jsonify({'ok': False, 'error': 'invalid_artifact_path'}), 400
    if not artifact_path.exists() or not artifact_path.is_file():
        return jsonify({'ok': False, 'error': 'artifact_not_found'}), 404
    return send_file(artifact_path, conditional=True)


@osint_bp.route('/jobs/<int:job_id>', methods=['GET', 'DELETE'])
def api_osint_job_detail(job_id: int):
    user = _require_admin()
    service = _get_osint_service()
    if request.method == 'DELETE':
        deleted = service.delete_job(job_id)
        if not deleted:
            return jsonify({'ok': False, 'error': 'Задача не найдена'}), 404
        try:
            _log_user_action(user, 'osint_search_delete', 'osint', job_id)
        except Exception:
            pass
        return jsonify({'ok': True})
    snapshot = service.get_job(job_id)
    if not snapshot:
        return jsonify({'ok': False, 'error': 'Задача не найдена'}), 404
    return jsonify({'ok': True, 'job': snapshot})


@osint_bp.route('/jobs/<int:job_id>/export', methods=['GET'])
def api_osint_job_export(job_id: int):
    _require_admin()
    service = _get_osint_service()
    format_param = (request.args.get('format') or 'markdown').strip().lower()
    if format_param not in {'markdown', 'json'}:
        return jsonify({'ok': False, 'error': 'Поддерживаются только форматы markdown/json'}), 400
    snapshot = service.get_job(job_id)
    if not snapshot:
        return jsonify({'ok': False, 'error': 'Задача не найдена'}), 404
    if format_param == 'json':
        payload = json.dumps(snapshot, ensure_ascii=False, indent=2)
        response = make_response(payload)
        response.headers['Content-Type'] = 'application/json; charset=utf-8'
        response.headers['Content-Disposition'] = f'attachment; filename="osint-job-{job_id}.json"'
        return response
    try:
        markdown = service.export_job_markdown(job_id)
    except ValueError:
        return jsonify({'ok': False, 'error': 'Задача не найдена'}), 404
    response = make_response(markdown)
    response.headers['Content-Type'] = 'text/markdown; charset=utf-8'
    response.headers['Content-Disposition'] = f'attachment; filename="osint-job-{job_id}.md"'
    return response


@osint_bp.route('/jobs/<int:job_id>/sources/<source_id>/retry', methods=['POST'])
def api_osint_job_retry_source(job_id: int, source_id: str):
    _require_admin()
    payload = request.get_json(silent=True) or {}
    force_refresh = bool(payload.get('force_refresh', True))
    service = _get_osint_service()
    try:
        snapshot = service.retry_source(job_id, source_id, force_refresh=force_refresh)
    except ValueError as exc:
        message = str(exc)
        if message == "osint_job_not_found":
            return jsonify({'ok': False, 'error': 'Задача не найдена'}), 404
        if message == "osint_source_not_found":
            return jsonify({'ok': False, 'error': 'Источник не найден'}), 404
        return jsonify({'ok': False, 'error': message}), 400
    return jsonify({'ok': True, 'job': snapshot})


@osint_bp.route('/browser/<source_id>', methods=['GET', 'POST'])
def api_osint_browser(source_id: str):
    if source_id is None:
        abort(404)
    target_url = request.args.get('next') or request.args.get('url')
    if not target_url:
        return jsonify({'ok': False, 'error': 'missing_target_url'}), 400
    method = request.method
    params = None
    data = None
    if method == 'GET':
        params = {key: value for key, value in request.args.items() if key not in {'next', 'url'}}
    else:
        data = request.form.to_dict(flat=False)
    try:
        proxied = fetch_proxied_page(
            source_id,
            target_url,
            method=method,
            params=params or None,
            data=data or None,
            allow_redirects=True,
            timeout=30,
        )
    except Exception as exc:
        return jsonify({'ok': False, 'error': str(exc)}), 502
    html_content = rewrite_html_for_proxy(source_id, target_url, proxied.text)
    response = make_response(html_content, proxied.status_code)
    response.headers['Content-Type'] = proxied.headers.get('Content-Type', 'text/html')
    response.headers['X-Frame-Options'] = 'ALLOWALL'
    response.headers['Content-Security-Policy'] = "frame-ancestors *"
    return response


@osint_bp.route('/browser-session', methods=['POST'])
def api_osint_browser_session():
    _require_admin()
    data = request.get_json(silent=True) or {}
    source_id = (data.get('source_id') or '').strip()
    url = (data.get('url') or '').strip()
    if not source_id or not url:
        return jsonify({'ok': False, 'error': 'source_id_and_url_required'}), 400
    try:
        session = remote_browser_manager.create(source_id, url)
    except Exception as exc:
        app.logger.exception("remote browser creation failed: %s", exc)
        return jsonify({'ok': False, 'error': 'remote_browser_failed'}), 500
    return jsonify({'ok': True, 'viewport': session.viewport()})


@osint_bp.route('/browser-session/<source_id>/snapshot', methods=['GET'])
def api_osint_browser_session_snapshot(source_id: str):
    _require_admin()
    session = remote_browser_manager.get(source_id)
    if not session:
        return jsonify({'ok': False, 'error': 'session_not_found'}), 404
    image = session.screenshot()
    if not image:
        return jsonify({'ok': False, 'error': 'snapshot_unavailable'}), 204
    return Response(image, mimetype='image/png')


@osint_bp.route('/browser-session/<source_id>/action', methods=['POST'])
def api_osint_browser_session_action(source_id: str):
    _require_admin()
    session = remote_browser_manager.get(source_id)
    if not session:
        return jsonify({'ok': False, 'error': 'session_not_found'}), 404
    data = request.get_json(silent=True) or {}
    action_type = (data.get('type') or '').strip().lower()
    viewport = session.viewport()
    try:
        if action_type == 'click':
            rel_x = float(data.get('x', 0))
            rel_y = float(data.get('y', 0))
            x = rel_x * viewport.get('width', 1366)
            y = rel_y * viewport.get('height', 768)
            session.click(x, y)
        elif action_type == 'type':
            text = str(data.get('text') or '')
            session.type_text(text)
        else:
            return jsonify({'ok': False, 'error': 'unsupported_action'}), 400
    except Exception as exc:
        app.logger.exception("remote browser action failed: %s", exc)
        return jsonify({'ok': False, 'error': 'action_failed'}), 500
    return jsonify({'ok': True})


@osint_bp.route('/browser-session/<source_id>', methods=['DELETE'])
def api_osint_browser_session_close(source_id: str):
    _require_admin()
    remote_browser_manager.close(source_id)
    return jsonify({'ok': True})

def _admin_count(exclude_id: int | None = None) -> int:
    q = User.query.filter(User.role == 'admin')
    if exclude_id is not None:
        q = q.filter(User.id != exclude_id)
    try:
        return q.count()
    except Exception:
        return 0

@users_bp.route('/', methods=['GET', 'POST'])
def api_users_collection():
    _require_admin()
    if request.method == 'GET':
        users = User.query.order_by(func.lower(User.username)).all()
        return jsonify({'ok': True, 'users': [_user_to_payload(u) for u in users]})
    data = request.get_json(silent=True) or {}
    username = (data.get('username') or '').strip()
    password = (data.get('password') or '').strip()
    full_name = (data.get('full_name') or '').strip()
    role = _normalize_user_role(data.get('role'), default=ROLE_EDITOR)
    if len(username) < 3:
        return jsonify({'ok': False, 'error': 'Логин должен содержать минимум 3 символа'}), 400
    if len(password) < 6:
        return jsonify({'ok': False, 'error': 'Пароль должен содержать минимум 6 символов'}), 400
    if len(full_name) < 5:
        return jsonify({'ok': False, 'error': 'Укажите ФИО (минимум 5 символов)'}), 400
    if User.query.filter(func.lower(User.username) == username.lower()).first():
        return jsonify({'ok': False, 'error': 'Пользователь с таким логином уже существует'}), 409
    user = User(username=username, role=role, full_name=full_name)
    user.set_password(password)
    db.session.add(user)
    db.session.commit()
    _log_user_action(
        _load_current_user(),
        'user_create',
        'user',
        user.id,
        detail=json.dumps({'username': username, 'role': role, 'full_name': full_name}, ensure_ascii=False)
    )
    return jsonify({'ok': True, 'user': _user_to_payload(user)}), 201


@users_bp.route('/search')
@require_admin
def api_users_search():
    q = (request.args.get('q') or '').strip()
    try:
        limit = min(max(int(request.args.get('limit', '20')), 1), 50)
    except Exception:
        limit = 20
    qry = User.query
    if q:
        like = f"%{q}%"
        qry = qry.filter(or_(User.username.ilike(like), User.full_name.ilike(like)))
    users = qry.order_by(func.lower(User.username)).limit(limit).all()
    return jsonify({'ok': True, 'users': [{
        'id': u.id,
        'username': u.username,
        'full_name': u.full_name,
        'role': _normalize_user_role(u.role, default=ROLE_EDITOR),
    } for u in users]})


@users_bp.route('/roles', methods=['GET'])
@require_admin
def api_users_roles():
    roles = [
        {'id': ROLE_ADMIN, 'label': 'Администратор'},
        {'id': ROLE_EDITOR, 'label': 'Редактор'},
        {'id': ROLE_VIEWER, 'label': 'Наблюдатель'},
    ]
    return jsonify({'ok': True, 'roles': roles})

@users_bp.route('/<int:user_id>', methods=['PATCH', 'DELETE'])
def api_users_detail(user_id: int):
    admin = _require_admin()
    user = User.query.get_or_404(user_id)
    if request.method == 'DELETE':
        if user.id == admin.id:
            return jsonify({'ok': False, 'error': 'Нельзя удалить собственную учётную запись'}), 400
        if _user_role(user) == ROLE_ADMIN and _admin_count(exclude_id=user.id) == 0:
            return jsonify({'ok': False, 'error': 'Нельзя удалить последнего администратора'}), 400
        db.session.delete(user)
        db.session.commit()
        _log_user_action(admin, 'user_delete', 'user', user_id)
        return jsonify({'ok': True})
    data = request.get_json(silent=True) or {}
    updated = False
    new_role = data.get('role')
    if new_role is not None:
        role = _normalize_user_role(str(new_role), default=_user_role(user))
        current_role = _user_role(user)
        if role != current_role:
            if current_role == ROLE_ADMIN and role != ROLE_ADMIN and _admin_count(exclude_id=user.id) == 0:
                return jsonify({'ok': False, 'error': 'Должен остаться хотя бы один администратор'}), 400
            user.role = role
            updated = True
    new_password = data.get('password')
    if new_password:
        pw = str(new_password).strip()
        if len(pw) < 6:
            return jsonify({'ok': False, 'error': 'Пароль должен содержать минимум 6 символов'}), 400
        user.set_password(pw)
        updated = True
    if not updated:
        return jsonify({'ok': False, 'error': 'Нет изменений'}), 400
    db.session.commit()
    _log_user_action(admin, 'user_update', 'user', user.id, detail=json.dumps({'role': _user_role(user), 'password_changed': bool(new_password)}))
    return jsonify({'ok': True, 'user': _user_to_payload(user)})


@admin_bp.route('/users', methods=['GET', 'POST'])
def api_admin_users():
    return api_users_collection()


@admin_bp.route('/users/search')
@require_admin
def api_admin_users_search():
    return api_users_search()


@admin_bp.route('/users/<int:user_id>', methods=['PATCH', 'DELETE'])
def api_admin_user_detail(user_id: int):
    return api_users_detail(user_id)

# ------------------- Инициализация и миграции коллекций -------------------
def _slugify(s: str) -> str:
    s = (s or '').strip().lower()
    s = re.sub(r"[^a-z0-9_\-а-яё]+", "-", s)
    s = re.sub(r"-+", "-", s).strip('-')
    if not s:
        s = 'collection'
    return s

def ensure_collections_schema():
    """Проверить наличие модели/таблицы коллекций и поля files.collection_id, создать базовую коллекцию.
    Привязать базовую коллекцию к файлам, у которых она отсутствует.
    """
    try:
        with app.app_context():
            # Создать таблицы, если их ещё нет
            db.create_all()
            # Добавить столбец collection_id в files, если его нет
            try:
                from sqlalchemy import text as _text
                with db.engine.begin() as conn:
                    insp = sa_inspect(conn)

                    def _columns(table_name: str) -> set[str]:
                        try:
                            return {str(col.get("name")) for col in insp.get_columns(table_name)}
                        except Exception:
                            return set()

                    files_cols = _columns("files")
                    if "collection_id" not in files_cols:
                        conn.execute(_text("ALTER TABLE files ADD COLUMN collection_id INTEGER"))

                    collections_cols = _columns("collections")
                    if "owner_id" not in collections_cols:
                        conn.execute(_text("ALTER TABLE collections ADD COLUMN owner_id INTEGER"))
                    if "is_private" not in collections_cols:
                        conn.execute(_text("ALTER TABLE collections ADD COLUMN is_private BOOLEAN NOT NULL DEFAULT FALSE"))

                    users_cols = _columns("users")
                    if "full_name" not in users_cols:
                        conn.execute(_text("ALTER TABLE users ADD COLUMN full_name TEXT"))

                    tasks_cols = _columns("task_records")
                    if "user_id" not in tasks_cols:
                        conn.execute(_text("ALTER TABLE task_records ADD COLUMN user_id INTEGER"))
                    if "collection_id" not in tasks_cols:
                        conn.execute(_text("ALTER TABLE task_records ADD COLUMN collection_id INTEGER"))
            except Exception:
                pass
            # Убедиться, что базовая коллекция существует
            base = Collection.query.filter_by(slug='base').first()
            if not base:
                base = Collection(name='Базовая коллекция', slug='base', searchable=True, graphable=True)
                db.session.add(base)
                db.session.commit()
            # Назначить базовую коллекцию всем файлам с NULL в collection_id
            try:
                File.query.filter(File.collection_id.is_(None)).update({File.collection_id: base.id})
                db.session.commit()
            except Exception:
                db.session.rollback()
    except Exception:
        # по возможности продолжаем дальше
        pass


def ensure_llm_schema():
    """Добавить дополнительные поля, используемые LLM-эндпоинтами."""
    try:
        with app.app_context():
            db.create_all()
            try:
                from sqlalchemy import text as _text
                with db.engine.begin() as conn:
                    insp = sa_inspect(conn)
                    cols = {str(col.get("name")) for col in insp.get_columns("llm_endpoints")}
                    if "provider" not in cols:
                        conn.execute(_text("ALTER TABLE llm_endpoints ADD COLUMN provider TEXT DEFAULT 'openai'"))
                    if "context_length" not in cols:
                        conn.execute(_text("ALTER TABLE llm_endpoints ADD COLUMN context_length INTEGER"))
                    if "instances" not in cols:
                        conn.execute(_text("ALTER TABLE llm_endpoints ADD COLUMN instances INTEGER NOT NULL DEFAULT 1"))
            except Exception:
                pass
    except Exception:
        pass


def _ensure_llm_schema_once():
    global LLM_SCHEMA_READY
    if LLM_SCHEMA_READY:
        return
    ensure_llm_schema()
    LLM_SCHEMA_READY = True


def _lmstudio_base_url(base_url: str) -> str:
    base = str(base_url or "").strip().rstrip("/")
    if not base:
        return ""
    if base.endswith("/v1"):
        return base[:-3]
    return base


def _lmstudio_headers(api_key: str | None) -> dict:
    headers = {"Content-Type": "application/json"}
    token = str(api_key or "").strip()
    if token:
        headers["Authorization"] = f"Bearer {token}"
    return headers


def _lmstudio_extract_models(payload: object) -> list[dict]:
    """Extract model list from LM Studio responses across versions/formats."""

    def _iter_candidate_nodes(node: object):
        if isinstance(node, list):
            for item in node:
                yield item
                yield from _iter_candidate_nodes(item)
        elif isinstance(node, dict):
            # prioritize common top-level containers
            for key in ("data", "models", "available", "loaded", "items", "results"):
                value = node.get(key)
                if isinstance(value, (list, dict)):
                    yield value
            for value in node.values():
                if isinstance(value, (list, dict)):
                    yield from _iter_candidate_nodes(value)

    def _pick_id_from_dict(d: dict) -> str:
        for key in (
            "id",
            "key",
            "model",
            "model_id",
            "modelId",
            "model_key",
            "modelKey",
            "identifier",
            "slug",
            "uid",
            "path",
            "name",
        ):
            value = d.get(key)
            if isinstance(value, str) and value.strip():
                return value.strip()
        model_node = d.get("model")
        if isinstance(model_node, dict):
            for key in ("id", "key", "identifier", "name", "modelKey"):
                value = model_node.get(key)
                if isinstance(value, str) and value.strip():
                    return value.strip()
        return ""

    def _normalize_item(item: object) -> dict | None:
        if isinstance(item, str):
            token = item.strip()
            if not token:
                return None
            return {"id": token, "name": token, "loaded": None}
        if not isinstance(item, dict):
            return None
        ident = _pick_id_from_dict(item)
        if not ident:
            return None
        loaded_raw = item.get("loaded")
        if loaded_raw is None:
            state = str(item.get("state") or item.get("status") or "").strip().lower()
            if state:
                loaded_raw = state in {"loaded", "active", "ready"}
        loaded = bool(loaded_raw) if loaded_raw is not None else None
        model_node = item.get("model") if isinstance(item.get("model"), dict) else None
        name = str(
            item.get("name")
            or item.get("display_name")
            or item.get("displayName")
            or item.get("title")
            or (model_node.get("name") if model_node else None)
            or ident
        ).strip() or ident
        return {"id": ident, "name": name, "loaded": loaded}

    out: list[dict] = []
    seen: set[str] = set()
    roots: list[object] = [payload]
    roots.extend(list(_iter_candidate_nodes(payload)))
    for root in roots:
        if isinstance(root, list):
            iterable = root
        elif isinstance(root, dict):
            iterable = [root]
        else:
            continue
        for raw in iterable:
            norm = _normalize_item(raw)
            if not norm:
                continue
            ident = norm["id"]
            if ident in seen:
                continue
            seen.add(ident)
            out.append(norm)
    return out


def _lmstudio_mark_activity() -> None:
    global LMSTUDIO_LAST_ACTIVITY_TS, LMSTUDIO_IDLE_UNLOADED
    with LMSTUDIO_IDLE_LOCK:
        LMSTUDIO_LAST_ACTIVITY_TS = time.time()
        LMSTUDIO_IDLE_UNLOADED = False


def _lmstudio_unload_all_if_idle() -> None:
    global LMSTUDIO_IDLE_UNLOADED
    idle_minutes = max(0, int(getattr(_rt(), "lmstudio_idle_unload_minutes", 0) or 0))
    if idle_minutes <= 0:
        return
    with LMSTUDIO_IDLE_LOCK:
        elapsed = time.time() - LMSTUDIO_LAST_ACTIVITY_TS
        if elapsed < idle_minutes * 60:
            return
        if LMSTUDIO_IDLE_UNLOADED:
            return
        LMSTUDIO_IDLE_UNLOADED = True
    try:
        endpoints = LlmEndpoint.query.all()
    except Exception:
        endpoints = []
    seen: set[tuple[str, str]] = set()
    for ep in endpoints:
        if (ep.provider or "openai") not in {"openai", "openrouter", "azure_openai"}:
            continue
        base = _lmstudio_base_url(ep.base_url)
        model = str(ep.model or "").strip()
        if not base or not model:
            continue
        ident = (base, model)
        if ident in seen:
            continue
        seen.add(ident)
        try:
            http_request(
                "POST",
                f"{base}/api/v1/models/unload",
                headers=_lmstudio_headers(ep.api_key),
                json={"model": model},
                timeout=(HTTP_CONNECT_TIMEOUT, max(10, HTTP_DEFAULT_TIMEOUT)),
                logger=app.logger,
            )
        except Exception:
            app.logger.debug("LM Studio idle unload failed for %s / %s", base, model)


def _start_lmstudio_idle_unload_scheduler() -> None:
    global LMSTUDIO_IDLE_UNLOAD_THREAD_STARTED
    if LMSTUDIO_IDLE_UNLOAD_THREAD_STARTED:
        return

    def loop():
        while True:
            try:
                with app.app_context():
                    _lmstudio_unload_all_if_idle()
            except Exception:
                app.logger.exception("LM Studio idle unload loop error")
            time.sleep(30)

    threading.Thread(target=loop, name="lmstudio-idle-unloader", daemon=True).start()
    LMSTUDIO_IDLE_UNLOAD_THREAD_STARTED = True

# ------------------- Вспомогательные функции безопасности путей -------------------
def _resolve_under_base(rel_path: str) -> tuple[Path, Path | None]:
    """Безопасно разрешать rel_path внутри UPLOAD_FOLDER. Возвращает (base_dir, abs_path или None, если путь вне папки)."""
    base_dir = Path(app.config.get('UPLOAD_FOLDER') or '.').resolve()
    try:
        # Нормализовать разделители и убрать ведущие слэши
        rel = str(rel_path).replace('\\', '/').lstrip('/')
        abs_path = (base_dir / rel).resolve()
        abs_path.relative_to(base_dir)
        return base_dir, abs_path
    except Exception:
        return base_dir, None

# ------------------- Routes -------------------

@app.after_request
def _force_utf8(resp):
    try:
        ct = (resp.headers.get('Content-Type') or '').lower()
        if ct.startswith('text/html'):
            resp.headers['Content-Type'] = 'text/html; charset=utf-8'
        elif ct.startswith('application/json'):
            # обеспечить кодировку utf-8 и для JSON
            resp.headers['Content-Type'] = 'application/json; charset=utf-8'
        return resp
    except Exception:
        return resp


@app.route('/health', methods=['GET'])
def health_check():
    checks: dict[str, dict[str, object]] = {}
    overall_status = 'ok'

    # Проверка базы данных
    try:
        db.session.execute(text('SELECT 1'))
        db.session.commit()
        checks['database'] = {'status': 'ok'}
    except Exception as exc:
        db.session.rollback()
        checks['database'] = {'status': 'error', 'detail': str(exc)}
        overall_status = 'degraded'

    # Проверка очереди задач
    queue = get_task_queue()
    queue_stats = queue.stats()
    queue_status = 'ok' if queue_stats.get('started') and not queue_stats.get('shutdown') else 'warning'
    if queue_status != 'ok':
        overall_status = 'degraded'
    checks['task_queue'] = {'status': queue_status, 'stats': queue_stats}
    llm_stats = get_llm_pool().stats()
    llm_status = 'ok'
    if int(llm_stats.get('queued') or 0) >= int(llm_stats.get('max_queue') or 1):
        llm_status = 'warning'
        overall_status = 'degraded'
    checks['llm_queue'] = {'status': llm_status, 'stats': llm_stats}

    # Проверка фонового шедулера очистки кэша
    if CACHE_CLEANUP_INTERVAL_HOURS > 0 and not _CLEANUP_THREAD_STARTED:
        checks['cache_cleanup'] = {'status': 'warning', 'detail': 'cleanup thread not started'}
        overall_status = 'degraded'
    else:
        checks['cache_cleanup'] = {'status': 'ok'}

    return jsonify({'status': overall_status, 'checks': checks})


@app.route('/metrics', methods=['GET'])
def metrics():
    lines: list[str] = []
    queue_stats = get_task_queue().stats()
    llm_pool_stats = get_llm_pool().stats()
    lines.append(f"task_queue_queued {queue_stats.get('queued', 0)}")
    lines.append(f"task_queue_workers {queue_stats.get('workers', 0)}")
    lines.append(f"task_queue_started {int(bool(queue_stats.get('started')))}")
    lines.append(f"task_queue_shutdown {int(bool(queue_stats.get('shutdown')))}")
    lines.append(f"llm_queue_queued {llm_pool_stats.get('queued', 0)}")
    lines.append(f"llm_queue_running {llm_pool_stats.get('running', 0)}")
    lines.append(f"llm_pool_workers {llm_pool_stats.get('workers', 0)}")

    try:
        totals = (
            db.session.query(TaskRecord.status, func.count(TaskRecord.id))
            .group_by(TaskRecord.status)
            .all()
        )
        for status, count in totals:
            metric_name = f"tasks_total_status_{(status or 'unknown').replace('-', '_')}"
            lines.append(f"{metric_name} {count}")
        db_status = '1'
    except Exception:
        db.session.rollback()
        db_status = '0'
    lines.append(f"database_available {db_status}")

    payload = "\n".join(lines) + "\n"
    return Response(payload, mimetype='text/plain; version=0.0.4; charset=utf-8')


# Загрузка файлов через веб-интерфейс
@app.route("/upload", methods=["GET", "POST"])
def upload_file():
    user = _load_current_user()
    if not user:
        abort(401)
    is_admin = _user_role(user) == ROLE_ADMIN
    if request.method == "GET":
        allowed = getattr(g, 'allowed_collection_ids', set())
        cols_q = Collection.query.order_by(Collection.name.asc())
        if allowed is not None:
            if not allowed:
                cols_q = cols_q.filter(Collection.id == -1)
            else:
                cols_q = cols_q.filter(Collection.id.in_(allowed))
        cols = cols_q.all()
        return jsonify({
            "allowed_extensions": sorted(ALLOWED_EXTS),
            "collections": [{"id": c.id, "name": c.name} for c in cols],
        })

    file = request.files.get("file")
    if not file or not file.filename:
        return json_error("Файл не выбран", 400)

    try:
        collection = _resolve_upload_collection(user, request.form, is_admin)
    except PermissionError:
        abort(403)
    except ValueError as exc:
        return json_error(str(exc), 400)
    except Exception as exc:
        return json_error(f"Не удалось подготовить коллекцию: {exc}", 500)

    try:
        save_result = _save_upload_to_disk(file)
    except ValueError as exc:
        return json_error(str(exc), 415)
    except Exception as exc:
        return json_error(f"Не удалось сохранить файл: {exc}", 500)

    try:
        import_result = _import_file_record(
            save_result['path'],
            save_result['rel_path'],
            collection,
            user_id=user.id if user else None,
            generate_metadata=True,
        )
    except Exception as exc:
        return json_error(f"Ошибка обработки файла: {exc}", 500)

    try:
        _log_user_action(
            _load_current_user(),
            'file_upload',
            'file',
            import_result.get('file_id'),
            detail=json.dumps({'filename': import_result.get('preview', {}).get('filename'), 'collection_id': getattr(collection, 'id', None)}),
        )
    except Exception:
        pass
    return jsonify({
        "status": "ok",
        "file_id": import_result.get('file_id'),
        "rel_path": import_result.get('rel_path'),
        "preview": import_result.get('preview'),
    })


@app.route('/api/import/jobs', methods=['POST'])
def api_import_jobs_create():
    user = _load_current_user()
    if not user:
        abort(401)
    if not _user_can_upload(user):
        abort(403)

    file = request.files.get('file')
    if not file or not file.filename:
        return json_error('Файл не выбран', 400)

    is_admin = _user_role(user) == ROLE_ADMIN
    try:
        collection = _resolve_upload_collection(user, request.form, is_admin)
    except PermissionError:
        abort(403)
    except ValueError as exc:
        return json_error(str(exc), 400)
    except Exception as exc:
        return json_error(f"Не удалось подготовить коллекцию: {exc}", 500)

    try:
        save_result = _save_upload_to_disk(file)
    except ValueError as exc:
        return json_error(str(exc), 415)
    except Exception as exc:
        return json_error(f"Не удалось сохранить файл: {exc}", 500)

    initial_preview = _build_initial_preview(save_result, collection)
    payload = {
        'kind': 'import_file',
        'submitted_by': user.id,
        'collection_id': getattr(collection, 'id', None),
        'filename': save_result.get('filename'),
        'rel_path': save_result.get('rel_path'),
        'initial_preview': initial_preview,
    }
    try:
        task = TaskRecord(
            name='import_file',
            status='queued',
            payload=json.dumps(payload, ensure_ascii=False),
            progress=0.0,
            user_id=user.id,
            collection_id=collection_id,
        )
        db.session.add(task)
        db.session.commit()
    except Exception as exc:
        db.session.rollback()
        return json_error(f"Не удалось создать задачу импорта: {exc}", 500)

    spec = {
        'path': str(save_result['path']),
        'rel_path': save_result['rel_path'],
        'collection_id': getattr(collection, 'id', None),
        'user_id': user.id,
        'filename': save_result.get('filename'),
    }
    _enqueue_import_job(task.id, spec, description=save_result.get('filename'))
    return jsonify({'ok': True, 'task_id': task.id, 'initial_preview': initial_preview})


def _task_payload_json(task: TaskRecord) -> dict:
    if not task.payload:
        return {}
    try:
        return json.loads(task.payload)
    except Exception:
        return {}


@app.route('/api/import/jobs', methods=['GET'])
def api_import_jobs_list():
    user = _load_current_user()
    if not user:
        abort(401)
    role = _user_role(user)
    tasks_query = TaskRecord.query.filter(TaskRecord.name == 'import_file').order_by(TaskRecord.created_at.desc())
    tasks = tasks_query.limit(50 if role == ROLE_ADMIN else 20).all()
    items = []
    for task in tasks:
        payload = _task_payload_json(task)
        if role != ROLE_ADMIN and payload.get('submitted_by') not in (user.id, None):
            continue
        info = _task_to_dict(task)
        info['payload_json'] = payload
        items.append(info)
    return jsonify({'ok': True, 'tasks': items})


@app.route('/api/import/jobs/<int:task_id>', methods=['GET'])
def api_import_jobs_status(task_id: int):
    user = _load_current_user()
    if not user:
        abort(401)
    task = TaskRecord.query.get_or_404(task_id)
    payload = _task_payload_json(task)
    if _user_role(user) != ROLE_ADMIN and payload.get('submitted_by') != user.id:
        abort(403)
    info = _task_to_dict(task)
    info['payload_json'] = payload
    return jsonify({'ok': True, 'task': info})


def _resolve_upload_collection(user: User | None, form, is_admin: bool) -> Collection | None:
    allowed = getattr(g, 'allowed_collection_ids', set())
    col: Collection | None = None
    new_col = (form.get('new_collection') or '').strip()
    col_id = form.get('collection_id')

    try:
        if new_col:
            slug = _slugify(new_col)
            col = Collection.query.filter((Collection.slug == slug) | (Collection.name == new_col)).first()
            if not col:
                requested_private = str(form.get('private', '')).lower() in ('1', 'true', 'yes', 'on')
                is_private_flag = requested_private or not is_admin
                col = Collection(
                    name=new_col,
                    slug=slug,
                    searchable=True,
                    graphable=True,
                    owner_id=user.id if user else None,
                    is_private=is_private_flag,
                )
                db.session.add(col)
                db.session.commit()
                if user:
                    try:
                        if not CollectionMember.query.filter_by(collection_id=col.id, user_id=user.id).first():
                            db.session.add(CollectionMember(collection_id=col.id, user_id=user.id, role='owner'))
                            db.session.commit()
                    except Exception:
                        db.session.rollback()
                if allowed is not None:
                    try:
                        allowed = set(allowed)
                        allowed.add(col.id)
                        g.allowed_collection_ids = allowed
                    except Exception:
                        pass
        elif col_id:
            try:
                col = Collection.query.get(int(col_id))
            except Exception:
                col = None
    except Exception:
        db.session.rollback()
        raise

    if not col and user:
        try:
            col = Collection.query.filter_by(owner_id=user.id, is_private=True).first()
        except Exception:
            col = None
    if not col:
        col = Collection.query.filter_by(slug='base').first() or Collection.query.first()

    if col:
        if allowed is not None and col.id not in allowed:
            if not _has_collection_access(col.id, write=True):
                raise PermissionError('Недостаточно прав для загрузки в коллекцию')
        else:
            if not _has_collection_access(col.id, write=True):
                raise PermissionError('Недостаточно прав для загрузки в коллекцию')
    return col


def _save_upload_to_disk(file_storage) -> dict:
    safe_name = secure_filename(file_storage.filename or '')
    ext = Path(safe_name).suffix.lower()
    if ext not in ALLOWED_EXTS:
        raise ValueError(f"Недопустимый тип файла: {ext}")

    scan_root = _scan_root_path()
    subdir = _import_subdir_value()
    base_dir = scan_root / subdir if subdir else scan_root
    base_dir.mkdir(parents=True, exist_ok=True)
    target = base_dir / safe_name
    stem = Path(safe_name).stem
    counter = 1
    while target.exists():
        target = base_dir / f"{stem}_{counter}{ext}"
        counter += 1
    file_storage.save(target)

    try:
        rel_path = str(target.relative_to(scan_root))
    except Exception:
        rel_path = target.name

    stat = target.stat()
    return {
        'path': target,
        'rel_path': rel_path,
        'ext': ext,
        'size': stat.st_size,
        'filename': target.name,
    }


def _build_initial_preview(save_result: dict, collection: Collection | None) -> dict:
    return {
        'filename': save_result.get('filename'),
        'ext': save_result.get('ext'),
        'size': int(save_result.get('size', 0) or 0),
        'collection_id': getattr(collection, 'id', None),
        'collection_name': getattr(collection, 'name', None) if collection else None,
    }


def _guess_language(text: str) -> str | None:
    if not text:
        return None
    cyr = sum(1 for ch in text if 'а' <= ch.lower() <= 'я' or ch.lower() == 'ё')
    lat = sum(1 for ch in text if 'a' <= ch.lower() <= 'z')
    total = cyr + lat
    if total < 20:
        return None
    return 'ru' if cyr >= lat else 'en'


def _extract_text_sample(path: Path, ext: str, limit: int = 6000) -> str:
    ext = ext.lower()
    try:
        if ext in {'.txt', '.md'}:
            return path.read_text(encoding='utf-8', errors='ignore')[:limit]
        if ext == '.pdf':
            return extract_text_pdf(path, limit_chars=limit)
        if ext == '.docx':
            return extract_text_docx(path, limit_chars=limit)
        if ext == '.rtf':
            return extract_text_rtf(path, limit_chars=limit)
        if ext == '.epub':
            return extract_text_epub(path, limit_chars=limit)
        if ext == '.djvu':
            return extract_text_djvu(path, limit_chars=limit)
    except Exception:
        return ''
    return ''


def _import_file_record(
    save_path: Path,
    rel_path: str,
    collection: Collection | None,
    user_id: int | None = None,
    *,
    generate_metadata: bool = True,
) -> dict:
    ext = save_path.suffix.lower()
    size = save_path.stat().st_size
    mtime = save_path.stat().st_mtime
    try:
        sha1 = sha1_of_file(save_path)
    except Exception:
        sha1 = None

    fobj = File.query.filter_by(path=str(save_path)).first()
    if not fobj:
        fobj = File(path=str(save_path))
        db.session.add(fobj)

    fobj.rel_path = rel_path
    fobj.filename = save_path.stem
    fobj.ext = ext
    fobj.size = size
    fobj.mtime = mtime
    fobj.sha1 = sha1
    if collection:
        fobj.collection_id = collection.id

    try:
        db.session.flush()
    except Exception as exc:
        db.session.rollback()
        raise exc

    text_sample = _extract_text_sample(save_path, ext)
    if text_sample:
        fobj.text_excerpt = text_sample[:4000]

    metadata_summary: dict[str, object] | None = None
    if generate_metadata:
        metadata = None
        sample_for_llm = text_sample or ''
        if not sample_for_llm and ext in AUDIO_EXTS:
            sample_for_llm = ''
        if sample_for_llm:
            try:
                metadata = call_lmstudio_for_metadata(sample_for_llm[:15000], save_path.name)
            except Exception as exc:
                logger.debug('LLM metadata generation failed for %s: %s', save_path.name, exc)
        if isinstance(metadata, dict) and metadata:
            metadata_summary = metadata
            mt_meta = normalize_material_type(metadata.get("material_type"))
            if ext in AUDIO_EXTS:
                fobj.material_type = 'audio'
            elif ext in IMAGE_EXTS:
                fobj.material_type = 'image'
            else:
                merged_type = _merge_llm_material_type(
                    fobj.material_type,
                    mt_meta,
                    allow_override=bool(TYPE_LLM_OVERRIDE),
                )
                if merged_type:
                    fobj.material_type = merged_type
            title_meta = (metadata.get('title') or '').strip()
            if title_meta:
                fobj.title = title_meta
            author_meta = _normalize_author(metadata.get('author'))
            if author_meta:
                fobj.author = author_meta
            year_meta = _normalize_year(metadata.get('year'))
            if year_meta:
                fobj.year = year_meta
            advisor = metadata.get('advisor')
            if advisor:
                advisor_str = str(advisor).strip()
                if advisor_str:
                    fobj.advisor = advisor_str
            kws = metadata.get('keywords') or []
            if isinstance(kws, list) and kws:
                fobj.keywords = ", ".join([str(x) for x in kws][:50])
            try:
                db.session.flush()
                if metadata.get('novelty'):
                    upsert_tag(fobj, 'научная новизна', str(metadata['novelty']))
                for extra_key in ('literature', 'organizations', 'classification'):
                    extra_val = metadata.get(extra_key)
                    if isinstance(extra_val, list) and extra_val:
                        upsert_tag(fobj, extra_key, "; ".join([str(x) for x in extra_val]))
            except Exception:
                db.session.rollback()
                db.session.add(fobj)
    try:
        _apply_metadata_quality_rules(
            fobj,
            ext=ext,
            text_excerpt=fobj.text_excerpt or "",
            filename=save_path.stem,
        )
    except Exception:
        pass

    try:
        _sync_file_to_fts(fobj)
    except Exception as exc:
        logger.debug('FTS sync failed for file %s: %s', getattr(fobj, 'id', None), exc)
    _invalidate_facets_cache('import')
    db.session.commit()

    language = _guess_language(fobj.text_excerpt or '') if getattr(fobj, 'text_excerpt', None) else None

    preview = {
        'filename': save_path.name,
        'ext': ext,
        'size': size,
        'collection_id': getattr(collection, 'id', None),
        'collection_name': getattr(collection, 'name', None) if collection else None,
        'title': fobj.title,
        'author': fobj.author,
        'material_type': fobj.material_type,
        'language': language,
        'text_excerpt': (fobj.text_excerpt or '')[:400],
    }
    if metadata_summary:
        preview['metadata'] = metadata_summary

    return {
        'file_id': fobj.id,
        'rel_path': rel_path,
        'preview': preview,
    }


def _enqueue_import_job(task_id: int, spec: dict, description: str | None = None) -> None:
    @copy_current_app_context
    def _runner():
        _run_import_job(task_id, spec)

    owner_id = None
    try:
        if spec.get('user_id') is not None:
            owner_id = int(spec.get('user_id'))
    except Exception:
        owner_id = None
    get_task_queue().submit(
        _runner,
        description=description or f"import-{spec.get('filename', '')}",
        owner_id=owner_id,
    )


def _run_import_job(task_id: int, spec: dict) -> None:
    task = TaskRecord.query.get(task_id)
    if not task:
        return
    try:
        task.status = 'running'
        task.started_at = datetime.utcnow()
        task.progress = 0.05
        db.session.commit()
    except Exception:
        db.session.rollback()

    path = Path(spec.get('path'))
    rel_path = spec.get('rel_path')
    collection = _get_collection_instance(spec.get('collection_id'))
    user_id = spec.get('user_id')

    try:
        if not path.exists():
            raise FileNotFoundError(f"Файл {path} недоступен")
        info = _import_file_record(path, rel_path, collection, user_id=user_id, generate_metadata=True)
        payload = {
            'kind': 'import_file',
            'submitted_by': user_id,
            'collection_id': getattr(collection, 'id', None),
            'filename': spec.get('filename'),
            'rel_path': rel_path,
            'result': info,
        }
        task = TaskRecord.query.get(task_id)
        if task:
            task.status = 'completed'
            task.progress = 1.0
            task.finished_at = datetime.utcnow()
            task.payload = json.dumps(payload, ensure_ascii=False)
            db.session.commit()
        if user_id:
            try:
                actor = User.query.get(user_id)
                if actor:
                    _log_user_action(actor, 'file_import_async', 'file', info.get('file_id'), detail=json.dumps({'filename': spec.get('filename'), 'collection_id': getattr(collection, 'id', None)}))
            except Exception:
                pass
    except Exception as exc:
        db.session.rollback()
        task = TaskRecord.query.get(task_id)
        if task:
            task.status = 'error'
            task.error = str(exc)
            task.finished_at = datetime.utcnow()
            payload = {
                'kind': 'import_file',
                'submitted_by': user_id,
                'collection_id': getattr(collection, 'id', None),
                'filename': spec.get('filename'),
                'rel_path': rel_path,
                'error': str(exc),
            }
            try:
                task.payload = json.dumps(payload, ensure_ascii=False)
            except Exception:
                task.payload = None
            db.session.commit()
        app.logger.warning('Import job %s failed: %s', task_id, exc)

# Утилита: безопасная относительная дорожка (для загрузки папок)
def _sanitize_relpath(p: str) -> str:
    p = (p or '').replace('\\', '/').lstrip('/')
    parts = [seg for seg in p.split('/') if seg not in ('', '.', '..')]
    return '/'.join(parts)


def _get_collection_instance(collection_ref) -> Collection | None:
    if isinstance(collection_ref, Collection):
        return collection_ref
    try:
        cid = int(collection_ref)
    except Exception:
        cid = None
    if cid:
        try:
            return Collection.query.get(cid)
        except Exception:
            return None
    return None


def _collection_dir_slug(col: Collection | None) -> str:
    if not col:
        return ''
    slug = (col.slug or '').strip()
    if slug:
        return slug
    name = (col.name or '').strip()
    if name:
        return _slugify(name)
    return f"collection-{col.id or 'unknown'}"


def _collection_root_dir(col: Collection | None, *, ensure: bool = True) -> Path:
    root = _scan_root_path()
    if not _collections_in_separate_dirs() or col is None:
        if ensure:
            root.mkdir(parents=True, exist_ok=True)
        return root
    slug = _collection_dir_slug(col) or f"collection-{col.id or 'unknown'}"
    target = root / 'collections' / slug
    if ensure:
        target.mkdir(parents=True, exist_ok=True)
    return target


def _import_base_dir_for_collection(col: Collection | None) -> Path:
    base_root = _collection_root_dir(col)
    sub = _import_subdir_value()
    return (base_root / sub) if sub else base_root

@app.route('/import', methods=['GET', 'POST'])
def import_files():
    """Импорт нескольких файлов и папок (через webkitdirectory) с возможным автосканом."""
    user = _load_current_user()
    if not user:
        abort(401)
    is_admin = _user_role(user) == ROLE_ADMIN
    scan_root = _scan_root_path()
    if request.method == 'GET':
        allowed = getattr(g, 'allowed_collection_ids', set())
        q = Collection.query.order_by(Collection.name.asc())
        if allowed is not None:
            if not allowed:
                q = q.filter(Collection.id == -1)
            else:
                q = q.filter(Collection.id.in_(allowed))
        cols = q.all()
        return jsonify({
            "allowed_extensions": sorted(ALLOWED_EXTS),
            "collections": [{"id": c.id, "name": c.name} for c in cols],
        })

    files = request.files.getlist('files')
    if not files:
        return json_error('Файлы не выбраны', 400)

    col_id = request.form.get('collection_id')
    new_col = (request.form.get('new_collection') or '').strip()
    col: Collection | None = None
    allowed = getattr(g, 'allowed_collection_ids', set())
    try:
        if new_col:
            slug = _slugify(new_col)
            col = Collection.query.filter((Collection.slug == slug) | (Collection.name == new_col)).first()
            if not col:
                requested_private = str(request.form.get('private', '')).lower() in ('1','true','yes','on')
                is_private_flag = requested_private or not is_admin
                col = Collection(name=new_col, slug=slug, searchable=True, graphable=True,
                                 owner_id=user.id, is_private=is_private_flag)
                db.session.add(col)
                db.session.commit()
                if user:
                    try:
                        if not CollectionMember.query.filter_by(collection_id=col.id, user_id=user.id).first():
                            db.session.add(CollectionMember(collection_id=col.id, user_id=user.id, role='owner'))
                            db.session.commit()
                    except Exception:
                        db.session.rollback()
                if allowed is not None:
                    try:
                        allowed = set(allowed)
                        allowed.add(col.id)
                        g.allowed_collection_ids = allowed
                    except Exception:
                        pass
        elif col_id:
            try:
                col = Collection.query.get(int(col_id))
            except Exception:
                col = None
    except Exception as e:
        db.session.rollback()
        return json_error(f"Не удалось создать коллекцию: {e}", 500)
    if not col and user:
        try:
            col = Collection.query.filter_by(owner_id=user.id, is_private=True).first()
        except Exception:
            col = None
    if not col:
        col = Collection.query.filter_by(slug='base').first() or Collection.query.first()
    if col:
        if allowed is not None and col.id not in allowed:
            allowed = None
        if not _has_collection_access(col.id, write=True):
            abort(403)

    base_dir = _import_base_dir_for_collection(col)
    saved = 0
    skipped = 0
    duplicates_found = 0
    saved_rel_paths: list[str] = []
    base_dir.mkdir(parents=True, exist_ok=True)

    for fs in files:
        raw_name = fs.filename or ''
        rel = _sanitize_relpath(raw_name)
        ext = Path(rel).suffix.lower()
        if not rel or ext not in ALLOWED_EXTS:
            skipped += 1
            continue
        dest = base_dir / rel
        if dest.exists():
            skipped += 1
            duplicates_found += 1
            app.logger.info(f'Skip duplicate by path: {dest}')
            continue
        try:
            dest.parent.mkdir(parents=True, exist_ok=True)
            fs.save(dest)
        except Exception as e:
            app.logger.warning(f'Failed to save %s: %s', rel, e)
            skipped += 1
            continue

        try:
            try:
                sha1 = sha1_of_file(dest)
            except Exception:
                sha1 = None

            duplicate_rec = None
            if sha1:
                try:
                    q = File.query.filter(File.sha1 == sha1)
                    if col:
                        q = q.filter(File.collection_id == col.id)
                    else:
                        q = q.filter(File.collection_id.is_(None))
                    duplicate_rec = q.first()
                except Exception:
                    duplicate_rec = None
            if duplicate_rec:
                try:
                    existing_path = Path(duplicate_rec.path) if duplicate_rec.path else None
                except Exception:
                    existing_path = None
                try:
                    if existing_path is None or existing_path.resolve() != dest.resolve():
                        try:
                            dest.unlink()
                        except Exception:
                            pass
                        duplicates_found += 1
                        skipped += 1
                        app.logger.info(f'Skip duplicate by hash: {dest}')
                        continue
                except Exception:
                    pass

            try:
                relp = str(dest.relative_to(scan_root))
            except Exception:
                relp = dest.name
            fobj = File.query.filter_by(path=str(dest)).first()
            if not fobj:
                fobj = File(
                    path=str(dest),
                    rel_path=relp,
                    filename=dest.stem,
                    ext=dest.suffix.lower(),
                    size=dest.stat().st_size,
                    mtime=dest.stat().st_mtime,
                    sha1=sha1,
                    collection_id=(col.id if col else None),
                )
                db.session.add(fobj)
            else:
                if col:
                    fobj.collection_id = col.id
            db.session.commit()
            saved_rel_paths.append(relp)
            saved += 1
        except Exception as e:
            db.session.rollback()
            app.logger.warning(f'Failed to register %s in DB: %s', rel, e)
            skipped += 1
            try:
                if dest.exists():
                    dest.unlink()
            except Exception:
                pass

    start_scan = request.form.get('start_scan') == 'on'
    extract_text = request.form.get('extract_text', 'on') == 'on'
    use_llm = request.form.get('use_llm') == 'on'
    prune = request.form.get('prune', 'on') == 'on'
    scan_started = False
    if start_scan and saved_rel_paths:
        if not is_admin:
            try:
                _log_user_action('scan_start_denied', 'collection', col.id if col else None,
                                 detail=json.dumps({'requested': len(saved_rel_paths)}))
            except Exception:
                pass
        else:
            paths_for_scan = [str((scan_root / p).resolve()) for p in saved_rel_paths]
            get_task_queue().submit(
                _run_scan_with_progress,
                extract_text,
                use_llm,
                prune,
                0,
                paths_for_scan,
                description="scan-import-batch",
                owner_id=getattr(_load_current_user(), 'id', None),
            )
            scan_started = True

    try:
        _log_user_action(
            _load_current_user(),
            'import_batch',
            'collection',
            col.id if col else None,
            detail=json.dumps({'saved': saved, 'skipped': skipped, 'duplicates': duplicates_found})
        )
    except Exception:
        pass
    return jsonify({
        "status": "ok",
        "saved": saved,
        "skipped": skipped,
        "duplicates": duplicates_found,
        "scan_started": scan_started,
        "imported": saved_rel_paths,
    })

# Просмотр и скачивание файлов
@app.route("/download/<path:rel_path>")
def download_file(rel_path):
    base_dir, abs_path = _resolve_under_base(rel_path)
    # Пытаемся заранее найти запись в БД и подготовить резерв, если файл переместили
    try:
        rp = str(rel_path)
        rp_alt = rp.replace('/', '\\') if ('/' in rp) else rp.replace('\\', '/')
        f = File.query.filter(or_(File.rel_path == rp, File.rel_path == rp_alt)).first()
    except Exception:
        f = None
    if (abs_path is None or not abs_path.exists()) and f is not None:
        try:
            filename_only = Path(f.rel_path or rp).name
            type_dirs = app.config.get('TYPE_DIRS') or {}
            sub = type_dirs.get((f.material_type or '').strip().lower())
            candidates = []
            if sub:
                candidates.append((base_dir / sub / filename_only).resolve())
            for v in set(type_dirs.values()):
                candidates.append((base_dir / v / filename_only).resolve())
            candidates.append((base_dir / filename_only).resolve())
            for cand in candidates:
                try:
                    cand.relative_to(base_dir)
                except Exception:
                    continue
                if cand.exists():
                    abs_path = cand
                    try:
                        f.path = str(cand)
                        f.rel_path = str(cand.relative_to(base_dir))
                        db.session.commit()
                    except Exception:
                        db.session.rollback()
                    # обновляем rel_path для дальнейших URL
                    rel_path = f.rel_path
                    break
            # В крайнем случае ищем по имени файла во всём base_dir (может быть медленно)
            if not (abs_path and abs_path.exists()):
                try:
                    for cand in base_dir.rglob(filename_only):
                        if cand.is_file():
                            abs_path = cand.resolve()
                            try:
                                abs_path.relative_to(base_dir)
                            except Exception:
                                continue
                            try:
                                f.path = str(cand)
                                f.rel_path = str(cand.relative_to(base_dir))
                                db.session.commit()
                            except Exception:
                                db.session.rollback()
                            rel_path = f.rel_path
                            break
                except Exception:
                    pass
        except Exception:
            pass
    if abs_path is None or not abs_path.exists():
        return ("Файл не найден.", 404)
    return send_from_directory(str(base_dir), str(Path(rel_path).as_posix()), as_attachment=True)

@app.route("/media/<path:rel_path>")
def media_file(rel_path):
    """Отдавать файл встраиваемо (для аудиоплеера в предпросмотре), без принудительного скачивания."""
    base_dir, abs_path = _resolve_under_base(rel_path)
    if abs_path is None or not abs_path.exists():
        # Встроенный 404 упрощает работу iframe/плеера
        return ("Not Found", 404)
    return send_from_directory(str(base_dir), str(Path(rel_path).as_posix()), as_attachment=False)

@app.route("/view/<path:rel_path>")
def view_file(rel_path):
    base_dir, abs_path = _resolve_under_base(rel_path)
    if abs_path is None or not abs_path.exists():
        return ("Файл не найден.", 404)
    ext = abs_path.suffix.lower()
    if ext == ".pdf":
        return send_from_directory(str(base_dir), str(Path(rel_path).as_posix()))
    elif ext in {".txt", ".md"}:
        try:
            content = abs_path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            content = "Не удалось прочитать файл."
        return Response(content, mimetype="text/plain; charset=utf-8")
    return json_error("Просмотр этого типа файлов не поддерживается", 415)


@app.route('/preview/<path:rel_path>')
def preview_file(rel_path):
    base_dir, abs_path = _resolve_under_base(rel_path)
    if abs_path is None or not abs_path.exists():
        body = (
            "<div style=\"min-height:100vh;display:flex;align-items:center;justify-content:center;"
            "font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;background:#0d1117;color:#c9d1d9;"
            "padding:24px;text-align:center;\">"
            "<div style=\"max-width:400px;background:#161b22;padding:28px;border-radius:18px;"
            "border:1px solid #30363d;box-shadow:0 20px 48px rgba(0,0,0,0.45);\">"
            "Файл не найден или недоступен.</div></div>"
        )
        return _html_page("Файл не найден", body, status=404)

    ext = abs_path.suffix.lower()
    is_pdf = ext == '.pdf'
    is_text = ext in {'.txt', '.md'}
    is_image = ext in IMAGE_EXTS
    is_audio = ext in AUDIO_EXTS
    is_docx = ext == '.docx'
    content = ''
    thumbnail_url = None
    abstract = ''
    audio_url = None
    duration = None
    image_url = None
    docx_html = None

    try:
        # Кэш‑каталог для текстовых фрагментов
        cache_dir = Path(app.static_folder) / 'cache' / 'text_excerpts'
        cache_dir.mkdir(parents=True, exist_ok=True)
        # по возможности используем sha1 в качестве ключа
        sha = None
        try:
            import hashlib
            with abs_path.open('rb') as fh:
                b = fh.read(1024*16)
                sha = hashlib.sha1(b).hexdigest()
        except Exception:
            sha = None

        cache_key = (sha or rel_path.replace('/', '_'))
        cache_file = cache_dir / (cache_key + '.txt')
        if cache_file.exists():
            content = cache_file.read_text(encoding='utf-8', errors='ignore')
        else:
            if is_pdf:
                content = extract_text_pdf(abs_path, limit_chars=4000)[:4000]
            elif ext == '.docx':
                content = extract_text_docx(abs_path, limit_chars=4000)
            elif ext == '.rtf':
                content = extract_text_rtf(abs_path, limit_chars=4000)
            elif ext == '.epub':
                content = extract_text_epub(abs_path, limit_chars=4000)
            elif ext == '.djvu':
                content = extract_text_djvu(abs_path, limit_chars=4000)
            elif is_text:
                content = abs_path.read_text(encoding='utf-8', errors='ignore')[:4000]
            # для аудио позднее предпочтём данные из БД
            try:
                cache_file.write_text(content, encoding='utf-8')
            except Exception:
                pass
        if is_docx and mammoth is not None:
            html_cache_dir = Path(app.static_folder) / 'cache' / 'docx_html'
            html_cache_dir.mkdir(parents=True, exist_ok=True)
            html_cache_file = html_cache_dir / (cache_key + '.html')
            if html_cache_file.exists():
                try:
                    docx_html = html_cache_file.read_text(encoding='utf-8', errors='ignore')
                except Exception:
                    docx_html = None
            if docx_html is None:
                generated_html = _docx_to_html(abs_path)
                if generated_html:
                    docx_html = generated_html
                    try:
                        html_cache_file.write_text(docx_html, encoding='utf-8')
                    except Exception:
                        pass
    except Exception:
        content = ''

    # Пробуем найти запись File, чтобы получить id для ссылки «Подробнее» и ключевые слова
    # Находим запись в БД по rel_path; учитываем различия слэшей/обратных слэшей (Windows)
    rp = str(rel_path)
    rp_alt = rp.replace('/', '\\') if ('/' in rp) else rp.replace('\\', '/')
    f = File.query.filter(or_(File.rel_path == rp, File.rel_path == rp_alt)).first()
    file_id = f.id if f else None
    keywords_str = (f.keywords or '') if f else ''
    if f and (is_audio or (f.material_type or '') == 'audio'):
        is_audio = True
        abstract = (f.abstract or '')
        # Если нет кэша, берём фрагмент транскрипта из БД
        if not content:
            content = (f.text_excerpt or '')[:4000]
        # аудиоплеер ссылается на endpoint скачивания
        # используем rel_path с прямыми слэшами для URL
        audio_url = url_for('media_file', rel_path=str(rel_path).replace('\\','/'))
        try:
            duration = audio_duration_hhmmss(abs_path)
        except Exception:
            duration = None
    if f and (is_image or (f.material_type or '') == 'image'):
        is_image = True
        abstract = (f.abstract or '')
        image_url = url_for('media_file', rel_path=str(rel_path).replace('\\','/'))

    # Сгенерировать миниатюру для PDF (с кэшем)
    if is_pdf:
        try:
            thumb_path = Path(app.static_folder) / 'thumbnails' / (Path(rel_path).stem + '.png')
            if not thumb_path.exists():
                try:
                    import fitz
                    doc = fitz.open(str(abs_path))
                    pix = doc[0].get_pixmap(matrix=fitz.Matrix(1, 1))
                    thumb_path.parent.mkdir(parents=True, exist_ok=True)
                    pix.save(str(thumb_path))
                except Exception as e:
                    app.logger.warning(f"Thumbnail generation failed: {e}")
            if thumb_path.exists():
                thumbnail_url = url_for('static', filename=f'thumbnails/{thumb_path.name}')
        except Exception:
            thumbnail_url = None

    rel_url = str(rel_path).replace('\\','/')
    embedded = str(request.args.get('embedded', '')).strip().lower() in ('1', 'true', 'yes', 'on')
    return _render_preview(
        rel_url,
        is_pdf=is_pdf,
        is_text=is_text,
        is_audio=is_audio,
        is_image=is_image,
        is_docx=is_docx,
        content=content,
        thumbnail_url=thumbnail_url,
        abstract=abstract,
        audio_url=audio_url,
        duration=duration,
        image_url=image_url,
        keywords=keywords_str,
        embedded=embedded,
        docx_html=docx_html,
    )

FILENAME_PATTERNS = [
    # "Author - Title (2021).pdf"
    re.compile(r"^(?P<author>.+?)\s*-\s*(?P<title>.+?)\s*\((?P<year>\d{4})\)$"),
    # "Author_Title_2020.pdf" или "Author Title 2020.pdf"
    re.compile(r"^(?P<author>.+?)[_ ]+(?P<title>.+?)[_ ]+(?P<year>\d{4})$"),
]

def sha1_of_file(fp: Path, chunk=1<<20):
    h = hashlib.sha1()
    with fp.open("rb") as f:
        while True:
            b = f.read(chunk)
            if not b:
                break
            h.update(b)
    return h.hexdigest()

def _count_matches(text: str, pattern: str) -> int:
    try:
        return len(re.findall(pattern, text))
    except Exception:
        return 0


def _analyze_extracted_text(text: str) -> dict[str, Any]:
    normalized = re.sub(r"\s+", " ", str(text or "")).strip()
    if not normalized:
        return {"has_text": False, "suspicious_mojibake": False}
    letters = _count_matches(normalized, r"[A-Za-zА-Яа-яЁё]")
    cyrillic = _count_matches(normalized, r"[А-Яа-яЁё]")
    mojibake_markers = _count_matches(normalized, r"[ÐÑÃâ\uFFFD]")
    mojibake_pairs = _count_matches(normalized, r"(?:Ð.|Ñ.|Ã.|â.)")
    marker_ratio = (mojibake_markers / len(normalized)) if normalized else 0.0
    cyrillic_ratio = (cyrillic / letters) if letters else 0.0
    suspicious = (
        mojibake_pairs >= 6
        or (marker_ratio > 0.03 and cyrillic_ratio < 0.05 and letters >= 40)
        or (mojibake_markers >= 20 and cyrillic == 0)
    )
    return {"has_text": True, "suspicious_mojibake": suspicious}

def extract_text_pdf(fp: Path, limit_chars=40000, force_ocr_first_page: bool = False):
    """Извлечение текста из PDF.
    Fallback: если текста крайне мало — OCR до первых 5 страниц (если установлен pytesseract).
    Также логируем включение/отсутствие OCR и затраченное время в прогресс-лог (если он активен).
    """
    try:
        import time as _time
        import tempfile
        max_ocr_pages = int(os.getenv('PDF_OCR_PAGES', '5'))
        ocr_langs = os.getenv('OCR_LANGS', 'rus+eng')
        used_ocr_pages = 0
        ocr_time_total = 0.0

        def _ocr_page(page_obj) -> str:
            nonlocal used_ocr_pages, ocr_time_total
            t0 = _time.time()
            pix = page_obj.get_pixmap(matrix=fitz.Matrix(2, 2))
            with tempfile.NamedTemporaryFile(suffix=".png", delete=True) as tf:
                pix.save(tf.name)
                ocr_text = pytesseract.image_to_string(tf.name, lang=ocr_langs)
            ocr_time_total += (_time.time() - t0)
            if (ocr_text or '').strip():
                used_ocr_pages += 1
            return ocr_text or ""

        with fitz.open(fp) as doc:
            text_parts = []
            for idx, page in enumerate(doc):
                raw = page.get_text("text") or ""
                if idx == 0 and pytesseract is not None:
                    first_analysis = _analyze_extracted_text(raw)
                    should_ocr_first = (
                        force_ocr_first_page
                        or not first_analysis.get("has_text")
                        or bool(first_analysis.get("suspicious_mojibake"))
                    )
                    if should_ocr_first:
                        try:
                            ocr = _ocr_page(page)
                            if ocr.strip():
                                ocr_analysis = _analyze_extracted_text(ocr)
                                if first_analysis.get("suspicious_mojibake") and not ocr_analysis.get("suspicious_mojibake"):
                                    raw = ocr
                                elif not raw.strip():
                                    raw = ocr
                                elif ocr.strip() not in raw:
                                    raw = f"{ocr}\n{raw}"
                        except Exception as oe:
                            app.logger.info(f"OCR(first page) failed for {fp}: {oe}")
                if idx < max_ocr_pages and pytesseract is not None and len((raw or '').strip()) < 30:
                    try:
                        ocr = _ocr_page(page)
                        if ocr.strip():
                            raw = ocr
                    except Exception as oe:
                        app.logger.info(f"OCR failed for page {idx} {fp}: {oe}")
                text_parts.append(raw)
                if limit_chars > 0 and sum(len(x) for x in text_parts) >= limit_chars:
                    break

            text = "\n".join(text_parts)
            analysis = _analyze_extracted_text(text)
            should_run_bulk_ocr = (
                pytesseract is not None
                and (
                    len(text.strip()) < 200
                    or bool(analysis.get("suspicious_mojibake"))
                )
            )
            if should_run_bulk_ocr:
                try:
                    text_ocr = []
                    pages_to_ocr = min(len(doc), max_ocr_pages)
                    for idx in range(pages_to_ocr):
                        ocr = _ocr_page(doc[idx])
                        if ocr.strip():
                            text_ocr.append(ocr)
                    if text_ocr:
                        ocr_joined = "\n".join(text_ocr).strip()
                        ocr_analysis = _analyze_extracted_text(ocr_joined)
                        base = (text or "").strip()
                        if analysis.get("suspicious_mojibake") and not ocr_analysis.get("suspicious_mojibake"):
                            text = ocr_joined
                            try:
                                _scan_log("OCR: заменил искажённый текст извлечённым OCR")
                            except Exception:
                                pass
                        elif not base:
                            text = ocr_joined
                        elif ocr_joined not in base:
                            text = f"{ocr_joined}\n\n{base}"
                            try:
                                _scan_log("OCR: объединён с базовым текстом PDF")
                            except Exception:
                                pass
                except Exception as oe:
                    app.logger.info(f"OCR fallback failed {fp}: {oe}")

        try:
            if used_ocr_pages > 0:
                _scan_log(f"OCR: использовано страниц {used_ocr_pages}, время {int(ocr_time_total*1000)} мс")
            elif pytesseract is None:
                _scan_log("OCR недоступен (pytesseract не установлен)")
        except Exception:
            pass
        return _limit_text_length(text, limit_chars)
    except Exception as e:
        app.logger.warning(f"PDF extract failed for {fp}: {e}")
        return ""

def extract_text_docx(fp: Path, limit_chars=40000):
    if not docx:
        return ""
    try:
        d = docx.Document(str(fp))
        text = "\n".join([p.text for p in d.paragraphs])
        return _limit_text_length(text, limit_chars)
    except Exception as e:
        app.logger.warning(f"DOCX extract failed for {fp}: {e}")
        return ""

def extract_text_rtf(fp: Path, limit_chars=40000):
    if not rtf_to_text:
        return ""
    try:
        text = rtf_to_text(fp.read_text(encoding="utf-8", errors="ignore"))
        return _limit_text_length(text, limit_chars)
    except Exception as e:
        app.logger.warning(f"RTF extract failed for {fp}: {e}")
        return ""

def extract_text_epub(fp: Path, limit_chars=40000):
    if not epub:
        return ""
    try:
        book = epub.read_epub(str(fp))
        text = ""
        for item in book.get_items():
            if item.get_type() == epub.ITEM_DOCUMENT:
                text += item.get_content().decode(errors="ignore")
                if limit_chars > 0 and len(text) >= limit_chars:
                    break
        return _limit_text_length(text, limit_chars)
    except Exception as e:
        app.logger.warning(f"EPUB extract failed for {fp}: {e}")
        return ""

def extract_text_djvu(fp: Path, limit_chars=40000):
    if not djvu:
        return ""
    try:
        with djvu.decode.open(str(fp)) as d:
            text = ""
            for page in d.pages:
                text += page.get_text()
                if limit_chars > 0 and len(text) >= limit_chars:
                    break
        return _limit_text_length(text, limit_chars)
    except Exception as e:
        app.logger.warning(f"DjVu extract failed for {fp}: {e}")
        return ""

def _ffmpeg_available():
    return shutil.which('ffmpeg') is not None

def _convert_to_wav_pcm16(src: Path, dst: Path, rate=16000):
    if not _ffmpeg_available():
        raise RuntimeError("ffmpeg not found for audio conversion")
    subprocess.run(['ffmpeg','-y','-i',str(src),'-ac','1','-ar',str(rate),'-f','wav',str(dst)], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)

def _ffprobe_duration_seconds(src: Path) -> float:
    if shutil.which('ffprobe') is None:
        return 0.0
    try:
        out = subprocess.check_output(['ffprobe','-v','error','-show_entries','format=duration','-of','default=nw=1:nk=1',str(src)], stderr=subprocess.DEVNULL)
        return float(out.strip())
    except Exception:
        return 0.0

def audio_duration_hhmmss(src: Path) -> str:
    """Return duration string HH:MM:SS for common audio; best-effort using wave/ffprobe."""
    try:
        if src.suffix.lower() == '.wav':
            with wave.open(str(src), 'rb') as wf:
                frames = wf.getnframes(); rate = wf.getframerate() or 1
                secs = frames / float(rate)
        else:
            secs = _ffprobe_duration_seconds(src)
    except Exception:
        secs = 0.0
    secs = int(round(secs))
    h = secs // 3600; m = (secs % 3600) // 60; s = secs % 60
    return (f"{h:02d}:{m:02d}:{s:02d}") if h else (f"{m:02d}:{s:02d}")

def transcribe_audio(fp: Path, limit_chars=40000,
                     backend_override: str | None = None,
                     model_path_override: str | None = None,
                     lang_override: str | None = None,
                     vad_override: bool | None = None) -> str:
    # Позволяем диагностике работать даже при выключенном глобальном флаге, когда заданы overrides
    if not TRANSCRIBE_ENABLED and backend_override is None:
        return ""
    backend = (backend_override or TRANSCRIBE_BACKEND or '').lower()
    model_path = (model_path_override if model_path_override is not None else TRANSCRIBE_MODEL_PATH) or ''
    lang = (lang_override if (lang_override is not None) else TRANSCRIBE_LANGUAGE) or 'ru'
    try:
        if backend == 'faster-whisper' and FasterWhisperModel:
            # Определить путь к модели (каталог / алиас / repo id)
            resolved = ''
            if model_path:
                if Path(model_path).expanduser().exists():
                    resolved = str(Path(model_path).expanduser())
                else:
                    resolved = _ensure_faster_whisper_model(model_path)
            # Запасной вариант: используем small, если путь не определился
            if not resolved:
                resolved = _ensure_faster_whisper_model(os.getenv('FASTER_WHISPER_DEFAULT_MODEL', 'small'))
            if not resolved:
                app.logger.warning("Не удалось определить путь к модели faster-whisper — укажите TRANSCRIBE_MODEL_PATH или установите huggingface_hub")
                return ""
            model = FasterWhisperModel(resolved, device="cpu", compute_type="int8")

            def _fw_try(vad=True, lng=lang):
                try:
                    segs, info = model.transcribe(str(fp), language=lng, vad_filter=vad)
                    txt = " ".join(((s.text or '').strip()) for s in segs)
                    return (txt or '').strip()
                except Exception as _e:
                    app.logger.info(f"Попытка распознавания faster-whisper не удалась (vad={vad}, lang={lng}): {_e}")
                    return ''

            # Последовательность попыток (или принудительно заданный режим VAD)
            if vad_override is True or vad_override is False:
                text = _fw_try(vad=bool(vad_override), lng=lang)
                if not text:
                    text = _fw_try(vad=bool(vad_override), lng=None)
            else:
                text = _fw_try(vad=True, lng=lang)
                if not text:
                    text = _fw_try(vad=False, lng=lang)
                if not text:
                    text = _fw_try(vad=False, lng=None)
            return _limit_text_length(text or '', limit_chars)
    except Exception as e:
        app.logger.warning(f"Транскрибация не удалась для {fp}: {e}")
    return ""

def call_lmstudio_summarize(text: str, filename: str) -> str:
    text = (text or "")[: int(os.getenv("SUMMARY_TEXT_LIMIT", "12000"))]
    text = text[:_lm_max_input_chars()]
    if not text:
        return ""
    system = PROMPTS.get('summarize_audio_system') or DEFAULT_PROMPTS.get('summarize_audio_system', '')
    user = f"Файл: {filename}\nСтенограмма:\n{text}"
    last_error: Exception | None = None
    for choice in _llm_iter_choices('summary'):
        label = _llm_choice_label(choice)
        provider = _llm_choice_provider(choice)
        messages = [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ]
        try:
            provider_name = provider  # сохраняем имя провайдера для логов
            _provider, response = _llm_send_chat(
                choice,
                messages,
                temperature=0.2,
                max_tokens=min(400, _lm_max_output_tokens()),
                timeout=120,
                cache_bucket='summary',
            )
            if _llm_response_indicates_busy(response):
                app.logger.info(f"LLM summary endpoint занята ({label}), переключаемся")
                last_error = RuntimeError('busy')
                continue
            response.raise_for_status()
            data = response.json()
            content = _llm_extract_content(_provider, data)
            if content:
                return content
            try:
                preview = (response.text or '')[:200]
                app.logger.warning(
                    "LLM summary empty content (%s, provider=%s): %s",
                    label,
                    provider_name,
                    preview,
                )
            except Exception:
                pass
        except ValueError as ve:
            last_error = ve
            app.logger.warning(f"Суммаризация не удалась ({label}): {ve}")
            continue
        except Exception as e:
            last_error = e
            if isinstance(e, RuntimeError) and str(e) == 'llm_cache_only_mode':
                app.logger.info(f"Суммаризация пропущена (режим cache-only, {label})")
            else:
                app.logger.warning(f"Суммаризация не удалась ({label}, {provider_name}): {e}")
            continue
    if last_error and str(last_error) not in {'busy', 'llm_cache_only_mode'}:
        app.logger.warning(f"Суммаризация не удалась: {last_error}")
    return ""

def call_lmstudio_compose(system: str, user: str, *, temperature: float = 0.2, max_tokens: int = 400) -> str:
    last_error: Exception | None = None
    for choice in _llm_iter_choices('compose'):
        label = _llm_choice_label(choice)
        provider = _llm_choice_provider(choice)
        messages = [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ]
        try:
            provider_name = provider
            _provider, response = _llm_send_chat(
                choice,
                messages,
                temperature=float(temperature),
                max_tokens=min(int(max_tokens), _lm_max_output_tokens()),
                timeout=120,
                cache_bucket='compose',
            )
            if _llm_response_indicates_busy(response):
                app.logger.info(f"LLM compose endpoint занята ({label}), переключаемся")
                last_error = RuntimeError('busy')
                continue
            response.raise_for_status()
            data = response.json()
            content = _llm_extract_content(_provider, data)
            if content:
                return content
        except ValueError as ve:
            last_error = ve
            app.logger.warning(f"Compose LLM endpoint некорректен ({label}): {ve}")
            continue
        except Exception as e:
            last_error = e
            if isinstance(e, RuntimeError) and str(e) == 'llm_cache_only_mode':
                app.logger.info(f"Compose-запрос пропущен (режим cache-only, {label})")
            else:
                app.logger.warning(f"LM compose не удался ({label}, {provider_name}): {e}")
    if last_error and str(last_error) not in {'busy', 'llm_cache_only_mode'}:
        app.logger.warning(f"LLM compose failed: {last_error}")
    return ""


def call_doc_chat_llm(messages: list[dict[str, Any]], *, temperature: float = 0.2, max_tokens: int = 600) -> str:
    last_error: Exception | None = None
    for choice in _llm_iter_choices('doc_chat'):
        label = _llm_choice_label(choice)
        provider = _llm_choice_provider(choice)
        try:
            _provider, response = _llm_send_chat(
                choice,
                messages,
                temperature=float(temperature),
                max_tokens=min(int(max_tokens), _lm_max_output_tokens()),
                timeout=180,
                cache_bucket='doc_chat',
            )
            if _llm_response_indicates_busy(response):
                app.logger.info(f"LLM doc_chat endpoint занята ({label}), переключаемся")
                last_error = RuntimeError('busy')
                continue
            response.raise_for_status()
            data = response.json()
            content = _llm_extract_content(_provider, data)
            if content:
                return content
        except ValueError as ve:
            last_error = ve
            app.logger.warning(f"Doc chat LLM endpoint некорректен ({label}): {ve}")
            continue
        except Exception as exc:
            last_error = exc
            if isinstance(exc, RuntimeError) and str(exc) == 'llm_cache_only_mode':
                app.logger.info(f"Doc chat пропущен (режим cache-only, {label})")
            else:
                app.logger.warning(f"Doc chat LLM не удался ({label}, {provider}): {exc}")
            continue
    if last_error and str(last_error) not in {'busy', 'llm_cache_only_mode'}:
        app.logger.warning(f"Doc chat LLM окончательно не ответил: {last_error}")
    raise RuntimeError("LLM не вернул ответ для документа.")
def call_lmstudio_keywords(text: str, filename: str):
    """Извлечь короткий список ключевых слов из текста или поискового запроса через LM Studio."""
    text = (text or "").strip()
    if not text:
        return []
    text = text[:int(os.getenv("KWS_TEXT_LIMIT", "8000"))]
    text = text[:_lm_max_input_chars()]
    is_query = str(filename or '').strip().lower() == 'ai-search'
    if is_query:
        system = PROMPTS.get('ai_search_keywords_system') or DEFAULT_PROMPTS.get('ai_search_keywords_system', '')
        user = (
            "Поисковый запрос пользователя:\n"
            f"{text}\n\n"
            "Верни JSON-массив вида [\"тег1\", \"тег2\", ...] без пояснений."
        )
    else:
        system = PROMPTS.get('keywords_system') or DEFAULT_PROMPTS.get('keywords_system', '')
        user = f"Файл: {filename}\nСтенограмма:\n{text}"
    last_error: Exception | None = None
    for choice in _llm_iter_choices('keywords'):
        label = _llm_choice_label(choice)
        provider = _llm_choice_provider(choice)
        messages = [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ]
        try:
            provider_name = provider
            _provider, response = _llm_send_chat(
                choice,
                messages,
                temperature=0.0,
                max_tokens=min(200, _lm_max_output_tokens()),
                timeout=90,
                cache_bucket='keywords_query' if is_query else 'keywords_transcript',
            )
            if _llm_response_indicates_busy(response):
                app.logger.info(f"LLM keywords endpoint занята ({label}), переключаемся")
                last_error = RuntimeError('busy')
                continue
            response.raise_for_status()
            data = response.json()
            content = _llm_extract_content(_provider, data)
            try:
                obj = json.loads(content)
                if isinstance(obj, list):
                    return [str(x) for x in obj][:12]
            except Exception:
                pass
            m = re.search(r"```json\s*(\[.*?\])\s*```", content, flags=re.S)
            if m:
                try:
                    obj = json.loads(m.group(1))
                    if isinstance(obj, list):
                        return [str(x) for x in obj][:12]
                except Exception:
                    pass
            rough = re.split(r"[\n;,]", content)
            res = [w.strip(" \t\r\n-•") for w in rough if w.strip()]
            if res:
                return res[:12]
        except ValueError as ve:
            last_error = ve
            app.logger.warning(f"LLM keywords endpoint некорректен ({label}): {ve}")
            continue
        except Exception as e:
            last_error = e
            if isinstance(e, RuntimeError) and str(e) == 'llm_cache_only_mode':
                app.logger.info(f"Извлечение ключевых слов пропущено (cache-only, {label})")
            else:
                app.logger.warning(f"Извлечение ключевых слов (LLM) не удалось ({label}, {provider_name}): {e}")
    if last_error and str(last_error) not in {'busy', 'llm_cache_only_mode'}:
        app.logger.warning(f"Извлечение ключевых слов (LLM) не удалось: {last_error}")
    return []

def call_lmstudio_vision(image_path: Path, filename: str):
    """Распознавание и описание изображения через LM Studio (совместимый с OpenAI Vision)."""
    try:
        import base64
    except Exception as e:
        app.logger.warning(f"Визуальное распознавание недоступно: {e}")
        return {}

    mime = "image/png"
    suf = image_path.suffix.lower()
    if suf in (".jpg", ".jpeg"):
        mime = "image/jpeg"
    elif suf in (".webp",):
        mime = "image/webp"
    elif suf in (".bmp",):
        mime = "image/bmp"
    elif suf in (".tif", ".tiff"):
        mime = "image/tiff"
    try:
        raw = image_path.read_bytes()
    except Exception as e:
        app.logger.warning(f"Визуальное распознавание: не удалось прочитать файл {image_path}: {e}")
        return {}
    data_url = f"data:{mime};base64," + base64.b64encode(raw).decode('ascii')

    system = PROMPTS.get('vision_system') or DEFAULT_PROMPTS.get('vision_system', '')
    user_content = [
        {"type": "text", "text": f"Файл: {filename}. Опиши и укажи ключевые слова."},
        {"type": "image_url", "image_url": {"url": data_url}},
    ]

    def _extract_kws_from_text(txt: str):
        try:
            t = (txt or '').replace('*', ' ')
            m = re.search(r"(?i)ключ[^\n\r:]*?:\s*(.+)", t)
            if m:
                line = m.group(1).strip()
                line = re.split(r"\n|\r|\u2028|\u2029|\s{2,}", line)[0]
                parts = [p.strip(" \t\r\n,.;·•-") for p in line.split(',')]
                parts = [p for p in parts if p]
                cleaned = re.sub(r"(?i)ключ[^\n\r:]*?:.*", "", t)
                return parts[:16], cleaned.strip()
        except Exception:
            return [], txt
        return [], txt

    last_error: Exception | None = None
    for choice in _llm_iter_choices('vision'):
        label = _llm_choice_label(choice)
        provider = _llm_choice_provider(choice)
        if provider not in ('openai', 'openrouter', 'azure_openai'):
            app.logger.info(f"LLM vision endpoint пропускается ({label}): провайдер {provider} не поддерживается")
            continue
        messages = [
            {"role": "system", "content": system},
            {"role": "user", "content": user_content},
        ]
        try:
            _provider, response = _llm_send_chat(
                choice,
                messages,
                temperature=0.2,
                max_tokens=500,
                timeout=180,
                cache_bucket='vision',
            )
            if _llm_response_indicates_busy(response):
                app.logger.info(f"LLM vision endpoint занята ({label}), переключаемся")
                last_error = RuntimeError('busy')
                continue
            response.raise_for_status()
            data = response.json()
            content = _llm_extract_content(_provider, data)
            try:
                obj = json.loads(content)
                if isinstance(obj, dict):
                    kws_list = obj.get("keywords") if isinstance(obj.get("keywords"), list) else []
                    if not kws_list:
                        kws_list, cleaned = _extract_kws_from_text(obj.get("description") or "")
                        if kws_list:
                            obj["description"] = cleaned
                            obj["keywords"] = kws_list
                    if "keywords" in obj and isinstance(obj["keywords"], list):
                        obj["keywords"] = [str(x) for x in obj["keywords"]][:16]
                    return obj
            except Exception:
                pass
            m = re.search(r"```json\s*(\{.*?\})\s*```", content, flags=re.S)
            if m:
                try:
                    obj = json.loads(m.group(1))
                    if isinstance(obj, dict):
                        kws_list = obj.get("keywords") if isinstance(obj.get("keywords"), list) else []
                        if not kws_list:
                            kws_list, cleaned = _extract_kws_from_text(obj.get("description") or "")
                            if kws_list:
                                obj["description"] = cleaned
                                obj["keywords"] = kws_list
                        if "keywords" in obj and isinstance(obj["keywords"], list):
                            obj["keywords"] = [str(x) for x in obj["keywords"]][:16]
                        return obj
                except Exception:
                    pass
            kws_guess, cleaned = _extract_kws_from_text(content or '')
            return {"description": cleaned.strip()[:2000], "keywords": kws_guess}
        except ValueError as ve:
            last_error = ve
            app.logger.warning(f"LLM vision endpoint некорректен ({label}): {ve}")
            continue
        except Exception as e:
            last_error = e
            if isinstance(e, RuntimeError) and str(e) == 'llm_cache_only_mode':
                app.logger.info(f"Визуальное распознавание пропущено (cache-only, {label})")
            else:
                app.logger.warning(f"Визуальное распознавание не удалось ({label}): {e}")
    if last_error and str(last_error) not in {'busy', 'llm_cache_only_mode'}:
        app.logger.warning(f"Визуальное распознавание не удалось: {last_error}")
    return {}

def call_lmstudio_for_metadata(text: str, filename: str):
    """
    Вызов OpenAI-совместимого API (LM Studio) для извлечения метаданных.
    Возвращает dict. Терпимо относится к не-JSON ответам: пытается вытащить из ```json ...``` блока.
    """
    text = (text or "")[: int(os.getenv("LLM_TEXT_LIMIT", "15000"))]
    text = text[:_lm_max_input_chars()]

    system = PROMPTS.get('metadata_system') or DEFAULT_PROMPTS.get('metadata_system', '')
    type_list, type_hints = _material_type_prompt_context()
    if "{{TYPE_LIST}}" in system:
        system = system.replace("{{TYPE_LIST}}", type_list or "document")
    elif type_list:
        system = f"{system}\nДоступные типы материалов: {type_list}."
    if "{{TYPE_HINTS}}" in system:
        system = system.replace("{{TYPE_HINTS}}", type_hints or "")
    elif type_hints:
        system = f"{system}\nТипы и подсказки:\n{type_hints}"
    system = system.strip()

    last_error: Exception | None = None
    for choice in _llm_iter_choices('metadata'):
        label = _llm_choice_label(choice)
        provider = _llm_choice_provider(choice)
        provider_name = provider
        current_text = text

        max_retries = 3
        backoff = 1.0
        busy = False
        for attempt in range(1, max_retries + 1):
            messages = [
                {"role": "system", "content": system},
                {"role": "user", "content": f"Файл: {filename}\nФрагмент текста:\n{current_text}"},
            ]
            try:
                _provider, response = _llm_send_chat(
                    choice,
                    messages,
                    temperature=0.0,
                    max_tokens=min(800, _lm_max_output_tokens()),
                    timeout=120,
                    cache_bucket='metadata',
                )
                if _llm_response_indicates_busy(response):
                    app.logger.info(f"LLM metadata endpoint занята ({label}), переключаемся")
                    busy = True
                    last_error = RuntimeError('busy')
                    break
                text_snippet = (response.text or '')[:2000]
                if response.status_code != 200:
                    app.logger.warning(f"LLM HTTP {response.status_code} ({label}, попытка {attempt}): {text_snippet}")
                    response.raise_for_status()
                data = response.json()
                content = _llm_extract_content(_provider, data)
                try:
                    return json.loads(content)
                except Exception:
                    pass
                m = re.search(r"```json\s*(\{.*?\})\s*```", content, flags=re.S)
                if m:
                    try:
                        return json.loads(m.group(1))
                    except Exception:
                        app.logger.warning("Не удалось разобрать JSON внутри блока ```json из ответа LLM")
                m = re.search(r"(\{.*\})", content, flags=re.S)
                if m:
                    try:
                        return json.loads(m.group(1))
                    except Exception:
                        app.logger.warning("Не удалось разобрать JSON‑фрагмент из ответа LLM")
                app.logger.warning(f"LLM вернул не‑JSON контент ({label}, первые 300 символов): {content[:300]}")
                return {}
            except RuntimeError as e:
                last_error = e
                if str(e) == 'llm_cache_only_mode':
                    app.logger.info(f"LLM metadata пропущено (cache-only, {label})")
                else:
                    app.logger.warning(f"LLM metadata unexpected runtime error ({label}, {provider_name}): {e}")
                break
            except requests.exceptions.RequestException as e:
                last_error = e
                app.logger.warning(f"Исключение при запросе к LLM ({label}, {provider_name}, попытка {attempt}): {e}")
                if isinstance(e, requests.HTTPError):
                    resp = getattr(e, 'response', None)
                    detail = ''
                    if resp is not None:
                        try:
                            detail = resp.text or ''
                        except Exception:
                            detail = ''
                    if detail and 'number of tokens to keep' in detail.lower() and len(current_text) > 800:
                        current_text = current_text[: max(len(current_text) // 2, 800)]
                        app.logger.info(f"LLM metadata: сокращаем контекст до {len(current_text)} символов и повторяем")
                        continue
                if attempt < max_retries:
                    import time
                    time.sleep(backoff)
                    backoff *= 2
                    continue
            except ValueError as e:
                app.logger.warning(f"LLM metadata endpoint некорректен ({label}, {provider_name}, попытка {attempt}): {e}")
                last_error = e
                break
            except Exception as e:
                app.logger.warning(f"Unexpected error calling LLM ({label}, {provider_name}, попытка {attempt}): {e}")
                last_error = e
                break
        if busy:
            continue
    if last_error and str(last_error) not in {'busy', 'llm_cache_only_mode'}:
        app.logger.warning(f"Метаданные через LLM не получены: {last_error}")
    return {}

def upsert_tag(file_obj: File, key: str, value: str):
    key = (key or "").strip()
    value = (value or "").strip()
    if not key or not value:
        return
    t = Tag.query.filter_by(file_id=file_obj.id, key=key, value=value).first()
    if not t:
        t = Tag(file_id=file_obj.id, key=key, value=value)
        db.session.add(t)
        _invalidate_facets_cache('tag upsert')

def _upsert_keyword_tags(file_obj: File):
    """Разложить строку ключевых слов на отдельные теги 'ключевое слово'."""
    try:
        # Сначала удалим существующие теги ключевых слов, чтобы не плодить дубли
        try:
            Tag.query.filter_by(file_id=file_obj.id).filter(
                or_(Tag.key == 'ключевое слово', Tag.key == 'keywords')
            ).delete(synchronize_session=False)
            _invalidate_facets_cache('tag cleanup')
        except Exception:
            pass
        raw = (file_obj.keywords or '')
        if not raw:
            return
        parts = re.split(r"[\n,;]+", str(raw))
        seen = set()
        for kw in parts:
            w = (kw or '').strip()
            if not w:
                continue
            wl = w.lower()
            if wl in seen:
                continue
            seen.add(wl)
            upsert_tag(file_obj, 'ключевое слово', w)
        try:
            search_service.sync_file(file_obj)
        except Exception:
            pass
    except Exception:
        pass

def normalize_material_type(s: str) -> str:
    s = (s or '').strip().lower()
    mapping = {
        'диссертация': 'dissertation',
        'автореферат': 'dissertation_abstract',
        'автореферат диссертации': 'dissertation_abstract',
        'статья': 'article', 'article': 'article', 'paper': 'article',
        'журнал': 'journal', 'journal': 'journal', 'magazine': 'journal', 'вестник': 'journal',
        'учебник': 'textbook', 'пособие': 'textbook', 'учебное пособие': 'textbook',
        'монография': 'monograph',
        'отчет': 'report', 'отчёт': 'report', 'report': 'report',
        'патент': 'patent', 'patent': 'patent',
        'презентация': 'presentation', 'presentation': 'presentation',
        'тезисы': 'proceedings', 'proceedings': 'proceedings', 'труды': 'proceedings',
        'стандарт': 'standard', 'gost': 'standard', 'gost r': 'standard', 'standard': 'standard',
        'заметки': 'note', 'note': 'note',
        'document': 'document', 'документ': 'document',
        'image': 'image', 'изображение':'image', 'картинка':'image'
    }
    # пытаемся найти прямое соответствие
    if s in mapping:
        return mapping[s]
    # частичные совпадения
    for k, v in mapping.items():
        if k in s:
            return v
    return s or 'document'


def _allowed_material_type_keys() -> set[str]:
    keys = {'document', 'audio', 'image'}
    try:
        for entry in _material_type_profiles():
            key = str((entry or {}).get('key') or '').strip()
            if key:
                keys.add(key)
    except Exception:
        pass
    return keys


def _merge_llm_material_type(
    current_type: str | None,
    llm_type: str | None,
    *,
    allow_override: bool,
) -> str | None:
    current = normalize_material_type(current_type or '')
    candidate = normalize_material_type(llm_type or '')
    allowed = _allowed_material_type_keys()
    if current and current not in allowed:
        current = 'document'
    if candidate and candidate not in allowed:
        candidate = ''
    if not candidate:
        return current or None
    if current and not allow_override:
        return current
    if current and current != 'document' and candidate == 'document':
        return current
    return candidate


def _looks_like_author_line(line: str) -> bool:
    line = (line or '').strip()
    if not line or len(line.split()) > 16:
        return False
    if re.search(r"\d{2,}|стр\.|с\.\s*\d", line, flags=re.I):
        return False
    # Классические форматы: Иванов И.И.; Иванов Иван Иванович; Ivanov I.I.; John Smith
    patterns = [
        r"[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.[А-ЯЁ]\.",
        r"[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+",
        r"[A-Z][a-z]+\s+[A-Z]\.[A-Z]\.",
        r"[A-Z][a-z]+\s+[A-Z][a-z]+",
    ]
    return any(re.search(pat, line) for pat in patterns)


MIN_JOURNAL_TOC_ENTRIES = 5
MIN_JOURNAL_PAGES = 20


def _extract_journal_toc_entries(text: str, limit: int = 30) -> list[dict[str, str]]:
    """Извлекает записи оглавления журнала: авторы, название, страница."""
    if not text:
        return []
    lower = text.lower()
    markers = ["оглавлен", "содержани", "contents", "table of contents"]
    positions = [lower.find(m) for m in markers if lower.find(m) != -1]
    if not positions:
        return []
    start = min(positions)
    slice_end = min(len(text), start + 15000)
    lines = text[start:slice_end].splitlines()
    entries: list[dict[str, str]] = []
    pending_authors: str | None = None
    blank_streak = 0
    for raw in lines[1:]:
        if len(entries) >= limit:
            break
        line = raw.strip()
        if not line:
            blank_streak += 1
            if blank_streak >= 3 and entries:
                break
            continue
        blank_streak = 0
        cleaned = re.sub(r"^[\dIVXLCM\s\.\-–·•]+", "", line, flags=re.IGNORECASE)
        cleaned = cleaned.replace("…", " ")
        cleaned = re.sub(r"[·•]{2,}", " ", cleaned)
        cleaned = re.sub(r"\.{2,}", " ", cleaned)
        cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
        if len(cleaned) < 4:
            continue
        page_match = re.search(r"(\d{1,4})(?:\s*(?:стр\.|с\.)?)?$", cleaned, flags=re.IGNORECASE)
        if not page_match:
            if _looks_like_author_line(cleaned):
                pending_authors = cleaned
            continue
        page = page_match.group(1)
        body = cleaned[:page_match.start()].strip(" .•—-–\t")
        if not body:
            continue
        authors = None
        title = body
        # Попытки отделить авторов
        separators = [" — ", " - ", " – ", " —", ": ", "; "]
        for sep in separators:
            if sep in body:
                left, right = body.split(sep, 1)
                if _looks_like_author_line(left):
                    authors = left.strip()
                    title = right.strip()
                    break
        if authors is None:
            m = re.match(r"^((?:[A-ZА-ЯЁ][A-Za-zА-Яа-яё\-']+(?:\s+[A-ZА-ЯЁ]\.[A-ZА-ЯЁ]\.)?(?:,\s*)?)+)\s+(.+)$", body)
            if m and _looks_like_author_line(m.group(1)):
                authors = m.group(1).strip()
                title = m.group(2).strip()
            elif pending_authors and not _looks_like_author_line(body):
                authors = pending_authors.strip()
                title = body.strip()
                pending_authors = None
        entry = {
            'title': title.strip(),
            'page': page.strip(),
        }
        if authors:
            entry['authors'] = authors
        entries.append(entry)
        pending_authors = None
    return entries


def _material_type_profiles() -> List[Dict[str, Any]]:
    runtime = _rt()
    items = getattr(runtime, "material_types", None) or []
    cleaned: List[Dict[str, Any]] = []
    seen = set()
    for entry in items:
        if not isinstance(entry, dict):
            continue
        key = str(entry.get("key") or "").strip()
        if not key or key in seen:
            continue
        cleaned.append(entry)
        seen.add(key)
    if "document" not in seen:
        cleaned.append({
            "key": "document",
            "label": "Документ",
            "enabled": True,
            "priority": -100,
            "threshold": 0.0,
        })
    return cleaned


def _material_type_rank_key(entry: Dict[str, Any]) -> Tuple[float, int, int, int, int]:
    score = float(entry.get("score") or 0.0)
    priority = int(entry.get("priority") or 0)
    special_hit = 1 if entry.get("special_match") else 0
    signal = int(entry.get("raw_text_count") or 0) + int(entry.get("raw_filename_count") or 0)
    index = int(entry.get("index") or 0)
    return (score, priority, special_hit, signal, -index)


def _evaluate_material_type_profile(
    profile: Dict[str, Any],
    ext_clean: str,
    text_lower: str,
    filename_lower: str,
    journal_toc: List[dict[str, str]],
    idx: int,
) -> Optional[Dict[str, Any]]:
    key = str(profile.get("key") or "").strip()
    if not key:
        return None
    if not profile.get("enabled", True):
        return None

    flow = [str(v or "").strip().lower() for v in profile.get("flow", []) if str(v or "").strip()]
    allow_extension = (not flow) or any(token in {"extension", "any"} for token in flow)
    allow_filename = (not flow) or any(token in {"filename", "any"} for token in flow)
    allow_heuristics = (not flow) or any(token in {"heuristics", "text", "any"} for token in flow)

    priority = int(profile.get("priority") or 0)
    threshold = float(profile.get("threshold", 1.0))
    extension_weight = float(profile.get("extension_weight", 2.0))
    filename_weight = float(profile.get("filename_weight", 1.5))
    text_weight = float(profile.get("text_weight", 1.0))

    require_extension = bool(profile.get("require_extension", False))
    require_filename = bool(profile.get("require_filename", False))
    require_text = bool(profile.get("require_text", False))

    extensions = [
        str(item or "").strip().lower().lstrip(".")
        for item in profile.get("extensions", [])
        if str(item or "").strip()
    ]
    extension_match = bool(ext_clean and extensions and ext_clean in extensions)
    if require_extension and not extension_match:
        return None

    filename_keywords = [
        str(item or "").strip().lower()
        for item in profile.get("filename_keywords", [])
        if str(item or "").strip()
    ]
    filename_count = sum(1 for kw in filename_keywords if kw and kw in filename_lower)
    if require_filename and filename_keywords and filename_count == 0:
        return None

    text_keywords = [
        str(item or "").strip().lower()
        for item in profile.get("text_keywords", [])
        if str(item or "").strip()
    ]
    text_count = sum(1 for kw in text_keywords if kw and kw in text_lower)
    if require_text and text_keywords and text_count == 0:
        return None

    exclude_keywords = [
        str(item or "").strip().lower()
        for item in profile.get("exclude_keywords", [])
        if str(item or "").strip()
    ]
    if exclude_keywords and any(ex in text_lower or ex in filename_lower for ex in exclude_keywords):
        return None

    special = profile.get("special") or {}
    special_match = False
    special_weight = 0.0
    if special.get("journal_toc_required"):
        min_entries = int(special.get("min_toc_entries") or MIN_JOURNAL_TOC_ENTRIES)
        if len(journal_toc) >= max(1, min_entries):
            special_match = True
            special_weight = float(special.get("weight") or 2.0)
        else:
            return None

    score = 0.0
    if allow_extension and extension_match:
        score += extension_weight
    if allow_filename and filename_count:
        score += filename_weight * filename_count
    if allow_heuristics and text_count:
        score += text_weight * text_count
    if special_match:
        score += special_weight

    effective_threshold = max(0.0, threshold if threshold is not None else 0.0)
    if score < effective_threshold and effective_threshold > 0 and key != "document":
        return None

    extension_stage = allow_extension and extension_match
    filename_stage = allow_filename and filename_count > 0
    heuristics_stage = allow_heuristics and (text_count > 0 or filename_count > 0 or special_match)

    if not (extension_stage or filename_stage or heuristics_stage):
        if key == "document":
            heuristics_stage = True
        else:
            return None

    return {
        "key": key,
        "score": float(score),
        "priority": priority,
        "extension": bool(extension_stage),
        "filename": bool(filename_stage),
        "heuristics": bool(heuristics_stage),
        "special_match": bool(special_match),
        "raw_extension_match": bool(extension_match),
        "raw_filename_count": int(filename_count),
        "raw_text_count": int(text_count),
        "index": idx,
    }


def _collect_material_type_matches(ext: str, text_excerpt: str, filename: str) -> List[Dict[str, Any]]:
    profiles = _material_type_profiles()
    if not profiles:
        return []
    ext_clean = (ext or "").lower().lstrip(".")
    text_lower = (text_excerpt or "").lower()
    filename_lower = (filename or "").lower()
    needs_toc = any((profile.get("special") or {}).get("journal_toc_required") for profile in profiles)
    journal_toc = _extract_journal_toc_entries(text_excerpt or "") if needs_toc else []
    matches: List[Dict[str, Any]] = []
    for idx, profile in enumerate(profiles):
        evaluation = _evaluate_material_type_profile(
            profile,
            ext_clean,
            text_lower,
            filename_lower,
            journal_toc,
            idx,
        )
        if evaluation:
            matches.append(evaluation)
    return matches


def _material_type_prompt_context() -> Tuple[str, str]:
    profiles = _material_type_profiles()
    visible = [entry for entry in profiles if entry.get("enabled", True)]
    type_keys: List[str] = []
    hints: List[str] = []
    for entry in visible:
        key = str(entry.get("key") or "").strip()
        if not key:
            continue
        type_keys.append(key)
        label = str(entry.get("label") or "").strip()
        hint = str(entry.get("llm_hint") or entry.get("description") or label or key).strip()
        if hint:
            hints.append(f"{key}: {hint}")
    type_list = ", ".join(type_keys)
    type_hints = "\n".join(f"- {line}" for line in hints if line)
    return type_list, type_hints


def _journal_safety_check(material_type: str | None, text: str | None, page_count: int | None, entries: list[dict[str, str]] | None = None) -> str | None:
    mt = (material_type or '').strip().lower()
    if mt != 'journal':
        return material_type
    entries = entries if entries is not None else (_extract_journal_toc_entries(text or '') if text else [])
    if len(entries) < MIN_JOURNAL_TOC_ENTRIES:
        return 'article'
    if page_count is not None and page_count < MIN_JOURNAL_PAGES:
        return 'article'
    return material_type

def guess_material_type(ext: str, text_excerpt: str, filename: str = "") -> str:
    matches = _collect_material_type_matches(ext, text_excerpt, filename)
    heuristics_matches = [
        entry for entry in matches if entry.get("heuristics") and entry.get("key") != "document"
    ]
    if heuristics_matches:
        best = max(heuristics_matches, key=_material_type_rank_key)
        return best["key"]
    doc_entry = next((entry for entry in matches if entry.get("key") == "document"), None)
    if doc_entry:
        return doc_entry["key"]
    return "document"


def _detect_type_pre_llm(ext: str, text_excerpt: str, filename: str) -> str | None:
    flow = [p.strip() for p in (TYPE_DETECT_FLOW or '').split(',') if p.strip()]
    if not flow:
        flow = ['extension', 'filename', 'heuristics']
    matches = _collect_material_type_matches(ext, text_excerpt, filename)
    if not matches:
        return None
    doc_candidate: Optional[Dict[str, Any]] = None
    for stage in flow:
        stage_key = str(stage or "").strip().lower()
        if not stage_key:
            continue
        if stage_key == 'llm':
            continue
        if stage_key not in {'extension', 'filename', 'heuristics', 'text'}:
            stage_key = 'heuristics'
        actual_key = 'heuristics' if stage_key in {'heuristics', 'text'} else stage_key
        candidates = [entry for entry in matches if entry.get(actual_key)]
        if not candidates:
            continue
        non_doc = [entry for entry in candidates if entry.get("key") != "document"]
        target = non_doc or candidates
        best = max(target, key=_material_type_rank_key)
        if best["key"] == "document":
            if doc_candidate is None:
                doc_candidate = best
            continue
        return best["key"]
    if doc_candidate:
        return doc_candidate["key"]
    return None

def looks_like_dissertation_filename(fn: str) -> bool:
    fl = (fn or '').lower()
    if not fl:
        return False
    if any(tok in fl for tok in ["автореферат", "autoreferat", "автoref"]):
        return True
    if any(tok in fl for tok in ["диссер", "dissert", "thesis"]):
        return True
    return False

def extract_tags_for_type(material_type: str, text: str, filename: str = "") -> dict:
    """Извлечение тегов по типу материала простыми регулярками.
    Возвращает dict: {key: value}.
    """
    t = (text or "")
    tl = t.lower()
    tags = {}
    # Общие сведения
    # Приблизительная оценка языка: доля кириллицы против латиницы
    try:
        cyr = sum(1 for ch in t if ('а' <= ch.lower() <= 'я') or (ch in 'ёЁ'))
        lat = sum(1 for ch in t if 'a' <= ch.lower() <= 'z')
        if cyr + lat > 20:
            tags.setdefault('lang', 'ru' if cyr >= lat else 'en')
    except Exception:
        pass
    # Распространённые идентификаторы
    # Общие
    # DOI (цифровой идентификатор объекта)
    m = re.search(r"\b(10\.\d{4,9}\/[\w\-\.:;()\/[\]A-Za-z0-9]+)", t)
    if m:
        tags.setdefault("doi", m.group(1))
    # ISBN (международный книжный номер)
    m = re.search(r"\bISBN[:\s]*([0-9\- ]{10,20})", t, flags=re.I)
    if m:
        tags.setdefault("isbn", m.group(1).strip())

    if material_type in ("dissertation", "dissertation_abstract"):
        # Научный руководитель (варианты написания)
        m = re.search(r"научн[ыийо]{1,3}\s*[-–:]?\s*руководител[ья][\s:–]+(.{3,80})", tl, flags=re.I)
        if m:
            tags.setdefault("научный руководитель", m.group(1).strip().title())
        # Специальность/шифр ВАК вида 05.13.11 или 01.02.03
        m = re.search(r"(?:шифр|специальн)[^\n\r]{0,30}?(\d{2}\.\d{2}\.\d{2})", tl)
        if not m:
            m = re.search(r"\b(\d{2}\.\d{2}\.\d{2})\b", tl)
        if m:
            tags.setdefault("специальность", m.group(1))
        # Организация/вуз (ФГБОУ ВО, университет, институт, академия, НИУ, МГУ, СПбГУ)
        m = re.search(r"(фгбоу\s*во|университет|институт|академия|ниу|мгу|спбгу)[^\n\r]{0,100}", tl)
        if m:
            tags.setdefault("организация", m.group(0).strip().title())
        # Кафедра
        m = re.search(r"кафедра\s+([^\n\r]{3,80})", tl)
        if m:
            tags.setdefault("кафедра", m.group(1).strip().title())
        # Степень
        if re.search(r"кандидат[а]?\b", tl):
            tags.setdefault("степень", "кандидат")
        elif re.search(r"доктор[а]?\b", tl):
            tags.setdefault("степень", "доктор")
        # Новизна (короткий фрагмент)
        m = re.search(r"научн[ао]я\s+новизн[аы][^\n\r:]*[:\-]\s*(.{20,300})", tl)
        if m:
            tags.setdefault("научная новизна", m.group(1).strip().capitalize())
        # Цель исследования
        m = re.search(r"цель\s+исследован[ияи][^:\n\r]*[:\-]\s*(.{10,200})", tl)
        if m:
            tags.setdefault("цель", m.group(1).strip().capitalize())
        # Объект/предмет исследования
        m = re.search(r"объект\s+исследован[ияи][^:\n\r]*[:\-]\s*(.{10,200})", tl)
        if m:
            tags.setdefault("объект", m.group(1).strip().capitalize())
        m = re.search(r"предмет\s+исследован[ияи][^:\n\r]*[:\-]\s*(.{10,200})", tl)
        if m:
            tags.setdefault("предмет", m.group(1).strip().capitalize())
        # Задачи исследования
        m = re.search(r"задач[аи]\s+исследован[ияи][^:\n\r]*[:\-]\s*(.{10,200})", tl)
        if m:
            tags.setdefault("задачи", m.group(1).strip().capitalize())
        # Положения, выносимые на защиту
        m = re.search(r"вынос[иы]м[ае][^\n\r]{0,40}на\s+защит[уы][^:\n\r]*[:\-]\s*(.{20,300})", tl)
        if m:
            tags.setdefault("на защиту", m.group(1).strip().capitalize())
    elif material_type == "article":
        # Журнал/Вестник
        m = re.search(r"(журнал|вестник|труды|transactions|journal)[:\s\-]+([^\n\r]{3,80})", tl, flags=re.I)
        if m:
            tags.setdefault("журнал", m.group(2).strip().title())
        # Номер/том/страницы
        m = re.search(r"том\s*(\d+)\b.*?№\s*(\d+)\b.*?с\.?\s*(\d+)[\-–](\d+)", tl)
        if m:
            tags.setdefault("том/номер", f"{m.group(1)}/{m.group(2)}")
            tags.setdefault("страницы", f"{m.group(3)}–{m.group(4)}")
        else:
            m = re.search(r"№\s*(\d+)\b.*?с\.?\s*(\d+)[\-–](\d+)", tl)
            if m:
                tags.setdefault("номер", m.group(1))
                tags.setdefault("страницы", f"{m.group(2)}–{m.group(3)}")
    elif material_type == "textbook":
        # Дисциплина
        m = re.search(r"по\s+дисциплин[еы]\s*[:\-]?\s*([^\n\r]{3,80})", tl)
        if m:
            tags.setdefault("дисциплина", m.group(1).strip().title())
        # Издательство
        m = re.search(r"издательств[оа]\s*[:\-]?\s*([^\n\r]{3,80})", tl)
        if m:
            tags.setdefault("издательство", m.group(1).strip().title())
    elif material_type == "monograph":
        # Издательство/город/год часто в шапке
        m = re.search(r"(издательство|изд.)\s*[:\-]?\s*([^\n\r]{3,80})", tl)
        if m:
            tags.setdefault("издательство", m.group(2).strip().title())

    return tags

def prune_missing_files(allowed_collection_ids: set[int] | None = None):
    """Удаляет из БД файлы, которых нет на диске."""
    removed = 0
    query = File.query
    if allowed_collection_ids is not None:
        if not allowed_collection_ids:
            query = query.filter(File.collection_id == -1)
        else:
            query = query.filter(File.collection_id.in_(allowed_collection_ids))
    for f in query.all():
        try:
            if not Path(f.path).exists():
                fid = f.id
                db.session.delete(f)
                removed += 1
                try:
                    _delete_file_from_fts(fid)
                except Exception as exc:
                    logger.debug('fts delete failed for missing file %s: %s', fid, exc)
        except Exception:
            fid = getattr(f, 'id', None)
            db.session.delete(f)
            removed += 1
            try:
                _delete_file_from_fts(fid)
            except Exception as exc:
                logger.debug('fts delete failed for missing file %s: %s', fid, exc)
    db.session.commit()
    return removed

# ------------------- Routes -------------------

@app.route("/")
def index():
    return redirect(url_for('app_react'))

def _serve_spa() -> Response:
    """Отдаёт собранный React UI либо dev-index, если билд отсутствует."""
    try:
        dist = BASE_DIR / 'frontend' / 'dist'
        idx = dist / 'index.html'
        if idx.exists():
            return send_file(str(idx))
    except Exception:
        pass
    dev_idx = BASE_DIR / 'frontend' / 'index.html'
    if dev_idx.exists():
        return send_file(str(dev_idx))
    body = (
        "<div style=\"min-height:100vh;display:flex;align-items:center;justify-content:center;"
        "font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;background:#0d1117;color:#c9d1d9;"
        "padding:24px;text-align:center;\">"
        "<div style=\"max-width:520px;background:rgba(22,27,34,0.85);padding:32px;border-radius:18px;"
        "border:1px solid #30363d;box-shadow:0 20px 48px rgba(0,0,0,0.45);\">"
        "<h1 style=\"margin-top:0;font-size:22px;\">React-интерфейс не собран</h1>"
        "<p style=\"line-height:1.6;\">В каталоге <code>frontend/dist</code> не найден файл <code>index.html</code>."
        " Соберите клиент командой <code>npm install && npm run build</code> внутри директории <code>frontend</code> "
        "или запустите dev-сервер (<code>npm run dev</code>) и проксируйте его вручную.</p>"
        "</div></div>"
    )
    return _html_page("Agregator — UI не найден", body, status=404)


@app.route("/app")
def app_react():
    """Точка входа для React SPA."""
    return _serve_spa()


@app.route("/app/<path:_path>")
def app_react_catchall(_path: str):
    """Внутренний роутинг React: любые пути внутри /app/*."""
    return _serve_spa()


def _collect_login_backgrounds() -> list[dict[str, str]]:
    """Enumerate available HTML backgrounds for the login screen."""
    entries: list[dict[str, str]] = []
    try:
        directory = LOGIN_BACKGROUNDS_DIR
        if not directory.exists():
            return entries
        for item in sorted(directory.glob('*.html')):
            if not item.is_file():
                continue
            label = item.stem.replace('_', ' ').replace('-', ' ').strip().title()
            try:
                text = item.read_text(encoding='utf-8', errors='ignore')
            except Exception:
                text = ''
            match = re.search(r'<title[^>]*>(.*?)</title>', text, flags=re.IGNORECASE | re.DOTALL)
            if match:
                raw = match.group(1).strip()
                if raw:
                    label = re.sub(r'\s+', ' ', raw)
            entries.append({
                'name': item.name,
                'label': label,
                'url': f"/login-backgrounds/{item.name}",
            })
    except Exception as exc:
        try:
            current_app.logger.warning('login backgrounds scan failed: %s', exc)
        except Exception:
            pass
    return entries


@app.route('/api/login-backgrounds')
def api_login_backgrounds():
    """Expose available login background animations."""
    return jsonify(_collect_login_backgrounds())


@app.route('/login-backgrounds/<path:filename>')
def login_background_file(filename: str):
    """Serve individual login background HTML files."""
    if '/' in filename or '\\' in filename:
        abort(404)
    if not filename.lower().endswith('.html'):
        abort(404)
    directory = LOGIN_BACKGROUNDS_DIR
    if not directory.exists():
        abort(404)
    try:
        response = send_from_directory(str(directory), filename)
    except FileNotFoundError:
        abort(404)
    try:
        # Разрешаем встраивание в iframe на этом же домене.
        response.headers['X-Frame-Options'] = 'SAMEORIGIN'
        csp = response.headers.get('Content-Security-Policy')
        if csp:
            if 'frame-ancestors' not in csp:
                response.headers['Content-Security-Policy'] = csp.rstrip(';') + "; frame-ancestors 'self'"
        else:
            response.headers['Content-Security-Policy'] = "frame-ancestors 'self'"
    except Exception:
        pass
    response.headers['Cache-Control'] = 'no-store'
    return response

@app.route('/assets/<path:filename>')
def vite_assets(filename):
    """Serve Vite built assets if present (dist/assets)."""
    dist_assets = BASE_DIR / 'frontend' / 'dist' / 'assets'
    if dist_assets.exists():
        return send_from_directory(str(dist_assets), filename)
    return ("Not found", 404)

@app.route("/aiword")
def aiword_index():
    _ensure_aiword_access()
    """Serve AiWord app (built with Vite) under /aiword with asset path rewrite and pre‑bootstrap.
    Injects a small script to prefill localStorage with BibTeX from this app's DB and LM Studio config.
    """
    dist = BASE_DIR / 'AiWord' / 'dist'
    idx = dist / 'index.html'
    if not idx.exists():
        return ("AiWord/dist/index.html not found. Please build AiWord or include dist.", 404)
    try:
        html = idx.read_text(encoding='utf-8', errors='ignore')
    except Exception:
        return send_file(str(idx))

    # Переписываем абсолютные ссылки на ресурсы в /aiword/assets
    html = html.replace('src="/assets/', 'src="/aiword/assets/')
    html = html.replace('href="/assets/', 'href="/aiword/assets/')
    # Переписываем vite.svg, если он упоминается
    html = html.replace('href="/vite.svg', 'href="/aiword/vite.svg')

    # Готовим bootstrap-скрипт для предварительной загрузки BibTeX и настроек ЛМ
    lm_base = os.getenv('LMSTUDIO_API_BASE', 'http://localhost:1234/v1')
    lm_model = os.getenv('LMSTUDIO_MODEL', 'google/gemma-3n-e4b')
    bootstrap = f"""
    <script>
    (function(){{
      try {{
        // Синхронизируем настройки LM Studio с Agregator
        localStorage.setItem('llm-writer-base-url', {lm_base!r});
        localStorage.setItem('llm-writer-model', {lm_model!r});
      }} catch(e){{}}
      // Загружаем библиографию из базы Agregator
      fetch('/api/aiword/bibtex').then(r=>r.ok?r.text():Promise.resolve('')).then(txt=>{{
        if (!txt) return;
        try {{
          var saved = localStorage.getItem('llm-writer-autosave');
          var data = saved ? JSON.parse(saved) : null;
          if (!data || !data.project || !data.content) {{
            data = {{
              project: {{
                title: 'Новая статья',
                language: 'ru',
                style_guide: 'Пиши чётко и структурированно. Без воды. Сохраняй академический тон, «выводы» — коротко и по делу. Используй Markdown: заголовки, списки, таблицы. Формулы как $...$ или ```math``` при необходимости.',
                persona: 'Ты — научный редактор и соавтор. Помогаешь с планом, структурой, стилем, и аргументацией.'
              }},
              content: '# Заголовок\n\nВставьте/пишите текст здесь...',
              bibText: ''
            }};
          }}
          data.bibText = txt;
          localStorage.setItem('llm-writer-autosave', JSON.stringify(data));
        }} catch(e){{}}
      }}).catch(function(){{}});
    }})();
    </script>
    """
    # Вставляем bootstrap перед первым тегом module-скрипта, если он есть
    marker = '<script type="module"'
    if marker in html:
        html = html.replace(marker, bootstrap + "\n" + marker, 1)
    else:
        html = html.replace('</head>', bootstrap + '\n</head>')
    return html

@app.route('/aiword/assets/<path:filename>')
def aiword_assets(filename):
    _ensure_aiword_access()
    dist_assets = BASE_DIR / 'AiWord' / 'dist' / 'assets'
    if dist_assets.exists():
        return send_from_directory(str(dist_assets), filename)
    return ("Not found", 404)

@app.route('/aiword/vite.svg')
def aiword_vite_svg():
    _ensure_aiword_access()
    p = BASE_DIR / 'AiWord' / 'dist' / 'vite.svg'
    if p.exists():
        return send_file(str(p))
    return ("Not found", 404)


def _metadata_quality(file_obj: File | None) -> dict[str, Any]:
    if not file_obj:
        return {"score": 0.0, "bucket": "low", "filled": 0, "total": 8}
    checks = {
        "title": bool((file_obj.title or "").strip()),
        "author": bool((file_obj.author or "").strip()),
        "year": bool((file_obj.year or "").strip()),
        "material_type": bool((file_obj.material_type or "").strip()),
        "keywords": bool((file_obj.keywords or "").strip()),
        "abstract": bool((getattr(file_obj, "abstract", None) or "").strip()),
        "text_excerpt": bool((file_obj.text_excerpt or "").strip()),
        "tags": bool(getattr(file_obj, "tags", None)),
    }
    total = len(checks)
    filled = sum(1 for v in checks.values() if v)
    score = round((filled / total) if total else 0.0, 3)
    if score >= 0.75:
        bucket = "high"
    elif score >= 0.45:
        bucket = "medium"
    else:
        bucket = "low"
    return {"score": score, "bucket": bucket, "filled": filled, "total": total}


def _serialize_search_file(file_obj: File) -> dict[str, Any]:
    return {
        "id": file_obj.id,
        "title": file_obj.title,
        "author": file_obj.author,
        "year": file_obj.year,
        "material_type": file_obj.material_type,
        "path": file_obj.path,
        "rel_path": file_obj.rel_path,
        "text_excerpt": file_obj.text_excerpt,
        "abstract": getattr(file_obj, 'abstract', None),
        "tags": [{"key": t.key, "value": t.value} for t in file_obj.tags],
        "metadata_quality": _metadata_quality(file_obj),
    }


def _ranked_rows_with_feedback(
    rows: Sequence[File],
    ranked_pairs: Sequence[tuple[int, float]],
    *,
    feedback_scale: float = 0.25,
) -> List[File]:
    if not rows:
        return []
    if not ranked_pairs:
        return list(rows)
    rank_score_map = {int(fid): float(score) for fid, score in ranked_pairs}
    feedback = _get_feedback_weights()
    row_ids = [int(row.id) for row in rows if getattr(row, "id", None)]
    realtime_delta: Dict[int, float] = {}
    if row_ids:
        try:
            agg_rows = (
                db.session.query(
                    AiSearchKeywordFeedback.file_id,
                    AiSearchKeywordFeedback.action,
                    func.count(AiSearchKeywordFeedback.id),
                )
                .filter(AiSearchKeywordFeedback.file_id.in_(row_ids))
                .group_by(AiSearchKeywordFeedback.file_id, AiSearchKeywordFeedback.action)
                .all()
            )
            per_file: Dict[int, Dict[str, int]] = {}
            for fid, action, cnt in agg_rows:
                if fid is None:
                    continue
                bucket = per_file.setdefault(int(fid), {"relevant": 0, "irrelevant": 0, "click": 0})
                key = str(action or "").strip().lower()
                if key in bucket:
                    bucket[key] += int(cnt or 0)
            for fid, stats in per_file.items():
                pos = float(stats.get("relevant", 0) or 0) + (float(stats.get("click", 0) or 0) * FEEDBACK_CLICK_WEIGHT)
                neg = float(stats.get("irrelevant", 0) or 0)
                realtime_delta[fid] = max(-1.2, min(1.2, math.log((1.0 + pos) / (1.0 + neg))))
        except Exception:
            realtime_delta = {}
    feedback_scale = max(0.0, float(feedback_scale or 0.0))
    scored: List[tuple[float, float, File]] = []
    for row in rows:
        base = rank_score_map.get(int(row.id), 0.0)
        fb_weight = float((feedback.get(int(row.id), {}) or {}).get("weight", 0.0) or 0.0)
        fb_weight += float(realtime_delta.get(int(row.id), 0.0) or 0.0)
        meta_score = float(_metadata_quality(row).get("score", 0.0) or 0.0)
        final = base + (fb_weight * feedback_scale) + (meta_score * 0.05)
        scored.append((final, row.mtime or 0.0, row))
    scored.sort(key=lambda item: (item[0], item[1]), reverse=True)
    return [item[2] for item in scored]


def _filter_rows_by_metadata_quality(rows: Sequence[File], bucket: str) -> List[File]:
    wanted = (bucket or "").strip().lower()
    if wanted not in {"low", "medium", "high"}:
        return list(rows)
    out: List[File] = []
    for row in rows:
        if _metadata_quality(row).get("bucket") == wanted:
            out.append(row)
    return out


@app.route("/api/search")
def api_search():
    q = request.args.get("q", "").strip()
    smart = str(request.args.get('smart', '')).lower() in ('1','true','yes','on')
    rank_profile = (request.args.get('rank_profile') or 'balanced').strip().lower()
    material_type = request.args.get("type", "").strip()
    tag_filters = request.args.getlist("tag")

    # Дополнительные необязательные фильтры (обратная совместимость)
    year_from = (request.args.get("year_from") or "").strip()
    year_to = (request.args.get("year_to") or "").strip()
    size_min = (request.args.get("size_min") or "").strip()
    size_max = (request.args.get("size_max") or "").strip()
    metadata_quality_bucket = (request.args.get("metadata_quality") or "").strip().lower()
    collection_filter = _parse_collection_param(request.args.get('collection_id'))

    query = File.query.join(Collection, File.collection_id == Collection.id)
    query = _apply_file_access_filter(query)
    try:
        query = query.filter(Collection.searchable == True)
    except Exception:
        pass
    if collection_filter is not None:
        query = query.filter(File.collection_id == collection_filter)
    if material_type:
        query = query.filter(File.material_type == material_type)
    if q and not smart:
        query = _apply_text_search_filter(query, q)
    if year_from:
        query = query.filter(File.year >= year_from)
    if year_to:
        query = query.filter(File.year <= year_to)
    if size_min:
        try:
            query = query.filter(File.size >= int(size_min))
        except Exception:
            pass
    if size_max:
        try:
            query = query.filter(File.size <= int(size_max))
        except Exception:
            pass
    for tf in tag_filters:
        if "=" in tf:
            k, v = tf.split("=", 1)
            t = aliased(Tag)
            query = query.join(t, t.file_id == File.id).filter(and_(t.key == k, t.value.ilike(f"%{v}%")))
    query = query.distinct()

    ranked_ids: list[int] = []
    ranked_pairs: list[tuple[int, float]] = []
    if q and smart:
        ranked_pairs = _search_ranked_candidates(q, limit=10000, profile=rank_profile)
        ranked_ids = [fid for fid, _ in ranked_pairs]
        if ranked_ids:
            query = query.filter(File.id.in_(ranked_ids))
        else:
            candidates = _search_candidate_ids(q)
            if candidates:
                ranked_ids = list(candidates)
                query = query.filter(File.id.in_(candidates))
            elif candidates == []:
                return jsonify([])
    elif q:
        ranked_pairs = _search_ranked_candidates(q, limit=10000, profile=rank_profile)
        ranked_ids = [fid for fid, _ in ranked_pairs]
        if ranked_ids:
            query = query.filter(File.id.in_(ranked_ids))
        else:
            return jsonify([])

    rows = query.order_by(File.mtime.desc().nullslast()).limit(10000 if q else 200).all()
    if ranked_pairs:
        rows = _ranked_rows_with_feedback(rows, ranked_pairs)
    rows = rows[:200]
    if q and smart:
        qlem = list(_expand_synonyms(_lemmas(q)))
        def _match(f: File) -> bool:
            lab = set(_lemmas(_row_text_for_search(f)))
            return any(l in lab for l in qlem)
        rows = [r for r in rows if _match(r)]
    if metadata_quality_bucket in {"low", "medium", "high"}:
        rows = _filter_rows_by_metadata_quality(rows, metadata_quality_bucket)
    return jsonify([_serialize_search_file(r) for r in rows])

@app.route("/api/search_v2")
def api_search_v2():
    """Расширенный поиск с пагинацией и доп. фильтрами.
    Query: q, type, tag=key=val (multi), year_from, year_to, size_min, size_max, limit, offset
    Returns: { items: [...], total: n }
    """
    q = request.args.get("q", "").strip()
    smart = str(request.args.get('smart', '')).lower() in ('1','true','yes','on')
    rank_profile = (request.args.get('rank_profile') or 'balanced').strip().lower()
    material_type = request.args.get("type", "").strip()
    tag_filters = request.args.getlist("tag")
    year_from = (request.args.get("year_from") or "").strip()
    year_to = (request.args.get("year_to") or "").strip()
    size_min = (request.args.get("size_min") or "").strip()
    size_max = (request.args.get("size_max") or "").strip()
    metadata_quality_bucket = (request.args.get("metadata_quality") or "").strip().lower()
    try:
        limit = min(max(int(request.args.get("limit", 50)), 1), 200)
    except Exception:
        limit = 50
    try:
        offset = max(int(request.args.get("offset", 0)), 0)
    except Exception:
        offset = 0

    base = File.query.join(Collection, File.collection_id == Collection.id)
    base = _apply_file_access_filter(base)
    collection_filter = _parse_collection_param(request.args.get('collection_id'))
    if collection_filter is not None:
        base = base.filter(File.collection_id == collection_filter)
    try:
        base = base.filter(Collection.searchable == True)
    except Exception:
        pass
    if material_type:
        base = base.filter(File.material_type == material_type)
    if q and not smart:
        base = _apply_text_search_filter(base, q)
    if year_from:
        base = base.filter(File.year >= year_from)
    if year_to:
        base = base.filter(File.year <= year_to)
    if size_min:
        try:
            base = base.filter(File.size >= int(size_min))
        except Exception:
            pass
    if size_max:
        try:
            base = base.filter(File.size <= int(size_max))
        except Exception:
            pass
    qx = base
    for tf in tag_filters:
        if "=" in tf:
            k, v = tf.split("=", 1)
            if k == 'author':
                qx = qx.filter(File.author.ilike(f"%{v}%"))
            else:
                t = aliased(Tag)
                qx = qx.join(t, t.file_id == File.id).filter(and_(t.key == k, t.value.ilike(f"%{v}%")))
    qx = qx.distinct()

    if q:
        ranked_pairs = _search_ranked_candidates(q, limit=10000, profile=rank_profile)
        ranked_ids = [fid for fid, _ in ranked_pairs]
        if ranked_ids:
            qx = qx.filter(File.id.in_(ranked_ids))
        else:
            candidates = _search_candidate_ids(q)
            if candidates:
                ranked_ids = list(candidates)
                qx = qx.filter(File.id.in_(candidates))
            else:
                return jsonify({"items": [], "total": 0})
        cap = 10000
        cand = qx.order_by(File.mtime.desc().nullslast()).limit(cap).all()
        if ranked_pairs:
            cand = _ranked_rows_with_feedback(cand, ranked_pairs)
        if smart:
            # Морфологическая фильтрация на сервере с синонимами
            qlem = list(_expand_synonyms(_lemmas(q)))
            def _match(f: File) -> bool:
                lab = set(_lemmas(_row_text_for_search(f)))
                return any(l in lab for l in qlem)
            cand = [r for r in cand if _match(r)]
        if metadata_quality_bucket in {"low", "medium", "high"}:
            cand = _filter_rows_by_metadata_quality(cand, metadata_quality_bucket)
        total = len(cand)
        rows = cand[offset:offset+limit]
    else:
        if metadata_quality_bucket in {"low", "medium", "high"}:
            cand = qx.order_by(File.mtime.desc().nullslast()).limit(10000).all()
            cand = _filter_rows_by_metadata_quality(cand, metadata_quality_bucket)
            total = len(cand)
            rows = cand[offset:offset+limit]
        else:
            total = qx.count()
            rows = qx.order_by(File.mtime.desc().nullslast()).offset(offset).limit(limit).all()
    items = [_serialize_search_file(r) for r in rows]
    return jsonify({"items": items, "total": total})


@app.route("/api/metadata/low-quality")
def api_metadata_low_quality():
    try:
        limit = min(max(int(request.args.get("limit", 200)), 1), 1000)
    except Exception:
        limit = 200
    collection_filter = _parse_collection_param(request.args.get('collection_id'))
    query = File.query
    query = _apply_file_access_filter(query)
    if collection_filter is not None:
        query = query.filter(File.collection_id == collection_filter)
    rows = query.order_by(File.mtime.desc().nullslast()).limit(5000).all()
    low_rows = [row for row in rows if _metadata_quality(row).get("bucket") == "low"]
    payload = [_serialize_search_file(row) for row in low_rows[:limit]]
    return jsonify({"ok": True, "items": payload, "total": len(low_rows)})


@app.route("/api/metadata/normalize", methods=["POST"])
@require_admin
def api_metadata_normalize():
    data = request.get_json(silent=True) or {}
    only_low = bool(data.get("only_low", True))
    try:
        limit = min(max(int(data.get("limit", 2000)), 1), 20000)
    except Exception:
        limit = 2000
    collection_id = data.get("collection_id")
    payload = {
        "status": "queued",
        "only_low": only_low,
        "limit": limit,
        "collection_id": collection_id,
    }
    task = TaskRecord(
        name='metadata_normalize',
        status='queued',
        progress=0.0,
        payload=json.dumps(payload, ensure_ascii=False),
        user_id=getattr(_load_current_user(), 'id', None),
        collection_id=int(collection_id) if collection_id else None,
    )
    try:
        db.session.add(task)
        db.session.commit()
    except Exception as exc:
        db.session.rollback()
        return jsonify({"ok": False, "error": str(exc)}), 500
    options = {
        "only_low": only_low,
        "limit": limit,
        "collection_id": collection_id,
    }
    get_task_queue().submit(
        _run_metadata_normalization_job,
        int(task.id),
        options,
        description="metadata-normalize",
        owner_id=getattr(_load_current_user(), 'id', None),
    )
    return jsonify({"ok": True, "task_id": int(task.id)})


@app.route('/api/voice-search', methods=['POST'])
def api_voice_search():
    """Принимает короткий аудиофрагмент, распознаёт его и возвращает текст."""
    user = _load_current_user()
    if not user:
        return jsonify({'ok': False, 'error': 'Не авторизовано'}), 401
    if not TRANSCRIBE_ENABLED:
        return jsonify({'ok': False, 'error': 'Распознавание аудио отключено'}), 503
    if request.content_length and request.content_length > 15 * 1024 * 1024:
        return jsonify({'ok': False, 'error': 'Аудио слишком большое'}), 413

    audio = request.files.get('audio')
    if not audio or audio.filename is None:
        return jsonify({'ok': False, 'error': 'Не удалось получить аудио'}), 400

    filename = secure_filename(audio.filename) or 'voice-input.webm'
    suffix = Path(filename).suffix or '.webm'
    src_path: Path | None = None
    wav_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            audio.save(tmp)
            src_path = Path(tmp.name)

        transcript = ''
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as wav_tmp:
                wav_path = Path(wav_tmp.name)
            _convert_to_wav_pcm16(src_path, wav_path)
            transcript = transcribe_audio(wav_path, limit_chars=5000)
        except Exception as conv_exc:
            app.logger.info(f'[voice-search] wav convert failed, fallback to raw: {conv_exc}')
            transcript = transcribe_audio(src_path, limit_chars=5000)

        transcript = (transcript or '').strip()
        if not transcript:
            return jsonify({'ok': True, 'text': '', 'warning': 'Речь не распознана'})

        try:
            _log_user_action(user, 'voice_search', detail=json.dumps({'length': len(transcript)}))
        except Exception:
            pass

        return jsonify({'ok': True, 'text': transcript})
    except Exception as exc:
        app.logger.warning(f'[voice-search] failed: {exc}')
        return jsonify({'ok': False, 'error': 'Не удалось распознать аудио'}), 500
    finally:
        if src_path:
            try:
                src_path.unlink()
            except FileNotFoundError:
                pass
        if wav_path:
            try:
                wav_path.unlink()
            except FileNotFoundError:
                pass

@app.route('/api/settings', methods=['GET', 'POST'])
def api_settings():
    """JSON API для чтения/изменения основных настроек UI/сканирования."""
    _require_admin()
    runtime = _rt()
    provider_default = runtime.lm_default_provider or 'openai'
    runtime_snapshot = runtime_settings_store.snapshot()
    runtime_fields_payload = []
    for field_def in CONFIG_FIELDS:
        snapshot_key = field_snapshot_key(field_def)
        value = runtime_snapshot.get(snapshot_key)
        if value is None and field_def.visibility == "env_only" and field_def.env_key:
            value = os.getenv(field_def.env_key, app.config.get(field_def.env_key))
        runtime_fields_payload.append({
            "name": field_def.name,
            "api_key": field_api_key(field_def),
            "group": field_def.group,
            "type": field_def.type_hint,
            "description": field_def.doc,
            "env_key": field_def.env_key,
            "runtime_mutable": bool(field_def.runtime_mutable),
            "visibility": field_def.visibility,
            "restart_required": bool(field_def.restart_required),
            "constraints": field_def.constraints or {},
            "ui_component": field_def.ui_component,
            "depends_on": field_def.depends_on,
            "risk_level": field_def.risk_level,
            "aliases": list(field_def.aliases),
            "value": value,
        })
    if request.method == 'GET':
        _ensure_llm_schema_once()
        db_dialect = _db_dialect_name()
        db_uri_raw = _active_database_uri()
        db_uri_safe = db_uri_raw
        if db_uri_raw:
            try:
                db_uri_safe = make_url(db_uri_raw).render_as_string(hide_password=True)
            except Exception:
                db_uri_safe = db_uri_raw
        try:
            cols = Collection.query.order_by(Collection.name.asc()).all()
            counts = {}
            try:
                rows = db.session.query(File.collection_id, func.count(File.id)).group_by(File.collection_id).all()
                counts = {cid: int(cnt) for (cid, cnt) in rows}
            except Exception:
                counts = {}
            collections = [{
                'id': c.id,
                'name': c.name,
                'slug': c.slug,
                'searchable': bool(c.searchable),
                'graphable': bool(c.graphable),
                'is_private': bool(getattr(c, 'is_private', False)),
                'count': int(counts.get(c.id, 0)),
            } for c in cols]
        except Exception:
            collections = []
        try:
            llms = LlmEndpoint.query.order_by(LlmEndpoint.created_at.desc()).all()
            preferred = _llm_pick_default_endpoint()
            effective_base = preferred.base_url if preferred else runtime.lmstudio_api_base
            effective_model = preferred.model if preferred else runtime.lmstudio_model
            effective_key = (preferred.api_key if preferred else runtime.lmstudio_api_key) or ''
            effective_provider = (preferred.provider if preferred else provider_default) or provider_default
            llm_items = [{
                'id': ep.id,
                'name': ep.name,
                'base_url': ep.base_url,
                'model': ep.model,
                'weight': float(ep.weight or 0.0),
                'purpose': ep.purpose,
                'purposes': _llm_parse_purposes(ep.purpose),
                'provider': (ep.provider or provider_default),
                'context_length': int(ep.context_length) if getattr(ep, 'context_length', None) else None,
                'instances': int(getattr(ep, 'instances', 1) or 1),
            } for ep in llms]
        except Exception:
            effective_base = runtime.lmstudio_api_base
            effective_model = runtime.lmstudio_model
            effective_key = runtime.lmstudio_api_key
            effective_provider = provider_default
            llm_items = []
        try:
            ai_entries = AiWordAccess.query.join(User, AiWordAccess.user_id == User.id).all()
            ai_users = [{
                'user_id': entry.user_id,
                'username': entry.user.username if entry.user else None,
                'full_name': entry.user.full_name if entry.user else None,
            } for entry in ai_entries]
        except Exception:
            ai_users = []
        return jsonify({
            'scan_root': str(runtime.scan_root),
            'extract_text': bool(runtime.extract_text),
            'lm_base': effective_base,
            'lm_model': effective_model,
            'lm_key': effective_key,
            'lm_provider': effective_provider,
            'rag_embedding_backend': runtime.rag_embedding_backend,
            'rag_embedding_model': runtime.rag_embedding_model,
            'rag_embedding_dim': int(runtime.rag_embedding_dim),
            'rag_embedding_batch': int(runtime.rag_embedding_batch_size),
            'rag_embedding_device': runtime.rag_embedding_device,
            'rag_embedding_endpoint': runtime.rag_embedding_endpoint or runtime.lmstudio_api_base,
            'rag_embedding_api_key': runtime.rag_embedding_api_key or runtime.lmstudio_api_key,
            'transcribe_enabled': bool(runtime.transcribe_enabled),
            'transcribe_backend': runtime.transcribe_backend,
            'transcribe_model': runtime.transcribe_model_path,
            'transcribe_language': runtime.transcribe_language,
            'summarize_audio': bool(runtime.summarize_audio),
            'audio_keywords_llm': bool(runtime.audio_keywords_llm),
            'vision_images': bool(runtime.images_vision_enabled),
            'kw_to_tags': bool(runtime.keywords_to_tags_enabled),
            'doc_chat_chunk_max_tokens': int(runtime.doc_chat_chunk_max_tokens),
            'doc_chat_chunk_overlap': int(runtime.doc_chat_chunk_overlap),
            'doc_chat_chunk_min_tokens': int(runtime.doc_chat_chunk_min_tokens),
            'doc_chat_max_chunks': int(runtime.doc_chat_max_chunks),
            'doc_chat_fallback_chunks': int(runtime.doc_chat_fallback_chunks),
            'doc_chat_image_min_width': int(runtime.doc_chat_image_min_width),
            'doc_chat_image_min_height': int(runtime.doc_chat_image_min_height),
            'type_detect_flow': runtime.type_detect_flow,
            'type_llm_override': bool(runtime.type_llm_override),
            'import_subdir': runtime.import_subdir,
            'move_on_rename': bool(runtime.move_on_rename),
            'collections_in_dirs': bool(runtime.collections_in_separate_dirs),
            'collection_type_subdirs': bool(runtime.collection_type_subdirs),
            'type_dirs': runtime.type_dirs,
            'ocr_langs': runtime.ocr_langs_cfg,
            'pdf_ocr_pages': int(runtime.pdf_ocr_pages_cfg),
            'prompts': runtime.prompts,
            'prompt_defaults': dict(DEFAULT_PROMPTS),
            'ai_rerank_llm': bool(runtime.ai_rerank_llm),
            'ocr_first_page_dissertation': bool(runtime.always_ocr_first_page_dissertation),
            'llm_cache_enabled': bool(runtime.llm_cache_enabled),
            'llm_cache_ttl_seconds': int(runtime.llm_cache_ttl_seconds),
            'llm_cache_max_items': int(runtime.llm_cache_max_items),
            'llm_cache_only_mode': bool(runtime.llm_cache_only_mode),
            'search_cache_enabled': bool(runtime.search_cache_enabled),
            'search_cache_ttl_seconds': int(runtime.search_cache_ttl_seconds),
            'search_cache_max_items': int(runtime.search_cache_max_items),
            'lm_max_input_chars': int(runtime.lm_max_input_chars),
            'lm_max_output_tokens': int(runtime.lm_max_output_tokens),
            'azure_openai_api_version': runtime.azure_openai_api_version,
            'collections': collections,
            'llm_endpoints': llm_items,
            'llm_purposes': LLM_PURPOSES,
            'llm_providers': LLM_PROVIDER_OPTIONS,
            'aiword_users': ai_users,
            'default_use_llm': bool(runtime.default_use_llm),
            'default_prune': bool(runtime.default_prune),
            'material_types': runtime.material_types,
            'runtime_fields': runtime_fields_payload,
            'settings_hub_v2_enabled': str(os.getenv('SETTINGS_HUB_V2_ENABLED', '1')).strip().lower() in {'1', 'true', 'yes', 'on'},
            'database_dialect': db_dialect,
            'database_uri': db_uri_safe,
            'database_type_options': ['sqlite', 'postgresql'],
        })

    data = request.json or {}
    before = runtime_settings_store.snapshot()
    update_payload = {}
    field_errors = {}

    def _set(key: str, value):
        update_payload[key] = value

    def _validate_schema_field(field_def, raw_value, source_key: str):
        try:
            type_hint = (field_def.type_hint or "").lower()
            if "bool" in type_hint:
                if isinstance(raw_value, bool):
                    value = raw_value
                else:
                    value = str(raw_value).strip().lower() in {"1", "true", "yes", "on"}
            elif "int" in type_hint:
                value = int(raw_value)
            elif "float" in type_hint:
                value = float(raw_value)
            elif "list" in type_hint:
                if isinstance(raw_value, list):
                    value = raw_value
                elif isinstance(raw_value, str):
                    value = [part.strip() for part in raw_value.split(",") if part.strip()]
                else:
                    raise ValueError("Ожидается список")
            elif "dict" in type_hint:
                if not isinstance(raw_value, dict):
                    raise ValueError("Ожидается объект")
                value = raw_value
            else:
                value = raw_value

            constraints = field_def.constraints or {}
            if "enum" in constraints and value not in constraints["enum"]:
                raise ValueError(f"Допустимые значения: {', '.join(str(v) for v in constraints['enum'])}")
            if isinstance(value, (int, float)):
                if "min" in constraints and value < constraints["min"]:
                    raise ValueError(f"Минимум: {constraints['min']}")
                if "max" in constraints and value > constraints["max"]:
                    raise ValueError(f"Максимум: {constraints['max']}")
            _set(field_snapshot_key(field_def), value)
        except Exception as exc:
            field_errors[source_key] = str(exc)

    if 'scan_root' in data:
        value = data.get('scan_root') or str(runtime.scan_root)
        _set('SCAN_ROOT', str(value))
    if 'extract_text' in data:
        _set('EXTRACT_TEXT', data.get('extract_text'))
    if 'lm_base' in data:
        _set('LMSTUDIO_API_BASE', data.get('lm_base'))
    if 'lm_model' in data:
        _set('LMSTUDIO_MODEL', data.get('lm_model'))
    if 'lm_key' in data:
        _set('LMSTUDIO_API_KEY', data.get('lm_key'))
    if 'lm_provider' in data:
        candidate = (data.get('lm_provider') or '').strip().lower()
        if candidate not in LLM_PROVIDER_CHOICES:
            candidate = runtime.lm_default_provider if runtime.lm_default_provider in LLM_PROVIDER_CHOICES else 'openai'
        _set('LM_DEFAULT_PROVIDER', candidate)
    if 'rag_embedding_backend' in data:
        backend = (data.get('rag_embedding_backend') or '').strip().lower()
        _set('RAG_EMBEDDING_BACKEND', backend or runtime.rag_embedding_backend)
    if 'rag_embedding_model' in data:
        _set('RAG_EMBEDDING_MODEL', data.get('rag_embedding_model'))
    if 'rag_embedding_dim' in data:
        try:
            _set('RAG_EMBEDDING_DIM', int(data.get('rag_embedding_dim')))
        except Exception:
            pass
    if 'rag_embedding_batch' in data:
        try:
            _set('RAG_EMBEDDING_BATCH', int(data.get('rag_embedding_batch')))
        except Exception:
            pass
    if 'rag_embedding_device' in data:
        _set('RAG_EMBEDDING_DEVICE', (data.get('rag_embedding_device') or '').strip())
    if 'rag_embedding_endpoint' in data:
        _set('RAG_EMBEDDING_ENDPOINT', (data.get('rag_embedding_endpoint') or '').strip())
    if 'rag_embedding_api_key' in data:
        _set('RAG_EMBEDDING_API_KEY', data.get('rag_embedding_api_key') or '')
    if 'transcribe_enabled' in data:
        _set('TRANSCRIBE_ENABLED', data.get('transcribe_enabled'))
    if 'transcribe_backend' in data:
        _set('TRANSCRIBE_BACKEND', data.get('transcribe_backend'))
    if 'transcribe_model' in data:
        _set('TRANSCRIBE_MODEL_PATH', data.get('transcribe_model'))
    if 'transcribe_language' in data:
        _set('TRANSCRIBE_LANGUAGE', data.get('transcribe_language'))
    if 'summarize_audio' in data:
        _set('SUMMARIZE_AUDIO', data.get('summarize_audio'))
    if 'audio_keywords_llm' in data:
        _set('AUDIO_KEYWORDS_LLM', data.get('audio_keywords_llm'))
    if 'vision_images' in data:
        _set('IMAGES_VISION_ENABLED', data.get('vision_images'))
    if 'kw_to_tags' in data:
        _set('KEYWORDS_TO_TAGS_ENABLED', data.get('kw_to_tags'))
    if 'doc_chat_chunk_max_tokens' in data:
        try:
            value = max(16, int(data.get('doc_chat_chunk_max_tokens') or 0))
        except Exception:
            value = None
        if value is not None:
            _set('DOC_CHAT_CHUNK_MAX_TOKENS', value)
    if 'doc_chat_chunk_overlap' in data:
        try:
            value = max(0, int(data.get('doc_chat_chunk_overlap') or 0))
        except Exception:
            value = None
        if value is not None:
            _set('DOC_CHAT_CHUNK_OVERLAP', value)
    if 'doc_chat_chunk_min_tokens' in data:
        try:
            value = max(1, int(data.get('doc_chat_chunk_min_tokens') or 0))
        except Exception:
            value = None
        if value is not None:
            _set('DOC_CHAT_CHUNK_MIN_TOKENS', value)
    if 'doc_chat_max_chunks' in data:
        try:
            value = max(0, int(data.get('doc_chat_max_chunks') or 0))
        except Exception:
            value = None
        if value is not None:
            _set('DOC_CHAT_MAX_CHUNKS', value)
    if 'doc_chat_fallback_chunks' in data:
        try:
            value = max(0, int(data.get('doc_chat_fallback_chunks') or 0))
        except Exception:
            value = None
        if value is not None:
            _set('DOC_CHAT_FALLBACK_CHUNKS', value)
    if 'doc_chat_image_min_width' in data:
        try:
            value = max(0, int(data.get('doc_chat_image_min_width') or 0))
        except Exception:
            value = None
        if value is not None:
            _set('DOC_CHAT_IMAGE_MIN_WIDTH', value)
    if 'doc_chat_image_min_height' in data:
        try:
            value = max(0, int(data.get('doc_chat_image_min_height') or 0))
        except Exception:
            value = None
        if value is not None:
            _set('DOC_CHAT_IMAGE_MIN_HEIGHT', value)
    if 'type_detect_flow' in data:
        _set('TYPE_DETECT_FLOW', data.get('type_detect_flow'))
    if 'type_llm_override' in data:
        _set('TYPE_LLM_OVERRIDE', data.get('type_llm_override'))
    if 'import_subdir' in data:
        value = data.get('import_subdir') or runtime.import_subdir
        _set('IMPORT_SUBDIR', value)
    if 'move_on_rename' in data:
        _set('MOVE_ON_RENAME', data.get('move_on_rename'))
    if 'collections_in_dirs' in data:
        _set('COLLECTIONS_IN_SEPARATE_DIRS', data.get('collections_in_dirs'))
    if 'collection_type_subdirs' in data:
        _set('COLLECTION_TYPE_SUBDIRS', data.get('collection_type_subdirs'))
    if 'type_dirs' in data and isinstance(data.get('type_dirs'), dict):
        _set('TYPE_DIRS', data.get('type_dirs'))
    if 'material_types' in data and isinstance(data.get('material_types'), list):
        _set('MATERIAL_TYPES', data.get('material_types'))
    if 'default_use_llm' in data:
        _set('DEFAULT_USE_LLM', data.get('default_use_llm'))
    if 'default_prune' in data:
        _set('DEFAULT_PRUNE', data.get('default_prune'))
    if 'ocr_langs' in data:
        _set('OCR_LANGS_CFG', data.get('ocr_langs'))
    if 'pdf_ocr_pages' in data:
        _set('PDF_OCR_PAGES_CFG', data.get('pdf_ocr_pages'))
    if 'ocr_first_page_dissertation' in data:
        _set('ALWAYS_OCR_FIRST_PAGE_DISSERTATION', data.get('ocr_first_page_dissertation'))
    if 'prompts' in data and isinstance(data.get('prompts'), dict):
        _set('PROMPTS', data.get('prompts'))
    if 'ai_rerank_llm' in data:
        _set('AI_RERANK_LLM', data.get('ai_rerank_llm'))
    if 'ai_query_variants_max' in data:
        _set('AI_QUERY_VARIANTS_MAX', data.get('ai_query_variants_max'))
    if 'ai_rag_retry_enabled' in data:
        _set('AI_RAG_RETRY_ENABLED', data.get('ai_rag_retry_enabled'))
    if 'ai_rag_retry_threshold' in data:
        _set('AI_RAG_RETRY_THRESHOLD', data.get('ai_rag_retry_threshold'))
    if 'llm_cache_enabled' in data:
        _set('LLM_CACHE_ENABLED', data.get('llm_cache_enabled'))
    if 'llm_cache_ttl_seconds' in data:
        _set('LLM_CACHE_TTL_SECONDS', data.get('llm_cache_ttl_seconds'))
    if 'llm_cache_max_items' in data:
        _set('LLM_CACHE_MAX_ITEMS', data.get('llm_cache_max_items'))
    if 'llm_cache_only_mode' in data:
        _set('LLM_CACHE_ONLY_MODE', data.get('llm_cache_only_mode'))
    if 'search_cache_enabled' in data:
        _set('SEARCH_CACHE_ENABLED', data.get('search_cache_enabled'))
    if 'search_cache_ttl_seconds' in data:
        _set('SEARCH_CACHE_TTL_SECONDS', data.get('search_cache_ttl_seconds'))
    if 'search_cache_max_items' in data:
        _set('SEARCH_CACHE_MAX_ITEMS', data.get('search_cache_max_items'))
    if 'lm_max_input_chars' in data:
        _set('LM_MAX_INPUT_CHARS', data.get('lm_max_input_chars'))
    if 'lm_max_output_tokens' in data:
        _set('LM_MAX_OUTPUT_TOKENS', data.get('lm_max_output_tokens'))
    if 'azure_openai_api_version' in data:
        _set('AZURE_OPENAI_API_VERSION', data.get('azure_openai_api_version'))
    if 'rag_rerank_backend' in data:
        _set('RAG_RERANK_BACKEND', data.get('rag_rerank_backend'))
    if 'rag_rerank_model' in data:
        _set('RAG_RERANK_MODEL', data.get('rag_rerank_model'))
    if 'rag_rerank_device' in data:
        _set('RAG_RERANK_DEVICE', data.get('rag_rerank_device'))
    if 'rag_rerank_batch_size' in data:
        _set('RAG_RERANK_BATCH_SIZE', data.get('rag_rerank_batch_size'))
    if 'rag_rerank_max_length' in data:
        _set('RAG_RERANK_MAX_LENGTH', data.get('rag_rerank_max_length'))
    if 'rag_rerank_max_chars' in data:
        _set('RAG_RERANK_MAX_CHARS', data.get('rag_rerank_max_chars'))
    runtime_fields_data = data.get('runtime_fields')
    if isinstance(runtime_fields_data, dict):
        for field_def in CONFIG_FIELDS:
            field_name = field_def.name
            if field_name not in runtime_fields_data:
                continue
            if not field_def.runtime_mutable:
                field_errors[field_name] = "Поле доступно только через env и требует перезапуска"
                continue
            _validate_schema_field(field_def, runtime_fields_data.get(field_name), field_name)

    for field_def in CONFIG_FIELDS:
        api_key = field_api_key(field_def)
        if api_key not in data:
            continue
        if not field_def.runtime_mutable:
            field_errors[api_key] = "Поле доступно только через env и требует перезапуска"
            continue
        _validate_schema_field(field_def, data.get(api_key), api_key)

    if field_errors:
        return jsonify({"ok": False, "error": "validation_error", "field_errors": field_errors}), 400

    if update_payload:
        runtime_settings_store.apply_updates(update_payload)
        _refresh_runtime_globals()
        _ensure_rag_embedding_defaults()
        _reset_rag_rerank_cache()
    after = runtime_settings_store.snapshot()
    runtime = _rt()

    if any(before.get(key) != after.get(key) for key in ('LMSTUDIO_API_BASE', 'LMSTUDIO_MODEL', 'LMSTUDIO_API_KEY', 'LM_DEFAULT_PROVIDER')):
        _invalidate_llm_cache()

    configure_llm_cache(
        enabled=runtime.llm_cache_enabled,
        max_items=runtime.llm_cache_max_items,
        ttl_seconds=runtime.llm_cache_ttl_seconds,
    )
    configure_search_cache(
        enabled=runtime.search_cache_enabled,
        max_items=runtime.search_cache_max_items,
        ttl_seconds=runtime.search_cache_ttl_seconds,
    )
    configure_llm_pool(
        workers=runtime.llm_pool_global_concurrency,
        per_user_limit=runtime.llm_pool_per_user_concurrency,
        max_queue=runtime.llm_queue_max_size,
    )
    runtime.apply_to_flask_config(app)

    collections_payload = data.get('collections')
    if isinstance(collections_payload, list):
        try:
            for entry in collections_payload:
                try:
                    cid = int(entry.get('id'))
                except Exception:
                    continue
                col = Collection.query.get(cid)
                if not col:
                    continue
                if 'searchable' in entry:
                    col.searchable = bool(entry['searchable'])
                if 'graphable' in entry:
                    col.graphable = bool(entry['graphable'])
            db.session.commit()
        except Exception as exc:
            db.session.rollback()
            app.logger.warning('Не удалось обновить коллекции в настройках: %s', exc)

    aiword_payload = data.get('aiword_users')
    if aiword_payload is not None:
        try:
            AiWordAccess.query.delete()
            db.session.commit()
            added = set()
            current_user = _load_current_user()
            grant_id = current_user.id if current_user else None
            users_list = aiword_payload if isinstance(aiword_payload, list) else []
            for item in users_list:
                try:
                    uid = int(item)
                except Exception:
                    continue
                if uid in added:
                    continue
                u = User.query.get(uid)
                if not u:
                    continue
                db.session.add(AiWordAccess(user_id=uid, granted_by=grant_id))
                added.add(uid)
            db.session.commit()
            try:
                _log_user_action('aiword_access_update', 'aiword', None, detail=json.dumps(sorted(list(added))))
            except Exception:
                pass
        except Exception:
            db.session.rollback()

    _save_runtime_settings_to_disk()
    return jsonify({'ok': True})


@app.route('/api/material-types')
def api_material_types():
    user = _load_current_user()
    if not user:
        return jsonify({'ok': False, 'error': 'Не авторизовано'}), 401
    runtime = _rt()
    items = getattr(runtime, 'material_types', None) or []
    payload = copy.deepcopy(items)
    labels: Dict[str, str] = {}
    for entry in payload:
        if not isinstance(entry, dict):
            continue
        key = str(entry.get('key') or '').strip()
        if not key:
            continue
        label = str(entry.get('label') or key).strip() or key
        labels[key] = label
        for alias in entry.get('aliases') or []:
            alias_key = str(alias or '').strip()
            if alias_key and alias_key not in labels:
                labels[alias_key] = label
    return jsonify({'ok': True, 'items': payload, 'labels': labels})


@app.route('/api/facet-config', methods=['GET', 'POST'])
def api_facet_config():
    admin = _require_admin()

    def _payload() -> dict:
        return {
            'ok': True,
            'config': _facet_config_state(),
            'options': facet_service.key_options()
        }

    if request.method == 'GET':
        return jsonify(_payload())

    data = request.get_json(silent=True) or {}
    search_payload = data.get('search') if isinstance(data.get('search'), dict) else {}
    graph_payload = data.get('graph') if isinstance(data.get('graph'), dict) else {}

    def _parse_keys(value) -> list[str] | None:
        if value is None:
            return None
        if isinstance(value, (list, tuple)):
            cleaned = [str(v or '').strip() for v in value if str(v or '').strip()]
            return cleaned
        return []

    runtime = _rt()
    include_types = bool(search_payload.get('include_types', runtime.search_facet_include_types))
    search_keys = _parse_keys(search_payload.get('tag_keys'))
    graph_keys = _parse_keys(graph_payload.get('tag_keys'))

    runtime_settings_store.apply_updates({
        'SEARCH_FACET_INCLUDE_TYPES': include_types,
        'SEARCH_FACET_TAG_KEYS': search_keys,
        'GRAPH_FACET_TAG_KEYS': graph_keys,
    })
    _refresh_runtime_globals()
    runtime = _rt()
    runtime.apply_to_flask_config(current_app)
    _save_runtime_settings_to_disk()
    try:
        detail = json.dumps({
            'search': {'include_types': include_types, 'tag_keys': search_keys},
            'graph': {'tag_keys': graph_keys}
        }, ensure_ascii=False)
        _log_user_action(admin, 'facet_config_update', 'facet_config', None, detail)
    except Exception:
        pass
    facet_service.invalidate('facet_config_update')
    return jsonify(_payload())


@app.route('/api/facets')
def api_facets():
    """Фасеты для текущих фильтров: типы и теги (с учётом выбранных фильтров)."""
    q = request.args.get('q', '').strip()
    material_type = request.args.get('type', '').strip()
    tag_filters = request.args.getlist('tag')
    context = (request.args.get('context') or 'search').strip().lower()
    if context not in ('search', 'graph'):
        context = 'search'
    cfg_key = 'GRAPH_FACET_TAG_KEYS' if context == 'graph' else 'SEARCH_FACET_TAG_KEYS'
    allowed_cfg = current_app.config.get(cfg_key)
    if isinstance(allowed_cfg, (list, tuple)):
        allowed_keys_list = [str(v or '').strip() for v in allowed_cfg if str(v or '').strip()]
    elif allowed_cfg is None:
        allowed_keys_list = None
    else:
        allowed_keys_list = []
    include_types = bool(current_app.config.get('SEARCH_FACET_INCLUDE_TYPES', True)) if context == 'search' else False

    year_from = (request.args.get('year_from') or '').strip()
    year_to = (request.args.get('year_to') or '').strip()
    size_min = (request.args.get('size_min') or '').strip()
    size_max = (request.args.get('size_max') or '').strip()
    collection_filter = _parse_collection_param(request.args.get('collection_id'))
    allowed_scope = _current_allowed_collections()

    request_args_tuple = tuple(sorted((key, tuple(request.args.getlist(key))) for key in request.args))
    params = FacetQueryParams(
        query=q,
        material_type=material_type,
        context=context,
        include_types=include_types,
        tag_filters=tag_filters,
        collection_filter=collection_filter,
        allowed_scope=allowed_scope,
        allowed_keys_list=allowed_keys_list,
        year_from=year_from,
        year_to=year_to,
        size_min=size_min,
        size_max=size_max,
        sources={'tags': True, 'authors': True, 'years': True},
        request_args=request_args_tuple,
    )

    payload = facet_service.get_facets(
        params,
        search_candidate_fn=_search_candidate_ids,
        like_filter_fn=_apply_like_filter,
    )
    return jsonify(payload)

    base_query = _apply_file_access_filter(File.query)
    try:
        base_query = base_query.join(Collection, File.collection_id == Collection.id).filter(Collection.searchable == True)
    except Exception:
        pass
    if collection_filter is not None:
        base_query = base_query.filter(File.collection_id == collection_filter)
    if material_type:
        base_query = base_query.filter(File.material_type == material_type)
    if q:
        like = f'%{q}%'
        tag_like = exists().where(and_(
            Tag.file_id == File.id,
            or_(Tag.value.ilike(like), Tag.key.ilike(like))
        ))
        filters = [
            File.title.ilike(like),
            File.author.ilike(like),
            File.keywords.ilike(like),
            File.filename.ilike(like),
            File.text_excerpt.ilike(like),
        ]
        if hasattr(File, 'abstract'):
            filters.append(File.abstract.ilike(like))
        filters.append(tag_like)
        base_query = base_query.filter(or_(*filters))
    if year_from:
        base_query = base_query.filter(File.year >= year_from)
    if year_to:
        base_query = base_query.filter(File.year <= year_to)
    if size_min:
        try:
            base_query = base_query.filter(File.size >= int(size_min))
        except Exception:
            pass
    if size_max:
        try:
            base_query = base_query.filter(File.size <= int(size_max))
        except Exception:
            pass

    types_facet: list[list] = []
    if include_types:
        types = db.session.query(File.material_type, func.count(File.id))
        types = types.filter(File.id.in_(base_query.with_entities(File.id)))
        types = types.group_by(File.material_type).all()
        types_facet = [[mt, cnt] for (mt, cnt) in types]

    base_ids_subq = base_query.with_entities(File.id).subquery()
    base_keys = [row[0] for row in db.session.query(Tag.key).filter(Tag.file_id.in_(base_ids_subq)).distinct().all()]
    selected: dict[str, list[str]] = {}
    for tf in tag_filters:
        if '=' in tf:
            k, v = tf.split('=', 1)
            selected.setdefault(k, []).append(v)
            if k not in base_keys:
                base_keys.append(k)

    if allowed_keys_set is not None:
        base_keys = [k for k in base_keys if k in allowed_keys_set or k in selected]
    for key in selected:
        if key not in base_keys:
            base_keys.append(key)

    tag_facets: dict[str, list[list]] = {}
    for key in base_keys:
        if allowed_keys_set is not None and key not in allowed_keys_set and key not in selected:
            continue
        qk = base_query
        for tf in tag_filters:
            if '=' not in tf:
                continue
            k, v = tf.split('=', 1)
            if k == key:
                continue
            tk = aliased(Tag)
            qk = qk.join(tk, tk.file_id == File.id).filter(and_(tk.key == k, tk.value.ilike(f'%{v}%')))
        ids_subq = qk.with_entities(File.id).distinct().subquery()
        rows = db.session.query(Tag.value, func.count(Tag.id)) \
            .filter(and_(Tag.file_id.in_(ids_subq), Tag.key == key)) \
            .group_by(Tag.value) \
            .order_by(func.count(Tag.id).desc()) \
            .all()
        tag_facets[key] = [[val, cnt] for (val, cnt) in rows]
        if key in selected:
            present_vals = {val for (val, _c) in tag_facets[key]}
            for v in selected[key]:
                if v not in present_vals:
                    tag_facets[key].append([v, 0])

    return jsonify({
        'types': types_facet,
        'tag_facets': tag_facets,
        'include_types': include_types,
        'allowed_keys': allowed_keys_list,
        'context': context,
    })

@app.route("/settings", methods=["GET"])
def settings_redirect():
    """Совместимость: перенаправить на новый интерфейс настроек."""
    return redirect('/app/settings')

@app.route('/settings/collections', methods=['POST'])
@require_admin
def settings_collections():
    # Обновляем флаги у существующих коллекций и при необходимости добавляем новую
    try:
        cols = Collection.query.all()
        # обновляем флаги
        for c in cols:
            s_key = f'search_{c.id}'
            g_key = f'graph_{c.id}'
            c.searchable = (request.form.get(s_key) == 'on')
            c.graphable = (request.form.get(g_key) == 'on')
        new_name = (request.form.get('new_name') or '').strip()
        if new_name:
            slug = _slugify(new_name)
            exists = Collection.query.filter((Collection.slug==slug) | (Collection.name==new_name)).first()
            if not exists:
                db.session.add(Collection(name=new_name, slug=slug, searchable=True, graphable=True))
        db.session.commit()
        return jsonify({"status": "ok"})
    except Exception as e:
        db.session.rollback()
        return json_error(f'Ошибка обновления коллекций: {e}', 500)


def _db_dialect_name() -> str:
    try:
        name = str(getattr(db.engine.dialect, "name", "") or "").strip().lower()
    except Exception:
        name = ""
    if name == "postgres":
        return "postgresql"
    return name or "unknown"


def _active_database_uri() -> str:
    return str(
        current_app.config.get("SQLALCHEMY_DATABASE_URI")
        or os.getenv("SQLALCHEMY_DATABASE_URI")
        or os.getenv("DATABASE_URL")
        or ""
    ).strip()


def _normalize_db_type(value: str | None) -> str:
    raw = str(value or "").strip().lower()
    if raw in {"postgres", "postgresql"}:
        return "postgresql"
    if raw in {"sqlite", "sqlite3"}:
        return "sqlite"
    return ""


def _requested_db_type() -> str:
    payload = request.get_json(silent=True) if request.is_json else None
    requested = ""
    if isinstance(payload, dict):
        requested = _normalize_db_type(payload.get("db_type"))
    if not requested:
        requested = _normalize_db_type(request.form.get("db_type"))
    if not requested:
        requested = _normalize_db_type(request.args.get("db_type"))
    return requested or _db_dialect_name()


def _pg_cli_connection(db_uri: str) -> tuple[list[str], dict[str, str], str]:
    try:
        parsed = make_url(db_uri)
    except Exception as exc:
        raise ValueError(f"Некорректный URI базы данных: {exc}") from exc
    if not str(parsed.drivername or "").startswith("postgresql"):
        raise ValueError("Текущая база не PostgreSQL")
    dbname = str(parsed.database or "").strip()
    if not dbname:
        raise ValueError("В URI PostgreSQL не указано имя базы")

    args: list[str] = []
    if parsed.host:
        args.extend(["-h", str(parsed.host)])
    if parsed.port:
        args.extend(["-p", str(parsed.port)])
    if parsed.username:
        args.extend(["-U", str(parsed.username)])

    env = os.environ.copy()
    if parsed.password is not None:
        env["PGPASSWORD"] = str(parsed.password)
    return args, env, dbname


def _run_cli(command: list[str], env: dict[str, str]) -> tuple[int, str]:
    proc = subprocess.run(
        command,
        env=env,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        check=False,
    )
    stderr = (proc.stderr or "").strip()
    stdout = (proc.stdout or "").strip()
    detail = stderr or stdout
    return int(proc.returncode or 0), detail[:4000]


def _clear_file_caches() -> None:
    try:
        static_dir = Path(app.static_folder)
        txt_cache = static_dir / 'cache' / 'text_excerpts'
        if txt_cache.exists():
            for fp in txt_cache.glob('*.txt'):
                try:
                    fp.unlink()
                except Exception:
                    pass
        thumbs = static_dir / 'thumbnails'
        if thumbs.exists():
            for fp in thumbs.glob('*.png'):
                try:
                    fp.unlink()
                except Exception:
                    pass
    except Exception:
        pass

@app.route('/admin/backup-db', methods=['POST'])
@require_admin
def backup_db():
    # Создаём и отправляем резервную копию БД для выбранного/текущего типа
    try:
        requested_type = _requested_db_type()
        active_type = _db_dialect_name()
        if requested_type != active_type:
            return json_error(
                f"Выбран тип '{requested_type}', но приложение подключено к '{active_type}'. "
                "Смените тип БД в настройках UI или подключение приложения.",
                400,
            )
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        bdir = BASE_DIR / 'backups'
        bdir.mkdir(exist_ok=True)
        if active_type == 'sqlite':
            src = BASE_DIR / 'catalogue.db'
            if not src.exists():
                return json_error('Файл базы не найден', 404)
            dst = bdir / f'catalogue_{ts}.db'
            shutil.copy2(src, dst)
            return send_from_directory(bdir, dst.name, as_attachment=True)

        if active_type == 'postgresql':
            payload = request.get_json(silent=True) if request.is_json else {}
            fmt = str(
                (payload or {}).get("format")
                or request.form.get("format")
                or request.args.get("format")
                or "dump"
            ).strip().lower()
            if fmt not in {"dump", "sql"}:
                fmt = "dump"
            pg_format = "c" if fmt == "dump" else "p"
            suffix = ".dump" if fmt == "dump" else ".sql"
            dst = bdir / f'catalogue_{ts}{suffix}'
            db_uri = _active_database_uri()
            conn_args, env, dbname = _pg_cli_connection(db_uri)
            cmd = [
                "pg_dump",
                *conn_args,
                "--no-owner",
                "--no-privileges",
                "--clean",
                "--if-exists",
                f"--format={pg_format}",
                "--file",
                str(dst),
                dbname,
            ]
            code, detail = _run_cli(cmd, env)
            if code != 0:
                return json_error(f"Ошибка pg_dump: {detail or 'неизвестная ошибка'}", 500)
            return send_from_directory(bdir, dst.name, as_attachment=True)

        return json_error(f"Неподдерживаемый тип БД: {active_type}", 400)
    except Exception as e:
        return json_error(f'Ошибка резервного копирования: {e}', 500)

@app.route('/admin/clear-db', methods=['POST'])
@require_admin
def clear_db():
    # Опасно: удалить все записи из основных таблиц
    try:
        Tag.query.delete()
        ChangeLog.query.delete()
        File.query.delete()
        db.session.commit()
        # очищаем статические кэши (текстовые фрагменты и миниатюры)
        try:
            static_dir = Path(app.static_folder)
            txt_cache = static_dir / 'cache' / 'text_excerpts'
            if txt_cache.exists():
                for fp in txt_cache.glob('*.txt'):
                    try: fp.unlink()
                    except Exception: pass
            thumbs = static_dir / 'thumbnails'
            if thumbs.exists():
                for fp in thumbs.glob('*.png'):
                    try: fp.unlink()
                    except Exception: pass
        except Exception:
            pass
        return jsonify({"status": "ok"})
    except Exception as e:
        db.session.rollback()
        return json_error(f'Ошибка очистки: {e}', 500)

@app.route('/admin/import-db', methods=['POST'])
@require_admin
def import_db():
    """Import DB dump/file for the selected current backend (sqlite/postgresql)."""
    requested_type = _requested_db_type()
    active_type = _db_dialect_name()
    if requested_type != active_type:
        return json_error(
            f"Выбран тип '{requested_type}', но приложение подключено к '{active_type}'. "
            "Смените тип БД в настройках UI или подключение приложения.",
            400,
        )

    file = request.files.get('dbfile')
    if not file or not file.filename:
        return json_error('Файл базы не выбран', 400)
    filename = secure_filename(file.filename)
    tmp: Path | None = None
    try:
        if active_type == "sqlite":
            if not filename.lower().endswith('.db'):
                return json_error('Ожидался файл .db (SQLite)', 400)
            try:
                head = file.stream.read(16)
                if head[:15] != b'SQLite format 3':
                    file.stream.seek(0)
                    return json_error('Файл не похож на SQLite базу', 400)
            except Exception:
                pass
            finally:
                try:
                    file.stream.seek(0)
                except Exception:
                    pass

            dst = BASE_DIR / 'catalogue.db'
            if dst.exists():
                ts = datetime.now().strftime('%Y%m%d_%H%M%S')
                bdir = BASE_DIR / 'backups'
                bdir.mkdir(exist_ok=True)
                backup_path = bdir / f'catalogue_before_import_{ts}.db'
                shutil.copy2(dst, backup_path)

            tmp = BASE_DIR / f'.upload_import_tmp_{os.getpid()}_{int(time.time())}.db'
            file.save(tmp)

            try:
                con = sqlite3.connect(str(tmp))
                cur = con.cursor()
                cur.execute("SELECT name FROM sqlite_master WHERE type='table'")
                tables = {r[0] for r in cur.fetchall()}
                required = {'files', 'tags', 'tag_schemas', 'changelog'}
                missing = required - tables
                con.close()
                if missing:
                    try:
                        tmp.unlink()
                    except Exception:
                        pass
                    return json_error('В файле базы нет требуемых таблиц: ' + ', '.join(sorted(missing)), 400)
            except Exception as e:
                if tmp and tmp.exists():
                    tmp.unlink()
                return json_error(f'Не удалось проверить схему базы: {e}', 500)

            try:
                db.session.close()
                db.session.remove()
            except Exception:
                pass
            try:
                db.engine.dispose()
            except Exception:
                pass

            shutil.move(str(tmp), str(dst))
            tmp = None
            _clear_file_caches()
            return jsonify({"status": "ok", "dialect": "sqlite"})

        if active_type == "postgresql":
            ts = datetime.now().strftime('%Y%m%d_%H%M%S')
            db_uri = _active_database_uri()
            conn_args, env, dbname = _pg_cli_connection(db_uri)
            bdir = BASE_DIR / "backups"
            bdir.mkdir(exist_ok=True)
            backup_before = bdir / f"catalogue_before_import_{ts}.dump"
            backup_cmd = [
                "pg_dump",
                *conn_args,
                "--no-owner",
                "--no-privileges",
                "--clean",
                "--if-exists",
                "--format=c",
                "--file",
                str(backup_before),
                dbname,
            ]
            code, detail = _run_cli(backup_cmd, env)
            if code != 0:
                return json_error(f"Не удалось создать pre-import backup: {detail or 'ошибка pg_dump'}", 500)

            suffix = Path(filename).suffix.lower() or ".bin"
            tmp = BASE_DIR / f'.upload_import_tmp_{os.getpid()}_{int(time.time())}{suffix}'
            file.save(tmp)

            try:
                db.session.close()
                db.session.remove()
            except Exception:
                pass
            try:
                db.engine.dispose()
            except Exception:
                pass

            reset_cmd = [
                "psql",
                *conn_args,
                "-d",
                dbname,
                "-v",
                "ON_ERROR_STOP=1",
                "-c",
                "DROP SCHEMA IF EXISTS public CASCADE; CREATE SCHEMA public;",
            ]
            code, detail = _run_cli(reset_cmd, env)
            if code != 0:
                return json_error(f"Не удалось очистить схему PostgreSQL: {detail or 'ошибка psql'}", 500)

            is_sql = suffix in {".sql"}
            if is_sql:
                import_cmd = [
                    "psql",
                    *conn_args,
                    "-d",
                    dbname,
                    "-v",
                    "ON_ERROR_STOP=1",
                    "-f",
                    str(tmp),
                ]
            else:
                import_cmd = [
                    "pg_restore",
                    *conn_args,
                    "-d",
                    dbname,
                    "--clean",
                    "--if-exists",
                    "--no-owner",
                    "--no-privileges",
                    str(tmp),
                ]
            code, detail = _run_cli(import_cmd, env)
            if code != 0:
                return json_error(f"Ошибка импорта PostgreSQL: {detail or 'ошибка восстановления'}", 500)

            _clear_file_caches()
            return jsonify({"status": "ok", "dialect": "postgresql"})

        return json_error(f"Неподдерживаемый тип БД: {active_type}", 400)
    except Exception as e:
        return json_error(f'Ошибка импорта базы: {e}', 500)
    finally:
        if tmp and tmp.exists():
            try:
                tmp.unlink()
            except Exception:
                pass


@app.route('/admin/db/migrate-sqlite-to-postgres', methods=['POST'])
@require_admin
def migrate_sqlite_to_postgres():
    payload = request.get_json(silent=True) or {}
    sqlite_path_raw = str(payload.get('sqlite_path') or 'catalogue.db').strip()
    pg_url = str(payload.get('postgres_url') or '').strip()
    pg_host = str(payload.get('postgres_host') or '').strip()
    pg_port = int(payload.get('postgres_port') or 5432)
    pg_db = str(payload.get('postgres_db') or '').strip()
    pg_user = str(payload.get('postgres_user') or '').strip()
    pg_password = str(payload.get('postgres_password') or '').strip()
    mode_raw = str(payload.get('mode') or 'dry-run').strip().lower()
    mode = 'run' if mode_raw == 'run' else 'dry-run'
    if not pg_url:
        if not (pg_host and pg_db and pg_user):
            return json_error('Укажите postgres_url или набор полей host/db/user/password', 400)
        auth = quote_plus(pg_user)
        if pg_password:
            auth = f"{auth}:{quote_plus(pg_password)}"
        pg_url = f"postgresql://{auth}@{pg_host}:{pg_port}/{quote_plus(pg_db)}"

    base_dir = BASE_DIR.resolve()
    sqlite_path = Path(sqlite_path_raw)
    if not sqlite_path.is_absolute():
        sqlite_path = (base_dir / sqlite_path).resolve()
    if not sqlite_path.exists():
        return json_error(f'SQLite файл не найден: {sqlite_path}', 404)
    if not sqlite_path.is_file():
        return json_error(f'Неверный путь SQLite: {sqlite_path}', 400)

    script_path = (base_dir / 'scripts' / 'migrate_sqlite_to_postgres.sh').resolve()
    if not script_path.exists():
        return json_error('Скрипт миграции не найден: scripts/migrate_sqlite_to_postgres.sh', 404)

    cmd = [
        str(script_path),
        '--sqlite',
        str(sqlite_path),
        '--pg',
        pg_url,
        '--run' if mode == 'run' else '--dry-run',
    ]
    env = os.environ.copy()
    env['DATABASE_URL'] = pg_url
    try:
        proc = subprocess.run(
            cmd,
            cwd=str(base_dir),
            env=env,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            check=False,
            timeout=30 * 60,
        )
    except subprocess.TimeoutExpired as exc:
        output = (exc.stdout or '') if isinstance(exc.stdout, str) else ''
        return jsonify({
            'ok': False,
            'mode': mode,
            'exit_code': 124,
            'output': output[-12000:],
            'error': 'Таймаут выполнения миграции (30 мин)',
        }), 504
    except FileNotFoundError:
        return json_error('Не удалось запустить скрипт миграции', 500)
    except Exception as exc:
        return json_error(f'Ошибка запуска миграции: {exc}', 500)

    output = (proc.stdout or '').strip()
    return jsonify({
        'ok': proc.returncode == 0,
        'mode': mode,
        'exit_code': int(proc.returncode or 0),
        'output': output[-12000:],
    }), (200 if proc.returncode == 0 else 500)


def _task_to_dict(task: TaskRecord) -> dict:
    payload_json = {}
    if task.payload:
        try:
            payload_json = json.loads(task.payload)
        except Exception:
            payload_json = {}
    return {
        'id': task.id,
        'name': task.name,
        'status': task.status,
        'progress': float(task.progress or 0.0),
        'payload': task.payload,
        'payload_json': payload_json,
        'created_at': task.created_at.isoformat() if task.created_at else None,
        'started_at': task.started_at.isoformat() if task.started_at else None,
        'finished_at': task.finished_at.isoformat() if task.finished_at else None,
        'error': task.error,
    }


def _reset_inflight_tasks() -> int:
    """Mark tasks that were left in non-final statuses as failed on startup."""
    try:
        stuck = TaskRecord.query.filter(TaskRecord.status.in_(TASK_STUCK_STATUSES)).all()
    except Exception as exc:
        app.logger.warning(f"[tasks] failed to load stuck tasks: {exc}")
        return 0
    if not stuck:
        return 0
    now = datetime.utcnow()
    updated = 0
    for task in stuck:
        try:
            task.status = 'error'
            task.error = 'Прервано: перезапуск приложения'
            if not task.finished_at:
                task.finished_at = now
            if not task.started_at:
                task.started_at = task.created_at or now
            if task.progress is None:
                task.progress = 0.0
            updated += 1
        except Exception:
            app.logger.debug('[tasks] failed to mark task #%s as interrupted', getattr(task, 'id', '?'))
    try:
        db.session.commit()
    except Exception as exc:
        db.session.rollback()
        app.logger.warning(f"[tasks] failed to commit stuck task reset: {exc}")
        return 0
    if updated:
        app.logger.info(f"[tasks] reset {updated} stuck task(s) on startup")
    return updated


def _cleanup_old_tasks() -> None:
    cutoff = datetime.utcnow() - TASK_RETENTION_WINDOW
    try:
        deleted = TaskRecord.query.filter(
            TaskRecord.status.in_(TASK_FINAL_STATUSES),
            TaskRecord.created_at.isnot(None),
            TaskRecord.created_at < cutoff
        ).delete(synchronize_session=False)
        db.session.commit()
        if deleted:
            app.logger.info(f"[tasks] removed {deleted} tasks older than {cutoff.isoformat()} (final statuses)")
    except Exception as exc:
        db.session.rollback()
        app.logger.warning(f"[tasks] cleanup failed: {exc}")


@admin_bp.route('/tasks', methods=['GET', 'DELETE'])
@require_admin
def api_admin_tasks():
    if request.method == 'DELETE':
        payload = request.get_json(silent=True) or {}
        status_key = str(request.args.get('status') or payload.get('status') or 'final').strip().lower()
        before_raw = str(request.args.get('before') or payload.get('before') or '').strip()
        status_map: dict[str, tuple[str, ...] | None] = {
            'final': TASK_FINAL_STATUSES,
            'done': TASK_FINAL_STATUSES,
            'completed': ('completed',),
            'complete': ('completed',),
            'errors': ('error',),
            'error': ('error',),
            'failed': ('error',),
            'cancelled': ('cancelled',),
            'canceled': ('cancelled',),
            'all': None,
            'any': None,
        }
        statuses = status_map.get(status_key or 'final')
        if statuses is None and status_key not in ('all', 'any'):
            return jsonify({'ok': False, 'error': 'unsupported_status'}), 400
        cutoff: datetime | None = None
        if before_raw:
            try:
                if len(before_raw) <= 10:
                    cutoff = datetime.strptime(before_raw, '%Y-%m-%d') + timedelta(days=1)
                else:
                    cutoff = datetime.fromisoformat(before_raw)
            except Exception:
                return jsonify({'ok': False, 'error': 'invalid_datetime'}), 400
        query = TaskRecord.query
        if statuses:
            query = query.filter(TaskRecord.status.in_(statuses))
        elif status_key not in ('all', 'any'):
            query = query.filter(TaskRecord.status.in_(TASK_FINAL_STATUSES))
        if cutoff is not None:
            query = query.filter(TaskRecord.created_at.isnot(None)).filter(TaskRecord.created_at < cutoff)
        try:
            deleted = query.delete(synchronize_session=False)
            db.session.commit()
        except Exception as exc:
            db.session.rollback()
            app.logger.warning(f"[tasks] bulk delete failed: {exc}")
            return jsonify({'ok': False, 'error': 'delete_failed'}), 500
        actor = _load_current_user()
        detail_payload = {'deleted': deleted, 'status': status_key or 'final'}
        if cutoff is not None:
            detail_payload['before'] = cutoff.isoformat()
        try:
            _log_user_action(actor, 'task_cleanup', 'task', None, detail=json.dumps(detail_payload))
        except Exception:
            pass
        app.logger.info(f"[tasks] bulk delete {deleted} records (status={status_key or 'final'}, before={cutoff.isoformat() if cutoff else '—'})")
        return jsonify({'ok': True, 'deleted': deleted})
    try:
        limit = int(request.args.get('limit', '50'))
    except Exception:
        limit = 50
    limit = max(1, min(limit, 200))
    status_filter = (request.args.get('status') or '').strip().lower()
    name_filter = (request.args.get('name') or '').strip()
    status_mode = None
    status_values: tuple[str, ...] | None = None
    if status_filter:
        if status_filter in {'active'}:
            status_mode = 'active'
        elif status_filter in {'final', 'done', 'completed'}:
            status_mode = 'final'
        elif status_filter in {'all'}:
            status_mode = 'all'
        else:
            tokens = [
                token.strip().lower()
                for token in status_filter.replace(';', ',').split(',')
                if token.strip()
            ]
            if tokens:
                status_values = tuple(tokens)
    _cleanup_old_tasks()
    cutoff = datetime.utcnow() - TASK_RETENTION_WINDOW
    query = TaskRecord.query.order_by(TaskRecord.created_at.desc())
    if name_filter:
        like_pattern = f"%{name_filter}%"
        query = query.filter(TaskRecord.name.ilike(like_pattern))
    if status_values:
        query = query.filter(func.lower(TaskRecord.status).in_(status_values))
    elif status_mode == 'active':
        query = query.filter(TaskRecord.status.notin_(TASK_FINAL_STATUSES))
    elif status_mode == 'final':
        query = query.filter(TaskRecord.status.in_(TASK_FINAL_STATUSES))
    if status_mode != 'all':
        query = query.filter(
            or_(
                TaskRecord.status.notin_(TASK_FINAL_STATUSES),
                TaskRecord.status.is_(None),
                TaskRecord.created_at.is_(None),
                TaskRecord.created_at >= cutoff
            )
        )
    tasks = query.limit(limit).all()
    payload = [_task_to_dict(t) for t in tasks]
    if SCAN_PROGRESS.get('running'):
        progress = 0.0
        total = float(SCAN_PROGRESS.get('total') or 0)
        processed = float(SCAN_PROGRESS.get('processed') or 0)
        if total > 0:
            progress = min(1.0, processed / total)
        active_id = SCAN_TASK_ID or 0
        if not any(t.get('id') == active_id for t in payload):
            payload.insert(0, {
                'id': active_id,
                'name': 'scan',
                'status': 'running',
                'progress': progress,
                'payload': json.dumps({
                    'stage': SCAN_PROGRESS.get('stage'),
                    'current': SCAN_PROGRESS.get('current'),
                }),
                'created_at': None,
                'started_at': None,
                'finished_at': None,
                'error': SCAN_PROGRESS.get('error'),
            })
    return jsonify({'ok': True, 'tasks': payload})


@admin_bp.route('/llm-queue', methods=['GET'])
@require_admin
def api_admin_llm_queue():
    stats = get_llm_pool().stats()
    return jsonify({'ok': True, 'stats': stats})


@admin_bp.route('/tasks/<int:task_id>', methods=['PATCH', 'DELETE'])
@require_admin
def api_admin_task_detail(task_id: int):
    task = TaskRecord.query.get_or_404(task_id)
    if request.method == 'DELETE':
        db.session.delete(task)
        db.session.commit()
        return jsonify({'ok': True})
    data = request.get_json(silent=True) or {}
    status = str(data.get('status') or '').strip().lower()
    if status == 'cancel':
        if task.name == 'scan' and SCAN_PROGRESS.get('running') and SCAN_TASK_ID == task.id:
            scan_cancel()
            return jsonify({'ok': True, 'task': _task_to_dict(task)})
        task.status = 'cancelled'
        db.session.commit()
        return jsonify({'ok': True, 'task': _task_to_dict(task)})
    if status:
        task.status = status
    progress = data.get('progress')
    if progress is not None:
        try:
            task.progress = float(progress)
        except Exception:
            pass
    db.session.commit()
    return jsonify({'ok': True, 'task': _task_to_dict(task)})


@admin_bp.route('/system-logs', methods=['GET', 'DELETE'])
@require_admin
def api_admin_system_logs():
    if request.method == 'DELETE':
        payload = request.get_json(silent=True) or {}
        name = str(request.args.get('name') or payload.get('name') or LOG_FILE_PATH.name).strip()
        path = _resolve_log_name(name)
        if not path:
            return jsonify({'ok': False, 'error': 'log_not_found'}), 404
        try:
            with path.open('w', encoding='utf-8') as fh:
                fh.truncate(0)
        except Exception as exc:
            return jsonify({'ok': False, 'error': str(exc)}), 500
        handler = _get_rotating_log_handler()
        if handler and Path(getattr(handler, 'baseFilename', '')).resolve() == path.resolve():
            handler.acquire()
            try:
                handler.flush()
                stream = getattr(handler, 'stream', None)
                if stream:
                    try:
                        stream.close()
                    except Exception:
                        pass
                handler.stream = handler._open()
            finally:
                handler.release()
        actor = _load_current_user()
        try:
            detail = json.dumps({'name': path.name, 'action': 'truncate'})
        except Exception:
            detail = None
        _log_user_action(actor, 'system_log_clear', 'log', None, detail=detail)
        app.logger.info(f"[logs] {getattr(actor, 'username', 'admin')} truncated log {path.name}")
        return jsonify({'ok': True, 'files': _list_system_log_files()})

    try:
        limit = int(request.args.get('limit', '200'))
    except Exception:
        limit = 200
    limit = max(10, min(limit, 2000))
    name = str(request.args.get('name') or '').strip()
    path = _resolve_log_name(name) or LOG_FILE_PATH
    if not path.exists():
        return jsonify({'ok': False, 'error': 'log_not_found'}), 404
    available_files = _list_system_log_files()
    selected = next((entry for entry in available_files if entry.get('name') == path.name), None)
    if not selected:
        try:
            stat = path.stat()
            selected = {
                'name': path.name,
                'size': stat.st_size,
                'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'rotated': path.name != LOG_FILE_PATH.name,
            }
        except Exception:
            selected = {'name': path.name, 'rotated': path.name != LOG_FILE_PATH.name}
    lines = _tail_log_file(path, limit)
    return jsonify({'ok': True, 'file': selected, 'lines': lines, 'available': available_files, 'limit': limit})


@admin_bp.route('/system-logs/rotate', methods=['POST'])
@require_admin
def api_admin_system_logs_rotate():
    handler = _get_rotating_log_handler()
    if not handler:
        return jsonify({'ok': False, 'error': 'rotation_unavailable'}), 400
    handler.acquire()
    try:
        handler.flush()
        handler.doRollover()
        if handler.stream is None:
            handler.stream = handler._open()
    except Exception as exc:
        handler.release()
        app.logger.warning(f"[logs] manual rotation failed: {exc}")
        return jsonify({'ok': False, 'error': str(exc)}), 500
    handler.release()
    actor = _load_current_user()
    try:
        detail = json.dumps({'name': LOG_FILE_PATH.name, 'action': 'rotate'})
    except Exception:
        detail = None
    _log_user_action(actor, 'system_log_rotate', 'log', None, detail=detail)
    app.logger.info(f"[logs] {getattr(actor, 'username', 'admin')} rotated system log")
    return jsonify({'ok': True, 'files': _list_system_log_files()})


@admin_bp.route('/system-logs/download', methods=['GET'])
@require_admin
def api_admin_system_logs_download():
    name = str(request.args.get('name') or LOG_FILE_PATH.name).strip()
    path = _resolve_log_name(name)
    if not path:
        abort(404)
    return send_file(path, mimetype='text/plain', as_attachment=True, download_name=path.name, conditional=True)


@admin_bp.route('/actions', methods=['GET', 'DELETE'])
@require_admin
def api_admin_actions():
    if request.method == 'DELETE':
        payload = request.get_json(silent=True) or {}
        before_raw = str(request.args.get('before') or payload.get('before') or '').strip()
        if not before_raw:
            return jsonify({'ok': False, 'error': 'parameter "before" is required'}), 400
        try:
            if len(before_raw) <= 10:
                before_dt = datetime.strptime(before_raw, '%Y-%m-%d') + timedelta(days=1)
            else:
                before_dt = datetime.fromisoformat(before_raw)
        except Exception:
            return jsonify({'ok': False, 'error': 'invalid datetime format'}), 400
        try:
            deleted = UserActionLog.query.filter(UserActionLog.created_at < before_dt).delete(synchronize_session=False)
            db.session.commit()
            app.logger.info(f"[actions] deleted {deleted} log records older than {before_dt.isoformat()}")
            return jsonify({'ok': True, 'deleted': deleted})
        except Exception as exc:
            db.session.rollback()
            app.logger.warning(f"[actions] failed to delete logs: {exc}")
            return jsonify({'ok': False, 'error': 'delete_failed'}), 500

    try:
        limit = int(request.args.get('limit', '100'))
    except Exception:
        limit = 100
    limit = max(1, min(limit, 500))
    q = UserActionLog.query.order_by(UserActionLog.created_at.desc())
    user_id = request.args.get('user_id')
    if user_id:
        try:
            q = q.filter(UserActionLog.user_id == int(user_id))
        except Exception:
            pass
    action = request.args.get('action')
    if action:
        q = q.filter(UserActionLog.action == action)
    logs = q.limit(limit).all()
    return jsonify({'ok': True, 'actions': [
        {
            'id': log.id,
            'user_id': log.user_id,
            'username': log.user.username if getattr(log, 'user', None) else None,
            'full_name': log.user.full_name if getattr(log, 'user', None) else None,
            'action': log.action,
            'entity': log.entity,
            'entity_id': log.entity_id,
            'detail': log.detail,
            'created_at': log.created_at.isoformat() if log.created_at else None,
        }
        for log in logs
    ]})


@admin_bp.route('/llm-endpoints', methods=['GET', 'POST'])
@require_admin
def api_admin_llm_endpoints():
    _ensure_llm_schema_once()
    if request.method == 'GET':
        eps = LlmEndpoint.query.order_by(LlmEndpoint.created_at.desc()).all()
        return jsonify({'ok': True, 'items': [
            {
                'id': ep.id,
                'name': ep.name,
                'base_url': ep.base_url,
                'model': ep.model,
                'weight': float(ep.weight or 0.0),
                'purpose': ep.purpose,
                'purposes': _llm_parse_purposes(ep.purpose),
                'provider': (ep.provider or _rt().lm_default_provider),
                'context_length': int(ep.context_length) if getattr(ep, 'context_length', None) else None,
                'instances': int(getattr(ep, 'instances', 1) or 1),
                'created_at': ep.created_at.isoformat() if ep.created_at else None,
            }
            for ep in eps
        ], 'purposes_catalog': LLM_PURPOSES, 'providers_catalog': LLM_PROVIDER_OPTIONS})
    data = request.get_json(silent=True) or {}
    name = (data.get('name') or '').strip()
    base_url = (data.get('base_url') or '').strip()
    model = (data.get('model') or '').strip()
    provider = (data.get('provider') or _rt().lm_default_provider or 'openai').strip().lower()
    if provider not in LLM_PROVIDER_CHOICES:
        return jsonify({'ok': False, 'error': 'некорректный провайдер LLM'}), 400
    if not name or not base_url or not model:
        return jsonify({'ok': False, 'error': 'name, base_url и model обязательны'}), 400
    normalized_purpose = _llm_normalize_purposes(data.get('purposes') if 'purposes' in data else data.get('purpose'))
    ep = LlmEndpoint(
        name=name,
        base_url=base_url,
        model=model,
        api_key=(data.get('api_key') or '').strip() or None,
        weight=float(data.get('weight') or 1.0),
        purpose=normalized_purpose,
        provider=provider,
        context_length=int(data.get('context_length')) if str(data.get('context_length') or '').strip().isdigit() else None,
        instances=max(1, int(data.get('instances') or 1)),
    )
    db.session.add(ep)
    db.session.commit()
    _sync_runtime_default_route_from_llm_endpoints()
    _invalidate_llm_cache()
    _log_user_action(_load_current_user(), 'llm_add', 'llm_endpoint', ep.id, detail=json.dumps({'name': name, 'purpose': ep.purpose}))
    return jsonify({'ok': True, 'item': {
        'id': ep.id,
        'name': ep.name,
        'base_url': ep.base_url,
        'model': ep.model,
        'weight': float(ep.weight or 0.0),
        'purpose': ep.purpose,
        'purposes': _llm_parse_purposes(ep.purpose),
        'provider': ep.provider,
        'context_length': int(ep.context_length) if getattr(ep, 'context_length', None) else None,
        'instances': int(getattr(ep, 'instances', 1) or 1),
        'created_at': ep.created_at.isoformat() if ep.created_at else None,
    }}), 201


@admin_bp.route('/llm-endpoints/<int:endpoint_id>', methods=['PATCH', 'DELETE'])
@require_admin
def api_admin_llm_endpoint_detail(endpoint_id: int):
    _ensure_llm_schema_once()
    ep = LlmEndpoint.query.get_or_404(endpoint_id)
    if request.method == 'DELETE':
        db.session.delete(ep)
        db.session.commit()
        _sync_runtime_default_route_from_llm_endpoints()
        _invalidate_llm_cache()
        _log_user_action(_load_current_user(), 'llm_delete', 'llm_endpoint', endpoint_id)
        return jsonify({'ok': True})
    data = request.get_json(silent=True) or {}
    updated = False
    if 'name' in data:
        ep.name = str(data['name']).strip() or ep.name
        updated = True
    if 'base_url' in data:
        ep.base_url = str(data['base_url']).strip() or ep.base_url
        updated = True
    if 'model' in data:
        ep.model = str(data['model']).strip() or ep.model
        updated = True
    if 'api_key' in data:
        ep.api_key = str(data['api_key']).strip() or None
        updated = True
    if 'weight' in data:
        try:
            ep.weight = float(data['weight'])
            updated = True
        except Exception:
            pass
    if 'purposes' in data:
        ep.purpose = _llm_normalize_purposes(data.get('purposes'))
        updated = True
    elif 'purpose' in data:
        ep.purpose = _llm_normalize_purposes(data.get('purpose'))
        updated = True
    if 'provider' in data:
        prov = str(data['provider']).strip().lower()
        if prov in LLM_PROVIDER_CHOICES:
            ep.provider = prov
            updated = True
        else:
            return jsonify({'ok': False, 'error': 'некорректный провайдер LLM'}), 400
    if 'context_length' in data:
        raw_ctx = data.get('context_length')
        if raw_ctx in (None, "", 0, "0"):
            ep.context_length = None
            updated = True
        else:
            try:
                ep.context_length = max(256, int(raw_ctx))
                updated = True
            except Exception:
                return jsonify({'ok': False, 'error': 'некорректный context_length'}), 400
    if 'instances' in data:
        try:
            ep.instances = max(1, int(data.get('instances') or 1))
            updated = True
        except Exception:
            return jsonify({'ok': False, 'error': 'некорректный instances'}), 400
    if updated:
        db.session.commit()
        _sync_runtime_default_route_from_llm_endpoints()
        _invalidate_llm_cache()
        _log_user_action(_load_current_user(), 'llm_update', 'llm_endpoint', ep.id)
    return jsonify({'ok': True, 'item': {
        'id': ep.id,
        'name': ep.name,
        'base_url': ep.base_url,
        'model': ep.model,
        'weight': float(ep.weight or 0.0),
        'purpose': ep.purpose,
        'purposes': _llm_parse_purposes(ep.purpose),
        'provider': ep.provider,
        'context_length': int(ep.context_length) if getattr(ep, 'context_length', None) else None,
        'instances': int(getattr(ep, 'instances', 1) or 1),
        'created_at': ep.created_at.isoformat() if ep.created_at else None,
    }})


@admin_bp.route('/lmstudio/models', methods=['GET'])
@require_admin
def api_admin_lmstudio_models():
    base_url = str(request.args.get('base_url') or '').strip()
    if not base_url:
        return jsonify({'ok': False, 'error': 'base_url is required'}), 400
    endpoint_id = request.args.get('endpoint_id')
    api_key = ''
    if endpoint_id:
        try:
            ep = LlmEndpoint.query.get(int(endpoint_id))
            if ep:
                api_key = ep.api_key or ''
        except Exception:
            pass
    if not api_key:
        api_key = str(request.args.get('api_key') or '').strip()
    base = _lmstudio_base_url(base_url)
    if not base:
        return jsonify({'ok': False, 'error': 'invalid base_url'}), 400
    response = http_request(
        'GET',
        f'{base}/api/v1/models',
        headers=_lmstudio_headers(api_key),
        timeout=(HTTP_CONNECT_TIMEOUT, max(10, HTTP_DEFAULT_TIMEOUT)),
        logger=app.logger,
    )
    if getattr(response, 'status_code', 500) >= 400:
        detail = ''
        try:
            detail = (response.text or '')[:300]
        except Exception:
            detail = ''
        return jsonify({'ok': False, 'error': f'LM Studio error {response.status_code}', 'detail': detail}), 502
    payload = response.json() if hasattr(response, 'json') else {}
    items = _lmstudio_extract_models(payload)
    # Compatibility fallback: some setups expose models on OpenAI-compatible /v1/models
    if not items:
        try:
            response_v1 = http_request(
                'GET',
                f'{base}/v1/models',
                headers=_lmstudio_headers(api_key),
                timeout=(HTTP_CONNECT_TIMEOUT, max(10, HTTP_DEFAULT_TIMEOUT)),
                logger=app.logger,
            )
            if getattr(response_v1, 'status_code', 500) < 400:
                payload_v1 = response_v1.json() if hasattr(response_v1, 'json') else {}
                items = _lmstudio_extract_models(payload_v1)
        except Exception:
            pass
    return jsonify({'ok': True, 'items': items})


@admin_bp.route('/lmstudio/models/load', methods=['POST'])
@require_admin
def api_admin_lmstudio_models_load():
    data = request.get_json(silent=True) or {}
    endpoint_id = data.get('endpoint_id')
    ep = None
    if endpoint_id is not None:
        ep = LlmEndpoint.query.get(int(endpoint_id))
    base_url = str(data.get('base_url') or (ep.base_url if ep else '') or '').strip()
    model = str(data.get('model') or (ep.model if ep else '') or '').strip()
    api_key = str(data.get('api_key') or (ep.api_key if ep else '') or '').strip()
    if not base_url or not model:
        return jsonify({'ok': False, 'error': 'base_url and model are required'}), 400
    context_length = data.get('context_length')
    payload = {'model': model}
    if context_length not in (None, '', 0, '0'):
        try:
            ctx = max(256, int(context_length))
            payload['context_length'] = ctx
        except Exception:
            return jsonify({'ok': False, 'error': 'invalid context_length'}), 400
    base = _lmstudio_base_url(base_url)
    response = http_request(
        'POST',
        f'{base}/api/v1/models/load',
        headers=_lmstudio_headers(api_key),
        json=payload,
        timeout=(HTTP_CONNECT_TIMEOUT, max(30, HTTP_DEFAULT_TIMEOUT)),
        logger=app.logger,
    )
    if getattr(response, 'status_code', 500) >= 400:
        detail = ''
        try:
            detail = (response.text or '')[:300]
        except Exception:
            detail = ''
        return jsonify({'ok': False, 'error': f'LM Studio load error {response.status_code}', 'detail': detail}), 502
    _lmstudio_mark_activity()
    return jsonify({'ok': True, 'result': response.json() if hasattr(response, 'json') else {}})


@admin_bp.route('/lmstudio/models/unload', methods=['POST'])
@require_admin
def api_admin_lmstudio_models_unload():
    data = request.get_json(silent=True) or {}
    endpoint_id = data.get('endpoint_id')
    ep = None
    if endpoint_id is not None:
        ep = LlmEndpoint.query.get(int(endpoint_id))
    base_url = str(data.get('base_url') or (ep.base_url if ep else '') or '').strip()
    model = str(data.get('model') or (ep.model if ep else '') or '').strip()
    api_key = str(data.get('api_key') or (ep.api_key if ep else '') or '').strip()
    if not base_url or not model:
        return jsonify({'ok': False, 'error': 'base_url and model are required'}), 400
    base = _lmstudio_base_url(base_url)
    response = http_request(
        'POST',
        f'{base}/api/v1/models/unload',
        headers=_lmstudio_headers(api_key),
        json={'model': model},
        timeout=(HTTP_CONNECT_TIMEOUT, max(30, HTTP_DEFAULT_TIMEOUT)),
        logger=app.logger,
    )
    if getattr(response, 'status_code', 500) >= 400:
        detail = ''
        try:
            detail = (response.text or '')[:300]
        except Exception:
            detail = ''
        return jsonify({'ok': False, 'error': f'LM Studio unload error {response.status_code}', 'detail': detail}), 502
    return jsonify({'ok': True, 'result': response.json() if hasattr(response, 'json') else {}})


# ------------------- CLI helpers -------------------

# ------------------- Diagnostics -------------------
@app.route('/diagnostics/transcribe')
def diag_transcribe():
    """Диагностика транскрибации одного файла.
    Параметры:
      - path: абсолютный или относительный к SCAN_ROOT путь
      - file_id: альтернатива path, взять файл из БД по id
      - limit: ограничение символов транскрипта (по умолчанию 1000)
    Возвращает JSON: состояние бэкенда/модели, сведения о файле и образец транскрипта.
    """
    res = {
        "transcribe_enabled": bool(TRANSCRIBE_ENABLED),
        "backend": (TRANSCRIBE_BACKEND or '').lower(),
        "language": TRANSCRIBE_LANGUAGE,
        "model_path": TRANSCRIBE_MODEL_PATH,
        "model_exists": False,
        "ffmpeg": _ffmpeg_available(),
        "ffprobe": shutil.which('ffprobe') is not None,
        "file": {},
        "transcript": {
            "ok": False,
            "len": 0,
            "sample": "",
        },
        "warnings": [],
        "faster_whisper": {},
        "applied": {},
    }
    try:
        mp = TRANSCRIBE_MODEL_PATH or ''
        try:
            res["model_exists"] = bool(mp) and Path(mp).exists()
        except Exception:
            res["model_exists"] = False

        # Детали разрешения модели faster-whisper (опционально)
        try:
            fw = {
                "hf_available": bool(hf_snapshot_download),
                "cache_dir": str(FW_CACHE_DIR),
                "model_ref": (request.args.get('model') or TRANSCRIBE_MODEL_PATH or '').strip(),
                "repo": None,
                "target_dir": None,
                "target_exists": None,
                "downloaded": False,
            }
            mref = fw["model_ref"]
            if mref:
                # Если локальный каталог существует, считаем путь найденным
                p = Path(mref).expanduser()
                if p.exists() and p.is_dir():
                    fw["target_dir"] = str(p)
                    fw["target_exists"] = True
                else:
                    repo = _fw_alias_to_repo(mref) or (mref if '/' in mref else None)
                    fw["repo"] = repo
                    if repo:
                        safe_name = repo.replace('/', '__')
                        target_dir = FW_CACHE_DIR / safe_name
                        fw["target_dir"] = str(target_dir)
                        fw["target_exists"] = target_dir.exists() and any(target_dir.iterdir())
                        # Необязательная загрузка, даже если бэкенд сейчас не faster-whisper
                        if (request.args.get('download') or '').lower() in ('1','true','yes','on'):
                            if hf_snapshot_download is None:
                                res["warnings"].append("Пакет huggingface_hub не установлен — автозагрузка невозможна")
                            else:
                                FW_CACHE_DIR.mkdir(parents=True, exist_ok=True)
                                hf_snapshot_download(repo_id=repo, local_dir=str(target_dir), local_dir_use_symlinks=False, revision=None)
                                fw["downloaded"] = True
                                fw["target_exists"] = target_dir.exists() and any(target_dir.iterdir())
                    else:
                        res["warnings"].append("Неизвестная ссылка на модель faster-whisper (не каталог, не алиас и не repo id)")
            res["faster_whisper"] = fw
        except Exception as e:
            res["warnings"].append(f"fw_resolve:{e}")

        # Переопределения (backend/lang/vad/model) из query-параметров
        eff_backend = (request.args.get('backend') or '').strip().lower() or None
        eff_lang = (request.args.get('lang') or '').strip() or None
        vad_param = (request.args.get('vad') or '').strip().lower()
        if vad_param in ('1','true','yes','on'): eff_vad = True
        elif vad_param in ('0','false','no','off'): eff_vad = False
        else: eff_vad = None
        model_override = (request.args.get('model') or '').strip() or None
        res["applied"] = {"backend": eff_backend, "lang": eff_lang, "vad": eff_vad, "model": model_override}

        # Нормализация backend: поддерживаем только faster-whisper
        if eff_backend and eff_backend != 'faster-whisper':
            res["warnings"].append("Бэкенд не поддерживается (Vosk удалён); используются настройки faster-whisper")
            eff_backend = None

        # Разрешение пути к файлу
        p = None
        q_path = (request.args.get('path') or '').strip()
        q_id = (request.args.get('file_id') or '').strip()
        if q_path:
            pp = Path(q_path)
            if not pp.is_absolute():
                root = _scan_root_path()
                # Пробуем типовые варианты, чтобы избежать дублирования корневой папки в пути
                candidates = [
                    root / q_path,
                    root.parent / q_path,
                ]
                picked = None
                for cand in candidates:
                    try:
                        if cand.exists():
                            picked = cand
                            break
                    except Exception:
                        pass
                pp = picked or (root / q_path)
            p = pp
        elif q_id:
            try:
                f = File.query.get(int(q_id))
                if f: p = Path(f.path)
            except Exception:
                p = None
        # Разрешить диагностику только модели (без файла)
        if not p:
            return jsonify(res)

        # метаданные файла
        file_info = {
            "requested": q_path or q_id,
            "path": str(p) if p else None,
            "exists": bool(p and p.exists()),
            "ext": (p.suffix.lower() if p else None),
            "size": (p.stat().st_size if p and p.exists() else None),
        }
        if p and p.exists():
            try:
                file_info["duration_seconds"] = _ffprobe_duration_seconds(p)
                file_info["duration_str"] = audio_duration_hhmmss(p)
            except Exception:
                file_info["duration_seconds"] = None
                file_info["duration_str"] = None
        res["file"] = file_info

        # Выполнить транскрибацию при наличии файла и разрешённых настройках (для аудио)
        if not TRANSCRIBE_ENABLED and eff_backend is None:
            res["warnings"].append("TRANSCRIBE_ENABLED is off")
            return jsonify(res)
        if not (p and p.exists() and p.is_file()):
            return jsonify(res)

        limit = 0
        try:
            limit = int(request.args.get('limit', '1000'))
        except Exception:
            limit = 1000
        limit = max(200, min(20000, limit or 1000))

        try:
            if p.suffix.lower() in AUDIO_EXTS:
                tx = transcribe_audio(p, limit_chars=limit,
                                      backend_override=eff_backend,
                                      model_path_override=model_override,
                                      lang_override=eff_lang,
                                      vad_override=eff_vad)
                sample = (tx or '')[:200]
                res["transcript"].update({"ok": bool(tx), "len": len(tx or ''), "sample": sample})
            elif p.suffix.lower() in IMAGE_EXTS and IMAGES_VISION_ENABLED:
                vis = call_lmstudio_vision(p, p.name)
                desc = (vis.get('description') or '') if isinstance(vis, dict) else ''
                sample = desc[:200]
                res["transcript"].update({"ok": bool(sample), "len": len(desc), "sample": sample})
            else:
                res["warnings"].append("Файл не аудио/изображение или распознавание изображений выключено")
        except Exception as e:
            res["warnings"].append(f"transcribe_error:{e}")
        return jsonify(res)
    except Exception as e:
        res["warnings"].append(f"unexpected:{e}")
    return jsonify(res), 500


def _format_bytes_short(num: float | int | None) -> str | None:
    if num is None:
        return None
    try:
        value = float(num)
    except (TypeError, ValueError):
        return None
    if value < 0:
        value = 0.0
    units = ["B", "KB", "MB", "GB", "TB", "PB"]
    idx = 0
    while value >= 1024 and idx < len(units) - 1:
        value /= 1024.0
        idx += 1
    if idx == 0:
        return f"{int(value)} {units[idx]}"
    return f"{value:.1f} {units[idx]}"


def _format_duration_brief(seconds: float | int | None) -> str | None:
    if seconds is None:
        return None
    try:
        total = int(max(0, float(seconds)))
    except (TypeError, ValueError):
        return None
    days, rem = divmod(total, 86400)
    hours, rem = divmod(rem, 3600)
    minutes, secs = divmod(rem, 60)
    if days:
        return f"{days}d {hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"


def _isoformat_dt(value: datetime | None) -> str | None:
    if value is None:
        return None
    try:
        trimmed = value.replace(microsecond=0)
        if trimmed.tzinfo is None:
            return trimmed.isoformat() + "Z"
        return trimmed.astimezone(timezone.utc).isoformat().replace("+00:00", "Z")
    except Exception:
        try:
            return value.isoformat()
        except Exception:
            return str(value)


def _isoformat_ts(value: float | int | None) -> str | None:
    if value is None:
        return None
    try:
        ts = float(value)
    except (TypeError, ValueError):
        return None
    try:
        return datetime.fromtimestamp(ts, tz=timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")
    except Exception:
        return None


_OSINT_FINAL_PROGRESS_STATUSES = {'completed', 'error', 'blocked', 'cancelled'}


def _osint_job_summary(job: Mapping[str, Any] | None) -> dict[str, Any]:
    """Собирает краткую сводку по задаче OSINT для отображения в админке."""
    if not isinstance(job, Mapping):
        return {}
    progress_raw = job.get('progress')
    progress: dict[str, Any] = {}
    if isinstance(progress_raw, Mapping):
        for key, value in progress_raw.items():
            entry_key = str(key)
            if isinstance(value, Mapping):
                progress[entry_key] = dict(value)
            else:
                progress[entry_key] = {'status': value}
    sources_total_raw = job.get('sources_total')
    sources_total: int | None
    try:
        sources_total = int(sources_total_raw)
        if sources_total <= 0:
            sources_total = None
    except Exception:
        sources_total = None
    if sources_total is None:
        sources_total = len(progress) or len(job.get('source_specs') or [])
    sources_completed_raw = job.get('sources_completed')
    try:
        sources_completed = int(sources_completed_raw)
        if sources_completed < 0:
            sources_completed = 0
    except Exception:
        sources_completed = sum(
            1
            for value in progress.values()
            if str(value.get('status') or '').lower() in _OSINT_FINAL_PROGRESS_STATUSES
        )
    percent_complete: float | None = None
    if sources_total:
        try:
            percent_complete = round(
                min(100.0, max(0.0, (float(sources_completed) / float(sources_total)) * 100.0)),
                1,
            )
        except Exception:
            percent_complete = None
    schedule_payload = job.get('schedule')
    schedule = dict(schedule_payload) if isinstance(schedule_payload, Mapping) else schedule_payload
    return {
        'id': job.get('id'),
        'query': job.get('query'),
        'status': job.get('status'),
        'error': job.get('error'),
        'created_at': job.get('created_at'),
        'started_at': job.get('started_at'),
        'completed_at': job.get('completed_at'),
        'sources_total': sources_total,
        'sources_completed': sources_completed,
        'percent_complete': percent_complete,
        'progress': progress,
        'schedule': schedule,
    }


def _rag_job_summary(task: TaskRecord | None) -> dict[str, Any]:
    if task is None:
        return {}
    try:
        payload_raw = json.loads(task.payload or "{}")
    except Exception:
        payload_raw = {}
    options = payload_raw.get('options') if isinstance(payload_raw, Mapping) else None
    progress_value = 0.0
    try:
        progress_value = max(0.0, min(1.0, float(task.progress or 0.0)))
    except Exception:
        progress_value = 0.0
    return {
        'id': task.id,
        'status': task.status,
        'progress': progress_value,
        'created_at': _isoformat_dt(task.created_at),
        'started_at': _isoformat_dt(task.started_at),
        'finished_at': _isoformat_dt(task.finished_at),
        'error': task.error,
        'collection_id': payload_raw.get('collection_id'),
        'collection_name': payload_raw.get('collection_name'),
        'total_files': payload_raw.get('total_files'),
        'options': options if isinstance(options, Mapping) else None,
    }


@admin_bp.route('/status/overview', methods=['GET'])
@require_admin
def api_admin_status_overview():
    now = datetime.utcnow().replace(microsecond=0)
    queue_stats = get_task_queue().stats()
    llm_pool_stats = get_llm_pool().stats()
    warnings: list[str] = []
    errors: dict[str, str] = {}

    uptime_seconds = max(0.0, time.time() - APP_START_TIME)
    app_info = {
        'created_at': _isoformat_dt(now),
        'started_at': _isoformat_dt(APP_STARTED_AT),
        'uptime_seconds': uptime_seconds,
        'uptime_human': _format_duration_brief(uptime_seconds),
        'environment': current_app.config.get('ENV'),
        'debug': bool(current_app.debug),
        'version': current_app.config.get('APP_VERSION'),
    }

    system_info = {
        'host': platform.node(),
        'platform': platform.platform(),
        'python_version': platform.python_version(),
        'pid': os.getpid(),
        'timezone': time.tzname,
    }
    try:
        load_avg = os.getloadavg()
        system_info['load_average'] = [round(load_avg[0], 2), round(load_avg[1], 2), round(load_avg[2], 2)]
    except (AttributeError, OSError):
        system_info['load_average'] = None

    queue_info = {
        'name': queue_stats.get('name'),
        'workers': queue_stats.get('workers'),
        'max_workers': queue_stats.get('max_workers'),
        'queued': queue_stats.get('queued'),
        'started': bool(queue_stats.get('started')),
        'shutdown': bool(queue_stats.get('shutdown')),
    }
    queue_status = 'ok'
    if not queue_info['started']:
        queue_status = 'warning'
        warnings.append('Очередь фоновых задач не запущена')
    elif queue_info['shutdown']:
        queue_status = 'warning'
        warnings.append('Очередь фоновых задач помечена как остановленная')
    queue_info['status'] = queue_status
    llm_queue_info = {
        'workers': llm_pool_stats.get('workers'),
        'per_user_limit': llm_pool_stats.get('per_user_limit'),
        'max_queue': llm_pool_stats.get('max_queue'),
        'queued': llm_pool_stats.get('queued'),
        'running': llm_pool_stats.get('running'),
        'avg_queue_wait_sec': llm_pool_stats.get('avg_queue_wait_sec'),
        'p95_queue_wait_sec': llm_pool_stats.get('p95_queue_wait_sec'),
    }
    if int(llm_queue_info.get('queued') or 0) >= int(llm_queue_info.get('max_queue') or 0):
        warnings.append('LLM очередь переполнена')

    # Database health
    db_status = 'ok'
    try:
        db.session.execute(text('SELECT 1'))
        db.session.commit()
    except Exception as exc:
        db_status = 'error'
        errors['database'] = str(exc)
        warnings.append('Ошибка проверки соединения с базой данных')
        db.session.rollback()

    counts: dict[str, Optional[int]] = {
        'files': None,
        'collections': None,
        'users': None,
        'tags': None,
        'rag_documents': None,
    }
    user_roles: dict[str, int] = {}
    rag_status_counts: dict[str, int] = {}
    rag_last_indexed = None

    if db_status == 'ok':
        try:
            counts['files'] = int(db.session.query(func.count(File.id)).scalar() or 0)
        except Exception as exc:
            errors['files_count'] = str(exc)
            warnings.append('Не удалось получить количество файлов')
            db.session.rollback()
        try:
            counts['collections'] = int(db.session.query(func.count(Collection.id)).scalar() or 0)
        except Exception as exc:
            errors['collections_count'] = str(exc)
            warnings.append('Не удалось получить количество коллекций')
            db.session.rollback()
        try:
            counts['tags'] = int(db.session.query(func.count(Tag.id)).scalar() or 0)
        except Exception as exc:
            errors['tags_count'] = str(exc)
            warnings.append('Не удалось получить количество тегов')
            db.session.rollback()
        try:
            counts['users'] = int(db.session.query(func.count(User.id)).scalar() or 0)
        except Exception as exc:
            errors['users_count'] = str(exc)
            warnings.append('Не удалось получить количество пользователей')
            db.session.rollback()
        try:
            counts['rag_documents'] = int(db.session.query(func.count(RagDocument.id)).scalar() or 0)
        except Exception as exc:
            errors['rag_documents_count'] = str(exc)
            warnings.append('Не удалось получить количество RAG-документов')
            db.session.rollback()

        try:
            role_rows = db.session.query(User.role, func.count(User.id)).group_by(User.role).all()
            user_roles = {str(role or 'unknown'): int(count or 0) for role, count in role_rows}
        except Exception as exc:
            errors['user_roles'] = str(exc)
            warnings.append('Не удалось получить распределение ролей пользователей')
            db.session.rollback()

        try:
            rag_rows = (
                db.session.query(RagDocument.import_status, func.count(RagDocument.id))
                .group_by(RagDocument.import_status)
                .all()
            )
            rag_status_counts = {str(status or 'unknown'): int(count or 0) for status, count in rag_rows}
            rag_last_indexed = db.session.query(func.max(RagDocument.last_indexed_at)).scalar()
        except Exception as exc:
            errors['rag_status'] = str(exc)
            warnings.append('Не удалось получить статус RAG-индексации')
            db.session.rollback()

    db_dialect = db.engine.dialect.name if getattr(db, "engine", None) is not None else None
    db_uri = str(current_app.config.get("SQLALCHEMY_DATABASE_URI") or "")
    db_uri_safe = db_uri
    if db_uri:
        try:
            db_uri_safe = make_url(db_uri).render_as_string(hide_password=True)
        except Exception:
            db_uri_safe = db_uri
    db_path = getattr(CONFIG, 'catalogue_db_path', None) if db_dialect == 'sqlite' else None
    db_size_bytes: Optional[int] = None
    if isinstance(db_path, Path):
        try:
            if db_path.exists():
                db_size_bytes = db_path.stat().st_size
        except Exception as exc:
            errors['database_size'] = str(exc)
            warnings.append('Не удалось определить размер базы данных')

    database_info = {
        'status': db_status,
        'dialect': db_dialect,
        'uri': db_uri_safe,
        'path': str(db_path) if db_path else None,
        'size_bytes': db_size_bytes,
        'size_pretty': _format_bytes_short(db_size_bytes),
        'counts': counts,
        'user_roles': user_roles,
        'rag_status': {
            'counts': rag_status_counts,
            'last_indexed_at': _isoformat_dt(rag_last_indexed),
        },
    }

    # Tasks summary
    task_counts: dict[str, int] = {}
    active_oldest: dict[str, Any] | None = None
    recent_tasks: list[dict[str, Any]] = []
    if db_status == 'ok':
        try:
            rows = db.session.query(TaskRecord.status, func.count(TaskRecord.id)).group_by(TaskRecord.status).all()
            task_counts = {str(status or 'unknown'): int(count or 0) for status, count in rows}
        except Exception as exc:
            errors['task_counts'] = str(exc)
            warnings.append('Не удалось получить распределение задач')
            db.session.rollback()
        try:
            active = (
                TaskRecord.query.filter(TaskRecord.status.notin_(TASK_FINAL_STATUSES))
                .order_by(TaskRecord.created_at.asc())
                .first()
            )
            if active:
                active_oldest = _task_to_dict(active)
        except Exception as exc:
            errors['task_active'] = str(exc)
            warnings.append('Не удалось получить активные задачи')
            db.session.rollback()
        try:
            rows = TaskRecord.query.order_by(TaskRecord.created_at.desc()).limit(5).all()
            recent_tasks = [_task_to_dict(task) for task in rows]
        except Exception as exc:
            errors['task_recent'] = str(exc)
            warnings.append('Не удалось получить последние задачи')
            db.session.rollback()

    tasks_info = {
        'counts': task_counts,
        'oldest_active': active_oldest,
        'recent': recent_tasks,
    }

    # OSINT job snapshots
    osint_summary: dict[str, Any] | None = None
    osint_service: OsintSearchService | None = None
    try:
        osint_service = _get_osint_service()
    except Exception as exc:  # noqa: BLE001
        osint_service = None
        warnings.append('Сервис OSINT-поиска недоступен')
        errors['osint_service'] = str(exc)
    if osint_service is not None:
        try:
            recent_jobs_raw = osint_service.list_jobs(limit=6) or []
            recent_jobs = [_osint_job_summary(job) for job in recent_jobs_raw if job]
            active_jobs = [
                job
                for job in recent_jobs
                if str(job.get('status') or '').lower() not in _OSINT_FINAL_PROGRESS_STATUSES
            ]
            queue_snapshot = getattr(osint_service, 'queue', None)
            queue_meta = queue_snapshot.stats() if queue_snapshot else None
            osint_summary = {
                'queue': queue_meta,
                'active': active_jobs,
                'recent': recent_jobs,
            }
        except Exception as exc:  # noqa: BLE001
            warnings.append('Не удалось получить состояние OSINT-поиска')
            errors['osint_status'] = str(exc)

    # External integrations snapshot
    browser_info = None
    try:
        browser_manager = get_browser_manager()
        browser_settings = browser_manager.settings
        browser_info = {
            'headless': browser_settings.headless,
            'proxy': browser_settings.proxy,
            'viewport': list(browser_settings.default_viewport),
            'context_timeout_ms': browser_settings.context_timeout_ms,
            'navigation_timeout_ms': browser_settings.navigation_timeout_ms,
        }
    except Exception as exc:
        errors['browser_settings'] = str(exc)
        warnings.append('Не удалось получить параметры браузера для OSINT')
    http_conf = get_http_settings()
    http_info = {
        'timeout': getattr(http_conf, 'timeout', None),
        'connect_timeout': getattr(http_conf, 'connect_timeout', None),
        'retries': getattr(http_conf, 'retries', None),
        'backoff_factor': getattr(http_conf, 'backoff_factor', None),
    }
    osint_integration = None
    if osint_service is not None:
        serp_settings = getattr(osint_service.fetcher, 'settings', None)
        if serp_settings is not None:
            osint_integration = {
                'cache_enabled': bool(getattr(serp_settings, 'cache_enabled', False)),
                'cache_ttl_seconds': getattr(serp_settings, 'cache_ttl_seconds', None),
                'retry_user_agents': len(getattr(serp_settings, 'retry_user_agents', ()) or ()),
                'retry_proxies': len(getattr(serp_settings, 'retry_proxies', ()) or ()),
                'wait_after_load_ms': getattr(serp_settings, 'wait_after_load_ms', None),
                'navigation_timeout_ms': getattr(serp_settings, 'navigation_timeout_ms', None),
                'user_agent_override': getattr(serp_settings, 'user_agent_override', None),
            }
            if osint_summary and osint_summary.get('queue'):
                osint_integration['queue'] = osint_summary['queue']
    integrations_summary = {
        'browser': browser_info,
        'http': http_info,
        'osint': osint_integration,
    }

    # User activity snapshot
    users_activity = None
    if db_status == 'ok':
        try:
            total_users = User.query.count()
            active_cutoff = datetime.utcnow() - timedelta(hours=24)
            active_24h = (
                db.session.query(func.count(func.distinct(UserActionLog.user_id)))
                .filter(UserActionLog.action == 'login', UserActionLog.user_id.isnot(None))
                .filter(UserActionLog.created_at >= active_cutoff)
                .scalar() or 0
            )
            recent_logins = (
                UserActionLog.query.filter(UserActionLog.action == 'login')
                .order_by(UserActionLog.created_at.desc())
                .limit(8)
                .all()
            )
            users_activity = {
                'total': total_users,
                'active_24h': int(active_24h),
                'recent': [
                    {
                        'user_id': entry.user_id,
                        'username': getattr(entry.user, 'username', None),
                        'full_name': getattr(entry.user, 'full_name', None),
                        'at': _isoformat_dt(entry.created_at),
                    }
                    for entry in recent_logins
                ],
            }
        except Exception as exc:
            errors['user_activity'] = str(exc)
            warnings.append('Не удалось получить статистику активности пользователей')
            db.session.rollback()

    # AI metrics snapshot
    ai_metrics_info: dict[str, Any] = {
        'window_size': 0,
        'latency_avg_ms': None,
        'last_measurement': None,
    }
    ai_metric_rows: list[AiSearchMetric] = []
    if db_status == 'ok':
        try:
            ai_metric_rows = (
                AiSearchMetric.query.order_by(AiSearchMetric.created_at.desc()).limit(25).all()
            )
            ai_metrics_info['window_size'] = len(ai_metric_rows)
            if ai_metric_rows:
                totals = [row.total_ms for row in ai_metric_rows if row.total_ms is not None]
                if totals:
                    ai_metrics_info['latency_avg_ms'] = round(sum(totals) / len(totals), 1)
                last_row = ai_metric_rows[0]
                meta_payload: Any = last_row.meta
                if isinstance(meta_payload, str):
                    try:
                        meta_payload = json.loads(meta_payload)
                    except Exception:
                        meta_payload = {'raw': meta_payload}
                ai_metrics_info['last_measurement'] = {
                    'id': last_row.id,
                    'created_at': _isoformat_dt(last_row.created_at),
                    'total_ms': last_row.total_ms,
                    'keywords_ms': last_row.keywords_ms,
                    'candidate_ms': last_row.candidate_ms,
                    'deep_ms': last_row.deep_ms,
                    'llm_answer_ms': last_row.llm_answer_ms,
                    'llm_snippet_ms': last_row.llm_snippet_ms,
                    'meta': meta_payload,
                }
        except Exception as exc:
            errors['ai_metrics'] = str(exc)
            warnings.append('Не удалось получить окна AI-метрик')
            db.session.rollback()

    def _avg_metric(field: str) -> float | None:
        values = [getattr(row, field) for row in ai_metric_rows if getattr(row, field) is not None]
        if not values:
            return None
        try:
            return round(sum(values) / len(values), 1)
        except Exception:
            return None

    # Scan snapshot
    scan_scope = None
    scope_payload = SCAN_PROGRESS.get('scope')
    if isinstance(scope_payload, dict):
        scan_scope = {
            'type': scope_payload.get('type'),
            'label': scope_payload.get('label'),
            'count': scope_payload.get('count'),
        }
    processed = SCAN_PROGRESS.get('processed') or 0
    total = SCAN_PROGRESS.get('total') or 0
    percent = None
    try:
        total_val = float(total)
        processed_val = float(processed)
        if total_val > 0:
            percent = round(min(100.0, max(0.0, (processed_val / total_val) * 100)), 1)
    except Exception:
        percent = None
    scan_info = {
        'running': bool(SCAN_PROGRESS.get('running')),
        'stage': SCAN_PROGRESS.get('stage'),
        'updated_at': _isoformat_ts(SCAN_PROGRESS.get('updated_at')),
        'processed': processed,
        'total': total,
        'added': SCAN_PROGRESS.get('added'),
        'updated': SCAN_PROGRESS.get('updated'),
        'duplicates': SCAN_PROGRESS.get('duplicates'),
        'removed': SCAN_PROGRESS.get('removed'),
        'current': SCAN_PROGRESS.get('current'),
        'eta_seconds': SCAN_PROGRESS.get('eta_seconds'),
        'eta_human': _format_duration_brief(SCAN_PROGRESS.get('eta_seconds')),
        'percent': percent,
        'scope': scan_scope,
        'task_id': SCAN_PROGRESS.get('task_id'),
        'last_log': None,
        'error': SCAN_PROGRESS.get('error'),
    }
    history = SCAN_PROGRESS.get('history') or []
    if isinstance(history, list) and history:
        last_entry = history[-1]
        if isinstance(last_entry, dict):
            scan_info['last_log'] = {
                'at': _isoformat_ts(last_entry.get('t')),
                'level': last_entry.get('level'),
                'message': last_entry.get('msg'),
            }
    if scan_info['error']:
        warnings.append('Обнаружена ошибка в последнем сканировании')

    # Runtime snapshot
    rt = runtime_settings_store.current
    provider_default = (getattr(rt, 'lm_default_provider', None) or 'openai').strip() or 'openai'
    llm_endpoints: list[dict[str, Any]] = []
    try:
        rows = (
            LlmEndpoint.query
            .order_by(LlmEndpoint.created_at.desc())
            .limit(20)
            .all()
        )
        for ep in rows:
            endpoint_provider = (ep.provider or provider_default) or 'openai'
            llm_endpoints.append({
                'id': ep.id,
                'name': ep.name,
                'model': ep.model,
                'provider': endpoint_provider,
                'weight': float(ep.weight or 0.0),
                'purposes': _llm_parse_purposes(ep.purpose),
                'tokenizer_hint': _guess_tokenizer_for_model(ep.model, endpoint_provider),
            })
    except Exception as exc:
        errors['llm_endpoints'] = str(exc)
        warnings.append('Не удалось получить список LLM-эндпоинтов')
        db.session.rollback()

    runtime_info = {
        'scan_root': str(rt.scan_root),
        'import_subdir': rt.import_subdir,
        'default_use_llm': bool(rt.default_use_llm),
        'default_prune': bool(rt.default_prune),
        'keywords_to_tags_enabled': bool(rt.keywords_to_tags_enabled),
        'transcribe_enabled': bool(rt.transcribe_enabled),
        'images_vision_enabled': bool(rt.images_vision_enabled),
        'ai_query_variants_max': int(rt.ai_query_variants_max),
        'lm_default_provider': provider_default,
        'lm_model': getattr(rt, 'lmstudio_model', None),
        'lm_tokenizer': _guess_tokenizer_for_model(getattr(rt, 'lmstudio_model', None), provider_default),
        'lm_max_input_chars': _lm_max_input_chars(),
        'lm_max_output_tokens': _lm_max_output_tokens(),
        'rag_embedding_backend': getattr(rt, 'rag_embedding_backend', None),
        'rag_embedding_model': getattr(rt, 'rag_embedding_model', None),
        'rag_embedding_dim': getattr(rt, 'rag_embedding_dim', None),
        'rag_rerank_backend': getattr(rt, 'rag_rerank_backend', None),
        'rag_rerank_model': getattr(rt, 'rag_rerank_model', None),
        'rag_rerank_batch_size': getattr(rt, 'rag_rerank_batch_size', None),
        'rag_rerank_max_length': getattr(rt, 'rag_rerank_max_length', None),
        'rag_rerank_max_chars': getattr(rt, 'rag_rerank_max_chars', None),
        'ai_rerank_llm': bool(getattr(rt, 'ai_rerank_llm', False)),
        'llm_endpoints': llm_endpoints,
    }
    try:
        usage = shutil.disk_usage(rt.scan_root)
        runtime_info['disk_usage'] = {
            'total_bytes': usage.total,
            'used_bytes': usage.used,
            'free_bytes': usage.free,
            'total_pretty': _format_bytes_short(usage.total),
            'used_pretty': _format_bytes_short(usage.used),
            'free_pretty': _format_bytes_short(usage.free),
        }
    except Exception as exc:
        errors['disk_usage'] = str(exc)

    service_flags = {
        'cache_cleanup': {
            'enabled': CACHE_CLEANUP_INTERVAL_HOURS > 0,
            'running': _CLEANUP_THREAD_STARTED,
            'interval_hours': CACHE_CLEANUP_INTERVAL_HOURS,
        },
    }
    feedback_stats = _feedback_scheduler_status()
    service_flags['feedback_trainer'] = {
        'enabled': bool(feedback_stats.get('scheduler_enabled')),
        'running': bool(feedback_stats.get('thread_started')),
        'interval_hours': feedback_stats.get('interval_hours'),
        'last_trigger_at': feedback_stats.get('last_trigger_at'),
        'next_run_in_seconds': feedback_stats.get('next_run_in_seconds'),
        'total_weighted': feedback_stats.get('total_weighted'),
    }
    if service_flags['cache_cleanup']['enabled'] and not service_flags['cache_cleanup']['running']:
        warnings.append('Поток очистки кэша не запущен')
    if service_flags['feedback_trainer']['enabled'] and not service_flags['feedback_trainer']['running']:
        warnings.append('Планировщик обучения AI-фидбэка не запущен')

    cache_info = {
        'llm': llm_cache_stats(),
        'search': search_cache_stats(),
    }
    if not cache_info['llm'].get('enabled'):
        warnings.append('Кэш LLM выключен – время ответов может увеличиться')
    if not cache_info['search'].get('enabled'):
        warnings.append('Кэш поиска выключен – увеличится нагрузка на БД')

    llm_latency_summary = {
        'total_avg_ms': _avg_metric('total_ms'),
        'answer_avg_ms': _avg_metric('llm_answer_ms'),
        'snippet_avg_ms': _avg_metric('llm_snippet_ms'),
    }
    provider_counts: dict[str, int] = {}
    for endpoint in llm_endpoints:
        provider = str(endpoint.get('provider') or 'unknown').strip().lower()
        provider_counts[provider] = provider_counts.get(provider, 0) + 1
    llm_summary = {
        'cache': cache_info.get('llm'),
        'endpoints': {
            'total': len(llm_endpoints),
            'per_provider': provider_counts,
        },
        'latency_ms': llm_latency_summary,
    }

    rag_ready = None
    rag_pending = None
    rag_jobs: list[dict[str, Any]] = []
    if db_status == 'ok':
        try:
            rag_ready = int(
                db.session.query(func.count(RagDocument.id))
                .filter(RagDocument.is_ready_for_rag.is_(True))
                .scalar() or 0
            )
            rag_pending = int(
                db.session.query(func.count(RagDocument.id))
                .filter(RagDocument.import_status != 'ready')
                .scalar() or 0
            )
        except Exception as exc:
            errors['rag_overview'] = str(exc)
            db.session.rollback()
        try:
            rag_task_rows = (
                TaskRecord.query.filter(TaskRecord.name == 'rag_collection')
                .filter(TaskRecord.status.notin_(TASK_FINAL_STATUSES))
                .order_by(TaskRecord.created_at.asc())
                .limit(5)
                .all()
            )
            rag_jobs = [_rag_job_summary(row) for row in rag_task_rows if row]
        except Exception as exc:
            errors['rag_jobs'] = str(exc)
            warnings.append('Не удалось получить активные задачи построения RAG')
            db.session.rollback()
    rag_overview = {
        'ready': rag_ready,
        'pending': rag_pending,
        'jobs': rag_jobs,
    }

    status = 'ok'
    if db_status == 'error':
        status = 'error'
    elif warnings:
        status = 'degraded'

    payload = {
        'ok': status != 'error',
        'status': status,
        'generated_at': _isoformat_dt(now),
        'warnings': warnings,
        'errors': errors,
        'app': app_info,
        'system': system_info,
        'queue': queue_info,
        'llm_queue': llm_queue_info,
        'database': database_info,
        'tasks': tasks_info,
        'osint': osint_summary,
        'ai_search': ai_metrics_info,
        'llm': llm_summary,
        'scan': scan_info,
        'runtime': runtime_info,
        'services': service_flags,
        'feedback': feedback_stats,
        'cache': cache_info,
        'rag': rag_overview,
        'users': users_activity,
        'integrations': integrations_summary,
    }
    return jsonify(payload)


@admin_bp.route('/ai-search/metrics', methods=['GET', 'DELETE'])
@require_admin
def api_admin_ai_metrics():
    if request.method == 'DELETE':
        try:
            deleted = AiSearchMetric.query.delete(synchronize_session=False)
            db.session.commit()
            return jsonify({'ok': True, 'deleted': deleted or 0})
        except Exception as exc:
            db.session.rollback()
            app.logger.warning(f"[ai-metrics] cleanup failed: {exc}")
            return jsonify({'ok': False, 'error': 'Не удалось очистить метрики'}), 500

    try:
        limit = int(request.args.get('limit', '100'))
    except Exception:
        limit = 100
    limit = max(1, min(limit, 500))
    rows = AiSearchMetric.query.order_by(AiSearchMetric.created_at.desc()).limit(limit).all()
    payload = []
    for row in rows:
        user_obj = row.user
        payload.append({
            'id': row.id,
            'query_hash': row.query_hash,
            'user_id': row.user_id,
             'user_username': getattr(user_obj, 'username', None),
             'user_full_name': getattr(user_obj, 'full_name', None),
            'total_ms': row.total_ms,
            'keywords_ms': row.keywords_ms,
            'candidate_ms': row.candidate_ms,
            'deep_ms': row.deep_ms,
            'llm_answer_ms': row.llm_answer_ms,
            'llm_snippet_ms': row.llm_snippet_ms,
            'created_at': row.created_at.isoformat() if row.created_at else None,
            'meta': row.meta,
        })
    summary = {}
    if rows:
        fields = ['total_ms', 'keywords_ms', 'candidate_ms', 'deep_ms', 'llm_answer_ms', 'llm_snippet_ms']
        for field in fields:
            vals = [getattr(r, field) for r in rows if getattr(r, field) is not None]
            if vals:
                summary[field] = sum(vals) / len(vals)
    return jsonify({'ok': True, 'items': payload, 'summary': summary})


@admin_bp.route('/ai-search/feedback/train', methods=['POST'])
@require_admin
def api_admin_feedback_train():
    user = _load_current_user()
    payload = request.get_json(silent=True) or {}
    cutoff_days_raw = payload.get('cutoff_days')
    cutoff_days: Optional[int]
    if cutoff_days_raw is None:
        cutoff_days = FEEDBACK_TRAIN_CUTOFF_DAYS
    else:
        try:
            cutoff_days = int(cutoff_days_raw)
        except Exception:
            return jsonify({'ok': False, 'error': 'invalid_cutoff'}), 400
    allow_duplicate = bool(payload.get('allow_duplicate'))
    task_id, queued = _enqueue_feedback_training(
        cutoff_days=cutoff_days,
        trigger='manual',
        submitted_by=getattr(user, 'id', None),
        allow_duplicate=allow_duplicate,
    )
    return jsonify({'ok': True, 'async': True, 'task_id': task_id, 'queued': queued})


@admin_bp.route('/ai-search/feedback/model', methods=['GET'])
@require_admin
def api_admin_feedback_model():
    try:
        limit = int(request.args.get('limit', '50'))
    except Exception:
        limit = 50
    limit = max(1, min(limit, 200))
    try:
        feedback_lookup = _get_feedback_weights()
    except Exception as exc:
        app.logger.debug("feedback model inspect failed: %s", exc)
        feedback_lookup = {}
    if not feedback_lookup:
        return jsonify({'ok': True, 'total': 0, 'positive': [], 'negative': []})

    positive_items = [(fid, info) for fid, info in feedback_lookup.items() if float(info.get('weight', 0.0) or 0.0) > 0]
    negative_items = [(fid, info) for fid, info in feedback_lookup.items() if float(info.get('weight', 0.0) or 0.0) < 0]
    pos_sorted = sorted(positive_items, key=lambda item: float(item[1].get('weight', 0.0) or 0.0), reverse=True)[:limit]
    neg_sorted = sorted(negative_items, key=lambda item: float(item[1].get('weight', 0.0) or 0.0))[:limit]
    needed_ids = {fid for fid, _ in pos_sorted + neg_sorted}
    file_meta_map: Dict[int, File] = {}
    if needed_ids:
        try:
            rows = (
                File.query.outerjoin(Collection, File.collection_id == Collection.id)
                .filter(File.id.in_(needed_ids))
                .all()
            )
            file_meta_map = {int(row.id): row for row in rows if getattr(row, 'id', None)}
        except Exception as exc:
            app.logger.debug("feedback model meta failed: %s", exc)
            file_meta_map = {}

    def _serialize(fid: int, payload: Dict[str, Any]) -> Dict[str, Any]:
        row = file_meta_map.get(fid)
        title = None
        author_name = None
        collection_id = None
        collection_name = None
        year_val = None
        if row is not None:
            title = row.title or row.filename or row.rel_path or f"Документ {fid}"
            author_name = row.author
            collection_id = row.collection_id
            year_val = row.year
            try:
                if getattr(row, "collection", None):
                    collection_name = row.collection.name
            except Exception:
                collection_name = None
        else:
            title = f"Документ {fid}"
        return {
            'file_id': fid,
            'title': title,
            'author': author_name,
            'year': year_val,
            'collection_id': collection_id,
            'collection_name': collection_name,
            'weight': round(float(payload.get('weight', 0.0) or 0.0), 6),
            'positive': int(payload.get('positive', 0) or 0),
            'negative': int(payload.get('negative', 0) or 0),
            'clicks': int(payload.get('clicks', 0) or 0),
            'updated_at': payload.get('updated_at'),
        }

    return jsonify({
        'ok': True,
        'total': len(feedback_lookup),
        'positive': [_serialize(fid, payload) for fid, payload in pos_sorted],
        'negative': [_serialize(fid, payload) for fid, payload in neg_sorted],
    })


def _feedback_scheduler_status() -> dict[str, Any]:
    scheduler_enabled = FEEDBACK_TRAIN_INTERVAL_HOURS > 0
    try:
        total_weighted = db.session.query(func.count(AiSearchFeedbackModel.file_id)).scalar() or 0
    except Exception:
        total_weighted = 0
    try:
        last_updated = db.session.query(func.max(AiSearchFeedbackModel.updated_at)).scalar()
    except Exception:
        last_updated = None
    active_task = TaskRecord.query.filter(
        TaskRecord.name == 'feedback_train',
        TaskRecord.status.notin_(TASK_FINAL_STATUSES),
    ).order_by(TaskRecord.created_at.desc()).first()
    last_task = TaskRecord.query.filter(TaskRecord.name == 'feedback_train').order_by(TaskRecord.created_at.desc()).first()
    now_ts = time.time()
    try:
        interval_hours = max(0.0, float(FEEDBACK_TRAIN_INTERVAL_HOURS or 0.0))
    except Exception:
        interval_hours = 0.0
    next_run_seconds: Optional[float] = None
    if scheduler_enabled and interval_hours > 0.0:
        interval_seconds = max(60.0, interval_hours * 3600.0)
        if FEEDBACK_LAST_TRIGGER_AT > 0:
            elapsed = max(0.0, now_ts - FEEDBACK_LAST_TRIGGER_AT)
            next_run_seconds = max(0.0, interval_seconds - elapsed)
        else:
            next_run_seconds = interval_seconds
    return {
        'ok': True,
        'scheduler_enabled': scheduler_enabled,
        'interval_hours': interval_hours,
        'cutoff_days': FEEDBACK_TRAIN_CUTOFF_DAYS,
        'thread_started': _FEEDBACK_THREAD_STARTED,
        'total_weighted': int(total_weighted),
        'last_updated_at': last_updated.isoformat() if last_updated else None,
        'active_task': _task_to_dict(active_task) if active_task else None,
        'last_task': _task_to_dict(last_task) if last_task else None,
        'last_trigger_at': _isoformat_ts(FEEDBACK_LAST_TRIGGER_AT),
        'next_run_in_seconds': next_run_seconds,
    }


@admin_bp.route('/ai-search/feedback/status', methods=['GET'])
@require_admin
def api_admin_feedback_status():
    return jsonify(_feedback_scheduler_status())


@admin_bp.route('/ai-search/feedback/scheduler', methods=['POST'])
@require_admin
def api_admin_feedback_scheduler():
    payload = request.get_json(silent=True) or {}
    updates: Dict[str, Any] = {}

    interval_value: Optional[float] = None
    if 'interval_hours' in payload:
        try:
            interval_value = max(0.0, float(payload.get('interval_hours') or 0.0))
        except Exception:
            return jsonify({'ok': False, 'error': 'invalid_interval'}), 400

    cutoff_value: Optional[int] = None
    if 'cutoff_days' in payload:
        try:
            cutoff_value = max(1, int(payload.get('cutoff_days') or 1))
        except Exception:
            return jsonify({'ok': False, 'error': 'invalid_cutoff'}), 400

    enabled_flag = payload.get('enabled')
    if enabled_flag is not None:
        enabled_bool = bool(enabled_flag)
        if not enabled_bool:
            updates['AI_FEEDBACK_TRAIN_INTERVAL_HOURS'] = 0.0
        else:
            target_interval = interval_value
            if target_interval is None:
                current_interval = getattr(runtime_settings_store.current, 'feedback_train_interval_hours', 0.0)
                if current_interval <= 0:
                    current_interval = 6.0
                target_interval = current_interval
            updates['AI_FEEDBACK_TRAIN_INTERVAL_HOURS'] = max(0.25, float(target_interval))
    elif interval_value is not None:
        updates['AI_FEEDBACK_TRAIN_INTERVAL_HOURS'] = interval_value

    if cutoff_value is not None:
        updates['AI_FEEDBACK_TRAIN_CUTOFF_DAYS'] = cutoff_value

    updates_applied = False
    if updates:
        runtime_settings_store.apply_updates(updates)
        runtime_settings_store.current.apply_env()
        runtime_settings_store.current.apply_to_flask_config(current_app)
        _refresh_runtime_globals()
        updates_applied = True

    FEEDBACK_SCHEDULER_EVENT.set()

    run_now = bool(payload.get('run_now'))
    manual_task = None
    if run_now:
        current_user = _load_current_user()
        try:
            task_id, queued = _enqueue_feedback_training(
                cutoff_days=FEEDBACK_TRAIN_CUTOFF_DAYS,
                trigger='manual_toggle',
                submitted_by=getattr(current_user, 'id', None),
                allow_duplicate=False,
            )
            if queued:
                global FEEDBACK_LAST_TRIGGER_AT
                FEEDBACK_LAST_TRIGGER_AT = time.time()
            manual_task = {'task_id': task_id, 'queued': queued}
        except Exception as exc:
            return jsonify({'ok': False, 'error': f'queue_failed:{exc}'}), 500

    status = _feedback_scheduler_status()
    status['updated'] = updates_applied
    if manual_task is not None:
        status['manual_task'] = manual_task
    return jsonify(status)


@admin_bp.route('/cache/llm', methods=['DELETE'])
@require_admin
def api_admin_cache_llm_clear():
    before = llm_cache_stats()
    llm_cache_clear()
    after = llm_cache_stats()
    return jsonify({'ok': True, 'before': before, 'after': after})


@admin_bp.route('/cache/search', methods=['DELETE'])
@require_admin
def api_admin_cache_search_clear():
    before = search_cache_stats()
    search_cache_clear()
    after = search_cache_stats()
    return jsonify({'ok': True, 'before': before, 'after': after})

# ------------------- Statistics & Visualization -------------------
from collections import Counter, defaultdict

@app.route("/api/stats")
def api_stats():
    # Агрегация по авторам, годам, типам материалов
    allowed_scope = _current_allowed_collections()
    base_query = File.query.join(Collection, File.collection_id == Collection.id).filter(Collection.searchable == True)
    if allowed_scope is not None:
        if not allowed_scope:
            base_query = base_query.filter(File.collection_id == -1)
        else:
            base_query = base_query.filter(File.collection_id.in_(allowed_scope))
    try:
        files = base_query.all()
    except Exception:
        fallback = File.query
        if allowed_scope is not None:
            if not allowed_scope:
                fallback = fallback.filter(File.collection_id == -1)
            else:
                fallback = fallback.filter(File.collection_id.in_(allowed_scope))
        files = fallback.all()
    authors = Counter()
    years = Counter()
    types = Counter()
    exts = Counter()
    size_buckets = Counter()
    months = Counter()
    # новые распределения
    weekdays = Counter()  # 0..6 (Mon..Sun)
    hours = Counter()     # 0..23
    # дополнительные агрегации
    kw = Counter()
    tag_keys = Counter()
    tag_values = Counter()
    total_files = 0
    total_size = 0
    # средний размер по типам
    size_sum_by_type = Counter()
    size_cnt_by_type = Counter()
    # заполненность ключевых полей
    meta_presence = Counter()
    meta_missing = Counter()
    # дополнительные показатели
    tags_total = 0
    tagless_files = 0
    now = datetime.utcnow()
    recent_7_cut = now - timedelta(days=7)
    recent_30_cut = now - timedelta(days=30)
    recent_counts = {'7d': 0, '30d': 0}
    def bucket_size(sz):
        if sz is None or sz <= 0:
            return "неизв."
        mb = sz / (1024*1024)
        if mb < 1: return "< 1 МБ"
        if mb < 10: return "1–10 МБ"
        if mb < 50: return "10–50 МБ"
        if mb < 100: return "50–100 МБ"
        return "> 100 МБ"
    for f in files:
        total_files += 1
        try:
            total_size += int(f.size or 0)
        except Exception:
            pass
        if f.author:
            authors[f.author] += 1
        if f.year:
            years[f.year] += 1
        if f.material_type:
            types[f.material_type] += 1
        if f.ext:
            exts[f.ext.lower().lstrip('.')] += 1
        size_buckets[bucket_size(f.size or 0)] += 1
        if f.mtime:
            try:
                d = datetime.fromtimestamp(f.mtime)
                months[d.strftime('%Y-%m')] += 1
                weekdays[d.weekday()] += 1  # Понедельник=0
                hours[d.hour] += 1
                if d >= recent_30_cut:
                    recent_counts['30d'] += 1
                if d >= recent_7_cut:
                    recent_counts['7d'] += 1
            except Exception:
                pass
        # ключевые слова
        if f.keywords:
            for part in re.split(r"[\n,;]+", f.keywords):
                w = (part or '').strip()
                if w:
                    kw[w] += 1
        # ключи тегов
        tag_count = 0
        try:
            for t in f.tags:
                tag_count += 1
                if t.key:
                    tag_keys[t.key] += 1
                if t.value:
                    tag_values[str(t.value).strip()] += 1
        except Exception:
            tag_count = 0
        tags_total += tag_count
        # средний размер по типам
        if f.material_type and (f.size or 0) > 0:
            size_sum_by_type[f.material_type] += int(f.size or 0)
            size_cnt_by_type[f.material_type] += 1
        # заполненность полей
        if f.title:
            meta_presence['Название'] += 1
        else:
            meta_missing['Название'] += 1
        if f.author:
            meta_presence['Автор'] += 1
        else:
            meta_missing['Автор'] += 1
        if f.year:
            meta_presence['Год'] += 1
        else:
            meta_missing['Год'] += 1
        if f.keywords:
            meta_presence['Ключевые слова'] += 1
        else:
            meta_missing['Ключевые слова'] += 1
        if tag_count > 0:
            meta_presence['Теги'] += 1
        else:
            meta_missing['Теги'] += 1
            tagless_files += 1
    # подготовка выходных структур
    # средний размер по типам (в МБ, округляем до десятых)
    avg_size_type = []
    for mt in size_sum_by_type.keys():
        cnt = max(1, size_cnt_by_type[mt])
        avg_mb = (size_sum_by_type[mt] / cnt) / (1024*1024)
        avg_size_type.append((mt, round(avg_mb, 1)))
    avg_size_type.sort(key=lambda x: x[1], reverse=True)
    # недели: упорядочим Пн..Вс
    weekday_names = ['Пн','Вт','Ср','Чт','Пт','Сб','Вс']
    weekdays_list = []
    for i in range(7):
        weekdays_list.append((weekday_names[i], int(weekdays.get(i, 0))))
    # часы 0..23
    hours_list = [(str(h), int(hours.get(h, 0))) for h in range(24)]
    # коллекции: количество файлов по коллекциям (searchable)
    collections_counts = []
    try:
        q = db.session.query(Collection.name, func.count(File.id)) \
            .join(File, File.collection_id == Collection.id) \
            .filter(Collection.searchable == True) \
            .group_by(Collection.id) \
            .order_by(func.count(File.id).desc())
        if allowed_scope is not None:
            if not allowed_scope:
                q = q.filter(File.collection_id == -1)
            else:
                q = q.filter(File.collection_id.in_(allowed_scope))
        collections_counts = [(name, int(cnt)) for (name, cnt) in q.all()]
    except Exception:
        pass
    # общий размер по коллекциям (только searchable)
    collections_total_size = []
    try:
        rows = db.session.query(Collection.name, func.coalesce(func.sum(File.size), 0)) \
            .join(File, File.collection_id == Collection.id) \
            .filter(Collection.searchable == True) \
            .group_by(Collection.id) \
            .order_by(func.coalesce(func.sum(File.size), 0).desc())
        if allowed_scope is not None:
            if not allowed_scope:
                rows = rows.filter(File.collection_id == -1)
            else:
                rows = rows.filter(File.collection_id.in_(allowed_scope))
        rows = rows.all()
        collections_total_size = [(name, int(total or 0)) for (name, total) in rows]
    except Exception:
        pass

    # самые крупные файлы (только searchable)
    largest_files = []
    try:
        rows = db.session.query(File.filename, File.size) \
            .join(Collection, File.collection_id == Collection.id) \
            .filter(Collection.searchable == True) \
            .order_by(File.size.desc().nullslast()) \
            .limit(15)
        if allowed_scope is not None:
            if not allowed_scope:
                rows = rows.filter(File.collection_id == -1)
            else:
                rows = rows.filter(File.collection_id.in_(allowed_scope))
        rows = rows.all()
        largest_files = [(fn or '(без имени)', int(sz or 0)) for (fn, sz) in rows]
    except Exception:
        pass

    tags_summary = {
        'avg_per_file': round(tags_total / total_files, 2) if total_files else 0.0,
        'with_tags': max(total_files - tagless_files, 0),
        'without_tags': tagless_files,
        'total_tags': tags_total,
    }

    authority_docs: List[Dict[str, Any]] = []
    authority_authors: List[Dict[str, Any]] = []
    authority_topics: List[Dict[str, Any]] = []
    authority_collections: List[Dict[str, Any]] = []
    try:
        snapshot = _compute_authority_snapshot(allowed_scope)
    except Exception as exc:
        app.logger.debug("authority stats computation failed: %s", exc)
        snapshot = None
    if snapshot:
        doc_scores = snapshot.get("doc_scores", {}) or {}
        if doc_scores:
            top_docs = heapq.nlargest(10, doc_scores.items(), key=lambda item: item[1])
            doc_ids = [int(fid) for fid, _ in top_docs]
            if doc_ids:
                doc_query = File.query.outerjoin(Collection, File.collection_id == Collection.id)
                doc_query = _apply_file_access_filter(doc_query)
                doc_rows = doc_query.filter(File.id.in_(doc_ids)).all()
                doc_map = {int(row.id): row for row in doc_rows if getattr(row, "id", None)}
                for fid, score in top_docs:
                    row = doc_map.get(int(fid))
                    if not row:
                        continue
                    try:
                        collection_name = row.collection.name if getattr(row, "collection", None) else None
                    except Exception:
                        collection_name = None
                    authority_docs.append({
                        "id": int(fid),
                        "score": round(float(score), 6),
                        "title": row.title or row.filename or row.rel_path or f"Документ {fid}",
                        "author": row.author,
                        "year": row.year,
                        "collection_id": row.collection_id,
                        "collection_name": collection_name,
                    })
        authority_authors = list(snapshot.get("author_entries", [])[:10])
        authority_topics = list(snapshot.get("topic_entries", [])[:12])
        authority_collections = list(snapshot.get("collection_entries", [])[:10])

    feedback_positive: List[Dict[str, Any]] = []
    feedback_negative: List[Dict[str, Any]] = []
    feedback_summary: Optional[Dict[str, Any]] = None
    try:
        feedback_lookup = _get_feedback_weights()
    except Exception as exc:
        app.logger.debug("feedback stats lookup failed: %s", exc)
        feedback_lookup = {}
    if feedback_lookup:
        positive_items = [(fid, info) for fid, info in feedback_lookup.items() if float(info.get('weight', 0.0) or 0.0) > 0]
        negative_items = [(fid, info) for fid, info in feedback_lookup.items() if float(info.get('weight', 0.0) or 0.0) < 0]
        pos_sorted = sorted(positive_items, key=lambda item: float(item[1].get('weight', 0.0) or 0.0), reverse=True)[:10]
        neg_sorted = sorted(negative_items, key=lambda item: float(item[1].get('weight', 0.0) or 0.0))[:10]
        needed_ids = {fid for fid, _ in pos_sorted + neg_sorted}
        file_meta_map: Dict[int, File] = {}
        if needed_ids:
            try:
                rows = (
                    File.query.outerjoin(Collection, File.collection_id == Collection.id)
                    .filter(File.id.in_(needed_ids))
                    .all()
                )
                file_meta_map = {int(row.id): row for row in rows if getattr(row, 'id', None)}
            except Exception as exc:
                app.logger.debug("feedback stats meta failed: %s", exc)
                file_meta_map = {}

        def _feedback_entry(fid: int, payload: Dict[str, Any]) -> Dict[str, Any]:
            row = file_meta_map.get(fid)
            title = None
            author_name = None
            collection_name = None
            collection_id = None
            year_val = None
            if row is not None:
                title = row.title or row.filename or row.rel_path or f"Документ {fid}"
                author_name = row.author
                year_val = row.year
                collection_id = row.collection_id
                try:
                    if getattr(row, "collection", None):
                        collection_name = row.collection.name
                except Exception:
                    collection_name = None
            else:
                title = f"Документ {fid}"
            return {
                "file_id": fid,
                "title": title,
                "author": author_name,
                "year": year_val,
                "collection_id": collection_id,
                "collection_name": collection_name,
                "weight": round(float(payload.get('weight', 0.0) or 0.0), 6),
                "positive": int(payload.get('positive', 0) or 0),
                "negative": int(payload.get('negative', 0) or 0),
                "clicks": int(payload.get('clicks', 0) or 0),
                "updated_at": payload.get('updated_at'),
            }

        feedback_positive = [_feedback_entry(fid, payload) for fid, payload in pos_sorted]
        feedback_negative = [_feedback_entry(fid, payload) for fid, payload in neg_sorted]
        pos_count = len([1 for _, info in feedback_lookup.items() if float(info.get('weight', 0.0) or 0.0) > 0])
        neg_count = len([1 for _, info in feedback_lookup.items() if float(info.get('weight', 0.0) or 0.0) < 0])
        feedback_summary = {
            "total_files": len(feedback_lookup),
            "positive": pos_count,
            "negative": neg_count,
        }

    return jsonify({
        "authors": authors.most_common(20),
        "authors_cloud": authors.most_common(100),
        "years": sorted(years.items()),
        "types": types.most_common(),
        "exts": exts.most_common(),
        "sizes": sorted(size_buckets.items(), key=lambda x: ["неизв.","< 1 МБ","1–10 МБ","10–50 МБ","50–100 МБ","> 100 МБ"].index(x[0]) if x[0] in ["неизв.","< 1 МБ","1–10 МБ","10–50 МБ","50–100 МБ","> 100 МБ"] else 999),
        "months": sorted(months.items()),
        "top_keywords": kw.most_common(30),
        "tag_keys": tag_keys.most_common(30),
        "tag_values_cloud": tag_values.most_common(150),
        "weekdays": weekdays_list,
        "hours": hours_list,
        "avg_size_type": avg_size_type,
        "meta_presence": sorted(meta_presence.items(), key=lambda x: x[0]),
        "meta_missing": sorted(meta_missing.items(), key=lambda x: x[0]),
        "total_files": total_files,
        "total_size_bytes": total_size,
        "collections_counts": collections_counts,
        "collections_total_size": collections_total_size,
        "largest_files": largest_files,
        "recent_counts": {
            "7d": int(recent_counts['7d']),
            "30d": int(recent_counts['30d']),
        },
        "tags_summary": tags_summary,
        "authority_docs": authority_docs,
        "authority_authors": authority_authors,
        "authority_topics": authority_topics,
        "authority_collections": authority_collections,
        "feedback_positive": feedback_positive,
        "feedback_negative": feedback_negative,
        "feedback_summary": feedback_summary,
    })

@app.route('/api/authority')
def api_authority_summary():
    user = _load_current_user()
    if not user:
        return jsonify({'ok': False, 'error': 'Не авторизовано'}), 401
    allowed_scope = _current_allowed_collections()
    try:
        snapshot = _compute_authority_snapshot(allowed_scope)
    except Exception as exc:
        app.logger.debug("authority summary failed: %s", exc)
        snapshot = None
    if not snapshot:
        return jsonify({
            "generated_at": None,
            "documents": [],
            "authors": [],
            "topics": [],
            "collections": [],
        })
    doc_scores = snapshot.get("doc_scores", {}) or {}
    top_docs = heapq.nlargest(20, doc_scores.items(), key=lambda item: item[1])
    doc_ids = [int(fid) for fid, _ in top_docs]
    doc_entries: List[Dict[str, Any]] = []
    if doc_ids:
        doc_query = File.query.outerjoin(Collection, File.collection_id == Collection.id)
        doc_query = _apply_file_access_filter(doc_query)
        rows = doc_query.filter(File.id.in_(doc_ids)).all()
        row_map = {int(row.id): row for row in rows if getattr(row, "id", None)}
        for fid, score in top_docs:
            row = row_map.get(int(fid))
            if not row:
                continue
            try:
                collection_name = row.collection.name if getattr(row, "collection", None) else None
            except Exception:
                collection_name = None
            doc_entries.append({
                "id": int(fid),
                "score": round(float(score), 6),
                "title": row.title or row.filename or row.rel_path or f"Документ {fid}",
                "author": row.author,
                "year": row.year,
                "collection_id": row.collection_id,
                "collection_name": collection_name,
            })
    generated_at = snapshot.get("generated_at")
    try:
        generated_iso = datetime.utcfromtimestamp(float(generated_at)).isoformat() if generated_at else None
    except Exception:
        generated_iso = None
    return jsonify({
        "generated_at": generated_iso,
        "documents": doc_entries,
        "authors": snapshot.get("author_entries", [])[:20],
        "topics": snapshot.get("topic_entries", [])[:40],
        "collections": snapshot.get("collection_entries", [])[:20],
    })

@app.route('/api/stats/tag-values')
def api_stats_tag_values():
    """Возвращает топ значений для заданного ключа тега.
    Query: key, limit=30
    Учитывает только коллекции с searchable=true.
    """
    key = (request.args.get('key') or '').strip()
    if not key:
        return jsonify({"items": []})
    allowed_scope = _current_allowed_collections()
    try:
        try:
            lim = min(max(int(request.args.get('limit', '30')), 1), 200)
        except Exception:
            lim = 30
        rows = db.session.query(Tag.value, func.count(Tag.id)) \
            .join(File, File.id == Tag.file_id) \
            .join(Collection, File.collection_id == Collection.id) \
            .filter(Collection.searchable == True) \
            .filter(Tag.key == key) \
            .group_by(Tag.value) \
            .order_by(func.count(Tag.id).desc()) \
            .limit(lim)
        if allowed_scope is not None:
            if not allowed_scope:
                rows = rows.filter(File.collection_id == -1)
            else:
                rows = rows.filter(File.collection_id.in_(allowed_scope))
        rows = rows.all()
        return jsonify({"items": [[v or '', int(c or 0)] for (v, c) in rows]})
    except Exception as e:
        return jsonify({"items": [], "error": str(e)}), 500

@app.route("/stats")
def stats_redirect():
    return redirect('/app/stats')


@app.route('/graph')
def graph_redirect():
    return redirect('/app/graph')


def _read_text_file_with_fallback(path: Path) -> str:
    for encoding in ("utf-8", "cp1251", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
        except Exception:
            break
    return path.read_text(encoding="utf-8", errors="ignore")


_CONTROL_CHAR_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")


def _sanitize_plain_text(text: str) -> str:
    if not text:
        return ""
    cleaned = _CONTROL_CHAR_RE.sub(" ", text)
    cleaned = "".join(ch if ch.isprintable() or ch in "\n\r\t" else " " for ch in cleaned)
    cleaned = cleaned.replace("\ufffd", " ")
    return cleaned


def _looks_like_binary_text(text: str) -> bool:
    if not text:
        return False
    sample = text[:4000]
    nulls = sample.count("\x00")
    if nulls > 0:
        return True
    control = sum(1 for ch in sample if ord(ch) < 32 and ch not in "\n\r\t")
    if control / max(1, len(sample)) > 0.1:
        return True
    return False


def _limit_text_length(text: str, limit_chars: int) -> str:
    """Trim text only when a positive limit is provided."""
    if not text:
        return ""
    if limit_chars and limit_chars > 0:
        return text[:limit_chars]
    return text


def _extract_text_for_rag(path: Path, *, limit_chars: int = 120_000) -> str:
    ext = path.suffix.lower().lstrip(".")
    if ext == "pdf":
        return extract_text_pdf(path, limit_chars=limit_chars, force_ocr_first_page=False)
    if ext == "docx":
        return extract_text_docx(path, limit_chars=limit_chars)
    if ext == "rtf":
        return extract_text_rtf(path, limit_chars=limit_chars)
    if ext == "epub":
        return extract_text_epub(path, limit_chars=limit_chars)
    if ext == "djvu":
        return extract_text_djvu(path, limit_chars=limit_chars)
    if ext in AUDIO_EXTS:
        return transcribe_audio(path, limit_chars=limit_chars)
    if ext in {"txt", "md", "markdown"}:
        return _read_text_file_with_fallback(path)
    try:
        return _read_text_file_with_fallback(path)
    except Exception:
        return ""


def _resolve_candidate_paths(file_obj: File) -> List[Path]:
    candidates: List[Path] = []
    if file_obj.path:
        candidates.append(Path(file_obj.path))
    if file_obj.rel_path:
        rel = Path(file_obj.rel_path)
        candidates.append(rel)
        try:
            scan_root = runtime_settings_store.current.scan_root
            candidates.append(scan_root / rel)
        except Exception:
            pass
        candidates.append(Path("library") / rel)
    seen: set[str] = set()
    resolved: List[Path] = []
    for candidate in candidates:
        try:
            normalized = candidate.expanduser()
        except Exception:
            normalized = candidate
        key = str(normalized)
        if key in seen:
            continue
        seen.add(key)
        resolved.append(normalized)
    return resolved


def _doc_chat_assess_answer_quality(
    answer: str,
    *,
    question_terms: Sequence[str],
    text_sources: Sequence[dict[str, Any]],
    image_sources: Sequence[dict[str, Any]],
) -> dict[str, Any]:
    normalized = (answer or "").strip()
    if not normalized:
        return {
            'status': 'error',
            'risk_score': 1.0,
            'risk_level': 'high',
            'messages': ['Ответ пустой или не был сгенерирован.'],
            'missing_citations': True,
            'unknown_citations': [],
            'citations': [],
        }
    text_count = len(text_sources)
    image_count = len(image_sources)
    citations: list[dict[str, Any]] = []
    unknown: list[dict[str, Any]] = []
    for match in DOC_CHAT_CITATION_RE.finditer(normalized):
        kind_raw, index_raw = match.groups()
        try:
            idx = int(index_raw)
        except Exception:
            continue
        kind_lower = kind_raw.lower()
        is_text = 'текст' in kind_lower or 'text' in kind_lower
        entry = {
            'type': 'text' if is_text else 'image',
            'index': idx,
            'label': f"{'Текст' if is_text else 'Изображение'} {idx}",
        }
        citations.append(entry)
        if (is_text and not (1 <= idx <= text_count)) or ((not is_text) and not (1 <= idx <= image_count)):
            unknown.append(entry)
    missing_citations = len(citations) == 0 and len(normalized) > 200
    digits = re.findall(r"\d{3,}", normalized)
    messages: list[str] = []
    risk = 0.0
    if missing_citations:
        messages.append("Ответ не содержит ссылок на «Текст N» или «Изображение N» — возможны галлюцинации.")
        risk += 0.55
    if unknown:
        labels = ", ".join(entry['label'] for entry in unknown)
        messages.append(f"Ответ ссылается на несуществующие фрагменты: {labels}.")
        risk += 0.2 * len(unknown)
    if digits and len(citations) == 0:
        messages.append("В ответе есть числовые данные без ссылок на источники.")
        risk += 0.15
    lower_answer = normalized.lower()
    informative_terms = [term for term in question_terms if len(term) >= 5][:4]
    if informative_terms:
        missing_terms = [term for term in informative_terms if term not in lower_answer]
        if len(missing_terms) == len(informative_terms):
            messages.append("Ключевые термины вопроса не упомянуты в ответе.")
            risk += 0.15
    risk = min(1.0, risk)
    if risk <= 0.2:
        status = 'ok'
        risk_level = 'low'
    elif risk <= 0.6:
        status = 'warning'
        risk_level = 'medium'
    else:
        status = 'error' if missing_citations or unknown else 'warning'
        risk_level = 'high'
    if not messages and status == 'ok':
        messages.append("Ссылки найдены, явные риски не обнаружены.")
    return {
        'status': status,
        'risk_score': round(risk, 3),
        'risk_level': risk_level,
        'messages': messages,
        'missing_citations': missing_citations,
        'unknown_citations': unknown,
        'citations': citations,
    }


def _collect_text_for_rag(file_obj: File, *, limit_chars: int = 120_000) -> tuple[str, Optional[Path]]:
    """Возвращает текст для RAG и путь к исходному файлу (если удалось определить)."""
    if not file_obj:
        return "", None
    last_error: Exception | None = None
    for candidate in _resolve_candidate_paths(file_obj):
        try:
            if not candidate.exists() or not candidate.is_file():
                continue
            raw_text = _extract_text_for_rag(candidate, limit_chars=limit_chars)
            if _looks_like_binary_text(raw_text or ""):
                continue
            text = _sanitize_plain_text(raw_text)
            if not text.strip():
                continue
            if text:
                return text, candidate
        except Exception as exc:
            last_error = exc
    fallback = (file_obj.text_excerpt or "").strip()
    if fallback:
        return fallback, None
    if last_error:
        raise RuntimeError(f"Не удалось извлечь текст: {last_error}") from last_error
    return "", None


def _embed_missing_chunks_for_document(
    document_id: int,
    backend: EmbeddingBackend,
    *,
    batch_size: int = 32,
    commit: bool = True,
) -> int:
    """Создаёт эмбеддинги для чанков документа, у которых они отсутствуют."""
    if not document_id:
        return 0
    model_name = getattr(backend, "model_name", "unknown")
    model_version = getattr(backend, "model_version", "unknown")
    join_condition = and_(
        RagChunkEmbedding.chunk_id == RagDocumentChunk.id,
        RagChunkEmbedding.model_name == model_name,
        RagChunkEmbedding.model_version == model_version,
    )
    pending_chunks: List[RagDocumentChunk] = (
        db.session.query(RagDocumentChunk)
        .outerjoin(RagChunkEmbedding, join_condition)
        .filter(RagDocumentChunk.document_id == document_id, RagChunkEmbedding.id.is_(None))
        .order_by(RagDocumentChunk.id.asc())
        .all()
    )
    if not pending_chunks:
        return 0
    total_created = 0
    batch_size = max(1, int(batch_size or 1))
    for start in range(0, len(pending_chunks), batch_size):
        batch = pending_chunks[start : start + batch_size]
        texts = [chunk.content or "" for chunk in batch]
        vectors = backend.embed_many(texts)
        if len(vectors) != len(batch):
            raise RuntimeError("Размера ответа backend эмбеддингов недостаточно для сохранения.")
        for chunk, vector in zip(batch, vectors):
            vector_bytes = vector_to_bytes(vector)
            checksum = hashlib.sha256(vector_bytes).hexdigest()
            embedding = RagChunkEmbedding(
                chunk_id=chunk.id,
                model_name=model_name,
                model_version=model_version,
                dim=len(vector),
                vector=vector_bytes,
                vector_checksum=checksum,
            )
            db.session.add(embedding)
            total_created += 1
        if commit:
            db.session.commit()
        else:
            db.session.flush()
    return total_created


def _run_rag_collection_job(task_id: int, collection_id: int, options: dict[str, object]) -> None:
    """Фоновая задача построения RAG-индекса для коллекции."""
    try:
        db.session.rollback()
    except Exception:
        db.session.remove()

    task = TaskRecord.query.get(task_id)
    if not task:
        return

    def _update_task(status: Optional[str] = None, progress: Optional[float] = None, error: Optional[str] = None, final: bool = False, payload_override: Optional[dict] = None) -> None:
        payload = payload_override or summary
        try:
            task_ref = TaskRecord.query.get(task_id)
            if not task_ref:
                return
            if status:
                task_ref.status = status
            if progress is not None:
                task_ref.progress = max(0.0, min(1.0, float(progress)))
            task_ref.payload = json.dumps(payload, ensure_ascii=False)
            if status == 'running' and task_ref.started_at is None:
                task_ref.started_at = datetime.utcnow()
            if status in ('completed', 'error', 'cancelled'):
                task_ref.finished_at = datetime.utcnow()
            if error:
                task_ref.error = error
            db.session.commit()
        except Exception:
            db.session.rollback()

    user_id = options.get("user_id")
    try:
        runtime = runtime_settings_store.current
    except Exception:
        runtime = None

    default_backend = (getattr(runtime, "rag_embedding_backend", "auto") or "auto").strip().lower()
    default_model = getattr(runtime, "rag_embedding_model", None) or "intfloat/multilingual-e5-large"
    default_dim = max(8, int(getattr(runtime, "rag_embedding_dim", 384) or 384))
    default_batch = max(1, int(getattr(runtime, "rag_embedding_batch_size", 32) or 32))
    default_device = getattr(runtime, "rag_embedding_device", None)
    default_endpoint = getattr(runtime, "rag_embedding_endpoint", None) or getattr(runtime, "lmstudio_api_base", "")
    default_api_key = getattr(runtime, "rag_embedding_api_key", None) or getattr(runtime, "lmstudio_api_key", "")

    def _log_result(action: str, detail_payload: dict[str, object]) -> None:
        try:
            actor = User.query.get(int(user_id)) if user_id else None
        except Exception:
            actor = None
        try:
            _log_user_action(
                actor,
                action,
                "collection",
                collection_id,
                detail=json.dumps(detail_payload, ensure_ascii=False),
            )
        except Exception:
            pass

    try:
        collection = Collection.query.get(collection_id)
        if not collection:
            raise RuntimeError("Коллекция не найдена.")
    except Exception as exc:
        summary = {
            "collection_id": collection_id,
            "collection_name": None,
            "total_files": 0,
            "ingested": 0,
            "skipped": 0,
            "embedded": 0,
            "failures": [
                {"reason": "collection_missing", "message": str(exc)},
            ],
            "status": "error",
        }
        _update_task(status='error', progress=0.0, error=str(exc), final=True, payload_override=summary)
        _log_result(
            "collection_rag_reindex_failed",
            {"reason": "collection_missing", "message": str(exc)},
        )
        return

    chunk_config = ChunkConfig(
        max_tokens=max(16, int(options.get("chunk_max_tokens", 700))),
        overlap=max(0, int(options.get("chunk_overlap", 120))),
        min_tokens=max(1, int(options.get("chunk_min_tokens", 80))),
    )
    skip_if_unchanged = bool(options.get("skip_if_unchanged", True))
    limit_chars = max(1000, int(options.get("limit_chars", 120_000)))
    embedding_backend_name = (
        str(options.get("embedding_backend") or default_backend or "auto").strip().lower() or default_backend or "auto"
    )
    embedding_model_name = str(options.get("embedding_model_name") or default_model).strip() or default_model
    embedding_dim = max(8, int(options.get("embedding_dim", default_dim)))
    embedding_batch = max(1, int(options.get("embedding_batch_size", default_batch)))
    embedding_normalize = bool(options.get("embedding_normalize", True))
    embedding_device = options.get("embedding_device")
    if not embedding_device and default_device:
        embedding_device = default_device
    embedding_endpoint = str(options.get("embedding_endpoint") or default_endpoint or "").strip()
    embedding_api_key = str(options.get("embedding_api_key") or default_api_key or "")

    try:
        files = (
            File.query.filter(File.collection_id == collection_id)
            .order_by(File.id.asc())
            .all()
        )
    except Exception as exc:
        summary = {
            "collection_id": collection_id,
            "collection_name": collection.name,
            "total_files": 0,
            "ingested": 0,
            "skipped": 0,
            "embedded": 0,
            "failures": [
                {"reason": "files_query_failed", "message": str(exc)},
            ],
            "status": "error",
        }
        _update_task(status='error', progress=0.0, error=str(exc), final=True, payload_override=summary)
        _log_result(
            "collection_rag_reindex_failed",
            {"reason": "files_query_failed", "message": str(exc)},
        )
        return

    indexer = RagIndexer(chunk_config=chunk_config, normalizer_version=str(options.get("normalizer_version") or "v1"))

    try:
        backend = load_embedding_backend(
            embedding_backend_name,
            model_name=embedding_model_name,
            dim=embedding_dim,
            normalize=embedding_normalize,
            batch_size=embedding_batch,
            device=embedding_device,
            base_url=embedding_endpoint or None,
            api_key=embedding_api_key or None,
        )
    except Exception as exc:
        summary = {
            "collection_id": collection_id,
            "collection_name": collection.name,
            "total_files": len(files),
            "ingested": 0,
            "skipped": 0,
            "embedded": 0,
            "failures": [
                {"reason": "backend_init_failed", "message": str(exc)},
            ],
            "status": "error",
        }
        _update_task(status='error', progress=0.0, error=str(exc), final=True, payload_override=summary)
        _log_result(
            "collection_rag_reindex_failed",
            {"reason": "backend_init_failed", "message": str(exc)},
        )
        return

    summary: dict[str, object] = {
        "collection_id": collection_id,
        "collection_name": collection.name,
        "total_files": len(files),
        "ingested": 0,
        "skipped": 0,
        "embedded": 0,
        "failures": [],
        "status": "running",
        "embedding_backend": {
            "name": getattr(backend, "name", embedding_backend_name),
            "model_name": getattr(backend, "model_name", embedding_model_name),
            "model_version": getattr(backend, "model_version", None),
            "dim": getattr(backend, "dim", embedding_dim),
            "endpoint": embedding_endpoint,
        },
        "options": {
            "skip_if_unchanged": skip_if_unchanged,
            "limit_chars": limit_chars,
            "chunk_max_tokens": chunk_config.max_tokens,
            "chunk_overlap": chunk_config.overlap,
            "chunk_min_tokens": chunk_config.min_tokens,
            "embedding_backend": embedding_backend_name,
            "embedding_model_name": embedding_model_name,
            "embedding_endpoint": embedding_endpoint,
        },
    }
    _update_task(status='running', progress=0.0)

    total_files = len(files) or 1
    failure_entries: list[dict[str, object]] = []

    for idx, file_obj in enumerate(files, start=1):
        ingest_result: dict[str, object] | None = None
        try:
            text, source_path = _collect_text_for_rag(file_obj, limit_chars=limit_chars)
            if not text.strip():
                raise RuntimeError("Пустой текст после извлечения.")
            metadata = {
                "source": "api.collection_rag_reindex",
                "source_path": str(source_path) if source_path else None,
                "file_sha1": file_obj.sha1,
                "collection_id": collection_id,
                "file_id": file_obj.id,
            }
            ingest_result = indexer.ingest_document(
                file_obj,
                text,
                metadata=metadata,
                skip_if_unchanged=skip_if_unchanged,
                commit=True,
            )
        except Exception as exc:
            db.session.rollback()
            failure_entries.append(
                {"file_id": file_obj.id, "reason": "ingest_failed", "message": str(exc)}
            )
            continue

        if ingest_result is None:
            continue

        if ingest_result.get("skipped"):
            summary["skipped"] = int(summary.get("skipped", 0)) + 1
        else:
            summary["ingested"] = int(summary.get("ingested", 0)) + 1

        document_id = int(ingest_result.get("document_id") or 0)
        if document_id:
            try:
                created = _embed_missing_chunks_for_document(
                    document_id,
                    backend,
                    batch_size=embedding_batch,
                    commit=True,
                )
                summary["embedded"] = int(summary.get("embedded", 0)) + created
            except Exception as exc:
                db.session.rollback()
                failure_entries.append(
                    {"file_id": file_obj.id, "reason": "embed_failed", "message": str(exc)}
                )

        progress = idx / total_files
        summary["failures"] = failure_entries[:20]
        _update_task(status='running', progress=progress)

    summary["status"] = "completed"
    try:
        backend.close()
    except Exception:
        pass

    summary["failures"] = failure_entries[:20]
    _update_task(status='completed', progress=1.0)

    _log_result(
        "collection_rag_reindex_done",
        {
            "ingested": summary.get("ingested"),
            "skipped": summary.get("skipped"),
            "embedded": summary.get("embedded"),
            "failures": len(failure_entries),
            "collection_name": collection.name,
            "backend": embedding_backend_name,
            "model": embedding_model_name,
        },
    )


def _run_rag_file_job(task_id: int, file_id: int, options: dict[str, object]) -> None:
    try:
        db.session.rollback()
    except Exception:
        db.session.remove()
    task = TaskRecord.query.get(task_id)
    if not task:
        return

    def _save(status: str, progress: float, payload: dict[str, Any], error: Optional[str] = None) -> None:
        try:
            row = TaskRecord.query.get(task_id)
            if not row:
                return
            row.status = status
            row.progress = max(0.0, min(1.0, float(progress)))
            row.payload = json.dumps(payload, ensure_ascii=False)
            if status == 'running' and row.started_at is None:
                row.started_at = datetime.utcnow()
            if status in TASK_FINAL_STATUSES:
                row.finished_at = datetime.utcnow()
            if error:
                row.error = error
            db.session.commit()
        except Exception:
            db.session.rollback()

    file_obj = File.query.get(file_id)
    if not file_obj:
        _save('error', 0.0, {"ok": False, "error": "file_not_found", "file_id": file_id}, "file_not_found")
        return

    runtime = runtime_settings_store.current
    backend_name = str(options.get("embedding_backend") or runtime.rag_embedding_backend or "auto").strip().lower() or "auto"
    model_name = str(options.get("embedding_model_name") or runtime.rag_embedding_model or "intfloat/multilingual-e5-large").strip()
    emb_dim = max(8, int(options.get("embedding_dim") or runtime.rag_embedding_dim or 384))
    emb_batch = max(1, int(options.get("embedding_batch_size") or runtime.rag_embedding_batch_size or 32))
    emb_device = options.get("embedding_device") or runtime.rag_embedding_device
    emb_endpoint = str(options.get("embedding_endpoint") or runtime.rag_embedding_endpoint or runtime.lmstudio_api_base or "").strip()
    emb_api_key = str(options.get("embedding_api_key") or runtime.rag_embedding_api_key or runtime.lmstudio_api_key or "")
    skip_if_unchanged = bool(options.get("skip_if_unchanged", True))
    limit_chars = max(1000, int(options.get("limit_chars") or 120_000))
    chunk_max_tokens = max(16, int(options.get("chunk_max_tokens") or 700))
    chunk_overlap = max(0, int(options.get("chunk_overlap") or 120))
    chunk_min_tokens = max(1, int(options.get("chunk_min_tokens") or 80))

    payload: dict[str, Any] = {
        "ok": False,
        "file_id": file_id,
        "status": "running",
    }
    _save('running', 0.05, payload)
    try:
        text, source_path = _collect_text_for_rag(file_obj, limit_chars=limit_chars)
        if not (text or "").strip():
            raise RuntimeError("Пустой текст после извлечения")
        indexer = RagIndexer(
            chunk_config=ChunkConfig(
                max_tokens=chunk_max_tokens,
                overlap=chunk_overlap,
                min_tokens=chunk_min_tokens,
            ),
            normalizer_version=str(options.get("normalizer_version") or "v1"),
        )
        metadata = {
            "source": "api.file_rag_reindex",
            "source_path": str(source_path) if source_path else None,
            "file_sha1": file_obj.sha1,
            "file_id": file_obj.id,
            "collection_id": file_obj.collection_id,
        }
        ingest_result = indexer.ingest_document(
            file_obj,
            text,
            metadata=metadata,
            skip_if_unchanged=skip_if_unchanged,
            commit=True,
        )
        _save('running', 0.65, {"ok": True, "stage": "ingested", "ingest": ingest_result, "file_id": file_id})
        backend = load_embedding_backend(
            backend_name,
            model_name=model_name,
            dim=emb_dim,
            batch_size=emb_batch,
            device=emb_device,
            base_url=emb_endpoint or None,
            api_key=emb_api_key or None,
        )
        try:
            embedded = _embed_missing_chunks_for_document(
                int(ingest_result.get("document_id") or 0),
                backend,
                batch_size=emb_batch,
                commit=True,
            )
        finally:
            try:
                backend.close()
            except Exception:
                pass
        final_payload = {
            "ok": True,
            "file_id": file_id,
            "status": "completed",
            "ingest": ingest_result,
            "embedded": int(embedded or 0),
            "backend": backend_name,
            "model": model_name,
        }
        _save('completed', 1.0, final_payload)
    except Exception as exc:
        db.session.rollback()
        _save('error', 1.0, {"ok": False, "file_id": file_id, "error": str(exc)}, str(exc))


def _run_metadata_normalization_job(task_id: int, options: dict[str, object]) -> None:
    try:
        db.session.rollback()
    except Exception:
        db.session.remove()
    task = TaskRecord.query.get(task_id)
    if not task:
        return

    def _save(status: str, progress: float, payload: dict[str, Any], error: Optional[str] = None) -> None:
        try:
            row = TaskRecord.query.get(task_id)
            if not row:
                return
            row.status = status
            row.progress = max(0.0, min(1.0, float(progress)))
            row.payload = json.dumps(payload, ensure_ascii=False)
            if status == 'running' and row.started_at is None:
                row.started_at = datetime.utcnow()
            if status in TASK_FINAL_STATUSES:
                row.finished_at = datetime.utcnow()
            if error:
                row.error = error
            db.session.commit()
        except Exception:
            db.session.rollback()

    only_low = bool(options.get("only_low", True))
    limit = max(1, int(options.get("limit") or 2000))
    collection_id = options.get("collection_id")
    query = File.query.order_by(File.id.asc())
    if collection_id is not None:
        try:
            query = query.filter(File.collection_id == int(collection_id))
        except Exception:
            pass
    files = query.limit(limit).all()
    total = len(files) or 1
    changed = 0
    low_before = 0
    low_after = 0
    _save('running', 0.02, {"status": "running", "total": len(files), "changed": 0})
    try:
        for idx, row in enumerate(files, start=1):
            quality_before = _metadata_quality(row)
            if quality_before.get("bucket") == "low":
                low_before += 1
            if only_low and quality_before.get("bucket") != "low":
                continue
            result = _apply_metadata_quality_rules(
                row,
                ext=row.ext or "",
                text_excerpt=row.text_excerpt or "",
                filename=row.filename or "",
            )
            if result.get("changed"):
                changed += 1
                try:
                    _sync_file_to_fts(row)
                except Exception:
                    pass
            quality_after = _metadata_quality(row)
            if quality_after.get("bucket") == "low":
                low_after += 1
            if idx == 1 or idx % 50 == 0:
                _save(
                    'running',
                    idx / total,
                    {
                        "status": "running",
                        "total": len(files),
                        "processed": idx,
                        "changed": changed,
                        "low_before": low_before,
                        "low_after": low_after,
                    },
                )
        db.session.commit()
        _invalidate_facets_cache('metadata_normalize')
        _save(
            'completed',
            1.0,
            {
                "status": "completed",
                "total": len(files),
                "changed": changed,
                "low_before": low_before,
                "low_after": low_after,
            },
        )
    except Exception as exc:
        db.session.rollback()
        _save(
            'error',
            1.0,
            {
                "status": "error",
                "total": len(files),
                "changed": changed,
                "low_before": low_before,
                "low_after": low_after,
                "error": str(exc),
            },
            str(exc),
        )


def _build_translation_hint(query_lang: Optional[str], section_lang: Optional[str]) -> str:
    q = (query_lang or "").lower()
    s = (section_lang or "").lower()
    if not q or not s or q == s or s == "mixed" or s == "unknown":
        return ""
    if q == "ru":
        return f"Фрагмент на {s}, ответ оставь на русском, цитаты — на оригинале."
    if s == "ru":
        return f"Fragment language ru; respond in {q or 'user language'}, keep quotes на русском."
    return f"Фрагмент на {s}, ответ оставь на {q or 'языке запроса'}, цитаты — на оригинале."


def _select_rag_embedding_variant(document_ids: Sequence[int]) -> Optional[tuple[str, Optional[str], int]]:
    if not document_ids:
        return None
    variant = (
        db.session.query(
            RagChunkEmbedding.model_name,
            RagChunkEmbedding.model_version,
            RagChunkEmbedding.dim,
        )
        .join(RagDocumentChunk, RagDocumentChunk.id == RagChunkEmbedding.chunk_id)
        .filter(RagDocumentChunk.document_id.in_(list(document_ids)))
        .order_by(RagChunkEmbedding.id.desc())
        .first()
    )
    if variant:
        return variant
    fallback_variant = (
        db.session.query(
            RagChunkEmbedding.model_name,
            RagChunkEmbedding.model_version,
            RagChunkEmbedding.dim,
        )
        .order_by(RagChunkEmbedding.id.desc())
        .first()
    )
    return fallback_variant


def _serialize_context_section(section: ContextSection) -> Dict[str, Any]:
    return {
        "doc_id": section.doc_id,
        "chunk_id": section.chunk_id,
        "title": section.title,
        "language": section.language,
        "section_path": section.extra.get("Раздел") or section.extra.get("section_path"),
        "translation_hint": section.translation_hint,
        "score_dense": section.score_dense,
        "score_sparse": section.score_sparse,
        "combined_score": section.combined_score,
        "reasoning_hint": section.reasoning_hint,
        "token_estimate": section.token_estimate,
        "preview": section.preview,
        "content": section.content,
        "url": section.url,
        "extra": section.extra,
    }


def _group_context_sections_by_document(sections: Sequence[ContextSection]) -> List[Dict[str, Any]]:
    groups: Dict[int, Dict[str, Any]] = {}
    for section in sections:
        doc_id = section.doc_id
        group = groups.get(doc_id)
        if group is None:
            group = {
                "doc_id": doc_id,
                "title": section.title,
                "url": section.url,
                "language": section.language,
                "combined_score": section.combined_score,
                "chunks": [],
            }
            groups[doc_id] = group
        chunk_payload = {
            "chunk_id": section.chunk_id,
            "section_path": section.extra.get("Раздел") or section.extra.get("section_path"),
            "preview": section.preview,
            "content": section.content,
            "score_dense": section.score_dense,
            "score_sparse": section.score_sparse,
            "combined_score": section.combined_score,
            "reasoning_hint": section.reasoning_hint,
            "translation_hint": section.translation_hint,
            "extra": copy.deepcopy(section.extra) if section.extra else {},
        }
        group["chunks"].append(chunk_payload)
        if section.combined_score is not None:
            group["combined_score"] = max(group.get("combined_score", 0.0), section.combined_score)
    for group in groups.values():
        group["chunks"].sort(key=lambda chunk: chunk.get("combined_score") or 0.0, reverse=True)
    return list(groups.values())


def _assess_rag_risk(
    answer_text: str,
    validation: "ValidationResult",
    sections: Sequence[ContextSection],
) -> Dict[str, Any]:
    risk_score = 0.0
    reasons: List[str] = []
    if not sections:
        risk_score += 0.5
        reasons.append("no_context")
    if validation.is_empty:
        risk_score += 0.5
        reasons.append("empty_answer")
    if validation.missing_citations:
        risk_score += 0.25
        reasons.append("missing_citations")
    if validation.unknown_citations:
        risk_score += 0.25
        reasons.append("unknown_citations")
    if validation.extra_citations:
        risk_score += 0.15
        reasons.append("extra_citations")
    if validation.hallucination_warning and "hallucination_warning" not in reasons:
        risk_score += 0.15
        reasons.append("hallucination_warning")
    citations_total = len(extract_citations(answer_text))
    if sections and citations_total < len(sections) // 2:
        risk_score += 0.1
        reasons.append("low_coverage")
    risk_score = min(1.0, risk_score)
    if risk_score >= 0.6:
        level = "high"
    elif risk_score >= 0.3:
        level = "medium"
    else:
        level = "low"
    flagged_refs = sorted({(int(doc), int(chunk)) for doc, chunk in validation.unknown_citations + validation.extra_citations})
    flagged = [
        {"doc_id": doc_id, "chunk_id": chunk_id}
        for doc_id, chunk_id in flagged_refs
    ]
    top_sections = [
        {
            "doc_id": sec.doc_id,
            "chunk_id": sec.chunk_id,
            "combined_score": round(sec.combined_score, 4),
            "reasoning_hint": sec.reasoning_hint,
        }
        for sec in sections[:5]
    ]
    return {
        "score": round(risk_score, 3),
        "level": level,
        "reasons": reasons,
        "flagged_refs": flagged,
        "hallucination_warning": bool(validation.hallucination_warning),
        "top_sections": top_sections,
    }


def _summarize_search_hits(hit_records: Sequence[Dict[str, Any]]) -> str:
    """Сводит типы совпадений классического поиска в компактную строку."""
    if not hit_records:
        return ""
    mapping = {
        "title": "название",
        "author": "автор",
        "keywords": "ключевые слова",
        "excerpt": "выдержка",
        "abstract": "аннотация",
    }
    summary: List[str] = []
    seen: set[str] = set()
    for hit in hit_records:
        hit_type = str((hit or {}).get("type") or "").lower()
        if not hit_type:
            continue
        if hit_type == "tag":
            tag_key = str((hit or {}).get("key") or "").strip()
            label = f"тег {tag_key}" if tag_key else "тег"
        else:
            label = mapping.get(hit_type, hit_type)
        if not label or label in seen:
            continue
        seen.add(label)
        summary.append(label)
        if len(summary) >= 5:
            break
    return ", ".join(summary)


def _collect_file_metadata_lines(file_obj: Optional["File"], *, max_tags: int = 3) -> List[str]:
    """Возвращает краткое описание файла для использования в контексте LLM."""
    if not file_obj:
        return []
    parts: List[str] = []
    if getattr(file_obj, "author", None):
        parts.append(f"Автор: {file_obj.author}")
    if getattr(file_obj, "year", None):
        parts.append(f"Год: {file_obj.year}")
    if getattr(file_obj, "material_type", None):
        parts.append(f"Тип: {file_obj.material_type}")
    if getattr(file_obj, "keywords", None):
        kw = str(file_obj.keywords).strip()
        if kw:
            parts.append(f"Ключевые слова: {kw}")
    tag_values: List[str] = []
    try:
        for tag in (getattr(file_obj, "tags", None) or []):
            key = str(getattr(tag, "key", "") or "").strip()
            value = str(getattr(tag, "value", "") or "").strip()
            if not key or not value:
                continue
            tag_values.append(f"{key}:{value}")
    except Exception:
        tag_values = []
    if tag_values:
        deduped: List[str] = []
        seen_tags: set[str] = set()
        for item in tag_values:
            if item in seen_tags:
                continue
            seen_tags.add(item)
            deduped.append(item)
        if deduped:
            parts.append(f"Теги: {', '.join(deduped[:max_tags])}")
    return parts


def _build_result_metadata_summary(
    file_obj: Optional["File"],
    result_entry: Optional[Dict[str, Any]],
) -> List[str]:
    """Комбинирует метаданные файла и сведения о совпадениях классического поиска."""
    summary = _collect_file_metadata_lines(file_obj)
    if not result_entry:
        return summary
    hits_summary = _summarize_search_hits(result_entry.get("hits") or [])
    if hits_summary:
        summary.append(f"Совпадения: {hits_summary}")
    matched_terms = result_entry.get("matched_terms") or []
    if matched_terms:
        terms_preview = ", ".join(str(term) for term in matched_terms[:5])
        if terms_preview:
            summary.append(f"Термины: {terms_preview}")
    sources = result_entry.get("snippet_sources") or []
    if sources:
        summary.append(f"Источники сниппетов: {', '.join(sources[:3])}")
    return summary


def _compose_snippet_based_answer(
    query: str,
    entries: Sequence[Dict[str, Any]],
    *,
    primary_terms: Sequence[str] | None = None,
    anchor_terms: Sequence[str] | None = None,
    term_idf: Optional[Mapping[str, float]] = None,
    max_items: int = 5,
    snippet_limit: int = 360,
) -> tuple[str, int, List[Dict[str, Any]]]:
    """Формирует безопасный ответ по сниппетам без генерации LLM."""
    noise_terms = {
        "кто",
        "что",
        "какой",
        "какая",
        "какие",
        "каков",
        "какова",
        "каково",
        "такой",
        "такая",
        "такие",
        "такое",
        "является",
        "являлся",
        "являлась",
        "являются",
        "есть",
        "это",
        "может",
        "нужно",
        "надо",
        "биография",
        "биографии",
    }
    usable_terms_raw = [
        term.lower()
        for term in (primary_terms or [])
        if term and len(term) >= 3 and term.lower() not in noise_terms
    ]
    usable_terms: List[str] = []
    if term_idf:
        for term in usable_terms_raw:
            if term_idf.get(term, 1.0) >= KEYWORD_IDF_MIN:
                usable_terms.append(term)
    else:
        usable_terms = usable_terms_raw[:]
    if not usable_terms:
        usable_terms = [term for term in usable_terms_raw if len(term) >= 5]
    anchor_terms_lower = [term.lower() for term in (anchor_terms or []) if term]
    def _anchor_in_text(lower_text: str) -> bool:
        if not anchor_terms_lower:
            return True
        for anchor in anchor_terms_lower:
            if anchor and anchor in lower_text:
                return True
            if len(anchor) >= 5:
                stem = anchor[:-2]
                if stem and stem in lower_text:
                    return True
            if len(anchor) >= 6:
                stem = anchor[:-3]
                if stem and stem in lower_text:
                    return True
        return False
    facts: List[str] = []
    sources: List[str] = []
    used_count = 0
    seen_snippets: set[str] = set()
    context_items: List[Dict[str, Any]] = []
    for idx, res in enumerate(entries, start=1):
        snippet_candidates: List[str] = []
        snippet_candidates.extend(res.get("snippets") or [])
        selected_snippets: List[str] = []
        for candidate in snippet_candidates:
            text = str(candidate or "").strip()
            if not text:
                continue
            normalized = re.sub(r"\s+", " ", text)
            if not normalized:
                continue
            lower = normalized.lower()
            if usable_terms and not any(term in lower for term in usable_terms):
                continue
            if not _anchor_in_text(lower):
                continue
            if lower in seen_snippets:
                continue
            seen_snippets.add(lower)
            trimmed = normalized[:max(snippet_limit, 360)].rstrip()
            selected_snippets.append(trimmed)
            if len(selected_snippets) >= 3:
                break
        if not selected_snippets:
            if not anchor_terms_lower and snippet_candidates:
                fallback_candidate = re.sub(r"\s+", " ", str(snippet_candidates[0] or "")).strip()
                if fallback_candidate:
                    selected_snippets.append(fallback_candidate[:max(snippet_limit, 360)].rstrip())
            else:
                continue
        if not selected_snippets:
            continue
        primary_snippet = selected_snippets[0][:snippet_limit].rstrip(". ")
        order = len(context_items) + 1
        facts.append(f"- {primary_snippet} [{order}]")
        title = (res.get("title") or res.get("rel_path") or f"file-{res.get('file_id')}") or ""
        title = str(title).strip()
        sources.append(f"- [{order}] {title}")
        used_count += 1
        context_items.append(
            {
                "order": order,
                "snippet": primary_snippet,
                "snippets": selected_snippets,
                "title": title,
                "file_id": int(res.get("file_id") or 0),
                "meta": res,
            }
        )
        if used_count >= max_items:
            break
    if not facts:
        return "", 0, []
    lines = ["Факты:", *facts, "Источники:", *sources]
    return "\n".join(lines), used_count, context_items


_SNIPPET_REF_RE = re.compile(r"\[(\d+)\]")


def _validate_snippet_llm_answer(answer: str, allowed_ids: Sequence[int]) -> Tuple[bool, List[str]]:
    text = (answer or "").strip()
    issues: List[str] = []
    if not text:
        issues.append("пустой ответ")
        return False, issues
    cited = sorted({int(match.group(1)) for match in _SNIPPET_REF_RE.finditer(text)})
    if not cited:
        issues.append("нет ссылок [n]")
    allowed_set = set(int(x) for x in allowed_ids)
    invalid = [cid for cid in cited if cid not in allowed_set]
    if invalid:
        issues.append(f"неизвестные ссылки {invalid}")
    return (not issues), issues


def _generate_snippet_llm_summary(
    query: str,
    context_items: Sequence[Dict[str, Any]],
    *,
    progress: Optional["_ProgressLogger"] = None,
    temperature: float = 0.1,
    max_tokens: int = 320,
) -> str:
    if not context_items:
        return ""
    allowed_ids = [item["order"] for item in context_items]
    context_lines: List[str] = []
    for item in context_items:
        title = item.get("title") or f"Источник {item.get('order')}"
        snippets = item.get("snippets") or [item.get("snippet") or ""]
        file_id = item.get("file_id")
        meta = item.get("meta") or {}
        hits = ", ".join(
            sorted({hit.get("term") for hit in (meta.get("hits") or []) if hit.get("term")})
        )
        header = f"[{item['order']}] Источник: {title}"
        if file_id:
            header += f" (file_id={file_id})"
        if hits:
            header += f" | ключи: {hits}"
        body_lines = []
        for snip in snippets:
            text = str(snip or "").strip()
            if not text:
                continue
            body_lines.append(f"- {text}")
        if not body_lines:
            body_lines.append("- (фрагмент отсутствует)")
        block = "\n".join([header] + body_lines)
        context_lines.append(block)
    system = (
        "Ты отвечаешь на вопросы поиска. Используй ТОЛЬКО предоставленные сниппеты, "
        "каждый помечен номером в квадратных скобках. Любой факт должен ссылаться на соответствующий номер. "
        "Если фрагменты содержат факты — перечисли их. Фразу «Недостаточно информации» используй "
        "только если ни один сниппет не даёт полезных сведений. Не придумывай новых фактов."
    )
    user = (
        f"Вопрос: {query}\n"
        "Доступные фрагменты:\n"
        + "\n".join(context_lines)
        + "\n\n"
        "Сформулируй краткий ответ на русском языке. Каждый факт сопровождай ссылкой [n]. "
        "Если информации недостаточно, ответь: 'Недостаточно информации в найденных источниках.'"
    )
    try:
        answer = call_lmstudio_compose(system, user, temperature=temperature, max_tokens=max_tokens) or ""
    except Exception:
        return ""
    answer = answer.strip()
    if not answer:
        return ""
    answer, translated = _ensure_russian_text(answer, label='snippet-answer')
    valid, issues = _validate_snippet_llm_answer(answer, allowed_ids)
    if not valid:
        if progress:
            progress.add(f"LLM ответ по сниппетам отклонён: {', '.join(issues)}")
        return ""
    if progress:
        note = "LLM ответ по сниппетам готов"
        if translated:
            note += ", выполнен перевод"
        progress.add(note)
    return answer


def _prepare_rag_context(
    query: str,
    results: Sequence[Dict[str, Any]],
    *,
    language_filters: Optional[Sequence[str]],
    max_chunks: int,
    progress: Optional["_ProgressLogger"] = None,
    variant_boost: int | float | None = None,
    dense_boost: float | int | None = None,
) -> tuple[Optional[Dict[str, Any]], list[str]]:
    notes: list[str] = []
    file_ids = [int(res.get('file_id')) for res in results if res.get('file_id')]
    if not file_ids:
        reason = "RAG: нет файлов с идентификаторами"
        if progress:
            progress.add(reason)
        notes.append(reason)
        return None, notes
    result_lookup: Dict[int, Dict[str, Any]] = {
        int(res.get("file_id")): res
        for res in results
        if res.get("file_id")
    }
    documents = (
        db.session.query(RagDocument)
        .filter(RagDocument.file_id.in_(file_ids), RagDocument.is_ready_for_rag == True)
        .all()
    )
    doc_map = {doc.id: doc for doc in documents}
    doc_ids = list(doc_map.keys())
    if not doc_ids:
        if progress:
            progress.add("RAG: нет документов, готовых для RAG-индекса")
        notes.append("Нет документов, подготовленных для RAG-индексации.")
        return None, notes
    allowed_scope = getattr(g, 'allowed_collection_ids', None)
    try:
        global_authority_scores = _compute_authority_scores(allowed_scope)
    except Exception:
        global_authority_scores = {}
    file_authority_scores: Dict[int, float] = {}
    if global_authority_scores:
        for doc in documents:
            fid = getattr(doc, 'file_id', None)
            if fid and fid in global_authority_scores:
                file_authority_scores[int(fid)] = float(global_authority_scores.get(fid, 0.0))
    file_feedback_scores: Dict[int, float] = {}
    try:
        feedback_weights_lookup = _get_feedback_weights()
    except Exception as exc:
        app.logger.debug("Feedback weights fetch failed in RAG context: %s", exc)
        feedback_weights_lookup = {}
    if feedback_weights_lookup:
        for doc in documents:
            fid = getattr(doc, 'file_id', None)
            if fid and fid in feedback_weights_lookup:
                file_feedback_scores[int(fid)] = float(feedback_weights_lookup[fid].get('weight', 0.0) or 0.0)
    variant = _select_rag_embedding_variant(doc_ids)
    if not variant:
        if progress:
            progress.add("RAG: эмбеддинги для документов не найдены")
        notes.append("Не найдены эмбеддинги для документов, используем fallback.")
        return None, notes
    model_name, model_version, dim = variant
    backend_choice = 'auto'
    version_token = (model_version or '').strip().lower()
    if 'lm-studio' in version_token or version_token in {'lmstudio', 'lm_studio'}:
        backend_choice = 'lm-studio'
    elif version_token in {'openai', 'openai-api'}:
        backend_choice = 'openai'
    try:
        backend = load_embedding_backend(
            backend_choice,
            model_name=model_name,
            dim=dim or 384,
        )
    except Exception as exc:
        if progress:
            progress.add(f"RAG: не удалось загрузить backend '{model_name}', fallback hash ({exc})")
        notes.append(f"Backend {model_name} недоступен ({exc}); используем hash fallback.")
        try:
            backend = load_embedding_backend(
                'hash',
                model_name=model_name,
                dim=dim or 384,
            )
        except Exception:
            backend = None
    if backend is None:
        if progress:
            progress.add("RAG: не удалось инициализировать backend эмбеддингов")
        notes.append("Не удалось инициализировать backend эмбеддингов.")
        return None, notes
    variant_limit = max(0, int(getattr(_rt(), "ai_query_variants_max", 0) or 0))
    if variant_boost:
        variant_limit += max(0, int(variant_boost))
    variant_limit = min(variant_limit, 6)
    query_variants: List[str] = [query]
    if variant_limit > 0:
        try:
            extra_variants = _generate_query_variants(query, max_variants=variant_limit)
        except Exception as exc:
            app.logger.debug("Query variants generation error: %s", exc)
            extra_variants = []
        seen_variants = {query.strip().lower()}
        for variant_text in extra_variants:
            normalized_variant = variant_text.strip()
            if not normalized_variant:
                continue
            lower = normalized_variant.lower()
            if lower in seen_variants:
                continue
            seen_variants.add(lower)
            query_variants.append(normalized_variant)
    backend_name_label = ""
    vector_map: Dict[str, Sequence[float]] = {}
    try:
        vectors = backend.embed_many(query_variants) or []
        for text, vec in zip(query_variants, vectors):
            vector_map[text] = vec
        query_vector = vector_map.get(query)
        if query_vector is None:
            query_vector = vectors[0] if vectors else [0.0] * (dim or 384)
        backend_name_label = getattr(backend, "name", getattr(backend, "model_name", "embedding-backend"))
    except Exception as exc:
        if progress:
            progress.add(f"RAG: не удалось вычислить эмбеддинг запроса ({exc})")
        notes.append(f"Не удалось вычислить эмбеддинг запроса ({exc}).")
        query_vector = [0.0] * (dim or 384)
        vector_map = {query: query_vector}
        query_variants = [query]
    finally:
        try:
            backend.close()
        except Exception:
            pass
    if backend_name_label:
        notes.append(f"Эмбеддинги: backend {backend_name_label}")
    language_filters = [lang.lower() for lang in (language_filters or [])] or None
    query_terms: List[str] = []
    for variant_text in query_variants:
        for token in _tokenize_query(variant_text):
            if token not in query_terms:
                query_terms.append(token)
    dense_multiplier = max(1.0, float(dense_boost or 1.0))
    vector_retriever = VectorRetriever(
        model_name=model_name,
        model_version=model_version,
        max_candidates=min(500, int(max_chunks * 40 * dense_multiplier)),
    )
    keyword_retriever = KeywordRetriever(
        limit=int(300 * dense_multiplier),
        search_service=search_service,
        expand_terms_fn=_expand_synonyms,
        lemma_fn=_lemma,
    )
    rag_reranker = _get_rag_reranker()
    selector = ContextSelector(
        vector_retriever=vector_retriever,
        keyword_retriever=keyword_retriever,
        dense_weight=1.0,
        sparse_weight=0.6,
        doc_penalty=0.2,
        max_per_document=2,
        dense_top_k=max(1, int(max_chunks * 3 * dense_multiplier)),
        sparse_limit=int(300 * dense_multiplier),
        max_total_tokens=_lm_safe_context_tokens(),
        rerank_fn=rag_reranker,
        feedback_scale=RAG_FEEDBACK_WEIGHT_SCALE,
    )
    dense_hits_override: Optional[List[RetrievedChunk]] = None
    if len(query_variants) > 1:
        variant_pairs: List[Tuple[str, Sequence[float]]] = []
        for text in query_variants:
            vec = vector_map.get(text)
            if vec:
                variant_pairs.append((text, vec))
        if len(variant_pairs) > 1:
            dense_hits_override = _rrf_combine_dense_hits(
                vector_retriever,
                variant_pairs,
                allowed_document_ids=doc_ids,
                top_k=max(1, int(max_chunks * 3 * dense_multiplier)),
            )
            if dense_hits_override:
                notes.append(f"RAG: использовано {len(variant_pairs)} формулировок запроса")
    contexts = selector.select(
        query,
        query_vector,
        top_k=max_chunks,
        languages=language_filters,
        allowed_document_ids=doc_ids,
        precomputed_dense_hits=dense_hits_override,
        query_terms=query_terms,
        authority_scores=file_authority_scores if file_authority_scores else None,
        feedback_scores=file_feedback_scores if file_feedback_scores else None,
    )
    if rag_reranker:
        cfg = getattr(rag_reranker, "_config", None)
        if cfg and getattr(cfg, "model_name", None):
            notes.append(f"RAG rerank: cross-encoder {cfg.model_name}")
        else:
            notes.append("RAG rerank: cross-encoder активен")
    if not contexts:
        if progress:
            progress.add("RAG: контекст не подобран")
        notes.append("Контекст RAG не подобран; используется классический режим.")
        return None, notes
    query_lang = detect_language(query)
    sections: List[ContextSection] = []
    allowed_refs: set[Tuple[int, int]] = set()
    chunk_ids: List[int] = []
    for idx, cand in enumerate(contexts, start=1):
        doc = cand.document or doc_map.get(cand.chunk.document_id)
        if doc is None:
            continue
        file_obj = getattr(doc, "file", None)
        title = (getattr(file_obj, "title", None) or getattr(file_obj, "filename", None) or getattr(file_obj, "rel_path", None) or f"Документ {doc.id}").strip()
        section_lang = cand.chunk.lang_primary or doc.lang_primary or ""
        translation_hint = _build_translation_hint(query_lang, section_lang)
        res_entry = None
        if file_obj is not None:
            res_entry = result_lookup.get(getattr(file_obj, "id", 0))
        preview_text = cand.preview or ""
        if not preview_text and res_entry:
            snippets = list(res_entry.get("snippets") or [])
            if res_entry.get("llm_snippet"):
                snippets.append(res_entry["llm_snippet"])
            if snippets:
                snippet_text = str(snippets[0]).strip()
                if snippet_text:
                    preview_text = snippet_text[:200]
        extra_fields: Dict[str, Any] = {
            "section_path": cand.section_path or cand.chunk.section_path or "",
            "keywords": cand.chunk.keywords_top or "",
        }
        if cand.section_path:
            extra_fields["Раздел"] = cand.section_path
        if cand.structure_bonus:
            extra_fields["Структурный вес"] = f"{cand.structure_bonus:.3f}"
        if cand.metadata_bonus:
            extra_fields["Метаданные вес"] = f"{cand.metadata_bonus:.3f}"
        if cand.metadata_hits:
            extra_fields["Совпадения метаданных"] = ", ".join(cand.metadata_hits[:5])
        if cand.authority_bonus:
            extra_fields["Вес авторитета"] = f"{cand.authority_bonus:.3f}"
        if getattr(cand, "feedback_bonus", 0.0):
            extra_fields["Вес фидбэка"] = f"{cand.feedback_bonus:.3f}"
        if file_obj is not None:
            if getattr(file_obj, "author", None):
                extra_fields["Автор"] = file_obj.author
            if getattr(file_obj, "year", None):
                extra_fields["Год"] = file_obj.year
            if getattr(file_obj, "material_type", None):
                extra_fields["Тип материала"] = file_obj.material_type
            if getattr(file_obj, "keywords", None):
                extra_fields["Ключевые слова файла"] = file_obj.keywords
            tag_lines = _collect_file_metadata_lines(file_obj)
            tag_strings = [line for line in tag_lines if line.startswith("Теги: ")]
            if tag_strings:
                extra_fields["Теги"] = tag_strings[0].replace("Теги: ", "", 1)
        if cand.matched_terms:
            extra_fields["Совпавшие термины"] = ", ".join(cand.matched_terms[:5])
        if res_entry:
            hits_summary = _summarize_search_hits(res_entry.get("hits") or [])
            if hits_summary:
                extra_fields["Совпадения поиска"] = hits_summary
            matched_terms = res_entry.get("matched_terms") or []
            if matched_terms:
                extra_fields.setdefault(
                    "Термины поиска",
                    ", ".join(str(term) for term in matched_terms[:5]),
                )
            snippet_sources = res_entry.get("snippet_sources") or []
            if snippet_sources:
                extra_fields["Источники сниппетов"] = ", ".join(snippet_sources[:3])
        llm_snippet = res_entry.get("llm_snippet") if res_entry else None
        section = ContextSection(
            doc_id=doc.id,
            chunk_id=cand.chunk.id,
            title=title,
            language=section_lang,
            token_estimate=cand.token_estimate,
            score_dense=cand.dense_score,
            score_sparse=cand.sparse_score,
            combined_score=cand.combined_score,
            reasoning_hint=cand.reasoning_hint,
            preview=preview_text,
            content=cand.chunk.content or "",
            url=getattr(file_obj, "rel_path", None),
            extra={
                **extra_fields,
                **({"LLM сниппет": llm_snippet} if llm_snippet else {}),
            },
            translation_hint=translation_hint,
        )
        sections.append(section)
        allowed_refs.add((doc.id, cand.chunk.id))
        chunk_ids.append(cand.chunk.id)
        if progress:
            progress.add(
                f"RAG контекст [{idx}/{len(contexts)}]: doc={doc.id} chunk={cand.chunk.id} combined={cand.combined_score:.3f} hint={cand.reasoning_hint or '—'}"
            )
    if not sections:
        notes.append("Контекст RAG не сформирован (нет подходящих секций).")
        return None, notes
    return {
        "sections": sections,
        "allowed_refs": allowed_refs,
        "query_lang": query_lang,
        "chunk_ids": chunk_ids,
        "model_name": model_name,
        "model_version": model_version,
        "notes": notes,
    }, notes


def _store_rag_session(
    *,
    query: str,
    query_lang: Optional[str],
    chunk_ids: Sequence[int],
    system_prompt: str,
    user_prompt: str,
    answer: str,
    validation: "ValidationResult",
    model_name: Optional[str],
    params: Optional[Dict[str, Any]],
    sources: Sequence[Dict[str, Any]],
    risk: Dict[str, Any],
) -> Optional[int]:
    try:
        validation_payload = {
            "validation": validation.as_dict(),
            "risk": risk,
            "sources": list(sources),
        }
        record = RagSession(
            query=query,
            query_lang=(query_lang or "")[:16],
            chunk_ids=",".join(str(cid) for cid in sorted(set(chunk_ids))),
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            answer=answer,
            validation=json.dumps(validation_payload, ensure_ascii=False),
            model_name=model_name,
            params=json.dumps(params or {}, ensure_ascii=False),
        )
        db.session.add(record)
        db.session.commit()
        return record.id
    except Exception as exc:
        db.session.rollback()
        app.logger.warning("Не удалось сохранить RAG-сессию: %s", exc)
        return None


def _estimate_section_tokens(section: ContextSection) -> int:
    estimate = getattr(section, "token_estimate", 0) or 0
    if estimate > 0:
        return int(estimate)
    text = (section.content or "").strip()
    if not text:
        return 120
    words = len(text.split())
    return max(80, int(words * 1.6) + 80)


def _total_section_tokens(sections: Sequence[ContextSection]) -> int:
    total = 0
    for section in sections:
        total += _estimate_section_tokens(section)
    return total


def _generate_rag_answer(
    query: str,
    bundle: Dict[str, Any],
    *,
    temperature: float,
    max_tokens: int,
    progress: Optional["_ProgressLogger"] = None,
) -> Tuple[str, "ValidationResult", Optional[int], str, str]:
    sections: List[ContextSection] = list(bundle.get("sections") or [])
    safe_limit = _lm_safe_context_tokens()
    trimmed_sections = sections[:]
    trimmed = False
    while len(trimmed_sections) > 1 and _total_section_tokens(trimmed_sections) > safe_limit:
        trimmed_sections = trimmed_sections[:-1]
        trimmed = True
    if trimmed and trimmed_sections:
        notes_list = bundle.setdefault("notes", [])
        note = f"Контекст RAG сокращён до {len(trimmed_sections)} секций (ограничение окна модели)."
        notes_list.append(note)
        if progress:
            progress.add(f"RAG: контекст сокращён до {len(trimmed_sections)} секций (ограничение окна модели)")
        allowed_refs_trimmed = {(sec.doc_id, sec.chunk_id) for sec in trimmed_sections}
        bundle["sections"] = trimmed_sections
        bundle["allowed_refs"] = allowed_refs_trimmed
        bundle["chunk_ids"] = [sec.chunk_id for sec in trimmed_sections]
        sections = trimmed_sections
    else:
        sections = bundle["sections"]
    system_prompt = build_system_prompt()
    user_prompt = build_user_prompt(query, sections)
    if progress:
        progress.add(f"RAG ответ: используем {len(sections)} чанков контекста")
    try:
        answer = call_lmstudio_compose(
            system_prompt,
            user_prompt,
            temperature=temperature,
            max_tokens=max_tokens,
        ).strip()
    except Exception as exc:
        if progress:
            progress.add(f"RAG ответ: ошибка генерации ({exc})")
        answer = ""
    if not answer:
        answer = fallback_answer()
        if progress:
            progress.add("RAG ответ: fallback без источников")
    validation = validate_answer(answer, sorted(bundle["allowed_refs"]))
    if validation.hallucination_warning and progress:
        progress.add("RAG валидация: обнаружены потенциальные проблемы с цитатами")
    if validation.missing_citations and progress:
        progress.add("RAG валидация: некоторые факты без ссылок")
    if validation.unknown_citations and progress:
        progress.add(f"RAG валидация: неизвестные ссылки {validation.unknown_citations}")
    if validation.extra_citations and progress:
        progress.add(f"RAG валидация: лишние ссылки {validation.extra_citations}")
    sources_payload = [
        {
            "doc_id": sec.doc_id,
            "chunk_id": sec.chunk_id,
            "title": sec.title,
            "section_path": sec.extra.get("Раздел") or sec.extra.get("section_path"),
            "combined_score": round(sec.combined_score, 4),
        }
        for sec in sections
    ]
    risk_info = _assess_rag_risk(answer, validation, sections)
    bundle["risk"] = risk_info
    bundle["sources"] = sources_payload
    session_id = _store_rag_session(
        query=query,
        query_lang=bundle.get("query_lang"),
        chunk_ids=bundle.get("chunk_ids") or [],
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        answer=answer,
        validation=validation,
        model_name=bundle.get("model_name"),
        params={
            "temperature": temperature,
            "max_tokens": max_tokens,
            "model_version": bundle.get("model_version"),
            "lm_model": getattr(_rt(), "lmstudio_model", None),
        },
        sources=sources_payload,
        risk=risk_info,
    )
    if session_id and progress:
        progress.add(f"RAG сессия сохранена (id={session_id})")
    return answer, validation, session_id, system_prompt, user_prompt


@app.cli.command("init-db")
def init_db():
    db.create_all()
    # Заполняем или обновляем TagSchema известными ключами (идемпотентно)
    seeds = [
            ("dissertation", "научный руководитель", "ФИО научного руководителя"),
            ("dissertation", "специальность", "Код/направление ВАК"),
            ("dissertation", "организация", "Базовая организация / вуз"),
            ("dissertation", "степень", "Кандидат / Доктор"),
            ("textbook", "дисциплина", "Учебная дисциплина"),
            ("textbook", "издательство", "Издательство"),
            ("article", "журнал", "Журнал / сборник"),
            ("article", "номер", "Номер выпуска"),
            ("article", "страницы", "Диапазон страниц"),
            ("article", "doi", "Digital Object Identifier"),
            ("monograph", "издательство", "Издательство"),
            ("any", "isbn", "Международный стандартный номер книги"),
        ]
    # Дополняем дополнительными ключами для новой таксономии тегов
    seeds += [
            ("any", "lang", "Язык текста (ru/en/...)"),
            ("any", "ext", "Расширение файла (без точки)"),
            ("any", "pages", "Число страниц (для PDF)"),
            ("any", "doi", "Digital Object Identifier"),
            ("any", "udk", "Универсальный десятичный классификатор (УДК)"),
            ("any", "bbk", "Библиотечно-библиографическая классификация (ББК)"),
            ("article", "journal", "Название журнала / сборника"),
            ("article", "volume_issue", "Том/номер журнала"),
            ("article", "pages", "Страницы в выпуске"),
            ("standard", "standard", "Код стандарта (ГОСТ/ISO/IEC/СТО/СП/СанПиН/ТУ)"),
            ("standard", "status", "Статус стандарта (active/replaced)"),
            ("proceedings", "conference", "Название конференции/симпозиума"),
            ("report", "doc_kind", "Вид документа (ТЗ/Пояснительная записка и т.п.)"),
            ("report", "organization", "Организация-издатель/разработчик"),
            ("patent", "patent_no", "Номер патента"),
            ("patent", "ipc", "Класс международной патентной классификации (IPC/МПК)"),
            ("presentation", "slides", "Признак презентации (слайды)"),
        ]
    added = 0
    for mt, k, d in seeds:
        exists = TagSchema.query.filter_by(material_type=mt, key=k).first()
        if not exists:
            db.session.add(TagSchema(material_type=mt, key=k, description=d))
            added += 1
    db.session.commit()
    print(f"DB initialized. Added {added} tag schema rows (existing preserved).")


def _delete_rag_index_for_collection(collection_id: int) -> Dict[str, int]:
    files: List[File] = list(File.query.filter(File.collection_id == collection_id).all())
    stats = {
        "affected_files": 0,
        "removed_documents": 0,
        "removed_chunks": 0,
        "removed_embeddings": 0,
        "removed_versions": 0,
        "removed_snippets": 0,
    }
    if not files:
        return stats
    file_ids = [f.id for f in files if getattr(f, "id", None)]
    for file_obj in files:
        doc = getattr(file_obj, "rag_document", None)
        if not doc:
            continue
        stats["affected_files"] += 1
        chunk_ids = [chunk.id for chunk in list(getattr(doc, "chunks", []) or []) if chunk.id]
        if chunk_ids:
            deleted_embeddings = (
                db.session.query(RagChunkEmbedding)
                .filter(RagChunkEmbedding.chunk_id.in_(chunk_ids))
                .delete(synchronize_session=False)
            )
            stats["removed_embeddings"] += int(deleted_embeddings or 0)
            deleted_chunks = (
                db.session.query(RagDocumentChunk)
                .filter(RagDocumentChunk.document_id == doc.id)
                .delete(synchronize_session=False)
            )
            stats["removed_chunks"] += int(deleted_chunks or 0)
        deleted_versions = (
            db.session.query(RagDocumentVersion)
            .filter(RagDocumentVersion.document_id == doc.id)
            .delete(synchronize_session=False)
        )
        stats["removed_versions"] += int(deleted_versions or 0)
        db.session.delete(doc)
        try:
            file_obj.rag_document = None  # type: ignore[attr-defined]
        except Exception:
            pass
        stats["removed_documents"] += 1
    if file_ids:
        deleted_snippets = (
            db.session.query(AiSearchSnippetCache)
            .filter(AiSearchSnippetCache.file_id.in_(file_ids))
            .delete(synchronize_session=False)
        )
        stats["removed_snippets"] += int(deleted_snippets or 0)
    return stats


@app.cli.command("rag-ingest-file")
@click.argument("file_id", type=int)
@click.option("--text-path", type=click.Path(exists=True, dir_okay=False), default=None, help="Путь к готовому тексту (если нужно переопределить извлечение).")
@click.option("--normalizer-version", default="v1", show_default=True, help="Версия нормализатора текста.")
@click.option("--max-tokens", type=int, default=700, show_default=True, help="Максимум токенов в чанке до overlap.")
@click.option("--overlap", type=int, default=120, show_default=True, help="Число токенов overlap между чанками.")
@click.option("--min-tokens", type=int, default=80, show_default=True, help="Минимум токенов в чанке, при меньшем объёме чанки сливаются.")
@click.option("--skip-if-unchanged/--force", default=True, show_default=True, help="Пропускать индексацию, если текст не изменился.")
@click.option("--commit/--no-commit", default=True, show_default=True, help="Коммитить изменения в базе данных.")
def rag_ingest_file_cli(
    file_id: int,
    text_path: Optional[str],
    normalizer_version: str,
    max_tokens: int,
    overlap: int,
    min_tokens: int,
    skip_if_unchanged: bool,
    commit: bool,
) -> None:
    """Индексирует файл в таблицах RAG (документ, чанки, версии)."""
    file_obj = File.query.filter_by(id=file_id).first()
    if not file_obj:
        raise click.ClickException(f"Файл с id={file_id} не найден.")

    cfg = ChunkConfig(
        max_tokens=max(16, max_tokens),
        overlap=max(0, min(overlap, max_tokens - 1)),
        min_tokens=max(1, min_tokens),
    )
    raw_text: Optional[str] = None
    source_path: Optional[Path] = None

    if text_path:
        source_path = Path(text_path)
        raw_text = _read_text_file_with_fallback(source_path)
    else:
        for candidate in _resolve_candidate_paths(file_obj):
            if not candidate.exists() or not candidate.is_file():
                continue
            extracted = _extract_text_for_rag(candidate)
            if extracted:
                source_path = candidate
                raw_text = extracted
                break
    if not raw_text:
        raw_text = file_obj.text_excerpt or ""
    if not raw_text.strip():
        raise click.ClickException("Не удалось получить текст для индексации (файл пустой или недоступен).")

    metadata = {
        "source": "cli.rag_ingest_file",
        "source_path": str(source_path) if source_path else None,
        "file_sha1": file_obj.sha1,
        "manual_text_path": bool(text_path),
    }

    indexer = RagIndexer(chunk_config=cfg, normalizer_version=normalizer_version)
    try:
        result = indexer.ingest_document(
            file_obj,
            raw_text,
            metadata=metadata,
            skip_if_unchanged=skip_if_unchanged,
            commit=commit,
        )
    except Exception as exc:
        app.logger.exception("RAG ingest failed for file_id=%s", file_id)
        raise click.ClickException(str(exc))

    if result.get("skipped"):
        click.echo(
            f"RAG индекс для файла {file_id} пропущен: {result.get('reason', 'unchanged')}. "
            f"Версия={result.get('version')}"
        )
    else:
        click.echo(
            f"RAG индекс готов: файл {file_id}, версия {result.get('version')} (id={result.get('version_id')}), "
            f"чанков={result.get('chunks')}, язык={result.get('language')}, commit={'yes' if commit else 'no'}"
        )


@app.cli.command("rag-embed-chunks")
@click.option("--backend", type=click.Choice(["auto", "hash", "sentence-transformers", "lm-studio", "openai"]), default="auto", show_default=True)
@click.option("--model-name", default="intfloat/multilingual-e5-large", show_default=True, help="Название модели/конфигурации.")
@click.option("--model-version", default=None, help="Версия модели; по умолчанию берётся из backend.")
@click.option("--dim", type=int, default=384, show_default=True, help="Размерность эмбеддингов (для hash backend).")
@click.option("--batch-size", type=int, default=32, show_default=True)
@click.option("--limit", type=int, default=256, show_default=True, help="Максимум чанков для обработки за один запуск.")
@click.option("--min-chunk-id", type=int, default=None, help="Обрабатывать чанки с id >= указанного.")
@click.option("--normalize/--no-normalize", default=True, show_default=True, help="Нормализовать векторы (L2).")
@click.option("--device", default=None, help="Устройство для sentence-transformers (cpu/cuda:0/... ).")
@click.option("--endpoint", default=None, help="URL OpenAI/LM Studio embeddings API (если отличается от настроек).")
@click.option("--api-key", default=None, help="API ключ для embeddings API (если требуется).")
@click.option("--timeout", type=float, default=None, help="Таймаут запросов к embeddings API, секунды.")
@click.option("--commit/--no-commit", default=True, show_default=True, help="Сохранять изменения в базе данных.")
def rag_embed_chunks_cli(
    backend: str,
    model_name: str,
    model_version: Optional[str],
    dim: int,
    batch_size: int,
    limit: int,
    min_chunk_id: Optional[int],
    normalize: bool,
    device: Optional[str],
    endpoint: Optional[str],
    api_key: Optional[str],
    timeout: Optional[float],
    commit: bool,
) -> None:
    """Генерирует эмбеддинги для чанков RAG и сохраняет их в БД."""
    try:
        runtime = runtime_settings_store.current
    except Exception:
        runtime = None
    resolved_endpoint = endpoint or (getattr(runtime, "rag_embedding_endpoint", None) if runtime else None) or (
        getattr(runtime, "lmstudio_api_base", None) if runtime else None
    )
    resolved_api_key = api_key or (getattr(runtime, "rag_embedding_api_key", None) if runtime else None) or (
        getattr(runtime, "lmstudio_api_key", None) if runtime else None
    )
    engine = load_embedding_backend(
        backend,
        model_name=model_name,
        dim=dim,
        normalize=normalize,
        batch_size=batch_size,
        device=device,
        base_url=resolved_endpoint,
        api_key=resolved_api_key,
        timeout=timeout,
    )
    resolved_model_name = getattr(engine, "model_name", model_name)
    resolved_model_version = model_version or getattr(engine, "model_version", "unknown")
    try:
        join_condition = and_(
            RagChunkEmbedding.chunk_id == RagDocumentChunk.id,
            RagChunkEmbedding.model_name == resolved_model_name,
            RagChunkEmbedding.model_version == resolved_model_version,
        )
        query = db.session.query(RagDocumentChunk).outerjoin(RagChunkEmbedding, join_condition)
        query = query.filter(RagChunkEmbedding.id.is_(None))
        if min_chunk_id is not None:
            query = query.filter(RagDocumentChunk.id >= min_chunk_id)
        query = query.order_by(RagDocumentChunk.id)
        if limit > 0:
            query = query.limit(limit)
        chunks: List[RagDocumentChunk] = query.all()
        if not chunks:
            click.echo("Нет чанков, требующих построения эмбеддингов.")
            return
        click.echo(f"Будет обработано чанков: {len(chunks)} (backend={getattr(engine, 'name', backend)})")
        processed = 0
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i : i + batch_size]
            texts = [c.content or "" for c in batch]
            vectors = engine.embed_many(texts)
            if len(vectors) != len(batch):
                raise RuntimeError("Размер ответа backend эмбеддингов не совпадает с размером батча.")
            for chunk_obj, vector in zip(batch, vectors):
                vector_bytes = vector_to_bytes(vector)
                vector_checksum = hashlib.sha256(vector_bytes).hexdigest()
                embedding = RagChunkEmbedding(
                    chunk_id=chunk_obj.id,
                    model_name=resolved_model_name,
                    model_version=resolved_model_version,
                    dim=len(vector),
                    vector=vector_bytes,
                    vector_checksum=vector_checksum,
                )
                db.session.add(embedding)
                processed += 1
        if commit:
            db.session.commit()
            click.echo(f"Сохранено эмбеддингов: {processed}")
        else:
            db.session.rollback()
            click.echo(f"[no-commit] Сгенерировано эмбеддингов: {processed}, изменения отменены.")
    finally:
        try:
            engine.close()
        except Exception:
            pass


@app.cli.command("rag-search")
@click.argument("query", type=str)
@click.option("--backend", type=click.Choice(["auto", "hash", "sentence-transformers"]), default="auto", show_default=True)
@click.option("--model-name", default="intfloat/multilingual-e5-large", show_default=True)
@click.option("--model-version", default=None)
@click.option("--dim", type=int, default=384, show_default=True)
@click.option("--normalize/--no-normalize", default=True, show_default=True)
@click.option("--batch-size", type=int, default=32, show_default=True)
@click.option("--device", default=None)
@click.option("--top-k", type=int, default=6, show_default=True)
@click.option("--max-candidates", type=int, default=1000, show_default=True)
def rag_search_cli(
    query: str,
    backend: str,
    model_name: str,
    model_version: Optional[str],
    dim: int,
    normalize: bool,
    batch_size: int,
    device: Optional[str],
    top_k: int,
    max_candidates: int,
) -> None:
    """Выполняет dense-поиск по чанкам RAG для текстового запроса."""
    engine = load_embedding_backend(
        backend,
        model_name=model_name,
        dim=dim,
        normalize=normalize,
        batch_size=batch_size,
        device=device,
    )
    try:
        vectors = engine.embed_many([query])
        resolved_model_name = getattr(engine, "model_name", model_name)
        resolved_model_version = model_version or getattr(engine, "model_version", "unknown")
    finally:
        try:
            engine.close()
        except Exception:
            pass
    if not vectors:
        click.echo("Не удалось получить эмбеддинг для запроса.")
        return
    retriever = VectorRetriever(
        model_name=resolved_model_name,
        model_version=resolved_model_version,
        max_candidates=max_candidates,
    )
    results = retriever.search_by_vector(vectors[0], top_k=top_k)
    if not results:
        click.echo("Совпадений не найдено.")
        return
    click.echo(f"Топ {len(results)} результатов:")
    for idx, item in enumerate(results, start=1):
        doc = item.document
        doc_info = f"doc_id={doc.id}" if doc else "doc_id=?"
        click.echo(
            f"{idx}. score={item.score:.4f} chunk_id={item.chunk.id} {doc_info} "
            f"lang={item.lang or '-'} keywords={item.keywords or '-'}"
        )


@app.cli.command("rag-context")
@click.argument("query", type=str)
@click.option("--backend", type=click.Choice(["auto", "hash", "sentence-transformers"]), default="auto", show_default=True)
@click.option("--model-name", default="intfloat/multilingual-e5-large", show_default=True)
@click.option("--model-version", default=None)
@click.option("--dim", type=int, default=384, show_default=True)
@click.option("--normalize/--no-normalize", default=True, show_default=True)
@click.option("--batch-size", type=int, default=32, show_default=True)
@click.option("--device", default=None)
@click.option("--top-k", type=int, default=6, show_default=True)
@click.option("--dense-top-k", type=int, default=12, show_default=True)
@click.option("--sparse-limit", type=int, default=100, show_default=True)
@click.option("--max-candidates", type=int, default=1000, show_default=True)
@click.option("--max-per-document", type=int, default=2, show_default=True)
@click.option("--doc-penalty", type=float, default=0.1, show_default=True)
@click.option("--dense-weight", type=float, default=1.0, show_default=True)
@click.option("--sparse-weight", type=float, default=0.6, show_default=True)
@click.option("--min-dense-score", type=float, default=0.0, show_default=True)
@click.option("--min-sparse-score", type=float, default=0.0, show_default=True)
@click.option("--min-combined-score", type=float, default=0.0, show_default=True)
@click.option("--max-tokens", type=int, default=3500, show_default=True, help="Ограничение по суммарным токенам контекста (0 = без ограничения).")
@click.option("--rerank-mode", type=click.Choice(["none", "combined", "dense", "sparse", "cross-encoder"]), default="none", show_default=True, help="Дополнительный реранж, например по dense-score.")
@click.option("--rerank-model", default=None, help="Название модели для cross-encoder rerank (например cross-encoder/ms-marco-MiniLM-L-6-v2).")
@click.option("--rerank-device", default=None, help="Устройство для cross-encoder rerank (по умолчанию как --device).")
@click.option("--rerank-batch-size", type=int, default=16, show_default=True, help="Batch size для cross-encoder rerank.")
@click.option("--rerank-max-length", type=int, default=512, show_default=True, help="Максимальная длина входа CrossEncoder (токены).")
@click.option("--rerank-max-chars", type=int, default=1200, show_default=True, help="Обрезка текста чанка перед cross-encoder rerank (символы).")
@click.option("--lang", multiple=True, help="Ограничить языки чанков (повторяемая опция).")
def rag_context_cli(
    query: str,
    backend: str,
    model_name: str,
    model_version: Optional[str],
    dim: int,
    normalize: bool,
    batch_size: int,
    device: Optional[str],
    top_k: int,
    dense_top_k: int,
    sparse_limit: int,
    max_candidates: int,
    max_per_document: int,
    doc_penalty: float,
    dense_weight: float,
    sparse_weight: float,
    min_dense_score: float,
    min_sparse_score: float,
    min_combined_score: float,
    max_tokens: int,
    rerank_mode: str,
    rerank_model: Optional[str],
    rerank_device: Optional[str],
    rerank_batch_size: int,
    rerank_max_length: int,
    rerank_max_chars: int,
    lang: Sequence[str],
) -> None:
    """Формирует комбинированный контекст (dense + keyword) для запроса."""
    languages = [item.strip() for item in lang if item.strip()] or None
    engine = load_embedding_backend(
        backend,
        model_name=model_name,
        dim=dim,
        normalize=normalize,
        batch_size=batch_size,
        device=device,
    )
    try:
        vectors = engine.embed_many([query])
        resolved_model_name = getattr(engine, "model_name", model_name)
        resolved_model_version = model_version or getattr(engine, "model_version", "unknown")
    finally:
        try:
            engine.close()
        except Exception:
            pass
    if not vectors:
        click.echo("Не удалось получить эмбеддинг для запроса.")
        return
    vector_retriever = VectorRetriever(
        model_name=resolved_model_name,
        model_version=resolved_model_version,
        max_candidates=max_candidates,
    )
    keyword_retriever = KeywordRetriever(
        limit=sparse_limit,
        search_service=search_service,
        expand_terms_fn=_expand_synonyms,
        lemma_fn=_lemma,
    )
    try:
        authority_scores_cli = _compute_authority_scores(None)
    except Exception:
        authority_scores_cli = {}
    try:
        feedback_scores_cli = {
            int(fid): float(entry.get("weight", 0.0) or 0.0)
            for fid, entry in _get_feedback_weights().items()
        }
    except Exception:
        feedback_scores_cli = {}
    query_terms_cli = _tokenize_query(query)
    mode = (rerank_mode or "none").lower()

    def _make_rerank(mode_name: str):
        if mode_name == "dense":
            return lambda _q, items: sorted(items, key=lambda c: (c.dense_score, c.combined_score), reverse=True)
        if mode_name == "sparse":
            return lambda _q, items: sorted(items, key=lambda c: (c.sparse_score, c.combined_score), reverse=True)
        if mode_name == "combined":
            return lambda _q, items: sorted(items, key=lambda c: c.combined_score, reverse=True)
        return None

    rerank_fn = None
    rerank_instance = None
    if mode == "cross-encoder":
        if not rerank_model:
            click.echo("Для cross-encoder режима требуется указать --rerank-model.")
            return
        try:
            from agregator.rag.rerank import CrossEncoderConfig, load_reranker

            rerank_instance = load_reranker(
                "cross-encoder",
                config=CrossEncoderConfig(
                    model_name=rerank_model,
                    device=rerank_device or device,
                    batch_size=rerank_batch_size,
                    max_length=rerank_max_length,
                    truncate_chars=rerank_max_chars,
                ),
            )
            rerank_fn = rerank_instance
        except Exception as exc:
            click.echo(f"Не удалось инициализировать cross-encoder rerank: {exc}")
            return
    else:
        rerank_fn = _make_rerank(mode)

    selector = ContextSelector(
        vector_retriever=vector_retriever,
        keyword_retriever=keyword_retriever,
        dense_weight=dense_weight,
        sparse_weight=sparse_weight,
        doc_penalty=doc_penalty,
        max_per_document=max_per_document,
        dense_top_k=dense_top_k,
        sparse_limit=sparse_limit,
        min_dense_score=min_dense_score,
        min_sparse_score=min_sparse_score,
        min_combined_score=min_combined_score,
        max_total_tokens=max_tokens if max_tokens > 0 else None,
        rerank_fn=rerank_fn,
        feedback_scale=RAG_FEEDBACK_WEIGHT_SCALE,
    )
    contexts = selector.select(
        query=query,
        query_vector=vectors[0],
        top_k=top_k,
        languages=languages,
        max_total_tokens=max_tokens if max_tokens > 0 else None,
        query_terms=query_terms_cli,
        authority_scores=authority_scores_cli if authority_scores_cli else None,
        feedback_scores=feedback_scores_cli if feedback_scores_cli else None,
    )
    if rerank_instance:
        try:
            rerank_instance.close()
        except Exception:
            pass
    if not contexts:
        click.echo("Контекст не подобран.")
        return
    click.echo(f"Отобрано чанков: {len(contexts)}")
    for idx, ctx in enumerate(contexts, start=1):
        doc = ctx.document
        doc_id = doc.id if doc else None
        title = None
        if doc and getattr(doc, "file", None):
            title = doc.file.title or doc.file.filename
        keywords_text = ", ".join(ctx.matched_terms[:5]) if ctx.matched_terms else "-"
        preview = (ctx.chunk.preview or ctx.chunk.content or "")[:200].replace("\n", " ").strip()
        click.echo(
            f"{idx}. doc={doc_id or '-'} chunk={ctx.chunk.id} adj={ctx.adjusted_score:.4f} "
            f"combined={ctx.combined_score:.4f} dense={ctx.dense_score:.4f} sparse={ctx.sparse_score:.4f} "
            f"lang={ctx.chunk.lang_primary or '-'} keywords={keywords_text}"
        )
        if title:
            click.echo(f"   title: {title}")
        if ctx.reasoning_hint:
            click.echo(f"   hint: {ctx.reasoning_hint}")
        if preview:
            click.echo(f"   preview: {preview}")


@app.cli.command("rag-generate")
@click.argument("query", type=str)
@click.option("--chunk-id", "chunk_ids", multiple=True, type=int, required=True, help="ID чанков, которые попадут в контекст.")
@click.option("--temperature", type=float, default=0.2, show_default=True)
@click.option("--max-tokens", type=int, default=400, show_default=True)
@click.option("--llm/--no-llm", default=True, show_default=True, help="Вызвать ли LLM, иначе вывести только промпты.")
@click.option("--custom-system", default=None, help="Переопределить системный промпт.")
@click.option("--store/--no-store", default=True, show_default=True, help="Сохранить результат в таблицу rag_sessions.")
def rag_generate_cli(
    query: str,
    chunk_ids: Sequence[int],
    temperature: float,
    max_tokens: int,
    llm: bool,
    custom_system: Optional[str],
    store: bool,
) -> None:
    """Собирает промпт для RAG-ответа и опционально вызывает LLM."""
    if not chunk_ids:
        click.echo("Не указаны chunk-id.")
        raise click.Abort()
    rows = (
        db.session.query(RagDocumentChunk, RagDocument, File)
        .join(RagDocument, RagDocumentChunk.document_id == RagDocument.id)
        .join(File, File.id == RagDocument.file_id)
        .filter(RagDocumentChunk.id.in_(chunk_ids))
        .order_by(RagDocumentChunk.id)
        .all()
    )
    if not rows:
        click.echo("Чанки не найдены.")
        return
    sections: List[ContextSection] = []
    allowed_refs = set()
    query_lang = detect_language(query)
    for chunk, document, file_obj in rows:
        extra_meta = {}
        if chunk.meta:
            try:
                extra_meta = json.loads(chunk.meta)
            except Exception:
                extra_meta = {}
        section_lang = (chunk.lang_primary or document.lang_primary or "").lower()
        raw_token_estimate = extra_meta.get("token_estimate")
        if raw_token_estimate is None:
            raw_token_estimate = chunk.token_count
        try:
            token_estimate = int(raw_token_estimate or 0)
        except (TypeError, ValueError):
            token_estimate = 0
        sections.append(
            ContextSection(
                doc_id=document.id,
                chunk_id=chunk.id,
                title=file_obj.title or file_obj.filename or f"Документ {file_obj.id}",
                language=section_lang,
                token_estimate=token_estimate,
                score_dense=float(extra_meta.get("dense_score") or 0.0),
                score_sparse=float(extra_meta.get("sparse_score") or 0.0),
                combined_score=float(extra_meta.get("combined_score") or 0.0),
                reasoning_hint=str(extra_meta.get("hint") or ""),
                preview=chunk.preview or str(extra_meta.get("preview") or ""),
                content=chunk.content or "",
                url=file_obj.rel_path,
                extra={"section_path": chunk.section_path or ""},
            )
        )
        allowed_refs.add((document.id, chunk.id))

    for section in sections:
        section.translation_hint = _build_translation_hint(query_lang, section.language)

    system_prompt = build_system_prompt(custom_system)
    user_prompt = build_user_prompt(query, sections)
    click.echo("=== System Prompt ===")
    click.echo(system_prompt)
    click.echo("\n=== User Prompt ===")
    click.echo(user_prompt)

    if not llm:
        click.echo("\nLLM вызов отключён (флаг --no-llm).")
        return

    answer = call_lmstudio_compose(
        system_prompt,
        user_prompt,
        temperature=temperature,
        max_tokens=max_tokens,
    ).strip()
    if not answer:
        answer = fallback_answer()
    click.echo("\n=== LLM Ответ ===")
    click.echo(answer)

    validation = validate_answer(answer, sorted(allowed_refs))
    click.echo("\n=== Валидация ===")
    click.echo(json.dumps(validation.as_dict(), ensure_ascii=False, indent=2))

    if store:
        try:
            session_record = RagSession(
                query=query,
                query_lang=query_lang,
                chunk_ids=",".join(str(cid) for cid in chunk_ids),
                system_prompt=system_prompt,
                user_prompt=user_prompt,
                answer=answer,
                validation=json.dumps(validation.as_dict(), ensure_ascii=False),
                model_name=getattr(_rt(), "lmstudio_model", None),
                params=json.dumps({
                    "temperature": temperature,
                    "max_tokens": max_tokens,
                    "llm": bool(llm),
                }, ensure_ascii=False),
            )
            db.session.add(session_record)
            db.session.commit()
            click.echo(f"\nSession saved: id={session_record.id}")
        except Exception as exc:
            db.session.rollback()
            click.echo(f"\n[warn] Не удалось сохранить сессию: {exc}")


# ------------------- App bootstrap -------------------

from routes import routes
app.register_blueprint(routes)

app.config['invalidate_facets'] = _invalidate_facets_cache
app.config['facet_cache'] = FACET_CACHE
app.config['facet_service'] = facet_service
app.config['search_service'] = search_service


def _run_refresh_task(task_id: int, file_id: int, params: dict[str, str]) -> None:
    with app.app_context():
        task = TaskRecord.query.get(task_id)
        if not task:
            return
        try:
            task.status = 'running'
            task.started_at = datetime.utcnow()
            task.progress = 0.1
            db.session.commit()
        except Exception:
            db.session.rollback()
        try:
            query = "&".join(
                f"{k}={quote_plus(str(v))}"
                for k, v in params.items()
                if v is not None and str(v) != ""
            )
            suffix = f"?{query}" if query else ""
            with app.test_request_context(
                f"/api/files/{int(file_id)}/refresh{suffix}",
                method="POST",
            ):
                response = api_file_refresh(int(file_id))
            status_code = 200
            payload: dict[str, Any] = {}
            if isinstance(response, tuple):
                flask_resp = response[0]
                status_code = int(response[1])
                payload = flask_resp.get_json(silent=True) if hasattr(flask_resp, "get_json") else {}
            else:
                flask_resp = response
                status_code = int(getattr(flask_resp, "status_code", 200) or 200)
                payload = flask_resp.get_json(silent=True) if hasattr(flask_resp, "get_json") else {}
            task = TaskRecord.query.get(task_id)
            if not task:
                return
            task.progress = 1.0
            task.payload = json.dumps(payload or {}, ensure_ascii=False)
            if status_code >= 400 or not (payload or {}).get("ok", False):
                task.status = 'error'
                task.error = str((payload or {}).get("error") or f"refresh_failed:{status_code}")
            else:
                task.status = 'completed'
            task.finished_at = datetime.utcnow()
            db.session.commit()
        except Exception as exc:
            db.session.rollback()
            task = TaskRecord.query.get(task_id)
            if task:
                task.status = 'error'
                task.error = str(exc)
                task.finished_at = datetime.utcnow()
                try:
                    db.session.commit()
                except Exception:
                    db.session.rollback()


# ------------------- Single-file Refresh API -------------------
@app.route("/api/files/<int:file_id>/refresh", methods=["POST"])
def api_file_refresh(file_id):
    """Re-extract text, refresh tags, and optionally re-run LLM for a single file.
    Respects runtime settings for LLM, audio summary and keywords. Always re-extracts text.
    """
    internal_call = str(request.args.get('internal') or '').strip().lower() in ('1', 'true', 'yes', 'on')
    async_requested = str(request.args.get('async') or '').strip().lower() in ('1', 'true', 'yes', 'on')
    user = _load_current_user()
    if async_requested and not internal_call:
        payload = {
            'file_id': int(file_id),
            'params': {
                'use_llm': request.args.get('use_llm'),
                'kws_audio': request.args.get('kws_audio'),
                'summarize': request.args.get('summarize'),
                'internal': '1',
            },
            'trigger_user_id': getattr(user, 'id', None),
        }
        task = TaskRecord(
            name='refresh',
            status='queued',
            progress=0.0,
            payload=json.dumps(payload, ensure_ascii=False),
            user_id=getattr(user, 'id', None),
            collection_id=getattr(File.query.get(file_id), 'collection_id', None),
        )
        try:
            db.session.add(task)
            db.session.commit()
        except Exception as exc:
            db.session.rollback()
            return jsonify({'ok': False, 'error': f'Не удалось создать задачу refresh: {exc}'}), 500
        params = {
            'use_llm': request.args.get('use_llm') or '1',
            'kws_audio': request.args.get('kws_audio') or '1',
            'summarize': request.args.get('summarize') or '1',
            'internal': '1',
        }
        get_task_queue().submit(
            _run_refresh_task,
            int(task.id),
            int(file_id),
            params,
            description=f"refresh-{file_id}",
            owner_id=getattr(user, 'id', None),
        )
        return jsonify({'ok': True, 'async': True, 'task_id': int(task.id)})
    f = File.query.get_or_404(file_id)
    refresh_task: TaskRecord | None = None
    if user and not internal_call:
        try:
            refresh_task = TaskRecord(
                name='refresh',
                status='running',
                progress=0.0,
                payload=json.dumps({'file_id': file_id, 'user_id': user.id}),
                user_id=user.id,
                collection_id=getattr(f, 'collection_id', None),
            )
            refresh_task.started_at = datetime.utcnow()
            db.session.add(refresh_task)
            db.session.commit()
        except Exception:
            db.session.rollback()
            refresh_task = None
    refresh_params = {k: request.args.get(k) for k in ('use_llm', 'kws_audio', 'summarize')}
    log_detail = {'file_id': file_id, 'params': {k: v for k, v in refresh_params.items() if v is not None}}
    try:
        detail_payload = json.dumps(log_detail, ensure_ascii=False)
    except Exception:
        detail_payload = str(log_detail)
    try:
        _log_user_action(user, 'file_refresh_start', 'file', file_id, detail=detail_payload[:2000])
    except Exception:
        pass
    try:
        app.logger.info(
            "[user-action] user=%s action=file_refresh_start file_id=%s params=%s",
            getattr(user, 'username', None) or 'anonymous',
            file_id,
            log_detail.get('params')
        )
    except Exception:
        pass
    log: list[str] = []
    def _log(m: str):
        log.append(m)
    try:
        p = Path(f.path)
        if not p.exists() or not p.is_file():
            return jsonify({"ok": False, "error": "file_not_found", "log": ["file not found on disk"]}), 404
        ext = p.suffix.lower()
        filename = p.stem

        # очищаем кэшированные артефакты для этого файла (миниатюра и текстовый фрагмент)
        try:
            # кэш текстового фрагмента
            cache_dir = Path(app.static_folder) / 'cache' / 'text_excerpts'
            key = (f.sha1 or (f.rel_path or '').replace('/', '_')) + '.txt'
            fp = cache_dir / key
            if fp.exists():
                fp.unlink()
        except Exception:
            pass
        try:
            # миниатюра PDF
            if ext == '.pdf':
                thumb = Path(app.static_folder) / 'thumbnails' / (p.stem + '.png')
                if thumb.exists():
                    thumb.unlink()
        except Exception:
            pass

        # извлекаем текст (всегда при обновлении)
        text_excerpt = ""
        _log(f"extract: start for {p.name}")
        if ext == ".pdf":
            force = _always_ocr_first_page_dissertation() and (
                ((f.material_type or '').lower() in ('dissertation', 'dissertation_abstract'))
                or looks_like_dissertation_filename(filename)
            )
            text_excerpt = extract_text_pdf(p, limit_chars=40000, force_ocr_first_page=force)
        elif ext == ".docx":
            text_excerpt = extract_text_docx(p, limit_chars=40000)
        elif ext == ".rtf":
            text_excerpt = extract_text_rtf(p, limit_chars=40000)
        elif ext == ".epub":
            text_excerpt = extract_text_epub(p, limit_chars=40000)
        elif ext == ".djvu":
            text_excerpt = extract_text_djvu(p, limit_chars=40000)
        elif ext in AUDIO_EXTS:
            text_excerpt = transcribe_audio(p, limit_chars=40000)
        if text_excerpt:
            f.text_excerpt = text_excerpt[:40000]
        _log(f"extract: done chars={len(text_excerpt or '')}")
        # Общие теги: расширение и страницы PDF
        try:
            upsert_tag(f, 'ext', ext.lstrip('.'))
        except Exception:
            pass
        page_count = None
        if ext == '.pdf':
            try:
                with fitz.open(str(p)) as _doc:
                    page_count = len(_doc)
                    upsert_tag(f, 'pages', str(len(_doc)))
            except Exception:
                pass

        # Теги, специфичные для аудио
        if ext in AUDIO_EXTS:
            try:
                upsert_tag(f, 'формат', ext.lstrip('.'))
                upsert_tag(f, 'длительность', audio_duration_hhmmss(p))
            except Exception:
                pass

        # Эвристики по имени файла (заполняем только пустые поля)
        title = author = year = None
        for pat in FILENAME_PATTERNS:
            m = pat.match(filename)
            if m:
                gd = m.groupdict()
                title = gd.get("title")
                author = gd.get("author")
                year = gd.get("year")
                break
        if title and not f.title:
            f.title = title
        if author and not f.author:
            f.author = author
        if year and not f.year:
            f.year = year

        # Определяем тип материала, если не указан; корректируем аудио
        if not f.material_type:
            # Явно проставим тип по расширению
            if ext in IMAGE_EXTS:
                f.material_type = 'image'
            elif ext in AUDIO_EXTS:
                f.material_type = 'audio'
            else:
                f.material_type = guess_material_type(ext, text_excerpt, filename)
        if ext in AUDIO_EXTS and (f.material_type or '') == 'document':
            f.material_type = 'audio'
        if ext in IMAGE_EXTS and (f.material_type or '') == 'document':
            f.material_type = 'image'

        journal_entries = _extract_journal_toc_entries(text_excerpt or '') if text_excerpt else []
        f.material_type = _journal_safety_check(f.material_type, text_excerpt, page_count, journal_entries)

        # Теги, зависящие от типа
        try:
            ttags = extract_tags_for_type(f.material_type or '', text_excerpt or '', filename)
            if ttags:
                db.session.flush()
                for k, v in ttags.items():
                    values = v if isinstance(v, (list, tuple, set)) else [v]
                    for single in values:
                        if single is None:
                            continue
                        upsert_tag(f, k, single)
            _log(f"type-tags: {len(ttags or {})}")
        except Exception:
            _log("type-tags: error")
        # Извлекаем встроенные ключевые слова («Ключевые слова:»), если есть в тексте
        try:
            import re as _re
            kw_m = _re.search(r"(?:Ключевые\s+слова|Keywords)\s*[:\-]\s*(.+)", (text_excerpt or ''), flags=_re.IGNORECASE)
            if kw_m:
                f.keywords = kw_m.group(1).strip()[:2000]
                if KEYWORDS_TO_TAGS_ENABLED:
                    db.session.flush(); _upsert_keyword_tags(f)
                _log("inline-keywords: ok")
        except Exception:
            _log("inline-keywords: error")
        # Дополнительные расширенные теги
        try:
            rtags = extract_richer_tags(f.material_type or '', text_excerpt or '', filename)
            if rtags:
                db.session.flush()
                for k, v in rtags.items():
                    values = v if isinstance(v, (list, tuple, set)) else [v]
                    for single in values:
                        if single is None:
                            continue
                        upsert_tag(f, k, single)
            _log(f"richer-tags: {len(rtags or {})}")
        except Exception:
            _log("richer-tags: error")

        # Необязательное обогащение через LLM по настройкам или флагам
        use_llm = (request.args.get('use_llm') in ('1','true','yes','on')) if ('use_llm' in request.args) else DEFAULT_USE_LLM
        do_summarize = (request.args.get('summarize') in ('1','true','yes','on')) if ('summarize' in request.args) else SUMMARIZE_AUDIO
        kws_audio_on = (request.args.get('kws_audio') in ('1','true','yes','on')) if ('kws_audio' in request.args) else AUDIO_KEYWORDS_LLM

        if use_llm and (text_excerpt or ext in {'.txt', '.md'}):
            llm_text = text_excerpt or ""
            if not llm_text and ext in {".txt", ".md"}:
                try:
                    llm_text = p.read_text(encoding="utf-8", errors="ignore")[:15000]
                except Exception:
                    pass
            _log(f"llm: metadata start (len={len(llm_text)})")
            meta = call_lmstudio_for_metadata(llm_text, p.name)
            if meta:
                mt_meta = normalize_material_type(meta.get("material_type"))
                # Не позволяем LLM менять тип для явных форматов
                if ext in AUDIO_EXTS:
                    f.material_type = 'audio'
                elif ext in IMAGE_EXTS:
                    f.material_type = 'image'
                else:
                    merged_type = _merge_llm_material_type(
                        f.material_type,
                        mt_meta,
                        allow_override=True,
                    )
                    if merged_type:
                        f.material_type = merged_type
                _t = (meta.get("title") or "").strip()
                if _t:
                    f.title = _t
                _a = _normalize_author(meta.get("author"))
                if _a:
                    f.author = _a
                _y = _normalize_year(meta.get("year"))
                if _y:
                    f.year = _y
                _adv = meta.get("advisor")
                if _adv is not None:
                    _adv_s = str(_adv).strip()
                    if _adv_s:
                        f.advisor = _adv_s
                kws = meta.get("keywords") or []
                if isinstance(kws, list) and kws:
                    f.keywords = ", ".join([str(x) for x in kws][:50])
                    if KEYWORDS_TO_TAGS_ENABLED:
                        db.session.flush()
                        _upsert_keyword_tags(f)
                if meta.get("novelty"):
                    db.session.flush()
                    upsert_tag(f, "научная новизна", str(meta.get("novelty")))
                for key in ("literature", "organizations", "classification"):
                    val = meta.get(key)
                    if isinstance(val, list) and val:
                        db.session.flush()
                        upsert_tag(f, key, "; ".join([str(x) for x in val]))
            else:
                _log("llm: metadata empty")

        # Резюме и ключевые слова для аудио
        if ext in AUDIO_EXTS:
            if kws_audio_on and (f.text_excerpt or '') and not (f.keywords or '').strip():
                try:
                    kws = call_lmstudio_keywords(f.text_excerpt, p.name)
                    if kws:
                        f.keywords = ", ".join(kws)
                        db.session.flush()
                        _upsert_keyword_tags(f)
                    _log(f"llm: audio keywords {len(kws or [])}")
                except Exception:
                    _log("llm: audio keywords error")
        # Анализ изображений (vision)
        if ext in IMAGE_EXTS and IMAGES_VISION_ENABLED:
            try:
                vis = call_lmstudio_vision(p, p.name)
                if isinstance(vis, dict):
                    desc = (vis.get('description') or '')
                    if desc:
                        f.abstract = desc[:8000]
                    kws = vis.get('keywords') or []
                    if isinstance(kws, list) and kws:
                        f.keywords = ", ".join([str(x) for x in kws][:50])
                        if KEYWORDS_TO_TAGS_ENABLED:
                            db.session.flush()
                            _upsert_keyword_tags(f)
                _log("vision: ok")
            except Exception:
                _log("vision: error")
            if do_summarize and (f.text_excerpt or ''):
                try:
                    summ = call_lmstudio_summarize(f.text_excerpt, p.name)
                    if summ:
                        f.abstract = summ[:2000]
                    _log("llm: summarize ok")
                except Exception:
                    _log("llm: summarize error")

        try:
            _apply_metadata_quality_rules(
                f,
                ext=ext,
                text_excerpt=text_excerpt or (f.text_excerpt or ""),
                filename=filename,
            )
        except Exception:
            pass
        db.session.flush()
        # Базовые теги из полей
        if f.material_type:
            upsert_tag(f, "тип", f.material_type)
        if f.author:
            upsert_tag(f, "автор", f.author)
        if f.year:
            upsert_tag(f, "год", str(f.year))
        if refresh_task:
            try:
                refresh_task.status = 'completed'
                refresh_task.progress = 1.0
                refresh_task.finished_at = datetime.utcnow()
                refresh_task.error = None
                db.session.add(refresh_task)
            except Exception:
                logger.warning('Не удалось обновить статус задачи refresh #%s', refresh_task.id if refresh_task else None)
        db.session.commit()

        data = file_to_dict(f)
        data["abstract"] = f.abstract
        data["text_excerpt"] = f.text_excerpt
        return jsonify({"ok": True, "file": data, "log": log})
    except Exception as e:
        db.session.rollback()
        log.append(f"exception: {e}")
        if refresh_task:
            try:
                task_obj = TaskRecord.query.get(refresh_task.id)
                if task_obj:
                    task_obj.status = 'error'
                    task_obj.error = str(e)
                    task_obj.finished_at = datetime.utcnow()
                    db.session.commit()
            except Exception:
                db.session.rollback()
        return jsonify({"ok": False, "error": str(e), "log": log}), 500

# ------------------- Filename Suggestion & Rename -------------------
def _safe_filename_component(s: str, max_len: int = 80) -> str:
    s = (s or '').strip()
    # Заменим небезопасные символы
    s = re.sub(r"[<>:\\/\\|?*\r\n\t\0]", " ", s)
    # Уберём кавычки
    s = s.replace('"', ' ').replace("'", ' ')
    # Превратим последовательности пробелов/точек в один пробел
    s = re.sub(r"[ ]+", " ", s)
    s = s.strip().strip('.')
    # Замена пробелов на подчёркивания
    s = s.replace(' ', '_')
    # Схлопнем повторные знаки подчёркивания
    s = re.sub(r"_+", "_", s)
    # Ограничение длины
    if len(s) > max_len:
        s = s[:max_len]
    return s or "file"

def _extract_lastname(full: str) -> str:
    a = (full or '').strip()
    if not a:
        return ''
    # несколько авторов — берём первого
    a = re.split(r"[,&;/]|\band\b", a, flags=re.I)[0]
    # удалим инициалы вида "И.И." и одиночные
    a = re.sub(r"\b[А-ЯA-Z]\.[А-ЯA-Z]\.?", "", a)
    a = re.sub(r"\b[А-ЯA-Z]\.\b", "", a)
    parts = [p for p in re.split(r"\s+", a) if p]
    if not parts:
        return ''
    # Русские ФИО часто в формате: Фамилия Имя Отчество — берём 1-ю
    cyr = all(re.match(r"^[А-Яа-я\-]+$", p) for p in parts)
    if cyr and len(parts) >= 2:
        # если 3 слова и последнее похоже на отчество (окончания -вич/-вна), используем первое
        if len(parts) >= 3 and re.search(r"(вич|вна|вны|ызы)$", parts[-1], flags=re.I):
            return parts[0]
        # две части: обычно Фамилия Имя — первое
        return parts[0]
    # Иначе — лучшее эвристическое: последнее слово
    return parts[-1]

def _mt_abbr(mt: str) -> str:
    m = (mt or '').lower()
    return {
        'dissertation': 'ДИС', 'dissertation_abstract': 'АВТ',
        'article': 'СТ', 'textbook': 'УЧ', 'monograph': 'МОНО',
        'report': 'ОТЧ', 'patent': 'ПАТ', 'presentation': 'ПРЕЗ',
        'proceedings': 'ТЕЗ', 'standard': 'СТД', 'note': 'ЗАМ',
        'document': 'ДОК', 'audio': 'АУД', 'image': 'ИЗО'
    }.get(m, 'ДОК')

def _degree_abbr(file_obj: File) -> str:
    # по тегу 'степень'
    deg = ''
    try:
        for t in file_obj.tags:
            if (t.key or '').lower() == 'степень':
                deg = (t.value or '').lower()
                break
    except Exception:
        pass
    if 'доктор' in deg:
        return 'ДН'
    if 'кандид' in deg:
        return 'КН'
    return ''

def _build_suggested_basename(file_obj: File) -> str:
    mt = (file_obj.material_type or '').lower()
    abbr = _mt_abbr(mt)
    ctx = {
        'abbr': abbr,
        'degree': _degree_abbr(file_obj),
        'title': (file_obj.title or '').strip(),
        'author_last': _extract_lastname(file_obj.author or ''),
        'year': (file_obj.year or '').strip(),
        'filename': file_obj.filename or '',
    }
    # Используем шаблоны из настроек, если заданы
    pattern = RENAME_PATTERNS.get(mt) or RENAME_PATTERNS.get('default')
    base = None
    if pattern:
        try:
            base = pattern.format(**ctx)
        except Exception:
            base = None
    if not base:
        # Запасной вариант (жёстко заданные правила)
        if mt in ('dissertation','dissertation_abstract'):
            base = f"{abbr}.{(ctx['degree'] + '.') if ctx['degree'] else ''}{ctx['title']}.{ctx['author_last']}"
        elif mt == 'article':
            base = f"СТ.{ctx['title']}.{ctx['author_last'] or ctx['year']}"
        elif mt == 'textbook':
            base = f"УЧ.{ctx['title']}.{ctx['author_last'] or ctx['year']}"
        elif mt == 'monograph':
            base = f"МОНО.{ctx['title']}.{ctx['author_last'] or ctx['year']}"
        elif mt == 'image':
            base = f"ИЗО.{ctx['title'] or ctx['filename']}"
        elif mt == 'audio':
            base = f"АУД.{ctx['title'] or ctx['filename']}"
        else:
            base = f"{abbr}.{ctx['title'] or ctx['filename']}.{ctx['author_last'] or ctx['year']}"

    parts = [p for p in [ _safe_filename_component(x) for x in (base or '').split('.') ] if p]
    name = '.'.join(parts) if parts else _safe_filename_component(file_obj.filename)
    return name[:120]


def _determine_target_dir_for_file(file_obj: File, current_path: Path, material_type: str) -> Path:
    col = getattr(file_obj, 'collection', None)
    if not isinstance(col, Collection):
        col = _get_collection_instance(file_obj.collection_id)
    material_type = (material_type or '').strip().lower()
    fallback_sub = TYPE_DIRS.get(material_type) or TYPE_DIRS.get('other', 'other')
    if COLLECTIONS_IN_SEPARATE_DIRS:
        base_root = _collection_root_dir(col)
        if col and COLLECTION_TYPE_SUBDIRS:
            return base_root / fallback_sub
        return base_root
    if MOVE_ON_RENAME:
            return _scan_root_path() / fallback_sub
    return current_path.parent


def _rename_file_record(file_obj: File, base_name: str | None = None) -> Path:
    p_old = Path(file_obj.path)
    if not p_old.exists():
        raise FileNotFoundError(f"Файл {file_obj.path} не найден")
    base = (base_name or _build_suggested_basename(file_obj)) or file_obj.filename or p_old.stem
    base = _safe_filename_component(base, max_len=120)
    ext = file_obj.ext or p_old.suffix
    mt = (file_obj.material_type or '').strip().lower()
    target_dir = _determine_target_dir_for_file(file_obj, p_old, mt)
    target_dir.mkdir(parents=True, exist_ok=True)
    p_new = target_dir / (base + (ext or ''))
    i = 1
    while p_new.exists() and p_new != p_old:
        p_new = target_dir / (f"{base}_{i}" + (ext or ''))
        i += 1
    renamed = p_new != p_old
    if renamed:
        p_old.rename(p_new)
        if (file_obj.ext or '').lower() == '.pdf':
            try:
                old_thumb = Path(app.static_folder) / 'thumbnails' / (p_old.stem + '.png')
                if old_thumb.exists():
                    old_thumb.unlink()
            except Exception:
                pass
    file_obj.path = str(p_new)
    try:
        file_obj.rel_path = str(p_new.relative_to(_scan_root_path()))
    except Exception:
        file_obj.rel_path = p_new.name
    file_obj.filename = p_new.stem
    try:
        file_obj.mtime = p_new.stat().st_mtime
    except Exception:
        pass
    if renamed:
        try:
            db.session.add(ChangeLog(file_id=file_obj.id, action='rename', field='filename', old_value=p_old.name, new_value=p_new.name))
        except Exception:
            pass
    return p_new

@app.route('/api/files/<int:file_id>/rename-suggest', methods=['GET'])
def api_rename_suggest(file_id):
    f = File.query.get_or_404(file_id)
    try:
        suggested = _build_suggested_basename(f)
        return jsonify({"ok": True, "suggested": suggested, "ext": f.ext or ''})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route('/api/files/<int:file_id>/rename', methods=['POST'])
def api_rename_apply(file_id):
    f = File.query.get_or_404(file_id)
    data = request.json or {}
    base = (data.get('base') or '').strip()
    if not base:
        base = _build_suggested_basename(f)
    base = _safe_filename_component(base, max_len=120)
    try:
        p_new = _rename_file_record(f, base)
        db.session.commit()
        return jsonify({"ok": True, "new_name": Path(p_new).name, "rel_path": f.rel_path})
    except Exception as e:
        db.session.rollback()
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route('/api/collections/<int:collection_id>/rename-all', methods=['POST'])
@require_admin
def api_collection_rename_all(collection_id: int):
    col = Collection.query.get_or_404(collection_id)
    files = File.query.filter(File.collection_id == collection_id).all()
    renamed = 0
    errors: list[dict[str, str]] = []
    for f in files:
        try:
            _rename_file_record(f)
            db.session.commit()
            renamed += 1
        except Exception as exc:
            db.session.rollback()
            errors.append({'id': str(f.id), 'error': str(exc)})
    try:
        _log_user_action(_load_current_user(), 'collection_rename_all', 'collection', col.id, detail=json.dumps({'renamed': renamed, 'errors': len(errors)}))
    except Exception:
        pass
    return jsonify({'ok': True, 'renamed': renamed, 'errors': errors})


@app.route('/api/collections/<int:collection_id>/clear', methods=['POST'])
@require_admin
def api_collection_clear(collection_id: int):
    col = Collection.query.get_or_404(collection_id)
    removed_files, errors = _delete_collection_files(collection_id)
    try:
        db.session.commit()
    except Exception as exc:
        db.session.rollback()
        return jsonify({'ok': False, 'error': str(exc)}), 500
    actor = _load_current_user()
    try:
        detail = json.dumps({'removed_files': removed_files, 'errors': len(errors)})
    except Exception:
        detail = None
    _log_user_action(actor, 'collection_clear', 'collection', col.id, detail=detail)
    if errors:
        app.logger.warning(f"[collections] clear reported {len(errors)} issues for collection {collection_id}: {errors[:3]}")
    return jsonify({'ok': True, 'removed_files': removed_files, 'errors': errors})


@app.route('/api/collections/<int:collection_id>/rag/delete', methods=['POST'])
@require_admin
def api_collection_rag_delete(collection_id: int):
    col = Collection.query.get_or_404(collection_id)
    try:
        stats = _delete_rag_index_for_collection(collection_id)
        db.session.commit()
    except Exception as exc:
        db.session.rollback()
        app.logger.exception("Failed to delete RAG index for collection %s", collection_id)
        return jsonify({'ok': False, 'error': f'Не удалось удалить RAG индекс: {exc}'}), 500
    actor = _load_current_user()
    try:
        detail = json.dumps(stats, ensure_ascii=False)
    except Exception:
        detail = None
    _log_user_action(actor, 'collection_rag_delete', 'collection', col.id, detail=detail)
    if stats.get('removed_documents'):
        message = f"Удалено RAG документов: {stats['removed_documents']}, чанков: {stats.get('removed_chunks', 0)}"
    else:
        message = 'RAG индекс для коллекции не найден.'
    stats.update({'ok': True, 'message': message})
    return jsonify(stats)


@app.route('/api/collections/<int:collection_id>/rag/reindex', methods=['POST'])
@require_admin
def api_collection_rag_reindex(collection_id: int):
    col = Collection.query.get_or_404(collection_id)
    try:
        file_count = db.session.query(func.count(File.id)).filter(File.collection_id == collection_id).scalar() or 0
    except Exception as exc:
        return jsonify({'ok': False, 'error': f'Не удалось подсчитать файлы: {exc}'}), 500
    if file_count == 0:
        return jsonify({'ok': False, 'error': 'В коллекции нет файлов для RAG', 'error_code': 'collection_empty'}), 400

    data = request.get_json(silent=True) or {}
    runtime = runtime_settings_store.current
    default_backend = (runtime.rag_embedding_backend or 'lm-studio').strip().lower() or 'lm-studio'
    default_model = runtime.rag_embedding_model or 'nomic-ai/nomic-embed-text-v1.5-GGUF'
    default_dim = max(8, int(runtime.rag_embedding_dim or 768))
    default_batch = max(1, int(runtime.rag_embedding_batch_size or 32))
    default_device = runtime.rag_embedding_device
    default_endpoint = runtime.rag_embedding_endpoint or runtime.lmstudio_api_base or ''
    default_api_key = runtime.rag_embedding_api_key or runtime.lmstudio_api_key or ''

    def _int_opt(value: object, default: int) -> int:
        try:
            return int(value)
        except (TypeError, ValueError):
            return default

    def _bool_opt(value: object, default: bool) -> bool:
        if value is None:
            return default
        if isinstance(value, bool):
            return value
        if isinstance(value, (int, float)):
            return bool(value)
        return str(value).strip().lower() in ('1', 'true', 'yes', 'on')

    chunk_max_tokens = max(16, _int_opt(data.get('max_tokens') or data.get('chunk_max_tokens'), 700))
    chunk_overlap = max(0, _int_opt(data.get('overlap') or data.get('chunk_overlap'), 120))
    chunk_min_tokens = max(1, _int_opt(data.get('min_tokens') or data.get('chunk_min_tokens'), 80))
    limit_chars = max(1000, _int_opt(data.get('limit_chars'), 120_000))
    skip_if_unchanged = _bool_opt(data.get('skip_if_unchanged'), True)
    normalizer_version = str(data.get('normalizer_version') or 'v1').strip() or 'v1'

    embedding_backend_raw = data.get('embedding_backend') or data.get('backend') or default_backend
    embedding_backend = str(embedding_backend_raw).strip().lower() or default_backend
    embedding_model_name = str(data.get('embedding_model_name') or data.get('model_name') or default_model).strip() or default_model
    embedding_dim = max(8, _int_opt(data.get('embedding_dim'), default_dim))
    embedding_batch_size = max(1, _int_opt(data.get('embedding_batch_size') or data.get('batch_size'), default_batch))
    embedding_normalize = _bool_opt(data.get('embedding_normalize'), True)
    device_value = data.get('embedding_device') or data.get('device')
    embedding_device = None
    if isinstance(device_value, str):
        trimmed = device_value.strip()
        embedding_device = trimmed or None
    if embedding_device is None and default_device:
        embedding_device = default_device
    endpoint_value = data.get('embedding_endpoint') or data.get('endpoint')
    if endpoint_value is None:
        endpoint_value = default_endpoint
    embedding_endpoint = str(endpoint_value or "").strip()
    api_key_raw = data.get('embedding_api_key') or data.get('api_key')
    if api_key_raw is None:
        api_key_raw = default_api_key
    embedding_api_key = str(api_key_raw or "")

    user = _load_current_user()

    options = {
        "chunk_max_tokens": chunk_max_tokens,
        "chunk_overlap": chunk_overlap,
        "chunk_min_tokens": chunk_min_tokens,
        "limit_chars": limit_chars,
        "skip_if_unchanged": skip_if_unchanged,
        "normalizer_version": normalizer_version,
        "embedding_backend": embedding_backend,
        "embedding_model_name": embedding_model_name,
        "embedding_dim": embedding_dim,
        "embedding_batch_size": embedding_batch_size,
        "embedding_normalize": embedding_normalize,
        "embedding_device": embedding_device,
        "embedding_endpoint": embedding_endpoint,
        "embedding_api_key": embedding_api_key,
        "user_id": getattr(user, 'id', None),
    }

    payload = {
        "collection_id": collection_id,
        "collection_name": col.name,
        "total_files": int(file_count),
        "status": "queued",
        "options": {
            "chunk_max_tokens": chunk_max_tokens,
            "chunk_overlap": chunk_overlap,
            "chunk_min_tokens": chunk_min_tokens,
            "limit_chars": limit_chars,
            "skip_if_unchanged": skip_if_unchanged,
            "embedding_backend": embedding_backend,
            "embedding_model": embedding_model_name,
            "embedding_endpoint": embedding_endpoint,
        },
    }

    task = TaskRecord(
        name='rag_collection',
        status='queued',
        payload=json.dumps(payload, ensure_ascii=False),
        progress=0.0,
        user_id=getattr(user, 'id', None),
        collection_id=collection_id,
    )
    try:
        db.session.add(task)
        db.session.commit()
    except Exception as exc:
        db.session.rollback()
        return jsonify({'ok': False, 'error': f'Не удалось создать задачу: {exc}'}), 500

    try:
        detail = json.dumps({"task_id": task.id, "files": int(file_count)}, ensure_ascii=False)
        _log_user_action(user, 'collection_rag_reindex_enqueued', 'collection', collection_id, detail=detail)
    except Exception:
        pass

    @copy_current_app_context
    def _runner():
        _run_rag_collection_job(task.id, collection_id, options)

    get_task_queue().submit(
        _runner,
        description=f"rag-collection-{collection_id}",
        owner_id=getattr(user, 'id', None),
    )

    return jsonify({'ok': True, 'task_id': task.id, 'files': int(file_count)})


@app.route('/api/files/<int:file_id>/rag/reindex', methods=['POST'])
@require_admin
def api_file_rag_reindex(file_id: int):
    file_obj = File.query.get_or_404(file_id)
    data = request.get_json(silent=True) or {}
    runtime = runtime_settings_store.current
    payload = {
        "file_id": int(file_id),
        "collection_id": file_obj.collection_id,
        "status": "queued",
        "options": {
            "skip_if_unchanged": bool(data.get("skip_if_unchanged", True)),
            "chunk_max_tokens": int(data.get("chunk_max_tokens") or 700),
            "chunk_overlap": int(data.get("chunk_overlap") or 120),
            "chunk_min_tokens": int(data.get("chunk_min_tokens") or 80),
            "embedding_backend": str(data.get("embedding_backend") or runtime.rag_embedding_backend or "auto"),
            "embedding_model_name": str(data.get("embedding_model_name") or runtime.rag_embedding_model or "intfloat/multilingual-e5-large"),
        },
    }
    task = TaskRecord(
        name='rag_file',
        status='queued',
        progress=0.0,
        payload=json.dumps(payload, ensure_ascii=False),
        user_id=getattr(_load_current_user(), 'id', None),
        collection_id=getattr(file_obj, 'collection_id', None),
    )
    try:
        db.session.add(task)
        db.session.commit()
    except Exception as exc:
        db.session.rollback()
        return jsonify({'ok': False, 'error': f'Не удалось создать задачу: {exc}'}), 500

    options = {
        "skip_if_unchanged": bool(data.get("skip_if_unchanged", True)),
        "limit_chars": int(data.get("limit_chars") or 120_000),
        "chunk_max_tokens": int(data.get("chunk_max_tokens") or 700),
        "chunk_overlap": int(data.get("chunk_overlap") or 120),
        "chunk_min_tokens": int(data.get("chunk_min_tokens") or 80),
        "normalizer_version": str(data.get("normalizer_version") or "v1"),
        "embedding_backend": str(data.get("embedding_backend") or runtime.rag_embedding_backend or "auto"),
        "embedding_model_name": str(data.get("embedding_model_name") or runtime.rag_embedding_model or "intfloat/multilingual-e5-large"),
        "embedding_dim": int(data.get("embedding_dim") or runtime.rag_embedding_dim or 384),
        "embedding_batch_size": int(data.get("embedding_batch_size") or runtime.rag_embedding_batch_size or 32),
        "embedding_device": data.get("embedding_device") or runtime.rag_embedding_device,
        "embedding_endpoint": str(data.get("embedding_endpoint") or runtime.rag_embedding_endpoint or runtime.lmstudio_api_base or ""),
        "embedding_api_key": str(data.get("embedding_api_key") or runtime.rag_embedding_api_key or runtime.lmstudio_api_key or ""),
    }
    get_task_queue().submit(
        _run_rag_file_job,
        int(task.id),
        int(file_id),
        options,
        description=f"rag-file-{file_id}",
        owner_id=getattr(_load_current_user(), 'id', None),
    )
    try:
        actor = _load_current_user()
        _log_user_action(actor, 'file_rag_reindex_enqueued', 'file', file_id, detail=json.dumps({"task_id": int(task.id)}))
    except Exception:
        pass
    return jsonify({'ok': True, 'task_id': int(task.id), 'file_id': int(file_id)})


def _delete_collection_files(col_id: int, remove_fs: bool = True) -> tuple[int, list[str]]:
    files = File.query.filter(File.collection_id == col_id).all()
    removed = 0
    errors: list[str] = []
    static_folder = Path(app.static_folder or (BASE_DIR / 'static'))
    thumbs_dir = static_folder / 'thumbnails'
    cache_dir = static_folder / 'cache' / 'text_excerpts'
    for f in files:
        path = Path(f.path) if f.path else None
        if remove_fs and path and path.exists():
            try:
                path.unlink()
                removed += 1
            except Exception as exc:
                errors.append(f"fs:{path}:{exc}")
        rel_path = Path(f.rel_path) if f.rel_path else None
        if rel_path:
            try:
                thumb = thumbs_dir / (rel_path.stem + '.png')
                if thumb.exists():
                    thumb.unlink()
            except Exception as exc:
                errors.append(f"thumb:{rel_path}:{exc}")
        try:
            key_base = f.sha1 or (f.rel_path or '').replace('/', '_')
            if key_base:
                cache_file = cache_dir / f"{key_base}.txt"
                if cache_file.exists():
                    cache_file.unlink()
        except Exception as exc:
            errors.append(f"cache:{f.id}:{exc}")
        fid = f.id
        db.session.delete(f)
        try:
            _delete_file_from_fts(fid)
        except Exception as exc:
            errors.append(f"fts:{fid}:{exc}")
    return removed, errors


def _remove_collection_directory_path(path: Path | None) -> None:
    if path is None:
        return
    try:
        if path.exists():
            shutil.rmtree(path)
    except Exception:
        pass


@app.route('/api/collections/<int:collection_id>', methods=['DELETE'])
@require_admin
def api_delete_collection(collection_id: int):
    col = Collection.query.get_or_404(collection_id)
    collection_root: Path | None = None
    if COLLECTIONS_IN_SEPARATE_DIRS:
        try:
            collection_root = _collection_root_dir(col, ensure=False)
        except Exception:
            collection_root = None
    removed_files, errors = _delete_collection_files(collection_id)
    try:
        CollectionMember.query.filter_by(collection_id=collection_id).delete()
        db.session.delete(col)
        db.session.commit()
    except Exception as exc:
        db.session.rollback()
        return jsonify({'ok': False, 'error': str(exc)}), 500
    _remove_collection_directory_path(collection_root)
    try:
        _log_user_action(_load_current_user(), 'collection_delete', 'collection', collection_id, detail=json.dumps({'removed_files': removed_files, 'errors': len(errors)}))
    except Exception:
        pass
    return jsonify({'ok': True, 'removed_files': removed_files, 'errors': errors})

# ------------------- Scan with Progress -------------------
import threading, time

SCAN_PROGRESS = {
    "running": False,
    "stage": "idle",
    "total": 0,
    "processed": 0,
    "added": 0,
    "updated": 0,
    "duplicates": 0,
    "removed": 0,
    "current": "",
    "use_llm": False,
    "error": None,
    "started_at": None,
    "updated_at": None,
    "eta_seconds": None,
    "history": [],
    "task_id": None,
}
SCAN_CANCEL = False
SCAN_TASK_ID: int | None = None
SCAN_STATE_LOCK = threading.RLock()

def _resolve_deleted_dir(root: Path) -> Path | None:
    deleted_cfg = app.config.get('DELETED_FOLDER')
    try:
        trash_dir = Path(deleted_cfg).expanduser() if deleted_cfg else Path('_deleted')
    except Exception:
        trash_dir = Path('_deleted')
    if not trash_dir.is_absolute():
        try:
            return (root / trash_dir).resolve()
        except Exception:
            return root / trash_dir
    try:
        return trash_dir.resolve()
    except Exception:
        return trash_dir


def _iter_files_for_scan(root: Path):
    trash_dir = _resolve_deleted_dir(root)
    for path in root.rglob("*"):
        if trash_dir is not None:
            try:
                path.relative_to(trash_dir)
                continue
            except ValueError:
                pass
        if not path.is_file():
            continue
        ext = path.suffix.lower()
        if ext not in ALLOWED_EXTS:
            continue
        yield path


def _bool_from_form(form, key: str, default: bool) -> bool:
    if form is None:
        return bool(default)
    if key in form:
        val = str(form.get(key) or '').strip().lower()
        return val in ('1', 'true', 'yes', 'on')
    return bool(default)

def _purge_cached_artifacts(records: list[tuple[File, Path]]):
    """Удаляет кэшированные превью и миниатюры для указанных файлов."""
    if not records:
        return
    try:
        static_dir = Path(app.static_folder or 'static').resolve()
    except Exception:
        return
    txt_cache = static_dir / 'cache' / 'text_excerpts'
    thumbs_dir = static_dir / 'thumbnails'
    for file_obj, path in records:
        try:
            if txt_cache.exists():
                key_base = file_obj.sha1 or (file_obj.rel_path or '').replace('/', '_')
                if key_base:
                    fp = txt_cache / f"{key_base}.txt"
                    if fp.exists():
                        fp.unlink()
        except Exception:
            pass
        try:
            if path.suffix.lower() == '.pdf' and thumbs_dir.exists():
                thumb = thumbs_dir / f"{path.stem}.png"
                if thumb.exists():
                    thumb.unlink()
        except Exception:
            pass

MAX_LOG_LINES = 200

def _scan_log(msg: str, level: str = "info"):
    with SCAN_STATE_LOCK:
        SCAN_PROGRESS.setdefault("history", [])
        entry = {"t": time.time(), "level": level, "msg": str(msg)}
        SCAN_PROGRESS["history"].append(entry)
        if len(SCAN_PROGRESS["history"]) > MAX_LOG_LINES:
            SCAN_PROGRESS["history"] = SCAN_PROGRESS["history"][-MAX_LOG_LINES:]
        SCAN_PROGRESS["updated_at"] = time.time()
    log_line = f"[scan] {msg}"
    try:
        if level == 'error':
            app.logger.error(log_line)
        elif level in ('warn', 'warning'):
            app.logger.warning(log_line)
        else:
            app.logger.info(log_line)
    except Exception:
        pass

def _update_eta():
    try:
        with SCAN_STATE_LOCK:
            st = SCAN_PROGRESS
            total = int(st.get("total") or 0)
            processed = int(st.get("processed") or 0)
            started = st.get("started_at") or time.time()
        if processed <= 0 or total <= 0:
            with SCAN_STATE_LOCK:
                st["eta_seconds"] = None
            return
        elapsed = max(0.001, time.time() - float(started))
        rate = processed / elapsed
        remain = max(0, total - processed)
        eta = remain / rate if rate > 0 else None
        with SCAN_STATE_LOCK:
            st["eta_seconds"] = int(eta) if eta is not None else None
    except Exception:
        with SCAN_STATE_LOCK:
            st = SCAN_PROGRESS
            st["eta_seconds"] = None

def _find_duplicate_file_by_sha1(sha1: str, current_file_id: int | None = None) -> File | None:
    key = (sha1 or "").strip()
    if not key:
        return None
    try:
        q = File.query.filter(File.sha1 == key)
        if current_file_id:
            q = q.filter(File.id != current_file_id)
        return q.order_by(File.id.asc()).first()
    except Exception:
        return None


def _hydrate_from_duplicate(target: File, source: File) -> None:
    if not target or not source:
        return
    for field in ("title", "author", "year", "advisor", "keywords", "abstract"):
        if not getattr(target, field, None) and getattr(source, field, None):
            setattr(target, field, getattr(source, field))
    if not getattr(target, "material_type", None):
        src_type = getattr(source, "material_type", None)
        if src_type:
            target.material_type = src_type
    elif (target.material_type or "").strip().lower() == "document":
        src_type = (getattr(source, "material_type", None) or "").strip().lower()
        if src_type and src_type != "document":
            target.material_type = src_type
    if not getattr(target, "text_excerpt", None):
        src_excerpt = getattr(source, "text_excerpt", None)
        if src_excerpt:
            target.text_excerpt = src_excerpt
    try:
        for tag in Tag.query.filter_by(file_id=source.id).all():
            upsert_tag(target, tag.key, tag.value)
    except Exception:
        pass

def _run_scan_with_progress(extract_text: bool, use_llm: bool, prune: bool, skip: int = 0, targets: list | None = None):
    global SCAN_PROGRESS, SCAN_CANCEL, SCAN_TASK_ID
    with app.app_context():
        try:
            try:
                db.session.rollback()
            except Exception:
                db.session.remove()
            task = None
            if SCAN_TASK_ID:
                try:
                    task = TaskRecord.query.get(SCAN_TASK_ID)
                    if task:
                        task.status = 'running'
                        task.started_at = datetime.utcnow()
                        task.progress = 0.0
                        db.session.commit()
                except Exception:
                    db.session.rollback()
                    task = None
            with SCAN_STATE_LOCK:
                SCAN_PROGRESS.update({
                    "running": True,
                    "stage": "counting",
                    "processed": 0,
                    "added": 0,
                    "updated": 0,
                    "duplicates": 0,
                    "removed": 0,
                    "error": None,
                    "use_llm": bool(use_llm),
                    "started_at": time.time(),
                    "updated_at": time.time(),
                    "eta_seconds": None,
                    "history": [],
                })
            _scan_log("Начало сканирования")
            root = _scan_root_path()
            # определяем набор для сканирования: явные цели или полный корневой обход
            if targets:
                try:
                    file_list = [Path(p) for p in targets]
                except Exception:
                    file_list = []
                # фильтруем по разрешённым путям и наличию
                file_list = [p for p in file_list if p.exists() and p.suffix.lower() in ALLOWED_EXTS]
                _scan_log(f"Сканирование только добавленных файлов: {len(file_list)}")
            else:
                file_list = list(_iter_files_for_scan(root))
            total = len(file_list)
            with SCAN_STATE_LOCK:
                SCAN_PROGRESS["total"] = total
            if task:
                try:
                    task.progress = 0.0
                    task.payload = json.dumps({
                        "total": total,
                        "use_llm": bool(use_llm),
                        "extract_text": bool(extract_text),
                        "prune": bool(prune),
                    })
                    db.session.commit()
                except Exception:
                    db.session.rollback()
            # поддержка продолжения: пропускаем первые N файлов, если задано
            skip = max(0, min(int(skip or 0), total))
            if skip:
                _scan_log(f"Продолжение: пропуск первых {skip} из {total}")
                file_list = file_list[skip:]
                with SCAN_STATE_LOCK:
                    SCAN_PROGRESS["processed"] = skip
            _scan_log(f"Найдено файлов: {len(file_list)}")

            added = updated = duplicates = 0
            cancelled = False
            for idx, path in enumerate(file_list, start=1):
                with SCAN_STATE_LOCK:
                    cancelled_requested = bool(SCAN_CANCEL)
                if cancelled_requested:
                    _scan_log("Отмена пользователем", level="warn")
                    cancelled = True
                    break
                with SCAN_STATE_LOCK:
                    SCAN_PROGRESS.update({
                        "stage": "processing",
                        "processed": (skip + idx),
                        "current": str(path.name),
                        "updated_at": time.time()
                    })
                _update_eta()
                if task and total:
                    try:
                        task.progress = min(1.0, float(skip + idx) / float(total))
                        if idx == 1 or idx % 10 == 0:
                            db.session.commit()
                    except Exception:
                        db.session.rollback()
                if idx == 1 or idx % 10 == 0:
                    _scan_log(f"Обработка: {path.name}")

                ext = path.suffix.lower()
                # по возможности вычисляем относительный путь к SCAN_ROOT
                try:
                    rel_path = str(path.relative_to(root))
                except Exception:
                    rel_path = path.name
                size = path.stat().st_size
                mtime = path.stat().st_mtime
                filename = path.stem
                page_count: int | None = None

                file_obj = File.query.filter_by(path=str(path)).first()
                if not file_obj:
                    sha1 = sha1_of_file(path)
                    file_obj = File(path=str(path), rel_path=rel_path, filename=filename,
                                    ext=ext, size=size, mtime=mtime, sha1=sha1)
                    db.session.add(file_obj)
                    added += 1
                    with SCAN_STATE_LOCK:
                        SCAN_PROGRESS["added"] = added
                else:
                    if (file_obj.size != size) or (file_obj.mtime != mtime):
                        sha1 = sha1_of_file(path)
                        file_obj.sha1 = sha1
                        file_obj.size = size
                        file_obj.mtime = mtime
                        file_obj.filename = filename
                    updated += 1
                    with SCAN_STATE_LOCK:
                        SCAN_PROGRESS["updated"] = updated

                duplicate_source: File | None = _find_duplicate_file_by_sha1(file_obj.sha1 or "", file_obj.id)
                skip_expensive_processing = duplicate_source is not None
                if duplicate_source:
                    _hydrate_from_duplicate(file_obj, duplicate_source)
                    duplicates += 1
                    with SCAN_STATE_LOCK:
                        SCAN_PROGRESS["duplicates"] = duplicates
                    if idx == 1 or idx % 10 == 0:
                        _scan_log(
                            f"Дубликат по хешу: {path.name} -> id={duplicate_source.id}; пропуск OCR/LLM",
                            level="info",
                        )

                # Извлечение текста (на основе текущей логики)
                text_excerpt = ""
                if extract_text and not skip_expensive_processing:
                    if ext == ".pdf":
                        text_excerpt = extract_text_pdf(path, limit_chars=40000)
                    elif ext == ".docx":
                        text_excerpt = extract_text_docx(path, limit_chars=40000)
                    elif ext == ".rtf":
                        text_excerpt = extract_text_rtf(path, limit_chars=40000)
                    elif ext == ".epub":
                        text_excerpt = extract_text_epub(path, limit_chars=40000)
                    elif ext == ".djvu":
                        text_excerpt = extract_text_djvu(path, limit_chars=40000)
                    elif ext in AUDIO_EXTS:
                        _scan_log(f"Транскрибация аудио: {path.name}")
                        text_excerpt = transcribe_audio(path, limit_chars=40000)
                    if text_excerpt:
                        file_obj.text_excerpt = text_excerpt[:40000]
                    # Общие теги: расширение и страницы PDF
                    try:
                        upsert_tag(file_obj, 'ext', ext.lstrip('.'))
                    except Exception:
                        pass
                    if ext == '.pdf':
                        try:
                            with fitz.open(str(path)) as _doc:
                                page_count = len(_doc)
                                upsert_tag(file_obj, 'pages', str(len(_doc)))
                        except Exception:
                            pass
                    # Теги, специфичные для аудио
                    if ext in AUDIO_EXTS:
                        try:
                            upsert_tag(file_obj, 'формат', ext.lstrip('.'))
                            upsert_tag(file_obj, 'длительность', audio_duration_hhmmss(path))
                        except Exception:
                            pass
                        # Лёгкие ключевые слова из транскрипта через LLM
                        try:
                            if AUDIO_KEYWORDS_LLM and (file_obj.text_excerpt or '') and not (file_obj.keywords or '').strip():
                                kws = call_lmstudio_keywords(file_obj.text_excerpt, path.name)
                                if kws:
                                    file_obj.keywords = ", ".join(kws)
                        except Exception as _e:
                            _scan_log(f"audio keywords llm failed: {_e}", level="warn")
                    # Теги, специфичные для изображений
                    if ext in IMAGE_EXTS:
                        try:
                            upsert_tag(file_obj, 'формат', ext.lstrip('.'))
                            if PILImage is not None:
                                with PILImage.open(str(path)) as im:
                                    w, h = im.size
                                upsert_tag(file_obj, 'разрешение', f'{w}x{h}')
                                orient = 'портрет' if h >= w else 'альбом'
                                upsert_tag(file_obj, 'ориентация', orient)
                        except Exception:
                            pass
                        # Лёгкие ключевые слова из транскрипта через LLM
                        try:
                            if AUDIO_KEYWORDS_LLM and (file_obj.text_excerpt or '') and not (file_obj.keywords or '').strip():
                                kws = call_lmstudio_keywords(file_obj.text_excerpt, path.name)
                                if kws:
                                    file_obj.keywords = ", ".join(kws)
                        except Exception as _e:
                            _scan_log(f"audio keywords llm failed: {_e}", level="warn")
                if skip_expensive_processing:
                    text_excerpt = file_obj.text_excerpt or ""

                # Встроенные ключевые слова в тексте для диссертаций/статей
                try:
                    import re as _re
                    kw_m = _re.search(r"(?:Ключевые\s+слова|Keywords)\s*[:\-]\s*(.+)", (text_excerpt or ''), flags=_re.IGNORECASE)
                    if kw_m:
                        file_obj.keywords = kw_m.group(1).strip()[:2000]
                        if KEYWORDS_TO_TAGS_ENABLED:
                            db.session.flush(); _upsert_keyword_tags(file_obj)
                except Exception:
                    pass

                # Эвристики на основе имени файла
                title, author, year = None, None, None
                for pat in FILENAME_PATTERNS:
                    m = pat.match(filename)
                    if m:
                        gd = m.groupdict()
                        title = gd.get("title") or title
                        author = gd.get("author") or author
                        year = gd.get("year") or year
                        break

                if title and not file_obj.title:
                    file_obj.title = title
                if author and not file_obj.author:
                    file_obj.author = author
                if year and not file_obj.year:
                    file_obj.year = year

                if not file_obj.material_type:
                    cand = _detect_type_pre_llm(ext, text_excerpt, filename)
                    if cand:
                        file_obj.material_type = cand
                if ext in AUDIO_EXTS and (file_obj.material_type or '') == 'document':
                    file_obj.material_type = 'audio'
                if ext in IMAGE_EXTS and (file_obj.material_type or '') == 'document':
                    file_obj.material_type = 'image'

                journal_entries = _extract_journal_toc_entries(text_excerpt or '') if text_excerpt else []
                file_obj.material_type = _journal_safety_check(file_obj.material_type, text_excerpt, page_count, journal_entries)

                # Типо-зависимые теги (до LLM)
                try:
                    ttags = extract_tags_for_type(file_obj.material_type or '', text_excerpt or '', filename)
                    if ttags:
                        db.session.flush()
                        for k, v in ttags.items():
                            values = v if isinstance(v, (list, tuple, set)) else [v]
                            for single in values:
                                if single is None:
                                    continue
                                upsert_tag(file_obj, k, single)
                except Exception as e:
                    _scan_log(f"type tags error: {e}", level="warn")
                # Дополнительные расширенные теги
                try:
                    rtags = extract_richer_tags(file_obj.material_type or '', text_excerpt or '', filename)
                    if rtags:
                        db.session.flush()
                        for k, v in rtags.items():
                            values = v if isinstance(v, (list, tuple, set)) else [v]
                            for single in values:
                                if single is None:
                                    continue
                                upsert_tag(file_obj, k, single)
                except Exception as e:
                    _scan_log(f"richer tags error: {e}", level="warn")

                # Обогащение через LLM (может быть медленным)
                if (not skip_expensive_processing) and use_llm and (text_excerpt or ext in {'.txt', '.md'} or ext in IMAGE_EXTS or ext in {'.pdf','.docx','.rtf','.epub','.djvu'}):
                    with SCAN_STATE_LOCK:
                        SCAN_PROGRESS["stage"] = "llm"
                        SCAN_PROGRESS["updated_at"] = time.time()
                    _scan_log(f"LLM-анализ: {path.name}")
                    llm_text = text_excerpt or ""
                    if not llm_text and ext in {".txt", ".md"}:
                        try:
                            llm_text = Path(path).read_text(encoding="utf-8", errors="ignore")[:15000]
                        except Exception:
                            pass
                    # Лёгкое извлечение для распространённых форматов документов, если extract_text выключен
                    if not llm_text and ext in {'.pdf','.docx','.rtf','.epub','.djvu'}:
                        try:
                            if ext == '.pdf':
                                llm_text = extract_text_pdf(
                                    path,
                                    limit_chars=12000,
                                    force_ocr_first_page=(
                                        _always_ocr_first_page_dissertation()
                                        and looks_like_dissertation_filename(filename)
                                    ),
                                )
                            elif ext == '.docx':
                                llm_text = extract_text_docx(path, limit_chars=12000)
                            elif ext == '.rtf':
                                llm_text = extract_text_rtf(path, limit_chars=12000)
                            elif ext == '.epub':
                                llm_text = extract_text_epub(path, limit_chars=12000)
                            elif ext == '.djvu':
                                llm_text = extract_text_djvu(path, limit_chars=12000)
                        except Exception as _e:
                            _scan_log(f"llm lightweight extract error: {_e}", level='warn')
                    meta = call_lmstudio_for_metadata(llm_text, path.name)
                    if meta:
                        mt_meta = normalize_material_type(meta.get("material_type"))
                        # Для очевидных форматов не даём LLM переопределять тип
                        if ext in AUDIO_EXTS:
                            file_obj.material_type = 'audio'
                        elif ext in IMAGE_EXTS:
                            file_obj.material_type = 'image'
                        else:
                            merged_type = _merge_llm_material_type(
                                file_obj.material_type,
                                mt_meta,
                                allow_override=bool(TYPE_LLM_OVERRIDE),
                            )
                            if merged_type:
                                file_obj.material_type = merged_type
                        _t = (meta.get("title") or "").strip()
                        if _t:
                            file_obj.title = _t
                        _a = _normalize_author(meta.get("author"))
                        if _a:
                            file_obj.author = _a
                        _y = _normalize_year(meta.get("year"))
                        if _y:
                            file_obj.year = _y
                        _adv = meta.get("advisor")
                        if _adv is not None:
                            _adv_s = str(_adv).strip()
                            if _adv_s:
                                file_obj.advisor = _adv_s
                        kws = meta.get("keywords") or []
                        if isinstance(kws, list):
                            file_obj.keywords = ", ".join([str(x) for x in kws][:50])
                        if meta.get("novelty"):
                            db.session.flush()
                            upsert_tag(file_obj, "научная новизна", str(meta.get("novelty")))
                        for key in ("literature", "organizations", "classification"):
                            val = meta.get(key)
                            if isinstance(val, list) and val:
                                db.session.flush()
                                upsert_tag(file_obj, key, "; ".join([str(x) for x in val]))
                        file_obj.material_type = _journal_safety_check(file_obj.material_type, llm_text or text_excerpt, page_count)
                    # Резюме для аудио
                    if ext in AUDIO_EXTS and SUMMARIZE_AUDIO and (file_obj.text_excerpt or ''):
                        summ = call_lmstudio_summarize(file_obj.text_excerpt, path.name)
                        if summ:
                            file_obj.abstract = summ[:2000]
                    # Анализ изображений через Vision
                    if ext in IMAGE_EXTS and IMAGES_VISION_ENABLED:
                        try:
                            vis = call_lmstudio_vision(path, path.name)
                            if isinstance(vis, dict):
                                desc = (vis.get('description') or '')
                                if desc:
                                    file_obj.abstract = desc[:8000]
                                kws = vis.get('keywords') or []
                                if isinstance(kws, list) and kws:
                                    file_obj.keywords = ", ".join([str(x) for x in kws][:50])
                        except Exception as e:
                            _scan_log(f"vision error: {e}", level="warn")
                    # Повторный прогон типовых тегов
                    try:
                        ttags = extract_tags_for_type(file_obj.material_type or '', text_excerpt or '', filename)
                        if ttags:
                            db.session.flush()
                            for k, v in ttags.items():
                                values = v if isinstance(v, (list, tuple, set)) else [v]
                                for single in values:
                                    if single is None:
                                        continue
                                    upsert_tag(file_obj, k, single)
                    except Exception as e:
                        _scan_log(f"type tags error(2): {e}", level="warn")
                    # Добавочные расширенные теги после LLM
                    try:
                        rtags = extract_richer_tags(file_obj.material_type or '', text_excerpt or '', filename)
                        if rtags:
                            db.session.flush()
                            for k, v in rtags.items():
                                values = v if isinstance(v, (list, tuple, set)) else [v]
                                for single in values:
                                    if single is None:
                                        continue
                                    upsert_tag(file_obj, k, single)
                    except Exception as e:
                        _scan_log(f"richer tags error(2): {e}", level="warn")

                try:
                    _apply_metadata_quality_rules(
                        file_obj,
                        ext=ext,
                        text_excerpt=text_excerpt or (file_obj.text_excerpt or ""),
                        filename=filename,
                    )
                except Exception:
                    pass
                db.session.flush()
                # Базовые теги
                if file_obj.material_type:
                    upsert_tag(file_obj, "тип", file_obj.material_type)
                if file_obj.author:
                    upsert_tag(file_obj, "автор", file_obj.author)
                if file_obj.year:
                    upsert_tag(file_obj, "год", str(file_obj.year))
                db.session.flush()
                try:
                    _sync_file_to_fts(file_obj)
                except Exception as sync_exc:
                    _scan_log(f"fts sync failed: {sync_exc}", level='warn')
                db.session.commit()

            removed = 0
            with SCAN_STATE_LOCK:
                should_prune = bool(prune and not SCAN_CANCEL)
            if should_prune:
                with SCAN_STATE_LOCK:
                    SCAN_PROGRESS["stage"] = "prune"
                    SCAN_PROGRESS["updated_at"] = time.time()
                _scan_log("Удаление отсутствующих файлов")
                removed = prune_missing_files()
                try:
                    _rebuild_files_fts()
                    _rebuild_tags_fts()
                except Exception as rebuild_exc:
                    _scan_log(f"fts rebuild failed: {rebuild_exc}", level='warn')
                db.session.commit()
            with SCAN_STATE_LOCK:
                SCAN_PROGRESS.update({"removed": removed, "stage": "done", "running": False, "updated_at": time.time()})
            _scan_log("Сканирование завершено")
            if task:
                try:
                    task.status = 'cancelled' if cancelled else 'completed'
                    task.finished_at = datetime.utcnow()
                    task.progress = 1.0 if not cancelled else task.progress
                    db.session.commit()
                except Exception:
                    db.session.rollback()
        except Exception as e:
            with SCAN_STATE_LOCK:
                SCAN_PROGRESS.update({"error": str(e), "running": False, "stage": "error", "updated_at": time.time()})
            _scan_log(f"Ошибка: {e}", level="error")
            if task:
                try:
                    task.status = 'error'
                    task.error = str(e)
                    task.finished_at = datetime.utcnow()
                    db.session.commit()
                except Exception:
                    db.session.rollback()
        finally:
            with SCAN_STATE_LOCK:
                SCAN_CANCEL = False
                SCAN_TASK_ID = None
                SCAN_PROGRESS["task_id"] = None
                SCAN_PROGRESS["scope"] = None
            try:
                db.session.remove()
            except Exception:
                pass

@app.route("/scan/start", methods=["POST"])
@require_admin
def scan_start():
    global SCAN_CANCEL, SCAN_TASK_ID
    with SCAN_STATE_LOCK:
        running = bool(SCAN_PROGRESS.get("running"))
    if running:
        return jsonify({"status": "busy"}), 409
    extract_text = _bool_from_form(request.form, "extract_text", EXTRACT_TEXT)
    use_llm = _bool_from_form(request.form, "use_llm", DEFAULT_USE_LLM)
    prune = _bool_from_form(request.form, "prune", DEFAULT_PRUNE)
    with SCAN_STATE_LOCK:
        SCAN_PROGRESS["scope"] = {
            "type": "library",
            "label": "Вся библиотека",
            "extract_text": bool(extract_text),
            "use_llm": bool(use_llm),
            "prune": bool(prune),
        }
        SCAN_CANCEL = False
    skip = 0
    try:
        skip = int(request.form.get('skip', '0') or 0)
    except Exception:
        skip = 0
    # Очищаем статические кэши перед полным сканом: текстовые фрагменты и миниатюры
    try:
        static_dir = Path(app.static_folder)
        txt_cache = static_dir / 'cache' / 'text_excerpts'
        if txt_cache.exists():
            for fp in txt_cache.glob('*.txt'):
                try: fp.unlink()
                except Exception: pass
        thumbs = static_dir / 'thumbnails'
        if thumbs.exists():
            for fp in thumbs.glob('*.png'):
                try: fp.unlink()
                except Exception: pass
    except Exception:
        pass
    payload = {
        "extract_text": extract_text,
        "use_llm": use_llm,
        "prune": prune,
        "skip": skip,
    }
    try:
        task = TaskRecord(
            name='scan',
            status='queued',
            payload=json.dumps(payload),
            progress=0.0,
            user_id=getattr(_load_current_user(), 'id', None),
            collection_id=None,
        )
        db.session.add(task)
        db.session.commit()
        with SCAN_STATE_LOCK:
            SCAN_TASK_ID = task.id
            SCAN_PROGRESS["task_id"] = task.id
    except Exception:
        db.session.rollback()
        with SCAN_STATE_LOCK:
            SCAN_TASK_ID = None
            SCAN_PROGRESS["task_id"] = None
    _log_user_action(_load_current_user(), 'scan_start', 'scan', SCAN_TASK_ID, detail=json.dumps(payload))
    get_task_queue().submit(
        _run_scan_with_progress,
        extract_text,
        use_llm,
        prune,
        skip,
        description="scan-full",
        owner_id=getattr(_load_current_user(), 'id', None),
    )
    return jsonify({"status": "started"})


@app.route("/scan/collection/<int:collection_id>", methods=["POST"])
@require_admin
def scan_collection(collection_id: int):
    global SCAN_CANCEL, SCAN_TASK_ID
    with SCAN_STATE_LOCK:
        running = bool(SCAN_PROGRESS.get("running"))
    if running:
        return jsonify({"status": "busy"}), 409
    collection = Collection.query.get(collection_id)
    if not collection:
        return jsonify({"status": "not_found", "error": "collection_not_found"}), 404
    extract_text = _bool_from_form(request.form, "extract_text", EXTRACT_TEXT)
    use_llm = _bool_from_form(request.form, "use_llm", DEFAULT_USE_LLM)
    prune = _bool_from_form(request.form, "prune", DEFAULT_PRUNE)
    files = File.query.filter(File.collection_id == collection_id).all()
    if not files:
        return jsonify({"status": "empty", "collection_id": collection_id, "files": 0})

    targets: list[str] = []
    records: list[tuple[File, Path]] = []
    missing = 0
    for file_obj in files:
        path_str = (file_obj.path or '').strip()
        path = Path(path_str) if path_str else None
        if not path or not path.exists():
            fallback = None
            if file_obj.rel_path:
                try:
                    fallback = (_scan_root_path() / Path(file_obj.rel_path)).resolve()
                except Exception:
                    fallback = None
            if fallback and fallback.exists():
                path = fallback
            else:
                missing += 1
                continue
        if not path.is_file():
            missing += 1
            continue
        ext = path.suffix.lower()
        if ext not in ALLOWED_EXTS:
            continue
        resolved = path.resolve()
        targets.append(str(resolved))
        records.append((file_obj, resolved))

    if not targets:
        return jsonify({"status": "empty", "collection_id": collection_id, "files": 0, "missing": missing})

    _purge_cached_artifacts(records)
    with SCAN_STATE_LOCK:
        SCAN_CANCEL = False
        SCAN_PROGRESS["scope"] = {
            "type": "collection",
            "collection_id": collection.id,
            "label": f"Коллекция «{collection.name}»",
            "count": len(targets),
            "extract_text": bool(extract_text),
            "use_llm": bool(use_llm),
            "prune": bool(prune),
        }

    payload = {
        "extract_text": extract_text,
        "use_llm": use_llm,
        "prune": prune,
        "collection_id": collection.id,
        "collection_name": collection.name,
        "targets": len(targets),
        "missing": missing,
    }
    try:
        task = TaskRecord(
            name='scan',
            status='queued',
            payload=json.dumps(payload),
            progress=0.0,
            user_id=getattr(_load_current_user(), 'id', None),
            collection_id=collection.id,
        )
        db.session.add(task)
        db.session.commit()
        with SCAN_STATE_LOCK:
            SCAN_TASK_ID = task.id
            SCAN_PROGRESS["task_id"] = task.id
    except Exception:
        db.session.rollback()
        with SCAN_STATE_LOCK:
            SCAN_TASK_ID = None
            SCAN_PROGRESS["task_id"] = None

    try:
        _log_user_action(_load_current_user(), 'scan_collection_start', 'collection', collection.id, detail=json.dumps(payload))
    except Exception:
        pass

    get_task_queue().submit(
        _run_scan_with_progress,
        extract_text,
        use_llm,
        prune,
        0,
        targets,
        description=f"scan-collection-{collection.id}",
        owner_id=getattr(_load_current_user(), 'id', None),
    )
    return jsonify({"status": "started", "files": len(targets), "missing": missing})

@app.route("/scan/status")
@require_admin
def scan_status():
    with SCAN_STATE_LOCK:
        snapshot = copy.deepcopy(SCAN_PROGRESS)
    return jsonify(snapshot)

@app.route("/scan/cancel", methods=["POST"])
@require_admin
def scan_cancel():
    global SCAN_CANCEL
    with SCAN_STATE_LOCK:
        SCAN_CANCEL = True
    try:
        if SCAN_TASK_ID:
            task = TaskRecord.query.get(SCAN_TASK_ID)
            if task and task.status not in ('completed','error','cancelled'):
                task.status = 'cancelling'
                db.session.commit()
    except Exception:
        db.session.rollback()
    _log_user_action(_load_current_user(), 'scan_cancel', 'scan', SCAN_TASK_ID)
    return jsonify({"status": "cancelling"})


# ------------------- AI Search (MVP) -------------------

def _normalize_keyword_candidate(raw: str) -> str:
    """Очистить строку, предназначенную для использования в качестве ключевого термина."""
    if not raw:
        return ""
    cleaned = re.sub(r"\s+", " ", str(raw)).strip()
    cleaned = cleaned.strip("\"'«»“”[](){}")
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    if not cleaned:
        return ""
    lowered = cleaned.lower()
    if len(lowered) < 2:
        return ""
    return lowered


def _extract_quoted_phrases(text: str) -> list[str]:
    """Извлечь фразы в кавычках и нормализовать их для использования как термов поиска."""
    if not text:
        return []
    pattern = r'"([^"\n]{2,})"|«([^»\n]{2,})»|“([^”\n]{2,})”|”([^“\n]{2,})“|ʼ([^ʼ\n]{2,})ʼ|\'([^\'\n]{2,})\''
    phrases: list[str] = []
    seen: set[str] = set()
    for match in re.finditer(pattern, str(text)):
        candidate = next((grp for grp in match.groups() if grp), "")
        normalized = _normalize_keyword_candidate(candidate)
        if not normalized:
            continue
        if normalized in seen:
            continue
        seen.add(normalized)
        phrases.append(normalized)
    return phrases


def _normalize_language_code(lang: str) -> str | None:
    if lang is None:
        return None
    code = re.sub(r"[^a-zA-Z\-]", "", str(lang)).lower()
    if not code:
        return None
    if "-" in code:
        code = code.split("-", 1)[0]
    if len(code) < 2:
        return None
    return code[:8]


def _guess_query_language(text: str) -> str | None:
    if not text:
        return None
    cyr = sum(1 for ch in text if ('а' <= ch.lower() <= 'я') or ch.lower() == 'ё')
    lat = sum(1 for ch in text if 'a' <= ch.lower() <= 'z')
    if cyr == 0 and lat == 0:
        return None
    return 'ru' if cyr >= lat else 'en'


def _ensure_russian_text(text: str, *, label: str = 'text') -> tuple[str, bool]:
    """
    Гарантировать русскую формулировку текста.

    Возвращает (new_text, translated_flag). Перевод применяется, только если в тексте
    отсутствует кириллица, но есть латиница — чтобы не трогать уже русские ответы и смешанные фрагменты.
    """
    original = (text or '').strip()
    if not original:
        return '', False
    cyr = sum(1 for ch in original if 'а' <= ch.lower() <= 'я' or ch.lower() == 'ё')
    lat = sum(1 for ch in original if 'a' <= ch.lower() <= 'z')
    if cyr > 0 or lat == 0:
        return original, False
    system = (
        "Ты профессиональный переводчик. Переведи приведённый текст на русский язык, "
        "сохраняя смысл, факты, числовые значения и ссылки. Не добавляй пояснений и комментариев."
    )
    user = (
        "Исходный текст:\n"
        f"{original}\n\n"
        "Переведи этот текст на русский язык, сохранив стиль и точность. "
        "Не добавляй ничего сверх перевода."
    )
    translated = ""
    try:
        translated = call_lmstudio_compose(
            system,
            user,
            temperature=0.0,
            max_tokens=min(max(200, len(original) + 120), _lm_max_output_tokens()),
        )
    except Exception as exc:
        app.logger.debug("Russian translation failed (%s): %s", label, exc)
    translated = (translated or '').strip()
    if not translated:
        return original, False
    trans_cyr = sum(1 for ch in translated if 'а' <= ch.lower() <= 'я' or ch.lower() == 'ё')
    if trans_cyr == 0:
        return original, False
    return translated, True


def _ai_expand_multilingual_terms(query: str, base_terms: Sequence[str], languages: Sequence[str]) -> tuple[list[str], list[tuple[str, int]]]:
    sanitized: list[str] = []
    seen_langs = set()
    for raw_lang in languages:
        normalized = _normalize_language_code(raw_lang)
        if not normalized:
            continue
        if normalized in seen_langs:
            continue
        seen_langs.add(normalized)
        sanitized.append(normalized)
    if not sanitized:
        return [], []
    base_lang = _guess_query_language(query)
    if base_lang:
        sanitized = [lang for lang in sanitized if lang != base_lang]
    if not sanitized:
        return [], []
    sanitized = sanitized[:5]
    base_preview = ", ".join([str(t) for t in list(base_terms)[:6] if t])
    system = (
        "Ты помогаешь поисковой системе расширить запрос. "
        "Для каждого указанного языка подбери 2-6 коротких ключевых фраз (1-3 слова) на этом языке. "
        "Фразы должны быть естественными, без транслитерации и повторов. "
        "Верни JSON-объект без пояснений: {\"en\": [\"term\"...], \"de\": [...]}. Ответ должен быть только JSON."
        " Используй только перечисленные языки в виде ключей."
    )
    user_parts = [
        f"Исходный запрос: {query}",
    ]
    if base_preview:
        user_parts.append(f"Базовые ключевые слова: {base_preview}")
    user_parts.append(f"Языки для расширения: {', '.join(sanitized)}")
    user_parts.append("Не добавляй другие поля и пояснения.")
    raw = call_lmstudio_compose(system, "\n".join(user_parts), temperature=0.15, max_tokens=280)
    if not raw:
        return [], []
    raw = raw.strip()
    if not raw:
        return [], []
    parsed: dict[str, list[str]] | None = None
    try:
        data = json.loads(raw)
        if isinstance(data, dict):
            parsed = data
    except Exception:
        match = re.search(r"```json\s*(\{.*?\})\s*```", raw, flags=re.S)
        if match:
            try:
                data = json.loads(match.group(1))
                if isinstance(data, dict):
                    parsed = data
            except Exception:
                parsed = None
    if not parsed:
        return [], []
    base_seen = {str(term).strip().lower() for term in base_terms if term}
    extras: list[str] = []
    lang_order: list[str] = []
    lang_counter: dict[str, int] = {}
    for lang, items in parsed.items():
        normalized_lang = _normalize_language_code(lang)
        if not normalized_lang or normalized_lang not in sanitized:
            continue
        if not isinstance(items, list):
            continue
        if normalized_lang not in lang_counter:
            lang_counter[normalized_lang] = 0
            lang_order.append(normalized_lang)
        for item in items:
            normalized_term = _normalize_keyword_candidate(item)
            if not normalized_term:
                continue
            key = normalized_term.lower()
            if key in base_seen:
                continue
            base_seen.add(key)
            extras.append(normalized_term)
            lang_counter[normalized_lang] += 1
    lang_stats = [(lang, lang_counter.get(lang, 0)) for lang in lang_order if lang_counter.get(lang, 0) > 0]
    return extras, lang_stats


def _ai_expand_keywords(
    query: str,
    *,
    multi_lang: bool = False,
    target_langs: Sequence[str] | None = None,
) -> tuple[list[str], list[tuple[str, int]]]:
    q = (query or "").strip()
    if not q:
        return [], []
    mandatory = _extract_quoted_phrases(q)
    sanitized_langs: list[str] = []
    if multi_lang and target_langs:
        seen_langs = set()
        for raw_lang in target_langs:
            normalized = _normalize_language_code(raw_lang)
            if not normalized:
                continue
            if normalized in seen_langs:
                continue
            seen_langs.add(normalized)
            sanitized_langs.append(normalized)
        if sanitized_langs:
            sanitized_langs = sanitized_langs[:5]
    base_lang = _guess_query_language(q)
    cross_langs: list[str] = []
    if base_lang == 'ru':
        cross_langs.append('en')
    elif base_lang == 'en':
        cross_langs.append('ru')
    active_langs: list[str] = []
    for lang in sanitized_langs + cross_langs:
        normalized = _normalize_language_code(lang)
        if not normalized:
            continue
        if normalized == base_lang:
            continue
        if normalized not in active_langs:
            active_langs.append(normalized)
    try:
        ttl_min = int(os.getenv("AI_EXPAND_TTL_MIN", "20") or 20)
    except Exception:
        ttl_min = 20
    cache_suffix = ""
    if active_langs:
        cache_suffix = f"|lang:{','.join(active_langs)}"
    key = _sha256(q + cache_suffix)
    now = _now()
    cached = AI_EXPAND_CACHE.get(key)
    if cached and (now - cached[0]) < ttl_min * 60:
        cached_terms = list(cached[1])
        for phrase in mandatory:
            if phrase and phrase not in cached_terms:
                cached_terms.insert(0, phrase)
        cached_langs = cached[2] if len(cached) > 2 else []
        return cached_terms, cached_langs
    kws = []
    try:
        kws = call_lmstudio_keywords(q, "ai-search") or []
    except Exception:
        kws = []
    if not kws:
        toks = [t.strip() for t in re.split(r"[\s,;]+", q) if t.strip()]
        kws = toks[:12]
    seen = set()
    res: list[str] = []
    for w in kws:
        normalized = _normalize_keyword_candidate(w)
        if not normalized:
            continue
        if normalized in STOP_WORDS:
            continue
        if normalized not in seen:
            seen.add(normalized)
            res.append(normalized)
    lang_details: list[tuple[str, int]] = []
    if active_langs:
        extra_terms, extra_stats = _ai_expand_multilingual_terms(q, res, active_langs)
        for term in extra_terms:
            if term and term not in seen:
                seen.add(term)
                res.append(term)
        lang_details.extend(extra_stats)
    for phrase in reversed(mandatory):
        if phrase and phrase not in seen:
            res.insert(0, phrase)
            seen.add(phrase)
    AI_EXPAND_CACHE[key] = (now, res, lang_details)
    return res, lang_details


def _plan_search_keywords(
    query: str,
    *,
    base_terms: Sequence[str] | None = None,
    ai_terms: Sequence[str] | None = None,
    ttl_minutes: int = 15,
) -> Dict[str, List[str]]:
    """Запрашивает у LLM структурированный набор термов для классического поиска."""
    q = (query or "").strip()
    if not q:
        return {"must": [], "phrases": [], "broad": []}
    plan_key = _sha256(f"kw-plan::{q.lower()}")
    now = _now()
    cached = AI_KEYWORD_PLAN_CACHE.get(plan_key)
    if cached and (now - cached[0]) < ttl_minutes * 60:
        return cached[1]

    base_preview = ", ".join(sorted({t.lower() for t in (base_terms or []) if t})[:8])
    ai_preview = ", ".join(sorted({t.lower() for t in (ai_terms or []) if t})[:8])
    system = (
        "Ты помощник поиска. Получив вопрос, ты возвращаешь JSON с ключами 'must', "
        "'phrases', 'broad'. Не добавляй комментариев. Каждое значение — список строк. "
        "'must' — обязательные термины (точные фамилии, ключевые слова), чистые токены без пунктуации. "
        "'phrases' — короткие фразы (2-4 слова), пригодные для полнотекстового поиска. "
        "'broad' — синонимы и расширения, включая перевод на другой язык. "
        "Если нечего добавить в категорию — верни пустой список."
    )
    user = (
        f"Запрос: {q}\n"
        f"Базовые токены: [{base_preview}]\n"
        f"AI-термины: [{ai_preview}]\n"
        "Ответь строго JSON-объектом, например "
        "{\"must\": [\"термин\"], \"phrases\": [\"фраза\"], \"broad\": [\"синоним\"]}."
    )
    raw_response = ""
    try:
        raw_response = (call_lmstudio_compose(system, user, temperature=0.0, max_tokens=200) or "").strip()
    except Exception:
        raw_response = ""
    data: Dict[str, List[str]] = {"must": [], "phrases": [], "broad": []}
    parsed: Optional[Dict[str, object]] = None
    if raw_response:
        try:
            parsed = json.loads(raw_response)
        except Exception:
            match = re.search(r"\{.*\}", raw_response, flags=re.S)
            if match:
                try:
                    parsed = json.loads(match.group(0))
                except Exception:
                    parsed = None
    if isinstance(parsed, dict):
        def _coerce_list(value: object) -> List[str]:
            if isinstance(value, list):
                return [str(item) for item in value if str(item).strip()]
            if isinstance(value, str):
                return [value]
            return []

        for key in ("must", "phrases", "broad"):
            items = _coerce_list(parsed.get(key))
            normalized: List[str] = []
            for item in items:
                norm = _normalize_keyword_candidate(item)
                if not norm:
                    continue
                if key == "phrases" and " " not in norm:
                    # фразовые термины должны содержать пробел
                    continue
                normalized.append(norm)
            data[key] = normalized
    if not data["must"]:
        data["must"] = [
            _normalize_keyword_candidate(tok)
            for tok in (base_terms or [])
            if _normalize_keyword_candidate(tok)
        ][:4]
    if not data["phrases"]:
        data["phrases"] = [
            _normalize_keyword_candidate(term)
            for term in (ai_terms or [])
            if term and " " in term and _normalize_keyword_candidate(term)
        ][:4]
    if not data["broad"]:
        data["broad"] = [
            _normalize_keyword_candidate(term)
            for term in (ai_terms or [])
            if term and _normalize_keyword_candidate(term)
        ][:6]
    AI_KEYWORD_PLAN_CACHE[plan_key] = (now, data)
    return data


def _generate_query_variants(query: str, *, max_variants: int = 2) -> List[str]:
    max_variants = max(0, int(max_variants or 0))
    query = (query or "").strip()
    if not query or max_variants <= 0:
        return []
    key = query.lower()
    now = time.monotonic()
    cached = QUERY_VARIANT_CACHE.get(key)
    if cached and (now - cached[0]) < QUERY_VARIANT_TTL:
        cached_list = cached[1]
        return cached_list[:max_variants]

    system = (
        "Ты помощник интеллектуального поиска. Получив исходный вопрос, "
        "верни JSON-массив из 1-3 альтернативных формулировок, которые сохраняют смысл, "
        "могут использовать синонимы и уточнять ключевые сущности. Без комментариев."
    )
    user = (
        f"Запрос: {query}\n"
        f"Сформируй до {max_variants} разных формулировок (если невозможно — верни пустой массив)."
    )
    variants: List[str] = []
    raw = ""
    try:
        raw = call_lmstudio_compose(system, user, temperature=0.15, max_tokens=160)
    except Exception as exc:
        app.logger.debug("Query variant generation failed: %s", exc)
        raw = ""
    parsed = None
    if raw:
        try:
            parsed = json.loads(raw)
        except Exception:
            match = re.search(r"\[.*\]", raw, flags=re.S)
            if match:
                try:
                    parsed = json.loads(match.group(0))
                except Exception:
                    parsed = None
    if isinstance(parsed, list):
        variants = [str(item).strip() for item in parsed if str(item).strip()]
    if not variants and raw:
        for line in raw.splitlines():
            candidate = line.strip().lstrip("-•0123456789. ")
            if candidate:
                variants.append(candidate)
    normalized: List[str] = []
    seen: set[str] = set()
    for candidate in variants:
        cand = candidate.strip()
        if not cand:
            continue
        lower = cand.lower()
        if lower == key or lower in seen:
            continue
        seen.add(lower)
        normalized.append(cand)
        if len(normalized) >= max_variants:
            break
    if not normalized and len(query.split()) > 2:
        # fallback: переставляем порядок слов и добавляем ключевую фразу
        tokens = query.split()
        alt = " ".join(sorted(tokens, key=str.lower))
        if alt.lower() != key:
            normalized.append(alt)
    QUERY_VARIANT_CACHE[key] = (now, normalized)
    return normalized


def _normalize_author_name(name: str) -> str:
    raw = (name or "").strip()
    if not raw:
        return ""
    cleaned = re.sub(r"\s+", " ", raw)
    return cleaned.lower()


def _normalize_tag_value(value: str) -> str:
    raw = (value or "").strip()
    if not raw:
        return ""
    cleaned = re.sub(r"\s+", " ", raw)
    return cleaned.lower()


def _compute_authority_snapshot(allowed_collections: Optional[Iterable[int]] = None) -> AuthorityCacheEntry:
    scope_key = "*"
    allowed_set: Optional[frozenset[int]] = None
    if allowed_collections is not None:
        allowed_set = frozenset(int(cid) for cid in allowed_collections)
        scope_key = ",".join(str(cid) for cid in sorted(allowed_set)) or "none"
    now_monotonic = time.monotonic()
    cached = AUTHORITY_GRAPH_CACHE.get(scope_key)
    if cached and (now_monotonic - float(cached.get("timestamp", 0.0))) < AUTHORITY_GRAPH_TTL:
        return cached

    try:
        file_query = File.query.outerjoin(Collection, File.collection_id == Collection.id)
        if allowed_set is not None:
            if not allowed_set:
                snapshot_empty = {
                    "timestamp": now_monotonic,
                    "generated_at": time.time(),
                    "doc_scores": {},
                    "author_entries": [],
                    "collection_entries": [],
                    "topic_entries": [],
                    "topic_index": {},
                }
                AUTHORITY_GRAPH_CACHE[scope_key] = snapshot_empty
                return snapshot_empty
            file_query = file_query.filter(File.collection_id.in_(allowed_set))
        else:
            file_query = file_query.filter(or_(Collection.searchable == True, Collection.id.is_(None)))
        files = file_query.all()
    except Exception as exc:
        app.logger.debug("Authority graph: failed to fetch files (%s)", exc)
        files = []

    if not files:
        snapshot_empty = {
            "timestamp": now_monotonic,
            "generated_at": time.time(),
            "doc_scores": {},
            "author_entries": [],
            "collection_entries": [],
            "topic_entries": [],
            "topic_index": {},
        }
        AUTHORITY_GRAPH_CACHE[scope_key] = snapshot_empty
        return snapshot_empty

    graph: Dict[int, Dict[int, float]] = defaultdict(dict)
    nodes: set[int] = set()
    collection_groups: Dict[int, List[int]] = defaultdict(list)
    author_groups: Dict[str, List[int]] = defaultdict(list)
    keyword_groups: Dict[str, List[int]] = defaultdict(list)
    tag_groups: Dict[Tuple[str, str], List[int]] = defaultdict(list)

    collection_ids_seen: set[int] = set()
    author_labels: Dict[str, str] = {}
    keyword_labels: Dict[str, str] = {}
    tag_labels: Dict[Tuple[str, str], str] = {}

    for f in files:
        if not getattr(f, "id", None):
            continue
        fid = int(f.id)
        nodes.add(fid)
        if f.collection_id:
            cid = int(f.collection_id)
            collection_groups[cid].append(fid)
            collection_ids_seen.add(cid)
        author_raw = (f.author or "").strip()
        author_clean = _normalize_author_name(author_raw)
        if author_clean:
            author_groups[author_clean].append(fid)
            author_labels.setdefault(author_clean, author_raw or author_clean)
        keywords_field = (f.keywords or "").strip()
        if keywords_field:
            for token in re.split(r"[,;\n]+", keywords_field):
                original = (token or "").strip()
                if not original:
                    continue
                kw = _normalize_tag_value(original)
                if kw:
                    keyword_groups[kw].append(fid)
                    keyword_labels.setdefault(kw, original)

    try:
        tag_query = db.session.query(Tag.file_id, Tag.key, Tag.value).join(File, Tag.file_id == File.id)
        if allowed_set is not None:
            tag_query = tag_query.filter(Tag.file_id.in_(nodes))
        rows = tag_query.all()
    except Exception:
        rows = []
    for fid, key, value in rows:
        fid = int(fid)
        if fid not in nodes:
            continue
        key_str = str(key or "").strip()
        normalized = _normalize_tag_value(value or "")
        if not key_str or not normalized:
            continue
        tag_groups[(key_str, normalized)].append(fid)
        tag_labels.setdefault((key_str, normalized), (value or "").strip() or normalized)

    def _add_group(group: List[int], weight: float) -> None:
        if len(group) < 2:
            return
        capped = group[:200]
        for i, a in enumerate(capped):
            for b in capped[i + 1:]:
                if a == b:
                    continue
                graph[a][b] = graph[a].get(b, 0.0) + weight
                graph[b][a] = graph[b].get(a, 0.0) + weight

    for members in collection_groups.values():
        _add_group(members, 0.4)
    for members in author_groups.values():
        _add_group(members, 0.8)
    for members in keyword_groups.values():
        _add_group(members, 0.3)
    for members in tag_groups.values():
        _add_group(members, 0.6)

    if not graph:
        snapshot_empty = {
            "timestamp": now_monotonic,
            "generated_at": time.time(),
            "doc_scores": {},
            "author_entries": [],
            "collection_entries": [],
            "topic_entries": [],
            "topic_index": {},
        }
        AUTHORITY_GRAPH_CACHE[scope_key] = snapshot_empty
        return snapshot_empty

    damping = 0.85
    max_iter = 20
    min_delta = 1e-6
    all_nodes = set(graph.keys())
    for neighbors in graph.values():
        all_nodes.update(neighbors.keys())
    N = len(all_nodes)
    if N == 0:
        snapshot_empty = {
            "timestamp": now_monotonic,
            "generated_at": time.time(),
            "doc_scores": {},
            "author_entries": [],
            "collection_entries": [],
            "topic_entries": [],
            "topic_index": {},
        }
        AUTHORITY_GRAPH_CACHE[scope_key] = snapshot_empty
        return snapshot_empty
    base = (1.0 - damping) / N
    scores = {node: 1.0 / N for node in all_nodes}
    for _ in range(max_iter):
        new_scores = {node: base for node in all_nodes}
        for node in all_nodes:
            neighbors = graph.get(node, {})
            if not neighbors:
                share = scores[node] * damping / N
                for target in new_scores:
                    new_scores[target] += share
                continue
            weight_sum = sum(neighbors.values()) or 1.0
            for target, weight in neighbors.items():
                new_scores[target] += damping * scores[node] * (weight / weight_sum)
        diff = sum(abs(new_scores[n] - scores.get(n, 0.0)) for n in all_nodes)
        scores = new_scores
        if diff < min_delta:
            break
    max_score = max(scores.values()) if scores else 1.0
    normalized_scores = {node: float(score) / max_score for node, score in scores.items() if score > 0}

    def _aggregate_group(members: List[int]) -> Tuple[float, int]:
        unique_members = list({int(m) for m in members})
        if not unique_members:
            return 0.0, 0
        values = [normalized_scores.get(mid, 0.0) for mid in unique_members]
        non_zero = [val for val in values if val > 0]
        if not non_zero:
            return 0.0, len(unique_members)
        return sum(non_zero) / len(non_zero), len(unique_members)

    collection_labels: Dict[int, str] = {}
    if collection_ids_seen:
        try:
            col_rows = db.session.query(Collection.id, Collection.name).filter(Collection.id.in_(collection_ids_seen)).all()
            collection_labels = {int(cid): (name or f"Коллекция {cid}") for cid, name in col_rows}
        except Exception:
            collection_labels = {cid: f"Коллекция {cid}" for cid in collection_ids_seen}

    author_entries: List[Dict[str, Any]] = []
    for normalized_name, members in author_groups.items():
        score, count = _aggregate_group(members)
        if count and score > 0:
            author_entries.append({
                "name": author_labels.get(normalized_name, normalized_name),
                "score": round(score, 6),
                "count": count,
            })

    collection_entries: List[Dict[str, Any]] = []
    for cid, members in collection_groups.items():
        score, count = _aggregate_group(members)
        if count and score > 0:
            collection_entries.append({
                "collection_id": cid,
                "name": collection_labels.get(cid, f"Коллекция {cid}"),
                "score": round(score, 6),
                "count": count,
            })

    topic_entries: List[Dict[str, Any]] = []
    topic_index: Dict[str, float] = {}

    for norm, members in keyword_groups.items():
        score, count = _aggregate_group(members)
        if count and score > 0:
            label = keyword_labels.get(norm, norm)
            topic_entries.append({
                "key": "keyword",
                "label": label,
                "score": round(score, 6),
                "count": count,
            })
            topic_index[f"keyword|||{norm}"] = score

    for (key_str, norm_val), members in tag_groups.items():
        score, count = _aggregate_group(members)
        if count and score > 0:
            label = tag_labels.get((key_str, norm_val), norm_val)
            topic_entries.append({
                "key": key_str,
                "label": label,
                "score": round(score, 6),
                "count": count,
            })
            topic_index[f"{key_str}|||{norm_val}"] = score

    author_entries.sort(key=lambda item: item["score"], reverse=True)
    collection_entries.sort(key=lambda item: item["score"], reverse=True)
    topic_entries.sort(key=lambda item: item["score"], reverse=True)

    snapshot = {
        "timestamp": now_monotonic,
        "generated_at": time.time(),
        "doc_scores": dict(normalized_scores),
        "author_entries": author_entries,
        "collection_entries": collection_entries,
        "topic_entries": topic_entries,
        "topic_index": topic_index,
    }
    AUTHORITY_GRAPH_CACHE[scope_key] = snapshot
    return snapshot


def _compute_authority_scores(allowed_collections: Optional[Iterable[int]] = None) -> Dict[int, float]:
    snapshot = _compute_authority_snapshot(allowed_collections)
    return dict(snapshot.get("doc_scores", {}))

def _invalidate_feedback_model_cache() -> None:
    global AI_FEEDBACK_MODEL_CACHE
    AI_FEEDBACK_MODEL_CACHE = (0.0, {})


def _get_feedback_weights() -> Dict[int, Dict[str, Any]]:
    global AI_FEEDBACK_MODEL_CACHE
    now = time.monotonic()
    cache_ts, cache_payload = AI_FEEDBACK_MODEL_CACHE
    if cache_payload and (now - cache_ts) < FEEDBACK_MODEL_TTL:
        return cache_payload
    try:
        rows = AiSearchFeedbackModel.query.all()
    except Exception as exc:
        app.logger.debug("Feedback model fetch failed: %s", exc)
        rows = []
    payload: Dict[int, Dict[str, Any]] = {}
    for row in rows:
        fid = int(getattr(row, "file_id", 0) or 0)
        if not fid:
            continue
        payload[fid] = {
            "weight": float(getattr(row, "weight", 0.0) or 0.0),
            "positive": int(getattr(row, "positive", 0) or 0),
            "negative": int(getattr(row, "negative", 0) or 0),
            "clicks": int(getattr(row, "clicks", 0) or 0),
            "updated_at": row.updated_at.isoformat() if getattr(row, "updated_at", None) else None,
        }
    AI_FEEDBACK_MODEL_CACHE = (now, payload)
    return payload


def _rebuild_feedback_model(*, cutoff_days: Optional[int] = None) -> dict:
    with app.app_context():
        cutoff_ts: Optional[datetime] = None
        if cutoff_days is not None:
            try:
                cutoff_days = int(cutoff_days)
                if cutoff_days > 0:
                    cutoff_ts = datetime.utcnow() - timedelta(days=cutoff_days)
            except Exception:
                cutoff_ts = None
        query = (
            db.session.query(
                AiSearchKeywordFeedback.file_id,
                AiSearchKeywordFeedback.action,
                func.count(AiSearchKeywordFeedback.id),
                func.max(AiSearchKeywordFeedback.created_at),
            )
            .filter(AiSearchKeywordFeedback.file_id.isnot(None))
        )
        if cutoff_ts is not None:
            query = query.filter(AiSearchKeywordFeedback.created_at >= cutoff_ts)
        rows = query.group_by(AiSearchKeywordFeedback.file_id, AiSearchKeywordFeedback.action).all()
        aggregates: Dict[int, Dict[str, Any]] = {}
        for fid, action, count, last_at in rows:
            if fid is None:
                continue
            fid = int(fid)
            entry = aggregates.setdefault(
                fid,
                {
                    "positive": 0,
                    "negative": 0,
                    "clicks": 0,
                    "last_positive": None,
                    "last_negative": None,
                    "last_click": None,
                },
            )
            cnt = int(count or 0)
            if action == 'relevant':
                entry["positive"] += cnt
                if last_at and (entry["last_positive"] is None or last_at > entry["last_positive"]):
                    entry["last_positive"] = last_at
            elif action == 'click':
                entry["clicks"] += cnt
                if last_at and (entry["last_click"] is None or last_at > entry["last_click"]):
                    entry["last_click"] = last_at
            elif action == 'irrelevant':
                entry["negative"] += cnt
                if last_at and (entry["last_negative"] is None or last_at > entry["last_negative"]):
                    entry["last_negative"] = last_at
        existing_rows = {
            row.file_id: row
            for row in db.session.query(AiSearchFeedbackModel).all()
        }
        updated = 0
        created = 0
        processed_ids: set[int] = set()
        now_dt = datetime.utcnow()
        for fid, entry in aggregates.items():
            processed_ids.add(fid)
            row = existing_rows.get(fid)
            if row is None:
                row = AiSearchFeedbackModel(file_id=fid)
                db.session.add(row)
                created += 1
            else:
                updated += 1
            positive = int(entry["positive"])
            negative = int(entry["negative"])
            clicks = int(entry["clicks"])
            row.positive = positive
            row.negative = negative
            row.clicks = clicks
            row.last_positive_at = entry.get("last_positive")
            row.last_negative_at = entry.get("last_negative")
            row.last_click_at = entry.get("last_click")
            pos_score = positive + (FEEDBACK_CLICK_WEIGHT * clicks)
            neg_score = negative
            if pos_score <= 0 and neg_score <= 0:
                weight = 0.0
            else:
                weight = math.log((FEEDBACK_POS_PRIOR + pos_score) / (FEEDBACK_NEG_PRIOR + max(0.0, neg_score)))
                weight = max(-FEEDBACK_MAX_WEIGHT, min(FEEDBACK_MAX_WEIGHT, weight))
            row.weight = weight
            row.updated_at = now_dt
        deleted = 0
        for fid, row in existing_rows.items():
            if fid not in processed_ids:
                db.session.delete(row)
                deleted += 1
        db.session.commit()
        _invalidate_feedback_model_cache()
        return {
            "files": len(aggregates),
            "created": created,
            "updated": updated,
            "deleted": deleted,
            "cutoff_days": cutoff_days,
            "timestamp": now_dt.isoformat(),
        }


def _enqueue_feedback_training(
    *,
    cutoff_days: Optional[int],
    trigger: str,
    submitted_by: Optional[int] = None,
    allow_duplicate: bool = False,
) -> tuple[int, bool]:
    active_task = TaskRecord.query.filter(
        TaskRecord.name == 'feedback_train',
        TaskRecord.status.notin_(TASK_FINAL_STATUSES),
    ).order_by(TaskRecord.created_at.desc()).first()
    if active_task and not allow_duplicate:
        return int(active_task.id), False
    payload = {
        'trigger': trigger,
        'cutoff_days': cutoff_days,
    }
    if submitted_by:
        payload['submitted_by'] = int(submitted_by)
    task = TaskRecord(
        name='feedback_train',
        status='queued',
        progress=0.0,
        payload=json.dumps(payload, ensure_ascii=False),
        user_id=int(submitted_by) if submitted_by else None,
    )
    db.session.add(task)
    db.session.commit()
    task_id = int(task.id)
    get_task_queue().submit(
        _feedback_training_job,
        task_id,
        cutoff_days,
        description='feedback_model_train',
        owner_id=int(submitted_by) if submitted_by else None,
    )
    return task_id, True


def _maybe_enqueue_feedback_training_online(*, trigger: str, submitted_by: Optional[int]) -> tuple[Optional[int], bool]:
    global FEEDBACK_ONLINE_TRAIN_LAST_AT
    now_ts = time.time()
    min_interval = max(30.0, float(FEEDBACK_ONLINE_TRAIN_MIN_INTERVAL_SEC or 0.0))
    if FEEDBACK_ONLINE_TRAIN_LAST_AT > 0 and (now_ts - FEEDBACK_ONLINE_TRAIN_LAST_AT) < min_interval:
        return None, False
    try:
        events = int(
            db.session.query(func.count(AiSearchKeywordFeedback.id))
            .scalar() or 0
        )
    except Exception:
        return None, False
    threshold = max(1, int(FEEDBACK_ONLINE_TRAIN_EVERY_EVENTS or 1))
    if events <= 0 or (events % threshold) != 0:
        return None, False
    task_id, queued = _enqueue_feedback_training(
        cutoff_days=FEEDBACK_TRAIN_CUTOFF_DAYS,
        trigger=trigger,
        submitted_by=submitted_by,
        allow_duplicate=False,
    )
    if queued:
        FEEDBACK_ONLINE_TRAIN_LAST_AT = now_ts
    return task_id, queued


def _feedback_training_job(task_id: int, cutoff_days: Optional[int]) -> None:
    with app.app_context():
        task = TaskRecord.query.get(task_id)
        if not task:
            return
        try:
            task.status = 'running'
            task.started_at = datetime.utcnow()
            task.progress = 0.0
            db.session.commit()
        except Exception:
            db.session.rollback()
        try:
            stats = _rebuild_feedback_model(cutoff_days=cutoff_days)
            task = TaskRecord.query.get(task_id)
            if task:
                original_payload: dict = {}
                if task.payload:
                    try:
                        original_payload = json.loads(task.payload)
                    except Exception:
                        original_payload = {}
                payload = {
                    'trigger': original_payload.get('trigger'),
                    'cutoff_days': cutoff_days,
                    'stats': stats,
                }
                task.status = 'completed'
                task.finished_at = datetime.utcnow()
                task.progress = 1.0
                task.payload = json.dumps(payload, ensure_ascii=False)
                db.session.commit()
        except Exception as exc:
            db.session.rollback()
            task = TaskRecord.query.get(task_id)
            if task:
                task.status = 'error'
                task.error = str(exc)
                task.finished_at = datetime.utcnow()
                try:
                    db.session.commit()
                except Exception:
                    db.session.rollback()
            app.logger.exception("[ai-feedback] training task #%s failed", task_id)

app.config['authority_snapshot_fn'] = _compute_authority_snapshot

def _rrf_combine_dense_hits(
    vector_retriever: VectorRetriever,
    variant_vectors: Sequence[Tuple[str, Sequence[float]]],
    *,
    allowed_document_ids: Optional[Sequence[int]],
    top_k: int,
) -> List[RetrievedChunk]:
    rrf_k = 60
    aggregate: Dict[int, Dict[str, object]] = {}
    for _variant, vector in variant_vectors:
        if not vector:
            continue
        hits = vector_retriever.search_by_vector(
            vector,
            top_k=top_k,
            allowed_document_ids=allowed_document_ids,
        )
        for rank, hit in enumerate(hits, start=1):
            weight = 1.0 / (rrf_k + rank)
            entry = aggregate.setdefault(
                hit.chunk.id,
                {
                    "chunk": hit.chunk,
                    "document": hit.document,
                    "lang": hit.lang,
                    "keywords": hit.keywords,
                    "preview": hit.preview,
                    "score": 0.0,
                    "best": hit.score,
                },
            )
            entry["score"] = float(entry.get("score", 0.0)) + weight
            entry["best"] = max(float(entry.get("best", 0.0)), float(hit.score or 0.0))
    if not aggregate:
        return []
    aggregated_items = sorted(
        aggregate.values(),
        key=lambda item: (float(item.get("score", 0.0)), float(item.get("best", 0.0))),
        reverse=True,
    )
    combined: List[RetrievedChunk] = []
    for item in aggregated_items:
        combined.append(
            RetrievedChunk(
                chunk=item["chunk"],
                document=item.get("document"),
                score=float(item.get("score", 0.0)),
                lang=item.get("lang"),
                keywords=item.get("keywords"),
                preview=item.get("preview"),
            )
        )
    return combined


def _read_cached_excerpt_for_file(f: File) -> str:
    try:
        # Предпочитаем фрагмент из базы данных
        if (f.text_excerpt or '').strip():
            return (f.text_excerpt or '')
        cache_dir = Path(app.static_folder) / 'cache' / 'text_excerpts'
        key = (f.sha1 or (f.rel_path or '').replace('/', '_')) + '.txt'
        fp = cache_dir / key
        if fp.exists():
            return fp.read_text(encoding='utf-8', errors='ignore')
    except Exception:
        pass
    return ''


SNIPPET_WINDOW_RADIUS = 220


def _collect_snippets(text: str, terms: list[str], max_snips: int = 2) -> list[str]:
    t = (text or '')
    if not t:
        return []
    tl = t.lower()
    outs: list[tuple[int, str]] = []
    windows = []
    for raw in terms:
        term = (raw or '').strip()
        if not term:
            continue
        ql = term.lower()
        pos = 0
        found_any = False
        for _i in range(3):  # максимум 3 позиционирования на термин
            idx = tl.find(ql, pos)
            if idx < 0:
                break
            found_any = True
            start = max(0, idx - SNIPPET_WINDOW_RADIUS)
            end = min(len(t), idx + len(term) + SNIPPET_WINDOW_RADIUS)
            windows.append((start, end))
            pos = idx + len(term)
        if not found_any and len(ql) >= 3:
            # пытаемся разделить термин на подслова
            for part in re.split(r"[\s\-_/]+", ql):
                if len(part) < 3:
                    continue
                idx = tl.find(part)
                if idx >= 0:
                    start = max(0, idx - SNIPPET_WINDOW_RADIUS)
                    end = min(len(t), idx + len(part) + SNIPPET_WINDOW_RADIUS)
                    windows.append((start, end))
    # объединяем перекрывающиеся окна
    windows.sort()
    merged = []
    for w in windows:
        if not merged or w[0] > merged[-1][1] + 40:
            merged.append(list(w))
        else:
            merged[-1][1] = max(merged[-1][1], w[1])
    for a, b in merged[:max_snips]:
        snip = t[a:b]
        # схлопываем переводы строк для компактности
        snip = re.sub(r"\s+", " ", snip).strip()
        outs.append((a, snip))
    outs.sort(key=lambda x: x[0])
    return [s for _pos, s in outs]


def _split_text_chunks(text: str, chunk_chars: int, max_chunks: int) -> list[str]:
    if not text:
        return []
    chunk_chars = max(256, int(chunk_chars or 1024))
    max_chunks = max(1, int(max_chunks or 1))
    limit = min(len(text), chunk_chars * max_chunks)
    trimmed = text[:limit]
    return [trimmed[i:i+chunk_chars] for i in range(0, len(trimmed), chunk_chars)]


def _iter_document_chunks(path: Path, chunk_chars: int = 6000, max_chunks: int = 30):
    if not path or not path.exists() or not path.is_file():
        return
    chunk_chars = max(512, int(chunk_chars or 2000))
    max_chunks = max(1, int(max_chunks or 1))
    ext = path.suffix.lower()
    yielded = 0
    try:
        if ext in {'.txt', '.md', '.markdown', '.csv', '.tsv', '.json', '.yaml', '.yml', '.log'}:
            with path.open('r', encoding='utf-8', errors='ignore') as fh:
                while yielded < max_chunks:
                    chunk = fh.read(chunk_chars)
                    if not chunk:
                        break
                    yielded += 1
                    yield chunk
            return
        if ext == '.pdf' and fitz is not None:
            with fitz.open(path) as doc:
                buffer = ''
                for page in doc:
                    buffer += page.get_text('text') or ''
                    while len(buffer) >= chunk_chars and yielded < max_chunks:
                        yield buffer[:chunk_chars]
                        buffer = buffer[chunk_chars:]
                        yielded += 1
                        if yielded >= max_chunks:
                            return
                if buffer and yielded < max_chunks:
                    yield buffer[:chunk_chars]
            return
        if ext == '.docx' and docx is not None:
            document = docx.Document(str(path))
            buffer = ''
            for para in document.paragraphs:
                buffer += (para.text or '') + '\n'
                while len(buffer) >= chunk_chars and yielded < max_chunks:
                    yield buffer[:chunk_chars]
                    buffer = buffer[chunk_chars:]
                    yielded += 1
                    if yielded >= max_chunks:
                        return
            if buffer and yielded < max_chunks:
                yield buffer[:chunk_chars]
            return
        if ext == '.rtf' and rtf_to_text is not None:
            text = rtf_to_text(path.read_text(encoding='utf-8', errors='ignore'))
            for chunk in _split_text_chunks(text, chunk_chars, max_chunks):
                yield chunk
            return
        if ext == '.epub' and epub is not None:
            book = epub.read_epub(str(path))
            buffer = ''
            for item in book.get_items():
                if item.get_type() != epub.ITEM_DOCUMENT:
                    continue
                buffer += item.get_content().decode(errors='ignore')
                while len(buffer) >= chunk_chars and yielded < max_chunks:
                    yield buffer[:chunk_chars]
                    buffer = buffer[chunk_chars:]
                    yielded += 1
                    if yielded >= max_chunks:
                        return
            if buffer and yielded < max_chunks:
                yield buffer[:chunk_chars]
            return
        if ext == '.djvu' and djvu is not None:
            with djvu.decode.open(str(path)) as d:
                buffer = ''
                for page in d.pages:
                    buffer += page.get_text()
                    while len(buffer) >= chunk_chars and yielded < max_chunks:
                        yield buffer[:chunk_chars]
                        buffer = buffer[chunk_chars:]
                        yielded += 1
                        if yielded >= max_chunks:
                            return
            if buffer and yielded < max_chunks:
                yield buffer[:chunk_chars]
            return
    except Exception:
        pass
    text = ''
    try:
        if ext == '.pdf':
            text = extract_text_pdf(path, limit_chars=chunk_chars * max_chunks)
        elif ext == '.docx':
            text = extract_text_docx(path, limit_chars=chunk_chars * max_chunks)
        elif ext == '.rtf':
            text = extract_text_rtf(path, limit_chars=chunk_chars * max_chunks)
        elif ext == '.epub':
            text = extract_text_epub(path, limit_chars=chunk_chars * max_chunks)
        elif ext == '.djvu':
            text = extract_text_djvu(path, limit_chars=chunk_chars * max_chunks)
        else:
            text = path.read_text(encoding='utf-8', errors='ignore')
    except Exception:
        text = ''
    for chunk in _split_text_chunks(text, chunk_chars, max_chunks):
        yield chunk


def _deep_scan_file(file_obj: File, terms: list[str], raw_query: str, *, chunk_chars: int = 5000, max_chunks: int = 40, max_snippets: int = 3):
    if not file_obj or not file_obj.path:
        return {"hit": False, "snippets": [], "matched_terms": set(), "boost": 0.0, "sources": []}
    path = Path(file_obj.path)
    if not path.exists() or not path.is_file():
        return {"hit": False, "snippets": [], "matched_terms": set(), "boost": 0.0, "sources": []}
    term_map = {str(t).lower(): t for t in terms if t}
    raw_lower = (raw_query or '').strip().lower()
    phrase_pat = re.compile(re.escape(raw_lower), flags=re.IGNORECASE) if raw_lower and len(raw_lower) >= 3 else None
    snippets: list[str] = []
    matched: set[str] = set()
    sources_used: set[str] = set()
    # Дополнительные текстовые данные (транскрипты, резюме, подписи) из БД/тегов
    extra_texts: list[tuple[str, str]] = []
    try:
        if (file_obj.text_excerpt or '').strip():
            extra_texts.append(('excerpt', file_obj.text_excerpt))
        if (getattr(file_obj, 'abstract', None) or '').strip():
            extra_texts.append(('abstract', getattr(file_obj, 'abstract')))
        for tag in getattr(file_obj, 'tags', []) or []:
            try:
                key = str(getattr(tag, 'key', '') or '')
                val = str(getattr(tag, 'value', '') or '')
            except Exception:
                continue
            key_l = key.lower()
            if not val or len(val.strip()) < 40:
                continue
            if any(token in key_l for token in ('transcript', 'стенограм', 'caption', 'описан', 'vision', 'summary', 'ocr')):
                extra_texts.append((f'tag:{key}', val))
    except Exception:
        pass
    try:
        for chunk in _iter_document_chunks(path, chunk_chars=chunk_chars, max_chunks=max_chunks):
            if not chunk:
                continue
            chunk_lower = chunk.lower()
            chunk_hits = {orig for low, orig in term_map.items() if low and low in chunk_lower}
            phrase_hit = phrase_pat.search(chunk) if phrase_pat else None
            if not chunk_hits and not phrase_hit:
                continue
            sn = _collect_snippets(chunk, terms, max_snips=1)
            if sn:
                snippets.extend(sn)
            else:
                cleaned = re.sub(r"\s+", " ", chunk).strip()
                if cleaned:
                    snippets.append(cleaned[:300])
            matched.update(chunk_hits)
            sources_used.add('full-text')
            if len(snippets) >= max_snippets:
                break
    except Exception:
        pass
    # При необходимости обогащаем дополнительными текстовыми источниками (транскрипты, подписи, резюме)
    if len(snippets) < max_snippets:
        for label, text in extra_texts:
            if not text:
                continue
            sn = _collect_snippets(text, terms, max_snips=1)
            cleaned = ''
            if sn:
                for item in sn:
                    if item not in snippets:
                        snippets.append(item)
            else:
                cleaned = re.sub(r"\s+", " ", text).strip()
                if cleaned:
                    snippets.append(cleaned[:300])
            if (sn and sn[0]) or cleaned:
                sources_used.add(label)
            if len(snippets) >= max_snippets:
                break
    if not snippets:
        return {"hit": False, "snippets": [], "matched_terms": matched, "boost": 0.0, "sources": sorted(sources_used)}
    boost = 1.5 * min(len(snippets), max_snippets) + 0.4 * len(matched)
    return {
        "hit": True,
        "snippets": snippets[:max_snippets],
        "matched_terms": matched,
        "boost": boost,
        "sources": sorted(sources_used),
    }


# Базовый список русских/английских стоп-слов для исключения неинформативных токенов в ИИ-поиске
STOP_WORDS = set([
    # Русские местоимения, частицы, предлоги, союзы, служебные слова
    'и','или','а','но','же','то','ли','не','ни','да','уж','вот','как','так','что','кто','где','когда','зачем','почему','какой','какая','какие','каков','это','эта','этот','эти','того','тому','этом','этих','тех','там','тут','здесь','бы','либо','пусть','дабы','быть','есть','нет','между','через','после','перед','около','возле','у','к','ко','от','до','для','по','под','над','о','об','обо','при','без','из','из‑за','изза','с','со','же','ну','в','во','на','над','ё','же','уж','еще','ещё','также','тоже','сам','сама','сами','само','свой','своя','свои','своё','мой','моя','мои','моё','твой','твоя','твои','твоё','наш','наша','наши','наше','ваш','ваша','ваши','ваше','тот','та','то','те','эт','прочее','другой','другая','другие','иной','иная','иное','иные',
    # Английские
    'the','a','an','and','or','but','if','then','else','when','where','why','how','what','who','whom','whose','this','that','these','those','is','are','was','were','be','been','being','to','of','in','on','for','with','at','by','from','as','about','into','through','after','over','between','out','against','during','without','before','under','around','among','it','its','we','you','they','he','she','them','his','her','their','our','your','my','me','us','do','does','did','not','no','yes'
])



class _ProgressLogger:
    def __init__(self, emitter=None):
        self.lines: list[str] = []
        self.emitter = emitter

    def add(self, line: str) -> None:
        text = str(line)
        self.lines.append(text)
        try:
            app.logger.info(f"[ai-search] {text}")
        except Exception:
            pass
        if self.emitter:
            try:
                self.emitter(text)
            except Exception:
                pass



def _tokenize_query(q: str) -> list[str]:
    s = (q or '').lower()
    # оставляем буквы, цифры, дефис, подчёркивание; разделяем по остальным символам
    parts = re.split(r"[^\w\-]+", s)
    filtered: list[str] = []
    for p in parts:
        if not p:
            continue
        if len(p) < 2:
            continue
        # отбрасываем токены, состоящие только из цифр
        if p.isdigit():
            continue
        # стоп-слова
        if p in STOP_WORDS:
            continue
        filtered.append(p)
    # удаляем дубликаты, сохраняя порядок
    seen = set()
    out: list[str] = []
    for p in filtered:
        if p not in seen:
            seen.add(p)
            out.append(p)
    return out[:16]


def _idf_for_terms(terms: list[str]) -> dict[str, float]:
    # вычисляем частоты документов по объединению полей файла и тегов
    idf: dict[str, float] = {}
    try:
        allowed = getattr(g, 'allowed_collection_ids', set())
        total_q = db.session.query(func.count(File.id))
        if allowed is not None:
            if not allowed:
                N = 1
            else:
                total_q = total_q.filter(File.collection_id.in_(allowed))
                N = total_q.scalar() or 1
        else:
            N = total_q.scalar() or 1
    except Exception:
        N = 1
    for w in terms:
        like = f"%{w}%"
        try:
            q = db.session.query(func.count(func.distinct(File.id))) \
                .outerjoin(Tag, Tag.file_id == File.id) \
                .filter(or_(
                    File.title.ilike(like),
                    File.author.ilike(like),
                    File.keywords.ilike(like),
                    File.text_excerpt.ilike(like),
                    File.abstract.ilike(like),
                    Tag.value.ilike(like),
                    Tag.key.ilike(like),
                ))
            if allowed is not None:
                if not allowed:
                    df = 0
                    idf[w] = 1.0
                    continue
                q = q.filter(File.collection_id.in_(allowed))
            df = int(q.scalar() or 0)
        except Exception:
            df = 0
        # сглаживание add-one
        val = float((1.0 + (N / (1.0 + df))))
        # логарифмическая шкала, минимум 1.0
        try:
            import math
            val = max(1.0, math.log(val + 1.0))
        except Exception:
            val = 1.0
        idf[w] = val
    return idf


def _ai_search_core(data: dict | None, progress_cb=None) -> dict:
    data = data or {}
    query = (data.get('query') or '').strip()
    if not query:
        raise ValueError("query is required")
    try:
        top_k = int(data.get('top_k') or 10)
    except Exception:
        top_k = 10
    original_top_k = top_k
    top_k = max(1, min(5, top_k))
    deep_search = bool(data.get('deep_search'))
    user = getattr(g, 'current_user', None)
    t_total_start = time.monotonic()
    stage_start = t_total_start
    durations: dict[str, float] = {}
    file_lookup: Dict[int, File] = {}
    try:
        max_candidates = int(data.get('max_candidates') or 15)
    except Exception:
        max_candidates = 15
    max_candidates = max(5, min(max_candidates, 400))
    full_text = bool(data.get('full_text'))
    llm_snippets = bool(data.get('llm_snippets'))
    try:
        chunk_chars = int(data.get('chunk_chars') or 5000)
    except Exception:
        chunk_chars = 5000
    chunk_chars = max(500, min(chunk_chars, 20000))
    try:
        max_chunks = int(data.get('max_chunks') or 40)
    except Exception:
        max_chunks = 40
    max_chunks = max(1, min(max_chunks, 200))
    try:
        max_snippets = int(data.get('max_snippets') or 3)
    except Exception:
        max_snippets = 3
    max_snippets = max(1, min(max_snippets, 10))
    all_languages = bool(data.get('all_languages'))
    try:
        requested_languages = [str(x).strip() for x in (data.get('languages') or []) if str(x).strip()]
    except Exception:
        requested_languages = []
    if requested_languages:
        dedup_langs: list[str] = []
        seen_langs = set()
        for raw_lang in requested_languages:
            key = raw_lang.lower()
            if key in seen_langs:
                continue
            seen_langs.add(key)
            dedup_langs.append(raw_lang)
        requested_languages = dedup_langs
    else:
        requested_languages = []
    rag_enabled = data.get('use_rag')
    if rag_enabled is None:
        rag_enabled = False
    else:
        rag_enabled = bool(rag_enabled)
    progress = _ProgressLogger(progress_cb)
    if not rag_enabled:
        progress.add("RAG: режим отключён настройками запроса")
    query_hash = _query_fingerprint(query, data)
    mandatory_terms = _extract_quoted_phrases(query)
    mandatory_set = set(mandatory_terms)
    progress.add(f"Запрос: {query}")
    progress.add(f"Отпечаток запроса: {query_hash[:12]}…")
    if top_k != original_top_k:
        progress.add(f"Запрошено Top K = {original_top_k}, ограничиваем до {top_k}")
    else:
        progress.add(f"Top K = {top_k}")
    progress.add("Режим: " + ("глубокий" if deep_search else "быстрый"))
    if mandatory_terms:
        preview = ', '.join(f'«{term}»' for term in mandatory_terms[:4])
        if len(mandatory_terms) > 4:
            preview += '…'
        progress.add(f"Обязательные фразы: {preview}")
    else:
        progress.add("Обязательные фразы: нет")
    sources = data.get('sources') or {}
    use_tags = sources.get('tags', True) if isinstance(sources, dict) else True
    use_text = sources.get('text', True) if isinstance(sources, dict) else True
    progress.add(f"Источники: теги {'вкл' if use_tags else 'выкл'}, метаданные {'вкл' if use_text else 'выкл'}")
    progress.add(f"Параметры: кандидатов ≤ {max_candidates}, сниппетов ≤ {max_snippets}")
    progress.add(f"Контент: чанк {chunk_chars} симв., чанков ≤ {max_chunks}, full-text {'вкл' if full_text else 'выкл'}")
    progress.add(f"LLM сниппеты: {'вкл' if llm_snippets else 'выкл'}")
    if all_languages:
        if requested_languages:
            preview_langs = ', '.join(requested_languages[:5])
            if len(requested_languages) > 5:
                preview_langs += '…'
            progress.add(f"Языки: все ({preview_langs})")
        else:
            progress.add("Языки: все (нет языковых тегов)")
    else:
        progress.add("Языки: по умолчанию")
    # Необязательный фильтр по коллекциям (список id)
    if 'collection_id' in data and 'collection_ids' not in data:
        data['collection_ids'] = [data['collection_id']]
    try:
        collection_ids = [int(x) for x in (data.get('collection_ids') or []) if str(x).strip()]
    except Exception:
        collection_ids = []
    allowed = getattr(g, 'allowed_collection_ids', set())
    if allowed is not None:
        if collection_ids:
            collection_ids = [cid for cid in collection_ids if cid in allowed]
        else:
            collection_ids = list(allowed)
    # Необязательный фильтр по типам материалов (список строк)
    material_types = []
    try:
        material_types = [str(x).strip().lower() for x in (data.get('material_types') or []) if str(x).strip()]
    except Exception:
        material_types = []
    # Необязательный диапазон лет
    year_from = (data.get('year_from') or '').strip()
    year_to = (data.get('year_to') or '').strip()
    # Необязательные фильтры тегов: ["key=value", ...]
    tag_filters = []
    try:
        tag_filters = [str(x) for x in (data.get('tag_filters') or []) if '=' in str(x)]
    except Exception:
        tag_filters = []

    # Расширяем и токенизируем
    keywords, lang_details = _ai_expand_keywords(query, multi_lang=all_languages, target_langs=requested_languages)
    base_tokens = _tokenize_query(query)
    keyword_plan = _plan_search_keywords(
        query,
        base_terms=base_tokens,
        ai_terms=keywords,
    )
    extra_tokens: list[str] = []
    ai_phrase_terms: list[str] = []
    if lang_details:
        preview_langs = ', '.join(f"{code}(+{count})" for code, count in lang_details if count > 0)
        if preview_langs:
            progress.add(f"Доп. языки: {preview_langs}")
    plan_must = [
        term for term in keyword_plan.get("must", [])
        if term and term not in mandatory_set
    ]
    plan_phrases = [
        term for term in keyword_plan.get("phrases", []) if term
    ]
    plan_broad = [
        term for term in keyword_plan.get("broad", []) if term
    ]
    plan_must_set = {term for term in plan_must}
    for term in plan_phrases:
        if term not in ai_phrase_terms and term not in mandatory_set:
            ai_phrase_terms.append(term)
    for term in plan_broad:
        extra_tokens.extend(_tokenize_query(term))
    for w in keywords:
        normalized_kw = _normalize_keyword_candidate(w)
        if normalized_kw:
            if ' ' in normalized_kw and normalized_kw not in mandatory_set:
                ai_phrase_terms.append(normalized_kw)
        extra_tokens.extend(_tokenize_query(w))
    if plan_must:
        progress.add("LLM ключи (строгие): " + ", ".join(plan_must[:6]) + ("…" if len(plan_must) > 6 else ""))
    if plan_phrases:
        progress.add("LLM фразы: " + ", ".join(plan_phrases[:6]) + ("…" if len(plan_phrases) > 6 else ""))
    if plan_broad:
        progress.add("LLM расширения: " + ", ".join(plan_broad[:6]) + ("…" if len(plan_broad) > 6 else ""))
    # уникальные термины (токены) с сохранением порядка, отдаём приоритет обязательным и LLM-фразам
    seen = set()
    terms: list[str] = []
    ordered_candidates = mandatory_terms + plan_must + ai_phrase_terms + base_tokens + extra_tokens + plan_broad
    for w in ordered_candidates:
        if not w:
            continue
        if w in seen:
            continue
        seen.add(w)
        terms.append(w)
    original_terms = list(terms)
    filtered_out_terms: list[str] = []
    if not terms and query:
        # запасной вариант: используем исходный запрос как термин
        terms = _tokenize_query(query) or [query.lower()]
    if terms:
        preview = ', '.join(terms[:6])
        if len(terms) > 6:
            preview += '…'
        progress.add(f"Ключевые термины: {preview}")
    else:
        progress.add("Ключевые термины: не выделены")

    # Предварительно считаем IDF по каждому термину
    idf = _idf_for_terms(terms)
    filtered_terms: list[str] = []
    removed_low_idf: list[str] = []
    for w in terms:
        if w in mandatory_set or w in plan_must_set:
            filtered_terms.append(w)
            continue
        if idf.get(w, 1.0) >= KEYWORD_IDF_MIN:
            filtered_terms.append(w)
        else:
            removed_low_idf.append(w)
    if removed_low_idf:
        filtered_out_terms.extend([w for w in removed_low_idf if w not in filtered_out_terms])
        progress.add(f"Фильтр ключевых слов: убрано {len(removed_low_idf)} общеупотребительных терминов")
    if not filtered_terms:
        fallback_terms = [w for w in terms if (w in mandatory_set) or len(w) >= 4]
        if fallback_terms:
            removed = [w for w in terms if w not in fallback_terms]
            if removed:
                filtered_out_terms.extend([w for w in removed if w not in filtered_out_terms])
            progress.add("Фильтр ключевых слов: оставляем только более длинные термины")
            filtered_terms = fallback_terms
    if filtered_terms:
        terms = filtered_terms
    try:
        feedback_rows = AiSearchKeywordFeedback.query.filter_by(query_hash=query_hash).all()
    except Exception:
        feedback_rows = []

    feedback_weights_lookup = _get_feedback_weights()
    feedback_applied_count = 0

    if terms:
        banned_terms = {str(row.keyword).lower() for row in feedback_rows if getattr(row, 'keyword', None) and row.action == 'irrelevant'}
        if banned_terms:
            before_len = len(terms)
            kept_terms: list[str] = []
            removed: list[str] = []
            for w in terms:
                if w in banned_terms and w not in mandatory_set:
                    removed.append(w)
                else:
                    kept_terms.append(w)
            terms = kept_terms
            if removed:
                filtered_out_terms.extend([w for w in removed if w not in filtered_out_terms])
            if len(terms) < before_len:
                progress.add(f"Фильтр обратной связи: игнорируем {before_len - len(terms)} термина по отзывам")
    durations['keywords'] = time.monotonic() - stage_start
    stage_start = time.monotonic()

    # Накапливаем кандидатов с оценками и попаданиями
    scores: dict[int, float] = {}
    hits: dict[int, list[dict]] = {}

    term_hits: dict[int, set[str]] = {}
    def add_score(fid: int, delta: float, hit: dict | None = None, term: str | None = None):
        scores[fid] = scores.get(fid, 0.0) + float(delta)
        if hit:
            hits.setdefault(fid, []).append(hit)
        if term:
            s = term_hits.setdefault(fid, set())
            if term:
                s.add(term)

    # Дополнительные кандидаты через FTS по ключевым терминам (учёт вопросительных слов)
    fts_candidates: List[int] = []
    fts_seed_terms = sorted(terms, key=lambda t: idf.get(t, 0.0), reverse=True)
    if not fts_seed_terms:
        backup_tokens = _tokenize_query(query)
        if backup_tokens:
            fts_seed_terms = sorted(
                backup_tokens,
                key=lambda t: idf.get(t, 0.0),
                reverse=True,
            )
    fts_seed_terms = [t for t in fts_seed_terms if len(t) >= 3][:5]
    fts_seed_query = " ".join(fts_seed_terms) if fts_seed_terms else query.strip()
    if fts_seed_query:
        try:
            fts_candidates = search_service.candidate_ids(fts_seed_query, limit=max(200, max_candidates * 10)) or []
        except Exception:
            fts_candidates = []
    if fts_candidates:
        base_term = fts_seed_terms[0] if fts_seed_terms else None
        if base_term:
            progress.add(f"FTS кандидаты: {len(fts_candidates)} (ядро: {base_term})")
        else:
            progress.add(f"FTS кандидаты: {len(fts_candidates)}")
        fts_cap = min(len(fts_candidates), max_candidates * 10)
        for fid in fts_candidates[:fts_cap]:
            if base_term:
                add_score(
                    fid,
                    AI_SCORE_FTS * idf.get(base_term, 1.0),
                    {"type": "fts", "term": base_term},
                    term=base_term,
                )
            else:
                add_score(
                    fid,
                    AI_SCORE_FTS,
                    {"type": "fts"},
                )
    else:
        progress.add("FTS кандидаты: не обнаружены")

    raw_query_clean = query.strip()
    raw_fts_candidates: List[int] = []
    if raw_query_clean:
        try:
            raw_fts_candidates = search_service.candidate_ids(raw_query_clean, limit=max(200, max_candidates * 10)) or []
        except Exception:
            raw_fts_candidates = []
    if raw_fts_candidates:
        primary_term = next((tok for tok in base_tokens if len(tok) >= 3), None)
        if not primary_term and terms:
            primary_term = next((tok for tok in terms if len(tok) >= 3), None)
        label_term = primary_term or raw_query_clean.lower()
        progress.add(f"FTS оригинальный запрос: {len(raw_fts_candidates)} кандидатов")
        fts_raw_cap = min(len(raw_fts_candidates), max_candidates * 10)
        for fid in raw_fts_candidates[:fts_raw_cap]:
            add_score(
                fid,
                AI_SCORE_FTS_RAW * (idf.get(label_term, 1.0) if isinstance(label_term, str) else 1.0),
                {"type": "fts_raw", "term": label_term},
                term=str(label_term) if isinstance(label_term, str) else None,
            )
    else:
        progress.add("FTS оригинальный запрос: не обнаружены")

    # Совпадения по тегам
    if use_tags:
        for w in terms:
            like = f"%{w}%"
            try:
                q = (
                    db.session.query(Tag.file_id, Tag.key, Tag.value)
                    .join(File, Tag.file_id == File.id)
                    .outerjoin(Collection, File.collection_id == Collection.id)
                    .filter(or_(Collection.searchable == True, Collection.id.is_(None)))
                    .filter(or_(Tag.value.ilike(like), Tag.key.ilike(like)))
                )
                if collection_ids:
                    q = q.filter(File.collection_id.in_(collection_ids))
                if material_types:
                    q = q.filter(func.lower(File.material_type).in_(material_types))
                if year_from:
                    q = q.filter(File.year >= year_from)
                if year_to:
                    q = q.filter(File.year <= year_to)
                rows = q \
                     .limit(4000).all()
                for fid, k, v in rows:
                    add_score(fid, AI_SCORE_TAG * idf.get(w, 1.0), {"type": "tag", "key": k, "value": v, "term": w}, term=w)
            except Exception:
                pass
        progress.add(f"Теги: найдено кандидатов {len(scores)}")
    else:
        progress.add("Теги: пропущено (источник отключён)")

    # Совпадения по полям файла
    if use_text:
        for w in terms:
            like = f"%{w}%"
            try:
                q = (
                    db.session.query(
                        File.id,
                        File.title,
                        File.author,
                        File.keywords,
                        File.text_excerpt,
                        File.abstract,
                    )
                    .outerjoin(Collection, File.collection_id == Collection.id)
                    .filter(or_(Collection.searchable == True, Collection.id.is_(None)))
                    .filter(
                        or_(
                            File.title.ilike(like),
                            File.author.ilike(like),
                            File.keywords.ilike(like),
                            File.text_excerpt.ilike(like),
                            File.abstract.ilike(like),
                        )
                    )
                )
                if collection_ids:
                    q = q.filter(File.collection_id.in_(collection_ids))
                if material_types:
                    q = q.filter(func.lower(File.material_type).in_(material_types))
                if year_from:
                    q = q.filter(File.year >= year_from)
                if year_to:
                    q = q.filter(File.year <= year_to)
                rows = q.limit(4000).all()
                for fid, title, author, kws, excerpt, abstract in rows:
                    if title and re.search(re.escape(w), title, flags=re.I):
                        add_score(fid, AI_SCORE_TITLE * idf.get(w, 1.0), {"type": "title", "term": w}, term=w)
                    if author and re.search(re.escape(w), author, flags=re.I):
                        add_score(fid, AI_SCORE_AUTHOR * idf.get(w, 1.0), {"type": "author", "term": w}, term=w)
                    if kws and re.search(re.escape(w), kws, flags=re.I):
                        add_score(fid, AI_SCORE_KEYWORDS * idf.get(w, 1.0), {"type": "keywords", "term": w}, term=w)
                    if excerpt and re.search(re.escape(w), excerpt, flags=re.I):
                        add_score(fid, AI_SCORE_EXCERPT * idf.get(w, 1.0), {"type": "excerpt", "term": w}, term=w)
                    if abstract and re.search(re.escape(w), abstract, flags=re.I):
                        add_score(fid, AI_SCORE_ABSTRACT * idf.get(w, 1.0), {"type": "abstract", "term": w}, term=w)
            except Exception:
                pass
        progress.add(f"Метаданные: найдено кандидатов {len(scores)}")
    else:
        progress.add("Метаданные: пропущено (источник отключён)")
    if plan_must:
        coverage_lines: List[str] = []
        for key in plan_must[:6]:
            count = sum(1 for matched in term_hits.values() if key in matched)
            coverage_lines.append(f"{key}:{count}")
        progress.add("Покрытие must-термов: " + (", ".join(coverage_lines) or "нет совпадений"))
    if mandatory_set and scores:
        kept_scores: dict[int, float] = {}
        kept_hits: dict[int, list[dict]] = {}
        kept_term_hits: dict[int, set[str]] = {}
        dropped = 0
        for fid, score_val in scores.items():
            matched = term_hits.get(fid, set())
            if mandatory_set.issubset(matched):
                kept_scores[fid] = score_val
                if fid in hits:
                    kept_hits[fid] = hits[fid]
                if matched:
                    kept_term_hits[fid] = matched
            else:
                dropped += 1
        if dropped:
            progress.add(f"Обязательные фразы: исключено {dropped} кандидатов без совпадений")
        scores = kept_scores
        hits = kept_hits
        term_hits = kept_term_hits
    durations['candidates'] = time.monotonic() - stage_start
    stage_start = time.monotonic()

    # Формируем результаты
    file_ids = list(scores.keys())
    results = []
    if file_ids:
        q_files = (
            File.query.outerjoin(Collection, File.collection_id == Collection.id)
            .filter(or_(Collection.searchable == True, Collection.id.is_(None)))
            .filter(File.id.in_(file_ids))
        )
        if collection_ids:
            q_files = q_files.filter(File.collection_id.in_(collection_ids))
        if material_types:
            q_files = q_files.filter(func.lower(File.material_type).in_(material_types))
        if year_from:
            q_files = q_files.filter(File.year >= year_from)
        if year_to:
            q_files = q_files.filter(File.year <= year_to)
        # Применяем фильтры тегов на финальном этапе
        for tf in tag_filters:
            try:
                k, v = tf.split('=', 1)
                k = (k or '').strip(); v = (v or '').strip()
                if not k or not v:
                    continue
                if k.lower() == 'author':
                    q_files = q_files.filter(File.author.ilike(f"%{v}%"))
                else:
                    tflt = aliased(Tag)
                    q_files = q_files.join(tflt, tflt.file_id == File.id).filter(and_(tflt.key == k, tflt.value.ilike(f"%{v}%")))
            except Exception:
                continue
        q_files = q_files.all()
        id2file = {f.id: f for f in q_files}
        file_lookup = id2file
        for fid, sc in scores.items():
            f = id2file.get(fid)
            if not f:
                continue
            # усиление за точную фразу (исходный запрос как фраза)
            phrase_boost = 0.0
            qraw = query.strip()
            if len(qraw) >= 3:
                try:
                    pat = re.escape(qraw)
                    if f.title and re.search(pat, f.title, flags=re.I):
                        phrase_boost += AI_BOOST_PHRASE
                    if f.keywords and re.search(pat, f.keywords, flags=re.I):
                        phrase_boost += AI_BOOST_PHRASE * 0.6
                except Exception:
                    pass
            # усиление за покрытие разных терминов
            n_terms = len(term_hits.get(fid, set()))
            coverage_boost = max(0, n_terms - 1) * AI_BOOST_MULTI
            # сниппеты из кэшированного фрагмента/аннотации
            snips = []
            snippet_sources: list[str] = []
            try:
                text_cache = _read_cached_excerpt_for_file(f)
                snips = _collect_snippets(text_cache, terms, max_snips=max_snippets) if text_cache else []
                if snips:
                    snippet_sources.append('excerpt-cache')
            except Exception:
                snips = []
            if (not snips or len(snips) < max_snippets) and getattr(f, 'abstract', None):
                try:
                    extra_count = max(1, max_snippets - len(snips))
                    extra = _collect_snippets(getattr(f, 'abstract'), terms, max_snips=extra_count)
                    added = False
                    for item in extra:
                        if item and item not in snips:
                            snips.append(item)
                            added = True
                    if added:
                        snippet_sources.append('abstract')
                except Exception:
                    pass
            # усиление за близость: несколько терминов в одном сниппете
            prox_boost = 0.0
            if snips:
                for s in snips:
                    present = 0
                    sl = s.lower()
                    for w in terms:
                        if w in sl:
                            present += 1
                    if present >= 2:
                        prox_boost += (present - 1) * AI_BOOST_SNIPPET_COOCCUR
            feedback_weight = 0.0
            feedback_boost = 0.0
            feedback_entry = feedback_weights_lookup.get(fid)
            if feedback_entry:
                feedback_weight = float(feedback_entry.get('weight', 0.0) or 0.0)
                if feedback_weight:
                    feedback_boost = feedback_weight * AI_FEEDBACK_WEIGHT_SCALE
                    feedback_applied_count += 1
            final_score = sc + phrase_boost + coverage_boost + prox_boost + feedback_boost
            results.append({
                "file_id": fid,
                "rel_path": f.rel_path,
                "title": f.title or f.filename,
                "score": round(final_score, 3),
                "hits": hits.get(fid, []),
                "snippets": snips,
                "snippet_sources": snippet_sources,
                "feedback_weight": round(feedback_weight, 6),
                "feedback_boost": round(feedback_boost, 3),
            })
        # сортируем по убыванию балла, затем по дате изменения
        results.sort(key=lambda x: (x.get('score') or 0.0, id2file.get(x['file_id']).mtime or 0.0), reverse=True)
        total_ranked = len(results)
        progress.add(f"Ранжирование: {total_ranked} кандидатов")
        if feedback_applied_count:
            progress.add(f"Фидбек: применён вес для {feedback_applied_count} документов")
        if total_ranked > max_candidates:
            progress.add(f"Ограничиваем до {max_candidates} лучших по баллу")
            results = results[:max_candidates]
        dismissed_file_ids = {int(row.file_id) for row in feedback_rows if getattr(row, 'file_id', None) and row.action == 'irrelevant'}
        if dismissed_file_ids:
            before_len = len(results)
            results = [res for res in results if int(res.get('file_id') or 0) not in dismissed_file_ids]
            after_len = len(results)
            if after_len < before_len:
                progress.add(f"Обратная связь: исключено {before_len - after_len} источников, помеченных как нерелевантные")

        for idx, res in enumerate(results, start=1):
            cand_title = (res.get('title') or res.get('rel_path') or f"file-{res.get('file_id')}").strip()
            progress.add(f"Кандидат [{idx}]: {cand_title}")
        if (deep_search or full_text) and results:
            if deep_search and full_text:
                scan_label = "Глубокий поиск + full-text"
                scan_step = "Глубокий поиск"
            elif deep_search:
                scan_label = "Глубокий поиск по контенту"
                scan_step = "Глубокий поиск"
            else:
                scan_label = "Полнотекстовый поиск"
                scan_step = "Full-text"
            scan_limit = len(results) if full_text else min(len(results), max(top_k * 2, 5))
            progress.add(f"{scan_label}: проверяем {scan_limit} файлов (чанк {chunk_chars}, чанков ≤ {max_chunks})")
            for idx, res in enumerate(results[:scan_limit], start=1):
                f = id2file.get(res['file_id'])
                if not f:
                    continue
                scan = _deep_scan_file(f, terms, query, chunk_chars=chunk_chars, max_chunks=max_chunks, max_snippets=max_snippets)
                boost = float(scan.get('boost') or 0.0)
                if boost:
                    res['score'] = round((res.get('score') or 0.0) + boost, 3)
                matched_terms = scan.get('matched_terms') or set()
                if matched_terms:
                    res['matched_terms'] = sorted({str(t) for t in matched_terms})
                if scan.get('hit'):
                    new_snips = scan.get('snippets') or []
                    if new_snips:
                        base_snips = res.get('snippets') or []
                        combined = base_snips + [sn for sn in new_snips if sn not in base_snips]
                        res['snippets'] = combined[:max_snippets]
                    sources = scan.get('sources') or []
                    if sources:
                        current_sources = set(res.get('snippet_sources') or [])
                        for src in sources:
                            if src not in current_sources:
                                current_sources.add(src)
                        res['snippet_sources'] = sorted(current_sources)
                title = (f.title or f.filename or f.rel_path or f"file-{f.id}").strip()
                srcs = ", ".join(scan.get('sources') or [])
                status = 'совпадения' if scan.get('hit') else 'нет'
                if srcs:
                    status += f"; источники: {srcs}"
                progress.add(f"{scan_step} [{idx}/{scan_limit}]: {title} — {status}")
            results.sort(key=lambda x: (x.get('score') or 0.0, id2file.get(x['file_id']).mtime or 0.0), reverse=True)
        durations['deep'] = time.monotonic() - stage_start
        stage_start = time.monotonic()
        results = results[:top_k]
        progress.add(f"Отобрано документов: {len(results)}")
        for idx, res in enumerate(results, start=1):
            title = (res.get('title') or res.get('rel_path') or f"file-{res.get('file_id')}").strip()
            sources = res.get('snippet_sources') or []
            if not sources:
                summary = 'не найдены'
            else:
                summary = ', '.join(sources)
            progress.add(f"Снипеты [{idx}] {title}: {summary}")
        if llm_snippets and results:
            total_targets = len(results)
            progress.add(f"LLM сниппеты: формируем для {total_targets} документов")
            for idx, res in enumerate(results, start=1):
                title = (res.get('title') or res.get('rel_path') or f"file-{res.get('file_id')}").strip()
                base_snippets = res.get('snippets') or []
                context = " ".join(base_snippets).strip()
                if not context:
                    progress.add(f"LLM сниппеты [{idx}/{total_targets}]: {title} — пропущено (нет контекста)")
                    continue
                cache_entry = _get_cached_snippet(res.get('file_id'), query_hash, True)
                if cache_entry and cache_entry.snippet:
                    translated_snippet, translated = _ensure_russian_text(cache_entry.snippet, label='snippet-cache')
                    res['llm_snippet'] = translated_snippet
                    progress.add(f"LLM сниппеты [{idx}/{total_targets}]: {title} — cache hit" + (' (перевод)' if translated else ''))
                    if translated:
                        _store_snippet_cache(res.get('file_id'), query_hash, True, translated_snippet, meta={'len': len(translated_snippet)}, ttl_hours=SNIPPET_CACHE_TTL_HOURS)
                    continue
                ctx = context[:1200]
                system = (
                    "Ты создаёшь короткий сниппет (до двух предложений) для поисковой выдачи. "
                    "Используй только предоставленный текст, не добавляй новых фактов. "
                    "Отвечай на русском языке."
                )
                user_msg = (
                    f"Запрос: {query}\n"
                    f"Фрагменты документа:\n{ctx}\n"
                    "Сформулируй 1-2 предложения с ключевыми фактами и сохрани смысл. Ответ дай на русском языке."
                )
                snippet_text = ""
                error_logged = False
                try:
                    snippet_text = (call_lmstudio_compose(system, user_msg, temperature=0.0, max_tokens=min(180, _lm_max_output_tokens())) or '').strip()
                except Exception as exc:
                    error_logged = True
                    progress.add(f"LLM сниппеты [{idx}/{total_targets}]: {title} — ошибка: {exc}")
                if snippet_text:
                    snippet_text, translated = _ensure_russian_text(snippet_text, label='snippet')
                    res['llm_snippet'] = snippet_text
                    note = f" ({len(snippet_text)} символов)"
                    if translated:
                        note += ", переведено"
                    progress.add(f"LLM сниппеты [{idx}/{total_targets}]: {title} — готово{note}")
                    _store_snippet_cache(res.get('file_id'), query_hash, True, snippet_text, meta={'len': len(snippet_text)})
                elif not error_logged:
                    progress.add(f"LLM сниппеты [{idx}/{total_targets}]: {title} — пустой ответ")
            durations['llm_snippets'] = time.monotonic() - stage_start
            stage_start = time.monotonic()
    else:
        progress.add("Совпадения не найдены")

    # Необязательный короткий ответ, используя сниппеты как контекст (поисковый промпт)
    answer = ""
    rag_context_sections: List[ContextSection] = []
    rag_validation: Optional[ValidationResult] = None
    rag_session_id: Optional[int] = None
    rag_bundle: Optional[Dict[str, Any]] = None
    rag_notes: list[str] = []
    rag_risk_info: Optional[Dict[str, Any]] = None
    rag_sources_payload: List[Dict[str, Any]] = []
    rag_auto_retry_attempted = False
    context_chunk_limit = max(4, min(8, top_k * 2))
    language_filters = requested_languages if requested_languages else None
    if rag_enabled:
        if results:
            rag_bundle, rag_notes = _prepare_rag_context(
                query,
                results,
                language_filters=language_filters,
                max_chunks=context_chunk_limit,
                progress=progress,
            )
        else:
            rag_notes.append("Нет результатов поиска, RAG контекст пропущен.")
    else:
        pass
    fallback_required = True
    if rag_enabled and rag_bundle and rag_bundle.get("sections"):
        fallback_required = False
        rag_notes.extend(rag_bundle.get("notes", []))
        stage_start = time.monotonic()
        answer, rag_validation, rag_session_id, _, _ = _generate_rag_answer(
            query,
            rag_bundle,
            temperature=0.1,
            max_tokens=min(350, _lm_max_output_tokens()),
            progress=progress,
        )
        durations['llm_answer'] = time.monotonic() - stage_start
        rag_context_sections = rag_bundle.get("sections", [])
        rag_risk_info = rag_bundle.get("risk")
        rag_sources_payload = rag_bundle.get("sources", [])
        retry_enabled = bool(getattr(_rt(), "ai_rag_retry_enabled", True))
        retry_threshold = float(getattr(_rt(), "ai_rag_retry_threshold", 0.6) or 0.6)
        initial_score = float(rag_risk_info.get("score", 0.0)) if rag_risk_info else 0.0
        if retry_enabled and initial_score >= retry_threshold:
            rag_auto_retry_attempted = True
            if progress:
                progress.add(f"RAG автоповтор: риск {initial_score:.2f} ≥ порога {retry_threshold:.2f}, расширяем контекст")
            rag_notes.append("Автоповтор RAG: риск ответа высокий, расширяем контекст.")
            retry_chunk_limit = min(max(6, context_chunk_limit + 2), max(6, top_k * 3))
            retry_bundle, retry_extra_notes = _prepare_rag_context(
                query,
                results,
                language_filters=language_filters,
                max_chunks=retry_chunk_limit,
                progress=progress,
                variant_boost=1,
                dense_boost=1.4,
            )
            if retry_bundle and retry_bundle.get("sections"):
                rag_notes.extend(retry_extra_notes)
                stage_start = time.monotonic()
                answer, rag_validation, rag_session_id, _, _ = _generate_rag_answer(
                    query,
                    retry_bundle,
                    temperature=0.1,
                    max_tokens=min(350, _lm_max_output_tokens()),
                    progress=progress,
                )
                durations['llm_answer'] = time.monotonic() - stage_start
                rag_bundle = retry_bundle
                rag_context_sections = retry_bundle.get("sections", [])
                rag_risk_info = retry_bundle.get("risk")
                rag_sources_payload = retry_bundle.get("sources", [])
                if rag_risk_info:
                    if progress:
                        progress.add(f"RAG автоповтор: новый риск {float(rag_risk_info.get('score', 0.0)):.2f}")
                    rag_notes.append("Автоповтор RAG выполнен: получен обновлённый ответ.")
                else:
                    rag_notes.append("Автоповтор RAG выполнен: риск не определён.")
            else:
                if progress:
                    progress.add("RAG автоповтор: не удалось сформировать дополнительный контекст")
                rag_notes.append("Автоповтор RAG: не удалось подобрать дополнительный контекст.")
    if fallback_required:
        stage_start = time.monotonic()
        if rag_notes:
            prefix = "RAG заметка" if not rag_enabled else "RAG fallback"
            for note in rag_notes:
                progress.add(f"{prefix}: {note}")
        topn = results[:10]
        snippet_context: List[Dict[str, Any]] = []
        if topn:
            core_terms = [term for term in base_tokens if len(term) >= 4] or [
                term for term in terms if len(term) >= 4
            ]
            anchor_terms = [term for term in base_tokens if len(term) >= 3] or [query.lower()]
            snippet_answer, used_count, snippet_context = _compose_snippet_based_answer(
                query,
                topn,
                primary_terms=core_terms,
                anchor_terms=anchor_terms,
                term_idf=idf,
                max_items=max(3, top_k),
            )
            if snippet_answer:
                llm_snippet_answer = _generate_snippet_llm_summary(
                    query,
                    snippet_context,
                    progress=progress,
                )
                if llm_snippet_answer:
                    if snippet_context and re.search(r"недостаточн[оы]", llm_snippet_answer, flags=re.IGNORECASE):
                        # Считаем, что LLM осторожничает — используем консервативный ответ по сниппетам
                        answer = snippet_answer
                        progress.add("LLM ответ по сниппетам отклонён: сообщает о нехватке данных")
                    else:
                        answer = llm_snippet_answer
                        progress.add(f"LLM ответ по сниппетам: использовано {len(snippet_context)} источников")
                else:
                    answer = snippet_answer
                    progress.add(f"Сниппет-ответ: использовано {used_count} источников (LLM не применён)")
            else:
                answer = fallback_answer()
                progress.add("Сниппет-ответ: подходящих фрагментов нет, возвращаем заглушку")
        else:
            answer = fallback_answer()
            progress.add("Сниппет-ответ: список результатов пуст")
        if snippet_context:
            detail_refs = []
            for item in snippet_context:
                title = item.get("title") or f"источник {item.get('order')}"
                detail_refs.append(f"[{item['order']}] {title}")
            if detail_refs:
                details_text = "Подробнее см.: " + "; ".join(detail_refs)
                answer = (answer + "\n\n" + details_text) if answer else details_text
        durations['llm_answer'] = time.monotonic() - stage_start
        stage_start = time.monotonic()

    # Необязательно: лёгкое реранжирование топ-15 через LLM с контекстом сниппетов
    if _rt().ai_rerank_llm and results:
        try:
            top = results[:15]
            prompt_lines = [f"[{i+1}] id={it['file_id']} :: { (it.get('snippets') or [''])[0] }" for i, it in enumerate(top)]
            prompt = "\n".join(prompt_lines)
            sys = "Ты ранжируешь источники по релевантности к запросу. Верни JSON-массив id в порядке убывания релевантности."
            user_prompt = f"Запрос: {query}\nИсточники:\n{prompt}\nОтвети только JSON массивом id."
            order = None
            last_error: Exception | None = None
            for choice in _llm_iter_choices('rerank'):
                label = _llm_choice_label(choice)
                url = _llm_choice_url(choice)
                if not url:
                    app.logger.warning(f"LLM rerank endpoint некорректен ({label}): пустой URL")
                    continue
                payload = {
                    "model": choice.get('model') or LMSTUDIO_MODEL,
                    "messages": [
                        {"role": "system", "content": sys},
                        {"role": "user", "content": user_prompt},
                    ],
                    "temperature": 0.0,
                    "max_tokens": 200,
                }
                try:
                    rr = http_request(
                        'POST',
                        url,
                        headers=_llm_choice_headers(choice),
                        json=payload,
                        timeout=(HTTP_CONNECT_TIMEOUT, max(HTTP_DEFAULT_TIMEOUT, 60)),
                        logger=app.logger,
                    )
                    if _llm_response_indicates_busy(rr):
                        app.logger.info(f"LLM rerank endpoint занята ({label}), переключаемся")
                        last_error = RuntimeError('busy')
                        continue
                    rr.raise_for_status()
                    dd = rr.json()
                    content = (dd.get("choices", [{}])[0].get("message", {}) or {}).get("content", "")
                    try:
                        order = json.loads(content)
                    except Exception:
                        m = re.search(r"\[(?:\s*\d+\s*,?\s*)+\]", content)
                        if m:
                            order = json.loads(m.group(0))
                    if order is not None:
                        break
                except Exception as e:
                    last_error = e
                    if isinstance(e, RuntimeError) and str(e) == 'llm_cache_only_mode':
                        app.logger.info(f"LLM rerank пропущен (cache-only, {label})")
                    else:
                        app.logger.warning(f"LLM rerank failed ({label}): {e}")
                    continue
            if order is None and last_error and str(last_error) not in {'busy', 'llm_cache_only_mode'}:
                app.logger.warning(f"LLM rerank failed: {last_error}")
            if isinstance(order, list) and all(isinstance(x, int) for x in order):
                pos = {int(fid): i for i, fid in enumerate(order)}
                results.sort(key=lambda x: (pos.get(int(x['file_id']), 10**6), -(x.get('score') or 0.0)))
                progress.add("Реранжирование LLM: порядок обновлён")
        except Exception:
            pass

    term_usage_counts: Dict[str, int] = {}
    for matched in term_hits.values():
        for term in matched:
            term_usage_counts[term] = term_usage_counts.get(term, 0) + 1
    final_keywords: List[str] = []
    seen_keywords: set[str] = set()
    base_tokens_limit = base_tokens[:4]
    for term in terms:
        if term in seen_keywords:
            continue
        if (
            term_usage_counts.get(term, 0) > 0
            or term in mandatory_set
            or term in base_tokens_limit
        ):
            final_keywords.append(term)
            seen_keywords.add(term)
    if not final_keywords:
        for term in base_tokens + list(terms) + plan_must:
            if term and term not in seen_keywords:
                final_keywords.append(term)
                seen_keywords.add(term)
            if len(final_keywords) >= 6:
                break
    filtered_keywords = sorted({w for w in filtered_out_terms if w})
    durations['total'] = time.monotonic() - t_total_start
    rag_context_payload = [_serialize_context_section(section) for section in rag_context_sections]
    extra_meta = {
        'query_preview': query[:200],
        'answer_preview': (answer or '')[:400],
        'keyword_count': len(final_keywords),
        'filtered_keywords': filtered_keywords,
        'mandatory_terms': mandatory_terms,
    }
    extra_meta.update({
        'rag_context_count': len(rag_context_sections),
        'rag_fallback': bool(rag_context_sections) is False and rag_enabled,
        'rag_hallucination_warning': bool(rag_validation and rag_validation.hallucination_warning),
        'rag_risk_score': float(rag_risk_info.get('score', 0.0)) if rag_risk_info else 0.0,
        'rag_auto_retry': bool(rag_auto_retry_attempted),
    })
    _record_search_metric(query_hash, durations, user, extra_meta)

    return {
        "query": query,
        "query_hash": query_hash,
        "keywords": final_keywords,
        "filtered_keywords": filtered_keywords,
        "mandatory_terms": mandatory_terms,
        "answer": answer,
        "items": results,
        "progress": progress.lines,
        "rag_context": rag_context_payload,
        "rag_context_groups": _group_context_sections_by_document(rag_context_sections),
        "rag_validation": rag_validation.as_dict() if rag_validation else None,
        "rag_risk": rag_risk_info,
        "rag_sources": rag_sources_payload,
        "rag_session_id": rag_session_id,
        "rag_notes": rag_notes,
        "rag_fallback": bool(rag_context_sections) is False and rag_enabled,
        "rag_enabled": rag_enabled,
        "rag_retry": bool(rag_auto_retry_attempted),
    }


def _resolve_problem_score(problem: Dict[str, Any]) -> Optional[float]:
    score_keys = (
        'score',
        'humanScore',
        'human_score',
        'remoteScore',
        'remote_score',
        'autoScore',
        'auto_score',
    )
    for key in score_keys:
        if key in problem:
            value = problem.get(key)
            try:
                return float(value) if value is not None else None
            except (TypeError, ValueError):
                continue
    return None


def _split_text_for_dataset(
    text: str,
    *,
    min_chars: int = 120,
    max_chars: int = 700,
    max_segments: int = 6,
) -> List[str]:
    cleaned = (text or '').strip()
    if not cleaned:
        return []

    # Сначала пробуем разбить по пустым строкам
    candidates = [segment.strip() for segment in re.split(r"\n\s*\n+", cleaned) if segment.strip()]

    segments: List[str] = []
    for block in candidates:
        if len(block) <= max_chars:
            if len(block) >= min_chars:
                segments.append(block)
            continue

        sentences = re.split(r"(?<=[.!?])\s+", block)
        buffer: List[str] = []
        current = ''
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            tentative = f"{current} {sentence}".strip() if current else sentence
            if len(tentative) > max_chars:
                if current:
                    segments.append(current.strip())
                if len(sentence) >= min_chars:
                    segments.append(sentence[:max_chars].strip())
                current = ''
            else:
                current = tentative
        if current and len(current) >= min_chars:
            segments.append(current.strip())

    if not segments:
        if len(cleaned) >= min_chars:
            segments = [cleaned[:max_chars]]

    return segments[:max_segments]


def _extract_pairs_from_response(raw: str) -> List[Dict[str, str]]:
    if not raw:
        return []
    start = raw.find('{')
    end = raw.rfind('}')
    if start == -1 or end == -1 or end <= start:
        return []
    snippet = raw[start:end + 1]
    try:
        payload = json.loads(snippet)
    except json.JSONDecodeError:
        try:
            snippet = re.sub(r".*?(\{.*\})", r"\1", raw, flags=re.S)
            payload = json.loads(snippet)
        except Exception:
            return []
    items = payload.get('pairs')
    if not isinstance(items, list):
        return []
    pairs: List[Dict[str, str]] = []
    for entry in items:
        if not isinstance(entry, dict):
            continue
        question = str(entry.get('question') or '').strip()
        answer = str(entry.get('answer') or '').strip()
        if question and answer:
            pairs.append({'question': question, 'answer': answer})
    return pairs


def _generate_pairs_from_snippet(
    snippet: str,
    *,
    desired_pairs: int,
    temperature: float,
    max_tokens: int,
    problem_question: Optional[str],
) -> List[Dict[str, str]]:
    paragraph = (snippet or '').strip()
    if not paragraph:
        return []

    system_prompt = (
        'Ты — помощник по подготовке данных для обучения. Отвечай на русском языке, '
        'формируй содержательные вопросы и ответы и соблюдай формат JSON без дополнительных комментариев.'
    )
    intro = (
        "Тебе дан абзац текста. Сформируй {count} пар \"вопрос-ответ\" для тонкой настройки модели. "
        "В каждом вопросе делай акцент на фактах, причинно-следственных связях, числах или важных утверждениях из абзаца. "
        "Не упоминай, что вопрос взят из текста. Ответ должен быть вытекающим только из абзаца.\n\n"
    )
    if problem_question:
        intro += f"Проблемный пользовательский вопрос: {problem_question.strip()}\n"

    user_prompt = (
        f"{intro}Абзац:\n\"\"\"\n{paragraph}\n\"\"\"\n\n"
        "Формат ответа:\n"
        "{\n  \"pairs\": [\n    { \"question\": \"...\", \"answer\": \"...\" }\n  ]\n}\n"
        f"Количество элементов в массиве должно быть ровно {desired_pairs}. Не добавляй других полей."
    )

    try:
        raw = call_lmstudio_compose(
            system_prompt,
            user_prompt,
            temperature=max(0.0, float(temperature)),
            max_tokens=max(128, int(max_tokens)),
        )
    except Exception as exc:
        app.logger.warning(f"QA generation failed: {exc}")
        return []

    pairs = _extract_pairs_from_response(raw)
    return pairs


def _launch_fine_tune_job(
    dataset: Sequence[Dict[str, str]],
    fine_tune_config: Dict[str, Any],
) -> Dict[str, Any]:
    server_url = (fine_tune_config.get('server_url') or '').strip()
    base_model_path = (fine_tune_config.get('base_model_path') or '').strip()
    config_payload = fine_tune_config.get('config') or {}
    if not server_url or not base_model_path or not isinstance(config_payload, dict):
        raise ValueError('server_url, base_model_path и config обязательны для запуска дообучения')

    payload: Dict[str, Any] = {
        'base_model_path': base_model_path,
        'dataset': [
            {
                'input': item['input'],
                'output': item['output'],
                'source': item.get('source'),
            }
            for item in dataset
        ],
        'config': config_payload,
    }

    if fine_tune_config.get('output_dir'):
        payload['output_dir'] = fine_tune_config['output_dir']

    endpoint = server_url.rstrip('/') + '/v1/fine-tunes'
    headers = fine_tune_config.get('headers')
    timeout = float(fine_tune_config.get('timeout', 900))
    response = http_request(
        'POST',
        endpoint,
        json=payload,
        headers=headers,
        timeout=(HTTP_CONNECT_TIMEOUT, max(timeout, HTTP_DEFAULT_TIMEOUT)),
        logger=app.logger,
    )
    if response.status_code >= 400:
        raise RuntimeError(f"Fine-tune request failed: {response.status_code} {response.text[:200]}")
    return response.json()


def _build_dataset_from_problems(
    payload: Dict[str, Any],
    logs: List[str],
) -> Tuple[List[Dict[str, str]], Dict[str, Any]]:
    evaluation = payload.get('evaluation') or {}
    items = evaluation.get('items') or []
    if not isinstance(items, list) or not items:
        raise ValueError('evaluation.items должен быть непустым списком')

    generation_opts = payload.get('generation') or {}
    search_opts = payload.get('search') or {}

    score_threshold = payload.get('score_threshold', evaluation.get('quality_gate'))
    try:
        threshold_value = float(score_threshold) if score_threshold is not None else None
    except (TypeError, ValueError):
        threshold_value = None

    max_questions = int(generation_opts.get('max_questions') or payload.get('max_questions') or len(items))
    target_pairs = int(payload.get('target_pairs') or generation_opts.get('target_pairs') or generation_opts.get('total_pairs') or 20)
    pairs_per_snippet = max(1, int(generation_opts.get('pairs_per_snippet') or 1))
    include_reference_pair = bool(generation_opts.get('include_reference_pair', True))
    min_paragraph_chars = max(30, int(generation_opts.get('min_paragraph_chars', 120)))
    max_paragraph_chars = max(min_paragraph_chars, int(generation_opts.get('max_paragraph_chars', 700)))
    llm_temperature = float(generation_opts.get('temperature', 0.2))
    llm_max_tokens = int(generation_opts.get('max_tokens', 512))

    scored_items: List[Tuple[Optional[float], int, Dict[str, Any]]] = []
    for idx, raw in enumerate(items):
        if not isinstance(raw, dict):
            continue
        score = _resolve_problem_score(raw)
        if threshold_value is None or (score is None) or (score <= threshold_value):
            scored_items.append((score, idx, raw))

    if not scored_items:
        raise ValueError('Не найдено вопросов, требующих внимания — уточните порог или список')

    scored_items.sort(key=lambda entry: (entry[0] if entry[0] is not None else -1.0, entry[1]))
    selected_items = scored_items[:max(1, max_questions)]

    dataset: List[Dict[str, str]] = []
    dataset_dedup: set[Tuple[str, str]] = set()
    processed = 0

    for rank, (score, _, problem) in enumerate(selected_items, start=1):
        question = str(problem.get('question') or '').strip()
        reference = str(problem.get('referenceAnswer') or problem.get('reference') or '').strip()
        logs.append(f"[problem {rank}] score={score!r} question='{question[:80]}'")

        if include_reference_pair and question and reference:
            key = (question, reference)
            if key not in dataset_dedup:
                dataset.append({'input': question, 'output': reference, 'source': 'evaluation'})
                dataset_dedup.add(key)
                logs.append(f"  добавлена пара из эталонного ответа")
                if len(dataset) >= target_pairs:
                    break

        search_parts = [question]
        if reference:
            search_parts.append(reference)
        search_query = ' '.join(part for part in search_parts if part).strip()
        if not search_query:
            logs.append('  пропущено: пустой поисковый запрос')
            continue

        search_payload: Dict[str, Any] = {
            'query': search_query,
            'top_k': search_opts.get('top_k', 5),
            'deep_search': bool(search_opts.get('deep_search', True)),
            'llm_snippets': False,
            'max_candidates': search_opts.get('max_candidates', 20),
            'max_snippets': search_opts.get('max_snippets', 3),
            'chunk_chars': search_opts.get('chunk_chars', 5000),
            'max_chunks': search_opts.get('max_chunks', 25),
        }
        if 'sources' in search_opts:
            search_payload['sources'] = search_opts['sources']
        if 'collection_ids' in search_opts:
            search_payload['collection_ids'] = search_opts['collection_ids']
        if 'material_types' in search_opts:
            search_payload['material_types'] = search_opts['material_types']

        progress_lines: List[str] = []
        result = _ai_search_core(search_payload, progress_lines.append)
        for line in progress_lines:
            logs.append(f"  [search] {line}")

        items_found = result.get('items') or []
        logs.append(f"  найдено кандидатов: {len(items_found)}")

        for cand_index, candidate in enumerate(items_found, start=1):
            file_id = candidate.get('file_id')
            if not file_id:
                continue
            file_obj = File.query.get(int(file_id))
            if not file_obj:
                continue

            snippets: List[str] = []
            for snippet in candidate.get('snippets') or []:
                snippet = (snippet or '').strip()
                if len(snippet) >= min_paragraph_chars:
                    snippets.append(snippet)

            if not snippets:
                text_excerpt = _read_cached_excerpt_for_file(file_obj)
                snippets = [
                    chunk
                    for chunk in _split_text_for_dataset(
                        text_excerpt,
                        min_chars=min_paragraph_chars,
                        max_chars=max_paragraph_chars,
                        max_segments=search_opts.get('max_segments', 4),
                    )
                ]

            if not snippets:
                continue

            for sn_index, snippet in enumerate(snippets, start=1):
                if len(dataset) >= target_pairs:
                    break
                pairs = _generate_pairs_from_snippet(
                    snippet,
                    desired_pairs=pairs_per_snippet,
                    temperature=llm_temperature,
                    max_tokens=llm_max_tokens,
                    problem_question=question,
                )
                if not pairs:
                    continue
                for pair in pairs:
                    q_text = pair['question'].strip()
                    a_text = pair['answer'].strip()
                    if not q_text or not a_text:
                        continue
                    key = (q_text, a_text)
                    if key in dataset_dedup:
                        continue
                    dataset.append({
                        'input': q_text,
                        'output': a_text,
                        'source': f"agregator:{file_obj.rel_path or file_obj.filename}#s{sn_index}"
                    })
                    dataset_dedup.add(key)
                    logs.append(
                        f"  добавлена пара из файла '{file_obj.rel_path or file_obj.filename}' (сниппет {sn_index})"
                    )
                    if len(dataset) >= target_pairs:
                        break
            if len(dataset) >= target_pairs:
                break

        processed += 1
        if len(dataset) >= target_pairs:
            break

    meta = {
        'selected_problems': len(selected_items),
        'processed_problems': processed,
        'target_pairs': target_pairs,
    }
    return dataset, meta


# CORS helpers now delegate to agregator.middleware.errors
def _add_pipeline_cors_headers(response: Response) -> Response:
    return _add_pipeline_cors_headers_impl(response)


def _pipeline_cors_preflight() -> Response:
    return _pipeline_cors_preflight_impl()


def _ai_search_stream_response(data: dict) -> Response:
    """
    Stream AI search progress via NDJSON events so the frontend can update in real time.
    """
    message_queue: queue.Queue = queue.Queue()

    def push(kind: str, payload: Any = None) -> None:
        message_queue.put((kind, payload))

    def encode(payload: dict[str, Any]) -> str:
        return json.dumps(payload, ensure_ascii=False) + "\n"

    @copy_current_app_context
    def run_search() -> None:
        try:
            result = _ai_search_core(
                data,
                progress_cb=lambda line: push('progress', str(line)),
            )
            push('complete', {"ok": True, **result})
        except ValueError as exc:
            push('error', {"ok": False, "error": str(exc), "status": 400})
        except Exception as exc:
            logger.exception("AI search failed", exc_info=True)
            push('error', {"ok": False, "error": str(exc), "status": 500})
        finally:
            push('done')

    threading.Thread(target=run_search, daemon=True).start()

    @stream_with_context
    def event_stream():
        while True:
            try:
                kind, payload = message_queue.get()
            except Exception:
                break
            if kind == 'done':
                break
            if kind == 'progress':
                yield encode({"type": "progress", "line": str(payload or "")})
                continue
            if kind == 'complete':
                payload = payload or {}
                yield encode({"type": "complete", "data": payload})
                continue
            if kind == 'error':
                payload = payload or {}
                yield encode({"type": "error", **payload})
                continue

    return Response(event_stream(), mimetype="application/x-ndjson")


@app.route('/api/training/problem-pipeline', methods=['POST', 'OPTIONS'])
def api_training_problem_pipeline():
    if request.method == 'OPTIONS':
        return _pipeline_cors_preflight()
    allowed, failure_response = _check_pipeline_access()
    if not allowed:
        return _add_pipeline_cors_headers(failure_response)
    payload = request.get_json(silent=True) or {}
    logs: List[str] = []

    try:
        dataset, meta = _build_dataset_from_problems(payload, logs)
    except ValueError as exc:
        logs.append(f"Ошибка: {exc}")
        resp = jsonify({'ok': False, 'error': str(exc), 'logs': logs})
        resp.status_code = 400
        return _add_pipeline_cors_headers(resp)
    except Exception as exc:
        app.logger.exception('problem pipeline failed')
        logs.append(f'Неожиданная ошибка: {exc}')
        resp = jsonify({'ok': False, 'error': 'Не удалось построить датасет', 'logs': logs})
        resp.status_code = 500
        return _add_pipeline_cors_headers(resp)

    if not dataset:
        logs.append('Датасет пуст — нечего отправлять на дообучение')
        resp = jsonify({'ok': False, 'error': 'Датасет пуст', 'logs': logs, 'meta': meta})
        resp.status_code = 400
        return _add_pipeline_cors_headers(resp)

    dry_run = bool(payload.get('dry_run'))
    include_dataset = bool(payload.get('include_dataset'))
    job_info: Optional[Dict[str, Any]] = None

    if not dry_run and payload.get('fine_tune'):
        try:
            job_info = _launch_fine_tune_job(dataset, payload['fine_tune'])
            logs.append(f"Запущено дообучение: задача {job_info.get('id')}")
        except Exception as exc:
            app.logger.exception('fine-tune launch failed')
            logs.append(f"Ошибка запуска дообучения: {exc}")
            resp = jsonify({
                'ok': False,
                'error': str(exc),
                'logs': logs,
                'meta': meta,
                'dataset_size': len(dataset),
            })
            resp.status_code = 502
            return _add_pipeline_cors_headers(resp)

    preview = [
        {
            'input': item['input'][:160],
            'output': item['output'][:160],
            'source': item.get('source'),
        }
        for item in dataset[:3]
    ]

    response: Dict[str, Any] = {
        'ok': True,
        'dataset_size': len(dataset),
        'meta': meta,
        'logs': logs,
        'preview': preview,
    }
    if job_info is not None:
        response['fine_tune_job'] = job_info
    if include_dataset:
        response['dataset'] = dataset
    elif dry_run:
        response['dataset_preview_full'] = dataset[: min(len(dataset), 10)]
    return _add_pipeline_cors_headers(jsonify(response))


@app.route('/api/doc-chat/documents')
def api_doc_chat_documents():
    user = _load_current_user()
    if not user:
        return jsonify({'ok': False, 'error': 'Не авторизовано'}), 401
    term = (request.args.get('q') or '').strip()
    try:
        query = _apply_file_access_filter(
            File.query.options(
                joinedload(File.collection),
                joinedload(File.rag_document),
            )
        )
        try:
            query = query.outerjoin(Collection, File.collection_id == Collection.id).filter(
                or_(Collection.searchable == True, Collection.id.is_(None))
            )
        except Exception:
            pass
        query = query.filter(or_(File.material_type.is_(None), File.material_type.notin_(['audio', 'image'])))
        if term:
            like = f"%{term}%"
            query = query.filter(or_(
                File.title.ilike(like),
                File.filename.ilike(like),
                File.author.ilike(like),
                File.keywords.ilike(like),
                File.rel_path.ilike(like),
            ))
        files = query.order_by(File.title.asc(), File.id.asc()).limit(400).all()
    except Exception as exc:
        app.logger.warning("doc-chat documents query failed: %s", exc)
        return jsonify({'ok': False, 'error': 'Не удалось получить список документов'}), 500
    cache_map: dict[int, DocChatCache] = {}
    try:
        file_ids = [f.id for f in files]
        if file_ids:
            cache_map = {
                cache.file_id: cache
                for cache in DocChatCache.query.filter(DocChatCache.file_id.in_(file_ids)).all()
            }
    except Exception:
        cache_map = {}

    items: list[dict[str, Any]] = []
    for file_obj in files:
        collection = getattr(file_obj, 'collection', None)
        doc = getattr(file_obj, 'rag_document', None)
        cache_record = cache_map.get(file_obj.id)
        items.append({
            'id': file_obj.id,
            'title': (file_obj.title or file_obj.filename or file_obj.rel_path or f"Файл {file_obj.id}").strip(),
            'author': file_obj.author,
            'year': file_obj.year,
            'collection': getattr(collection, 'name', None),
            'collection_id': getattr(collection, 'id', None),
            'material_type': file_obj.material_type,
            'rel_path': file_obj.rel_path,
            'has_rag': bool(doc and getattr(doc, 'is_ready_for_rag', False)),
            'mtime': file_obj.mtime,
            'doc_chat_ready': _doc_chat_cache_is_valid(file_obj, cache_record),
        })
    return jsonify({'ok': True, 'items': items})


@app.route('/api/doc-chat/collections/<int:collection_id>/prepare', methods=['POST'])
@require_admin
def api_doc_chat_prepare_collection(collection_id: int):
    user = _load_current_user()
    if not user:
        return jsonify({'ok': False, 'error': 'Не авторизовано'}), 401
    collection = Collection.query.get(collection_id)
    if not collection:
        return jsonify({'ok': False, 'error': 'Коллекция не найдена'}), 404

    active_tasks = TaskRecord.query.filter(TaskRecord.status.in_(('queued', 'running'))).count()
    if active_tasks > 0:
        return jsonify({'ok': False, 'error': 'Есть активные задачи, дождитесь их завершения перед запуском обработки.'}), 409

    queue_stats = get_task_queue().stats()
    if (queue_stats.get('queued') or 0) > 0:
        return jsonify({'ok': False, 'error': 'Очередь фоновых задач занята, попробуйте чуть позже.'}), 409

    files = File.query.filter(File.collection_id == collection_id).all()
    pending_ids = [f.id for f in files if not _doc_chat_cache_is_valid(f, getattr(f, 'doc_chat_cache', None))]
    if not pending_ids:
        return jsonify({'ok': True, 'pending': 0, 'message': 'Все документы коллекции уже подготовлены для чата.'})

    payload = {'collection_id': collection_id, 'pending': len(pending_ids)}
    task = TaskRecord(
        name='doc_chat_collection_prepare',
        status='queued',
        payload=json.dumps(payload, ensure_ascii=False),
        progress=0.0,
        user_id=getattr(user, 'id', None),
        collection_id=collection_id,
    )
    try:
        db.session.add(task)
        db.session.commit()
    except Exception as exc:
        db.session.rollback()
        return jsonify({'ok': False, 'error': f'Не удалось создать задачу: {exc}'}), 500

    allowed = _current_allowed_collections()
    allowed_ids = list(allowed) if allowed is not None else None
    user_id = getattr(user, 'id', None)

    @copy_current_app_context
    def _runner():
        _doc_chat_collection_job(task.id, pending_ids, user_id, allowed_ids)

    get_task_queue().submit(
        _runner,
        description=f"doc-chat-collection-{collection_id}",
        owner_id=getattr(user, 'id', None),
    )
    return jsonify({'ok': True, 'task_id': task.id, 'pending': len(pending_ids)})


@app.route('/api/doc-chat/prepare', methods=['POST'])
def api_doc_chat_prepare():
    user = _load_current_user()
    if not user:
        return jsonify({'ok': False, 'error': 'Не авторизовано'}), 401
    try:
        app.logger.info("[doc-chat] prepare endpoint hit by user=%s (v2)", getattr(user, 'id', None))
    except Exception:
        pass
    data = request.get_json(silent=True) or {}
    try:
        file_id = int(data.get('file_id'))
    except Exception:
        return jsonify({'ok': False, 'error': 'Некорректный идентификатор файла'}), 400
    force_prepare = bool(data.get('force'))
    file_obj = File.query.get(file_id)
    if not file_obj:
        return jsonify({'ok': False, 'error': 'Документ не найден'}), 404
    app.logger.info("[doc-chat] file resolved id=%s title=%s", file_obj.id, (file_obj.title or file_obj.filename))
    allowed = _current_allowed_collections()
    if allowed is not None and file_obj.collection_id not in allowed:
        return jsonify({'ok': False, 'error': 'Нет доступа к документу'}), 403

    existing_id: Optional[str] = None
    with DOC_CHAT_LOCK:
        for sid, sess in DOC_CHAT_SESSIONS.items():
            if sess.get('user_id') == user.id and sess.get('file_id') == file_obj.id and sess.get('status') in {'processing', 'ready'}:
                existing_id = sid
                break
    now_ts = time.time()
    if existing_id:
        try:
            existing = _doc_chat_internal_session(existing_id)
        except Exception:
            existing = None
        if existing:
            status = existing.get('status')
            percent = existing.get('percent')
            updated_at = float(existing.get('updated_at') or existing.get('created_at') or 0.0)
            age = now_ts - updated_at if updated_at else None
            app.logger.info(
                "[doc-chat:%s] reuse candidate status=%s percent=%s age=%s",
                existing_id[:8],
                status,
                percent,
                f"{age:.1f}s" if age is not None else None,
            )
            is_ready = status == 'ready'
            is_active = status == 'processing' and age is not None and age < 45
            if not force_prepare and (is_ready or is_active):
                snapshot = _doc_chat_public_session(existing_id)
                if snapshot:
                    return jsonify({'ok': True, 'session': snapshot})
            should_reset = force_prepare or (age is not None and age > 45)
            if should_reset:
                if not force_prepare and age is not None and age > 45:
                    app.logger.warning(
                        "[doc-chat:%s] stale session detected (status=%s age=%.1fs) — переинициализация",
                        existing_id[:8],
                        status,
                        age,
                    )
                cache_path = existing.get('cache_path')
                with DOC_CHAT_LOCK:
                    DOC_CHAT_SESSIONS.pop(existing_id, None)
                if cache_path:
                    try:
                        shutil.rmtree(Path(cache_path), ignore_errors=True)
                    except Exception:
                        pass
        else:
            with DOC_CHAT_LOCK:
                DOC_CHAT_SESSIONS.pop(existing_id, None)

    cached_payload = None
    if not force_prepare:
        try:
            cached_payload = _doc_chat_load_cache_payload(file_obj)
        except Exception as exc:
            app.logger.warning("[doc-chat] cache lookup failed for file %s: %s", file_obj.id, exc)

    if cached_payload and not force_prepare:
        try:
            session_payload = _doc_chat_create_session(user, file_obj)
        except Exception as exc:
            app.logger.exception("[doc-chat] failed to create cached session: %s", exc)
            return jsonify({'ok': False, 'error': 'Не удалось создать сессию документа'}), 500
        _doc_chat_store_data(session_payload['id'], cached_payload)
        _doc_chat_progress(session_payload['id'], "Документ загружен из кэша", percent=90)
        _doc_chat_set_status(session_payload['id'], 'ready', percent=100)
        snapshot = _doc_chat_public_session(session_payload['id'])
        if snapshot is None:
            return jsonify({'ok': False, 'error': 'Не удалось создать сессию документа'}), 500
        return jsonify({'ok': True, 'session': snapshot, 'cached': True})
    try:
        session_payload = _doc_chat_create_session(user, file_obj)
    except Exception as exc:
        app.logger.exception("[doc-chat] failed to create session: %s", exc)
        return jsonify({'ok': False, 'error': 'Не удалось создать сессию документа'}), 500
    app.logger.info("[doc-chat:%s] session payload created", session_payload['id'][:8])
    allowed_ids = list(allowed) if allowed is not None else None
    app.logger.info(
        "[doc-chat:%s] session created status=%s percent=%s",
        session_payload['id'][:8],
        session_payload.get('status'),
        session_payload.get('percent'),
    )
    fallback_thread = None
    fallback_reason: Optional[str] = None
    try:
        queue = get_task_queue()
        stats_before = queue.stats()
        app.logger.info(
            "[doc-chat:%s] submit task: before workers=%s started=%s queued=%s",
            session_payload['id'][:8],
            stats_before.get('workers'),
            stats_before.get('started'),
            stats_before.get('queued'),
        )
        task_id = queue.submit(
            _doc_chat_prepare_worker,
            session_payload['id'],
            file_obj.id,
            user.id,
            allowed_ids,
            description=f"doc-chat-prepare-{file_obj.id}",
        )
        stats_after = queue.stats()
        app.logger.info(
            "[doc-chat:%s] очередь подготовки поставлена (task=%s, workers=%s, started=%s, queued=%s)",
            session_payload['id'][:8],
            task_id,
            stats_after.get('workers'),
            stats_after.get('started'),
            stats_after.get('queued'),
        )
        if not stats_after.get('started') or not stats_after.get('workers'):
            fallback_reason = f"queue_inactive workers={stats_after.get('workers')} started={stats_after.get('started')}"
        def _queue_watchdog() -> None:
            app.logger.info("[doc-chat:%s] watchdog started", session_payload['id'][:8])
            try:
                time.sleep(3.0)
                session_probe = _doc_chat_internal_session(session_payload['id'])
                if not session_probe:
                    app.logger.warning("[doc-chat:%s] watchdog: session отсутствует — запускаем резерв", session_payload['id'][:8])
                    raise RuntimeError("session_missing")
                status = session_probe.get('status')
                if status in {'processing', 'ready', 'error'}:
                    app.logger.info("[doc-chat:%s] watchdog: очередь работает (status=%s)", session_payload['id'][:8], status)
                    return
                app.logger.warning("[doc-chat:%s] watchdog: очередь не стартовала (status=%s) — запускаем резерв", session_payload['id'][:8], status)
                raise RuntimeError(f"queue_idle:{status}")
            except Exception as watchdog_exc:
                try:
                    app.logger.warning(
                        "[doc-chat:%s] watchdog triggered fallback: %s",
                        session_payload['id'][:8],
                        watchdog_exc,
                    )
                    _doc_chat_prepare_worker(
                        session_payload['id'],
                        file_obj.id,
                        user.id,
                        allowed_ids,
                    )
                    app.logger.info("[doc-chat:%s] watchdog fallback completed", session_payload['id'][:8])
                except Exception as exc_inner:
                    app.logger.exception(
                        "[doc-chat:%s] watchdog fallback failed: %s",
                        session_payload['id'][:8],
                        exc_inner,
                    )

        threading.Thread(
            target=_queue_watchdog,
            name=f"doc-chat-watchdog-{session_payload['id'][:8]}",
            daemon=True,
        ).start()
    except Exception as exc:
        app.logger.exception(
            "[doc-chat:%s] не удалось поставить задачу: %s",
            session_payload['id'][:8],
            exc,
        )
        fallback_reason = f"queue_error:{exc}"

    if fallback_reason:
        try:
            app.logger.warning(
                "[doc-chat:%s] используем резервный поток подготовки (%s)",
                session_payload['id'][:8],
                fallback_reason,
            )
            fallback_thread = threading.Thread(
                target=_doc_chat_prepare_worker,
                args=(session_payload['id'], file_obj.id, user.id, allowed_ids),
                name=f"doc-chat-fallback-{session_payload['id'][:8]}",
                daemon=True,
            )
            fallback_thread.start()
            app.logger.info("[doc-chat:%s] fallback thread started", session_payload['id'][:8])
        except Exception as exc:
            app.logger.exception(
                "[doc-chat:%s] резервный запуск подготовки не удался: %s",
                session_payload['id'][:8],
                exc,
            )
            return jsonify({'ok': False, 'error': 'Не удалось запустить обработку документа'}), 500

    snapshot = _doc_chat_public_session(session_payload['id'])
    return jsonify({'ok': True, 'session': snapshot})


@app.route('/api/doc-chat/status/<session_id>')
def api_doc_chat_status(session_id: str):
    user = _load_current_user()
    if not user:
        return jsonify({'ok': False, 'error': 'Не авторизовано'}), 401
    owner_id = _doc_chat_owner_id(session_id)
    if owner_id is None:
        return jsonify({'ok': False, 'error': 'Сессия не найдена'}), 404
    if owner_id != user.id and _user_role(user) != ROLE_ADMIN:
        return jsonify({'ok': False, 'error': 'Нет доступа'}), 403
    snapshot = _doc_chat_public_session(session_id)
    if snapshot is None:
        return jsonify({'ok': False, 'error': 'Сессия не найдена'}), 404
    return jsonify({'ok': True, 'session': snapshot})


@app.route('/api/doc-chat/ask', methods=['POST'])
def api_doc_chat_ask():
    user = _load_current_user()
    if not user:
        return jsonify({'ok': False, 'error': 'Не авторизовано'}), 401
    payload = request.get_json(silent=True) or {}
    session_id = str(payload.get('session_id') or '').strip()
    question = str(payload.get('question') or '').strip()
    if not session_id:
        return jsonify({'ok': False, 'error': 'Не указана сессия'}), 400
    if not question:
        return jsonify({'ok': False, 'error': 'Вопрос не должен быть пустым'}), 400
    owner_id = _doc_chat_owner_id(session_id)
    if owner_id is None:
        return jsonify({'ok': False, 'error': 'Сессия не найдена'}), 404
    if owner_id != user.id and _user_role(user) != ROLE_ADMIN:
        return jsonify({'ok': False, 'error': 'Нет доступа'}), 403
    session_internal = _doc_chat_internal_session(session_id)
    if not session_internal:
        return jsonify({'ok': False, 'error': 'Сессия не найдена'}), 404
    if session_internal.get('status') != 'ready':
        return jsonify({'ok': False, 'error': 'Документ ещё обрабатывается'}), 409
    doc_data = session_internal.get('data') or {}
    history = session_internal.get('history') or []
    requested_mode_raw = str(payload.get('mode') or '').strip().lower()
    mode_key, mode_config = _doc_chat_resolve_mode(requested_mode_raw)
    max_chunks = max(1, int(mode_config.get('max_chunks', 4)))
    fallback_chunks = max(1, int(mode_config.get('fallback_chunks', min(2, max_chunks))))
    prefer_math_chunks = bool(mode_config.get('prefer_math_chunks'))
    require_terms = bool(mode_config.get('require_terms'))
    prefer_technical_images = bool(mode_config.get('prefer_technical_images'))
    image_limit = max(1, int(mode_config.get('image_limit', 3)))
    user_prefs = _doc_chat_get_user_preferences(user)
    tone_pref = DOC_CHAT_TONE_OPTIONS.get(user_prefs.get('tone', ''), DOC_CHAT_TONE_OPTIONS['neutral'])
    detail_pref = DOC_CHAT_DETAIL_OPTIONS.get(user_prefs.get('detail', ''), DOC_CHAT_DETAIL_OPTIONS['balanced'])
    language_pref = DOC_CHAT_LANGUAGE_OPTIONS.get(user_prefs.get('language', ''), DOC_CHAT_LANGUAGE_OPTIONS['ru'])
    history_context = _doc_chat_history_context(history, max_pairs=3, max_chars=1200)
    question_for_vector = question
    if history_context:
        question_for_vector = f"{history_context}\n\nТекущий вопрос: {question}"
    recent_chunk_ids = _doc_chat_recent_chunk_ids(history, limit=4)
    document_id = doc_data.get('document_id')
    embedding_info = doc_data.get('embedding') or {}
    if not document_id or not embedding_info:
        return jsonify({'ok': False, 'error': 'Сессия не готова к ответам'}), 409

    runtime = _rt()
    embedding_backend = (embedding_info.get('backend') or runtime.rag_embedding_backend or 'auto').strip().lower() or 'auto'
    model_name = embedding_info.get('model_name') or runtime.rag_embedding_model or 'intfloat/multilingual-e5-large'
    model_version = embedding_info.get('model_version')
    dim = int(embedding_info.get('dim') or getattr(runtime, 'rag_embedding_dim', 384) or 384)
    batch_size = max(1, int(getattr(runtime, 'rag_embedding_batch_size', 32) or 32))
    device = getattr(runtime, 'rag_embedding_device', None)
    endpoint = getattr(runtime, 'rag_embedding_endpoint', None) or runtime.lmstudio_api_base
    api_key = getattr(runtime, 'rag_embedding_api_key', None) or runtime.lmstudio_api_key

    images = doc_data.get('images') or []
    image_vector_indexes: dict[int, int] = {}
    try:
        backend = load_embedding_backend(
            embedding_backend,
            model_name=model_name,
            dim=dim,
            batch_size=batch_size,
            device=device,
            base_url=endpoint or None,
            api_key=api_key or None,
        )
        texts_to_embed = [question_for_vector]
        for idx, image in enumerate(images):
            description = str(image.get('description') or '').strip()
            keywords = image.get('keywords') or []
            kw_text = ", ".join(str(k) for k in keywords if k)
            before_ctx = str(image.get('context_before') or '').strip()
            after_ctx = str(image.get('context_after') or '').strip()
            ocr_snippet = str(image.get('ocr_text') or '')[:800].strip()
            parts: list[str] = []
            if description:
                parts.append(description)
            if kw_text:
                parts.append(f"Ключевые слова: {kw_text}")
            if before_ctx:
                parts.append(f"Предыдущий абзац: {before_ctx}")
            if after_ctx:
                parts.append(f"Следующий абзац: {after_ctx}")
            if ocr_snippet:
                parts.append(f"OCR: {ocr_snippet}")
            combined = "\n".join(parts).strip()
            if combined:
                image_vector_indexes[idx] = len(texts_to_embed)
                texts_to_embed.append(combined)
        vectors = backend.embed_many(texts_to_embed)
    except Exception as exc:
        try:
            backend.close()
        except Exception:
            pass
        app.logger.warning("doc-chat embedding failed: %s", exc)
        return jsonify({'ok': False, 'error': 'Не удалось обработать вопрос'}), 500
    try:
        backend.close()
    except Exception:
        pass
    if not vectors:
        return jsonify({'ok': False, 'error': 'Не удалось обработать вопрос'}), 500
    question_vector = vectors[0]

    retriever = VectorRetriever(
        model_name=model_name,
        model_version=model_version,
        max_candidates=200,
    )
    retrieved = retriever.search_by_vector(
        question_vector,
        top_k=max(6, max_chunks + 2),
        allowed_document_ids=[document_id],
    )

    question_terms = [term for term in re.findall(r"\w+", question_for_vector.lower()) if len(term) >= 3]
    max_score = max((float(item.score or 0.0) for item in retrieved), default=0.0)
    threshold_factor = float(mode_config.get('score_threshold_factor', 0.45))
    low_score_factor = float(mode_config.get('low_score_threshold_factor', max(threshold_factor, 0.6)))
    score_switch_point = float(mode_config.get('score_switch_point', 0.2))
    min_threshold = float(mode_config.get('min_threshold', 0.05))
    score_threshold = 0.0
    if max_score > 0.0:
        factor = threshold_factor if max_score >= score_switch_point else low_score_factor
        score_threshold = max(min_threshold, max_score * factor)

    filtered_chunks: list[tuple[Any, str, list[str], float]] = []
    deferred_terms: list[tuple[Any, str, list[str], float]] = []
    deferred_math: list[tuple[Any, str, list[str], float]] = []
    seen_snippets: set[str] = set()
    for item in retrieved:
        if len(filtered_chunks) >= max_chunks:
            break
        chunk = item.chunk
        snippet_full = (chunk.content or '').strip()
        if not snippet_full:
            continue
        snippet = snippet_full[:800]
        score = float(item.score or 0.0)
        snippet_norm = re.sub(r'\s+', ' ', snippet_full[:200]).lower()
        if snippet_norm in seen_snippets:
            continue
        has_terms = not question_terms or any(term in snippet_full.lower() for term in question_terms)
        passes_score = score >= score_threshold or not filtered_chunks
        if not has_terms and not passes_score:
            continue
        if len(snippet_full) < 60 and max_score > 0.0 and score < max_score * 0.8:
            continue
        chunk_has_math = _doc_chat_snippet_has_math(snippet_full)
        highlights = _doc_chat_extract_highlights(snippet_full, question_terms, limit=3)
        entry = (item, snippet, highlights, score)
        seen_snippets.add(snippet_norm)
        if require_terms and not has_terms:
            deferred_terms.append(entry)
            continue
        if prefer_math_chunks and not chunk_has_math:
            deferred_math.append(entry)
            continue
        filtered_chunks.append(entry)

    if len(filtered_chunks) < max_chunks:
        for pool in (deferred_terms, deferred_math):
            if len(filtered_chunks) >= max_chunks:
                break
            for entry in pool:
                filtered_chunks.append(entry)
                if len(filtered_chunks) >= max_chunks:
                    break

    if recent_chunk_ids:
        chunk_map: dict[int, Any] = {}
        for item in retrieved:
            chunk_obj = getattr(item, 'chunk', None)
            chunk_id = getattr(chunk_obj, 'id', None)
            try:
                chunk_id_int = int(chunk_id)
            except Exception:
                continue
            chunk_map[chunk_id_int] = item
        existing_ids: set[int] = set()
        for entry in filtered_chunks:
            chunk_obj = getattr(entry[0], 'chunk', None)
            chunk_id = getattr(chunk_obj, 'id', None)
            try:
                chunk_id_int = int(chunk_id)
            except Exception:
                continue
            existing_ids.add(chunk_id_int)
        for chunk_id in recent_chunk_ids:
            if chunk_id in existing_ids:
                continue
            candidate = chunk_map.get(chunk_id)
            if not candidate:
                continue
            chunk_obj = candidate.chunk
            snippet_full = (chunk_obj.content or '').strip()
            if not snippet_full:
                continue
            snippet = snippet_full[:800]
            highlights = _doc_chat_extract_highlights(snippet_full, question_terms, limit=2)
            entry = (candidate, snippet, highlights, float(candidate.score or 0.0))
            filtered_chunks.insert(0, entry)
            existing_ids.add(chunk_id)
            if len(filtered_chunks) >= max_chunks:
                break

    if not filtered_chunks:
        for item in retrieved[:fallback_chunks]:
            chunk = item.chunk
            snippet_full = (chunk.content or '').strip()
            if not snippet_full:
                continue
            snippet = snippet_full[:800]
            highlights = _doc_chat_extract_highlights(snippet_full, question_terms, limit=2)
            filtered_chunks.append((item, snippet, highlights, float(item.score or 0.0)))
            if len(filtered_chunks) >= fallback_chunks:
                break

    if len(filtered_chunks) > max_chunks:
        filtered_chunks = filtered_chunks[:max_chunks]

    context_sections: list[dict[str, Any]] = []
    text_sources: list[dict[str, Any]] = []
    for idx, (item, snippet, highlights, score) in enumerate(filtered_chunks, start=1):
        chunk = item.chunk
        section_path = chunk.section_path
        page = None
        if chunk.meta:
            try:
                meta = json.loads(chunk.meta)
                page = meta.get('page') or meta.get('page_number') or meta.get('page_index')
            except Exception:
                page = None
        label = f"Текст {idx}"
        chunk_lower = (chunk.content or '').lower()
        matched_terms = sorted({term for term in question_terms if term in chunk_lower})
        context_sections.append({
            'label': label,
            'content': snippet,
            'section_path': section_path,
            'page': page,
            'score': round(float(score), 4),
            'highlights': highlights,
        })
        chunk_id_val = getattr(chunk, 'id', None)
        try:
            chunk_id_int = int(chunk_id_val)
        except Exception:
            chunk_id_int = None
        text_sources.append({
            'label': label,
            'chunk_id': chunk_id_int,
            'section_path': section_path,
            'page': page,
            'score': round(float(score), 4),
            'preview': (chunk.preview or snippet[:300]),
            'highlights': highlights,
            'matched_terms': matched_terms,
        })

    if not context_sections:
        previews = doc_data.get('preview_chunks') or []
        for idx, preview in enumerate(previews, start=1):
            snippet = (preview.get('preview') or '')[:800]
            label = f"Текст {idx}"
            context_sections.append({
                'label': label,
                'content': snippet,
                'section_path': preview.get('section_path'),
                'page': None,
                'score': 0.0,
            })
            text_sources.append({
                'label': label,
                'chunk_id': None,
                'section_path': preview.get('section_path'),
                'page': None,
                'score': 0.0,
                'preview': snippet[:300],
            })

    image_sources: list[dict[str, Any]] = []
    if images:
        scored_images: list[dict[str, Any]] = []
        for idx, image in enumerate(images):
            vec_idx = image_vector_indexes.get(idx)
            vector = vectors[vec_idx] if vec_idx is not None and vec_idx < len(vectors) else None
            score = _doc_chat_cosine(question_vector, vector) if vector else 0.0
            label = f"Изображение {idx + 1}"
            scored_images.append({
                'label': label,
                'page': image.get('page'),
                'description': image.get('description'),
                'keywords': image.get('keywords') or [],
                'url': image.get('url'),
                'context_before': image.get('context_before'),
                'context_after': image.get('context_after'),
                'ocr_preview': image.get('ocr_preview') or (str(image.get('ocr_text') or '').strip()[:400] or None),
                'score': round(float(score), 4),
            })
        scored_images.sort(key=lambda item: item['score'], reverse=True)
        technical_images = [img for img in scored_images if _doc_chat_image_is_technical(img)] if prefer_technical_images else []
        relevant_images = [img for img in scored_images if img['score'] >= 0.05]
        candidates: list[dict[str, Any]] = []
        if prefer_technical_images and technical_images:
            candidates = technical_images
        elif relevant_images:
            candidates = relevant_images
        else:
            candidates = scored_images
        image_sources = candidates[:image_limit] if candidates else []

    doc_meta = session_internal.get('file_meta') or {}
    doc_title = doc_meta.get('title') or f"Документ {doc_meta.get('id')}"
    doc_author = doc_meta.get('author')
    doc_year = doc_meta.get('year')
    info_parts = [f"Документ: «{doc_title}»"]
    if doc_author:
        info_parts.append(f"Автор: {doc_author}")
    if doc_year:
        info_parts.append(f"Год: {doc_year}")
    doc_info_line = "; ".join(info_parts)

    context_parts: list[str] = []
    for section in context_sections:
        meta_line = []
        if section.get('page'):
            meta_line.append(f"стр. {section['page']}")
        if section.get('section_path'):
            meta_line.append(section['section_path'])
        meta_suffix = f" ({', '.join(meta_line)})" if meta_line else ""
        context_parts.append(f"[{section['label']}{meta_suffix}]\n{section['content']}".strip())
        if section.get('highlights'):
            bullets = "\n".join(f"- {line}" for line in section['highlights'][:3])
            context_parts.append(f"Ключевые мысли {section['label']}:\n{bullets}")
    if image_sources:
        context_parts.append("Изображения:")
        for image in image_sources:
            meta_line = []
            if image.get('page'):
                meta_line.append(f"стр. {image['page']}")
            meta_suffix = f" ({', '.join(meta_line)})" if meta_line else ""
            keywords = image.get('keywords') or []
            desc_parts: list[str] = []
            description = str(image.get('description') or '').strip()
            if description:
                desc_parts.append(description)
            if keywords:
                desc_parts.append(f"Ключевые слова: {', '.join(keywords)}.")
            before_ctx = image.get('context_before')
            if before_ctx:
                desc_parts.append(f"Предыдущий абзац: {before_ctx}")
            after_ctx = image.get('context_after')
            if after_ctx:
                desc_parts.append(f"Следующий абзац: {after_ctx}")
            ocr_preview = str(image.get('ocr_preview') or '').strip()
            if not ocr_preview:
                ocr_preview = str(image.get('ocr_text') or '')[:400].strip()
            if ocr_preview:
                desc_parts.append(f"OCR: {ocr_preview}")
            block_text = "\n".join(desc_parts).strip() or "Описание изображения недоступно."
            context_parts.append(f"[{image['label']}{meta_suffix}]\n{block_text}")
    if not context_parts:
        context_parts.append("Контекст: подходящие фрагменты не найдены, ответ может быть неполным.")
    context_blob = "\n\n".join(context_parts)
    max_chars = _lm_max_input_chars()
    if len(context_blob) > max_chars:
        context_blob = context_blob[:max_chars]

    base_instruction = (
        "Ты виртуальный эксперт по документам. Используй приведённые фрагменты текста и описания изображений, "
        "чтобы ответить на вопрос. Указывай ссылки на использованные части форматом «Текст N» или «Изображение N». "
        "Если данных недостаточно, прямо сообщи об этом."
    )
    persona_instruction_parts = [
        tone_pref.get('instruction'),
        detail_pref.get('instruction'),
        language_pref.get('instruction'),
    ]
    persona_instruction = "\n".join(part for part in persona_instruction_parts if part)
    if persona_instruction:
        base_instruction = f"{base_instruction}\n\n{persona_instruction}"
    mode_instruction = str(mode_config.get('instruction') or '').strip()
    if mode_instruction:
        base_instruction = f"{base_instruction}\n\n{mode_instruction}"
    history_section = f"История диалога:\n{history_context}\n\n" if history_context else ""
    system_prompt = f"{base_instruction}\n\n{history_section}{doc_info_line}\n\nКонтекст:\n{context_blob}"
    messages: list[dict[str, str]] = [
        {"role": "system", "content": system_prompt},
    ]
    for turn in history[-6:]:
        role = turn.get('role')
        content = turn.get('content')
        if role in {'user', 'assistant'} and content:
            messages.append({'role': role, 'content': str(content)})
    messages.append({'role': 'user', 'content': question})

    llm_max_tokens = int(detail_pref.get('max_tokens', 900) or 900)
    try:
        answer = call_doc_chat_llm(messages, temperature=0.2, max_tokens=llm_max_tokens)
    except Exception as exc:
        app.logger.warning("doc-chat LLM failed: %s", exc)
        return jsonify({'ok': False, 'error': 'Не удалось получить ответ'}), 500
    answer = (answer or '').strip()

    validation_info = _doc_chat_assess_answer_quality(
        answer,
        question_terms=question_terms,
        text_sources=text_sources,
        image_sources=image_sources,
    )

    _doc_chat_append_history(session_id, {'role': 'user', 'content': question, 'mode': mode_key})
    _doc_chat_append_history(session_id, {
        'role': 'assistant',
        'content': answer,
        'mode': mode_key,
        'sources': {'texts': text_sources, 'images': image_sources},
        'validation': validation_info,
    })

    try:
        detail = json.dumps({
            'session_id': session_id,
            'file_id': session_internal.get('file_id'),
            'question': question[:200],
            'mode': mode_key,
        }, ensure_ascii=False)
        _log_user_action(user, 'doc_chat_ask', 'file', session_internal.get('file_id'), detail=detail[:2000])
    except Exception:
        pass

    snapshot = _doc_chat_public_session(session_id)
    return jsonify({
        'ok': True,
        'answer': answer,
        'sources': {
            'texts': text_sources,
            'images': image_sources,
        },
        'validation': validation_info,
        'session': snapshot,
        'mode': mode_key,
    })


@app.route('/api/doc-chat/clear', methods=['POST'])
def api_doc_chat_clear():
    user = _load_current_user()
    if not user:
        return jsonify({'ok': False, 'error': 'Не авторизовано'}), 401
    payload = request.get_json(silent=True) or {}
    session_id = str(payload.get('session_id') or '').strip()
    if not session_id:
        return jsonify({'ok': False, 'error': 'Не указана сессия'}), 400
    owner_id = _doc_chat_owner_id(session_id)
    if owner_id is None:
        return jsonify({'ok': False, 'error': 'Сессия не найдена'}), 404
    if owner_id != user.id and _user_role(user) != ROLE_ADMIN:
        return jsonify({'ok': False, 'error': 'Нет доступа'}), 403
    now = _doc_chat_now()
    with DOC_CHAT_LOCK:
        session = DOC_CHAT_SESSIONS.get(session_id)
        if not session:
            return jsonify({'ok': False, 'error': 'Сессия не найдена'}), 404
        session['history'] = []
        session['updated_at'] = now
    snapshot = _doc_chat_public_session(session_id)
    return jsonify({'ok': True, 'session': snapshot})


@app.route('/api/doc-chat/preferences', methods=['GET', 'POST'])
def api_doc_chat_preferences():
    user = _load_current_user()
    if not user:
        return jsonify({'ok': False, 'error': 'Не авторизовано'}), 401
    if request.method == 'GET':
        prefs = _doc_chat_get_user_preferences(user)
        return jsonify({
            'ok': True,
            'preferences': prefs,
            'options': _doc_chat_pref_options_payload(),
        })
    data = request.get_json(silent=True) or {}
    updates: dict[str, Any] = {}
    for key in DOC_CHAT_PREFERENCE_OPTIONS:
        if key in data:
            updates[key] = data[key]
    if not updates:
        return jsonify({'ok': False, 'error': 'Нет параметров для обновления'}), 400
    try:
        prefs = _doc_chat_set_user_preferences(user, updates)
    except Exception as exc:
        app.logger.warning("[doc-chat] failed to update preferences for user %s: %s", getattr(user, 'id', None), exc)
        return jsonify({'ok': False, 'error': 'Не удалось сохранить настройки'}), 500
    return jsonify({
        'ok': True,
        'preferences': prefs,
        'options': _doc_chat_pref_options_payload(),
    })


@app.route('/api/ai-search', methods=['POST'])
def api_ai_search():
    data = request.get_json(silent=True) or {}
    user = _load_current_user()
    query_preview = str(data.get('query') or '').strip()[:200]
    detail_obj = {
        'query_preview': query_preview,
        'top_k': data.get('top_k'),
        'deep_search': data.get('deep_search'),
        'sources': data.get('sources'),
    }
    try:
        detail_payload = json.dumps(detail_obj, ensure_ascii=False)
    except Exception:
        detail_payload = str(detail_obj)
    try:
        _log_user_action(user, 'ai_search', 'search', None, detail=detail_payload[:2000])
    except Exception:
        pass
    try:
        app.logger.info(
            "[user-action] user=%s action=ai_search query_preview=%s top_k=%s deep_search=%s",
            getattr(user, 'username', None) or 'anonymous',
            query_preview,
            detail_obj.get('top_k'),
            detail_obj.get('deep_search')
        )
    except Exception:
        pass
    stream_flag = str(request.args.get('stream') or '').strip().lower()
    accept_header = (request.headers.get('Accept') or '').lower()
    if stream_flag in {'1', 'true', 'yes'} or 'application/x-ndjson' in accept_header:
        return _ai_search_stream_response(data)
    try:
        result = _ai_search_core(data)
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400
    except Exception as exc:
        logger.exception("AI search failed", exc_info=True)
        return jsonify({"ok": False, "error": str(exc)}), 500
    return jsonify({"ok": True, **result})
@app.route('/api/ai-search/feedback', methods=['POST'])
def api_ai_search_feedback():
    user = _load_current_user()
    if not user:
        return jsonify({'ok': False, 'error': 'Не авторизовано'}), 401
    data = request.get_json(silent=True) or {}
    query_hash = str(data.get('query_hash') or '').strip()
    action = str(data.get('action') or '').strip().lower()
    if not query_hash or action not in {'click', 'relevant', 'irrelevant', 'ignored'}:
        return jsonify({'ok': False, 'error': 'Некорректные параметры'}), 400
    keyword = (data.get('keyword') or '').strip() or None
    detail = data.get('detail')
    file_id = data.get('file_id')
    try:
        file_id_int = int(file_id) if file_id is not None else None
    except Exception:
        file_id_int = None
    score = data.get('score')
    try:
        score_val = float(score) if score is not None else None
    except Exception:
        score_val = None
    entry = AiSearchKeywordFeedback(
        user_id=user.id,
        file_id=file_id_int,
        query_hash=query_hash,
        keyword=keyword,
        action=action,
        score=score_val,
        detail=json.dumps(detail, ensure_ascii=False) if isinstance(detail, (dict, list)) else (detail or None),
    )
    try:
        db.session.add(entry)
        db.session.commit()
    except Exception as exc:
        db.session.rollback()
        app.logger.warning(f"[ai-feedback] failed to store feedback: {exc}")
        return jsonify({'ok': False, 'error': 'Не удалось сохранить отклик'}), 500
    train_task_id = None
    if action in {'click', 'relevant', 'irrelevant'}:
        try:
            train_task_id, _ = _maybe_enqueue_feedback_training_online(
                trigger='online',
                submitted_by=getattr(user, 'id', None),
            )
        except Exception:
            train_task_id = None
    return jsonify({'ok': True, 'train_task_id': train_task_id})


@app.route('/api/search/feedback', methods=['POST'])
def api_search_feedback():
    user = _load_current_user()
    if not user:
        return jsonify({'ok': False, 'error': 'Не авторизовано'}), 401
    data = request.get_json(silent=True) or {}
    action = str(data.get('action') or '').strip().lower()
    if action not in {'click', 'relevant', 'irrelevant', 'ignored'}:
        return jsonify({'ok': False, 'error': 'Некорректные параметры'}), 400
    query_hash = str(data.get('query_hash') or '').strip()
    query_text = str(data.get('query') or '').strip()
    if not query_hash:
        query_hash = _sha256(query_text or '')
    if not query_hash:
        return jsonify({'ok': False, 'error': 'Не указан query_hash/query'}), 400
    try:
        file_id = int(data.get('file_id'))
    except Exception:
        return jsonify({'ok': False, 'error': 'Не указан file_id'}), 400
    if not File.query.filter_by(id=file_id).first():
        return jsonify({'ok': False, 'error': 'Файл не найден'}), 404
    detail_payload = {
        'source': 'catalogue_search',
        'query': query_text[:500] if query_text else None,
        'meta': data.get('meta') if isinstance(data.get('meta'), (dict, list)) else None,
    }
    entry = AiSearchKeywordFeedback(
        user_id=user.id,
        file_id=file_id,
        query_hash=query_hash,
        keyword=None,
        action=action,
        score=None,
        detail=json.dumps(detail_payload, ensure_ascii=False),
    )
    try:
        db.session.add(entry)
        db.session.commit()
    except Exception as exc:
        db.session.rollback()
        app.logger.warning("[search-feedback] store failed: %s", exc)
        return jsonify({'ok': False, 'error': 'Не удалось сохранить отклик'}), 500
    train_task_id = None
    if action in {'click', 'relevant', 'irrelevant'}:
        try:
            train_task_id, _ = _maybe_enqueue_feedback_training_online(
                trigger='online-search',
                submitted_by=getattr(user, 'id', None),
            )
        except Exception:
            train_task_id = None
    return jsonify({'ok': True, 'train_task_id': train_task_id})


def _initialize_background_jobs():
    _start_cache_cleanup_scheduler()
    _start_feedback_training_scheduler()
    _start_lmstudio_idle_unload_scheduler()
    try:
        _get_osint_service().ensure_schedule_worker()
    except Exception:
        app.logger.exception("Failed to start OSINT schedule worker")


if hasattr(app, "before_serving"):
    app.before_serving(_initialize_background_jobs)


@app.before_request
def _ensure_background_jobs_started():
    if (not _CLEANUP_THREAD_STARTED) or (not _FEEDBACK_THREAD_STARTED) or (not LMSTUDIO_IDLE_UNLOAD_THREAD_STARTED):
        _initialize_background_jobs()


if __name__ == "__main__":
    setup_app()
    with app.app_context():
        ensure_collections_schema()
        ensure_llm_schema()
        ensure_default_admin()
    port = int(os.environ.get("AGREGATOR_PORT") or os.environ.get("PORT", 5050))
    app.run(host="0.0.0.0", port=port, debug=False, use_reloader=False, threaded=True)
